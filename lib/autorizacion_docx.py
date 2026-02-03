"""
Módulo para generar documentos de autorización en DOCX.
Rellena la plantilla de autorización con datos del expediente.
"""

import os
import logging
from datetime import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    # Intentar con win32com como alternativa (solo Windows)
    try:
        import win32com.client
        WIN32COM_AVAILABLE = True
    except ImportError:
        WIN32COM_AVAILABLE = False

logger = logging.getLogger(__name__)


def generar_autorizacion_docx(
    plantilla_path,
    ruta_destino,
    codigo_rma,
    cliente=None,
    persona_de_contacto=None,
    email_de_contacto=None,
    fecha_emision=None,
    motivo=None,
    observaciones=None,
    fecha_autorizacion=None,
    usar_cuno=False,
    cuno_path=None,
    ruta_firma=None
):
    """
    Genera un documento de autorización en formato DOCX usando python-docx.
    
    Args:
        plantilla_path: Ruta a la plantilla DOCX
        ruta_destino: Ruta donde guardar el documento generado
        codigo_rma: Código del RMA
        cliente: Nombre del cliente
        persona_de_contacto: Persona de contacto
        email_de_contacto: Email de contacto
        fecha_emision: Fecha de emisión del RMA
        motivo: Motivo del RMA
        observaciones: Observaciones adicionales
        fecha_autorizacion: Fecha de autorización
        usar_cuno: Si se debe incluir el cuño
        cuno_path: Ruta al archivo del cuño
        ruta_firma: Ruta al archivo de la firma
    
    Returns:
        bool: True si se generó correctamente, False en caso contrario
    """
    try:
        logger.info(f"Generando documento de autorización para {codigo_rma}")
        
        # Validar que existe la plantilla
        if not os.path.exists(plantilla_path):
            logger.error(f"No se encuentra la plantilla: {plantilla_path}")
            return False
        
        # Si ruta_destino termina en .pdf, cambiarla temporalmente a .docx
        # para el procesamiento DOCX → PDF
        ruta_pdf_final = None
        if ruta_destino.lower().endswith('.pdf'):
            ruta_pdf_final = ruta_destino
            ruta_destino = ruta_destino[:-4] + '.docx'  # Cambiar .pdf por .docx
            logger.info(f"Ruta ajustada para procesamiento: {ruta_destino}")
        
        # Cargar la plantilla
        document = docx.Document(plantilla_path)
        
        # Función auxiliar para convertir fecha a dd/mm/yyyy
        def formato_fecha(fecha_str):
            if not fecha_str:
                return ''
            try:
                # Intentar parsear varios formatos
                for fmt in ['%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y']:
                    try:
                        dt = datetime.strptime(fecha_str, fmt)
                        return dt.strftime('%d/%m/%Y')
                    except ValueError:
                        continue
                # Si no se pudo parsear, devolver el original
                return str(fecha_str)
            except:
                return str(fecha_str)
        
        # Preparar el mapeo de marcadores a valores
        mapeo = {
            '[[CODIGO_RMA]]': str(codigo_rma or ''),
            '[[CLIENTE]]': str(cliente or ''),
            '[[PERSONA_CONTACTO]]': str(persona_de_contacto or ''),
            '[[EMAIL_CONTACTO]]': str(email_de_contacto or ''),
            '[[FECHA_EMISION]]': formato_fecha(fecha_emision),
            '[[MOTIVO]]': str(motivo or ''),
            '[[OBSERVACIONES]]': str(observaciones or ''),
            '[[FECHA_AUTORIZACION]]': formato_fecha(fecha_autorizacion or datetime.now().strftime('%Y-%m-%d'))
        }
        
        logger.info(f"Campos a rellenar: {len(mapeo)}")
        for clave, valor in mapeo.items():
            if valor:
                logger.info(f"  {clave} = '{valor}'")
        
        # Función auxiliar para reemplazar texto preservando formato
        def reemplazar_texto_preservando_formato(document, mapeo):
            """Reemplaza texto manteniendo el formato original"""
            
            def reemplazar_en_parrafo(paragraph):
                """Reemplaza marcadores en un párrafo, manejando marcadores divididos en múltiples runs"""
                # Construir el texto completo del párrafo
                texto_completo = ''.join(run.text for run in paragraph.runs)
                
                # Verificar si hay algún marcador en este párrafo
                tiene_marcador = any(clave in texto_completo for clave in mapeo.keys())
                
                if not tiene_marcador:
                    return
                
                # Si hay marcadores, procesar run por run
                for run in paragraph.runs:
                    texto_modificado = run.text
                    for clave, valor in mapeo.items():
                        if clave in texto_modificado:
                            texto_modificado = texto_modificado.replace(clave, valor)
                    run.text = texto_modificado
                
                # Manejo especial: si un marcador está dividido en múltiples runs
                # Reconstruir el párrafo combinando runs
                texto_completo = ''.join(run.text for run in paragraph.runs)
                for clave, valor in mapeo.items():
                    if clave in texto_completo and not any(clave in run.text for run in paragraph.runs):
                        # El marcador está dividido entre runs
                        # Limpiar todos los runs y poner el texto en el primero
                        texto_reemplazado = texto_completo.replace(clave, valor)
                        if paragraph.runs:
                            paragraph.runs[0].text = texto_reemplazado
                            # Eliminar el texto de los demás runs
                            for run in paragraph.runs[1:]:
                                run.text = ''
                        break
            
            # Reemplazar en párrafos
            for paragraph in document.paragraphs:
                reemplazar_en_parrafo(paragraph)
            
            # Reemplazar en tablas
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            reemplazar_en_parrafo(paragraph)
        
        # Aplicar reemplazos
        reemplazar_texto_preservando_formato(document, mapeo)
        logger.info("Texto reemplazado en el documento")
        
        # Añadir imágenes en cuadros de texto
        # Buscar cuadros de texto con marcadores [[CUNO]] y [[FIRMA]]
        firma_agregada = False
        cuno_agregado = False
        
        # En python-docx, los cuadros de texto son shapes dentro del documento
        # Necesitamos acceder al XML del documento directamente
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Función auxiliar para obtener dimensiones del cuadro de texto
        def obtener_dimensiones_textbox(txbxContent_element):
            """Obtiene ancho y alto del cuadro de texto en pulgadas"""
            # Navegar hacia arriba para encontrar el shape que contiene el textbox
            # txbxContent -> wps:txbx -> wps:wsp (shape)
            try:
                txbx = txbxContent_element.getparent()
                wsp = txbx.getparent()
                
                # Buscar spPr (shape properties) que contiene las dimensiones
                spPr = wsp.find(qn('wps:spPr'))
                if spPr is not None:
                    # Buscar xfrm (transform) que contiene ext (extent - dimensiones)
                    xfrm = spPr.find(qn('a:xfrm'))
                    if xfrm is not None:
                        ext = xfrm.find(qn('a:ext'))
                        if ext is not None:
                            # Las dimensiones están en EMUs (914400 EMU = 1 pulgada)
                            cx = int(ext.get('cx'))  # ancho
                            cy = int(ext.get('cy'))  # alto
                            ancho_inches = cx / 914400.0
                            alto_inches = cy / 914400.0
                            return ancho_inches, alto_inches
            except Exception as e:
                logger.debug(f"No se pudieron obtener dimensiones del textbox: {e}")
            
            # Valores por defecto si no se pueden obtener
            return 1.5, 1.5
        
        # Función auxiliar para añadir imagen a un párrafo XML
        def anadir_imagen_a_parrafo_xml(paragraph_element, ruta_imagen, ancho_box=None, alto_box=None):
            """Añade una imagen a un elemento de párrafo XML del cuadro de texto, llenando el espacio disponible"""
            try:
                # Crear un párrafo temporal en el documento para generar la imagen
                temp_paragraph = document.add_paragraph()
                temp_run = temp_paragraph.add_run()
                
                # Si tenemos dimensiones del cuadro, usar esas dimensiones para llenar el cuadro
                if ancho_box and alto_box:
                    # Leer dimensiones de la imagen para calcular la mejor escala
                    from PIL import Image as PILImage
                    with PILImage.open(ruta_imagen) as img:
                        ancho_px, alto_px = img.size
                        aspect_ratio_img = ancho_px / alto_px
                        aspect_ratio_box = ancho_box / alto_box
                        
                        # Ajustar para llenar el cuadro al máximo manteniendo proporción
                        # Escalar por la dimensión menor para que llene más
                        if aspect_ratio_img > aspect_ratio_box:
                            # Imagen más ancha: ajustar por alto para llenar más verticalmente
                            temp_picture = temp_run.add_picture(ruta_imagen, height=Inches(alto_box))
                        else:
                            # Imagen más alta: ajustar por ancho para llenar más horizontalmente
                            temp_picture = temp_run.add_picture(ruta_imagen, width=Inches(ancho_box))
                else:
                    # Sin dimensiones del cuadro, usar tamaño por defecto
                    temp_picture = temp_run.add_picture(ruta_imagen, width=Inches(1.5))
                
                # Obtener el XML del run con la imagen
                imagen_run_xml = temp_run._element
                
                # Copiar el run con la imagen al párrafo del cuadro de texto
                paragraph_element.append(imagen_run_xml)
                
                # Eliminar el párrafo temporal del documento
                temp_p_element = temp_paragraph._element
                temp_p_element.getparent().remove(temp_p_element)
                
                return True
            except Exception as e:
                logger.error(f"Error al crear imagen en cuadro de texto: {e}")
                return False
        
        # Buscar todos los elementos de cuadro de texto en el documento
        for element in document.element.body.iter():
            # Buscar elementos de tipo textbox
            if element.tag == qn('w:txbxContent'):
                # Obtener dimensiones del cuadro de texto
                ancho_box, alto_box = obtener_dimensiones_textbox(element)
                
                # Este es el contenido de un cuadro de texto
                # Buscar párrafos dentro del cuadro de texto
                for paragraph in element.iter(qn('w:p')):
                    # Obtener el texto del párrafo
                    texto_completo = ''
                    for texto_elem in paragraph.iter(qn('w:t')):
                        if texto_elem.text:
                            texto_completo += texto_elem.text
                    
                    # Verificar si contiene el marcador de cuño
                    if '[[CUNO]]' in texto_completo:
                        # Limpiar el contenido del cuadro de texto (eliminar todos los runs)
                        for run in list(paragraph.iter(qn('w:r'))):
                            paragraph.remove(run)
                        
                        # Si se debe usar el cuño, añadir la imagen adaptada al tamaño del cuadro
                        if usar_cuno and cuno_path and os.path.exists(cuno_path):
                            if anadir_imagen_a_parrafo_xml(paragraph, cuno_path, ancho_box, alto_box):
                                cuno_agregado = True
                                logger.info(f"Cuño añadido en cuadro de texto ({ancho_box:.2f}x{alto_box:.2f} pulgadas)")
                        else:
                            logger.info("Cuadro de texto del cuño dejado vacío (no se solicitó cuño)")
                    
                    # Verificar si contiene el marcador de firma
                    elif '[[FIRMA]]' in texto_completo:
                        # Limpiar el contenido del cuadro de texto (eliminar todos los runs)
                        for run in list(paragraph.iter(qn('w:r'))):
                            paragraph.remove(run)
                        
                        # Si hay firma, añadir la imagen adaptada al tamaño del cuadro
                        if ruta_firma and os.path.exists(ruta_firma):
                            if anadir_imagen_a_parrafo_xml(paragraph, ruta_firma, ancho_box, alto_box):
                                firma_agregada = True
                                logger.info(f"Firma añadida en cuadro de texto")
                        else:
                            logger.info("Cuadro de texto de firma dejado vacío (no se proporcionó firma)")
        
        # Guardar el documento DOCX temporalmente
        document.save(ruta_destino)
        logger.info(f"Documento DOCX generado: {ruta_destino}")
        
        # Si se especificó PDF como destino final, convertir
        if ruta_pdf_final:
            ruta_pdf = ruta_pdf_final
        else:
            # Si no se especificó, asumir que queremos PDF
            ruta_pdf = ruta_destino.replace('.docx', '.pdf')
        
        # Convertir a PDF
        try:
            if DOCX2PDF_AVAILABLE:
                # Método 1: docx2pdf (más simple)
                docx2pdf_convert(ruta_destino, ruta_pdf)
                logger.info(f"Convertido a PDF con docx2pdf: {ruta_pdf}")
            elif WIN32COM_AVAILABLE:
                # Método 2: win32com (solo Windows, requiere Word instalado)
                word = win32com.client.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(ruta_destino))
                doc.SaveAs(os.path.abspath(ruta_pdf), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close()
                word.Quit()
                logger.info(f"Convertido a PDF con win32com: {ruta_pdf}")
            else:
                logger.warning("No hay herramienta de conversión a PDF disponible. Se guardará como DOCX.")
                # No hacer nada, mantener el DOCX
                return True
            
            # Eliminar el DOCX temporal
            try:
                os.remove(ruta_destino)
                logger.info(f"DOCX temporal eliminado")
            except Exception as e:
                logger.warning(f"No se pudo eliminar DOCX temporal: {e}")
            
            logger.info(f"Documento de autorización PDF generado: {ruta_pdf}")
            return True
            
        except Exception as e:
            logger.error(f"Error al convertir a PDF: {e}", exc_info=True)
            # Si falla la conversión, mantener el DOCX
            logger.warning("Se mantendrá el archivo DOCX")
            return True
        
    except Exception as e:
        logger.error(f"Error al generar documento de autorización: {e}", exc_info=True)
        return False

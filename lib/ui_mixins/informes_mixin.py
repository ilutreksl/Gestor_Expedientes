"""Mixin extraido automaticamente de VentanaPrincipal (app.py).

Estas clases NO son instanciables por si solas: solo aportan metodos que se
combinan con VentanaPrincipal via herencia multiple. Dependen de atributos de
instancia (self.conn, self.username, self.tree_rmas, etc.) inicializados en
VentanaPrincipal.__init__.
"""
from lib.app_core import *  # noqa: F401,F403 - helpers/constantes/imports compartidos con app.py
from lib.app_core import _get_cached_query, invalidate_cache  # nombres "privados" que el wildcard import no trae

class InformesMixin:
    @staticmethod
    def _reemplazar_marcador_en_parrafo(paragraph, clave, valor_str):
        """Sustituye todas las apariciones de `clave` dentro de un párrafo,
        aunque Word haya partido el marcador en varios runs (p.ej. por
        autocorrección o revisión ortográfica: "[[COD" + "IGO_RMA]]"). El
        formato del run donde empieza cada aparición se conserva; los runs
        que quedan completamente "consumidos" por el marcador se vacían.
        """
        while True:
            runs = paragraph.runs
            texto_completo = ''.join(r.text for r in runs)
            idx = texto_completo.find(clave)
            if idx == -1:
                return
            fin = idx + len(clave)

            pos = 0
            primer_run_usado = False
            for run in runs:
                run_ini, run_fin = pos, pos + len(run.text)
                pos = run_fin
                if run_fin <= idx or run_ini >= fin:
                    continue  # este run no se solapa con el marcador
                texto_antes   = run.text[:max(0, idx - run_ini)]
                texto_despues = run.text[max(0, fin - run_ini):]
                if not primer_run_usado:
                    run.text = texto_antes + valor_str + texto_despues
                    primer_run_usado = True
                else:
                    run.text = texto_antes + texto_despues

    def _reemplazar_marcadores_preservando_formato(self, document, mapeo):
        """Reemplaza marcadores [[CLAVE]] manteniendo el formato original de la
        plantilla (fuente, tamaño, negrita, etc.), incluso si un marcador está
        repartido en varios runs dentro del mismo párrafo.

        Sustituye a la asignación de `paragraph.text`, que en python-docx
        borra todos los runs del párrafo y los reemplaza por uno solo con el
        formato por defecto, perdiendo el de la plantilla.
        """
        def _procesar_parrafo(paragraph):
            for clave, valor in mapeo.items():
                valor_str = str(valor) if valor is not None else ""
                if clave in paragraph.text:
                    self._reemplazar_marcador_en_parrafo(paragraph, clave, valor_str)

        for paragraph in document.paragraphs:
            _procesar_parrafo(paragraph)

        # También procesar tablas si las hay
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _procesar_parrafo(paragraph)

    def generar_informe_dinamico(self):
        """
        Genera un informe dinámico usando python-docx, lo guarda en Backblaze B2 
        y lo registra en la base de datos.
        """
        # 1. Validaciones y Obtención de Datos
        if not self.current_rma_id:
            messagebox.showerror("Error", "Debe cargar un RMA guardado para generar el informe.")
            return

        # Asumimos que self.datos_rma_maestro contiene los datos del RMA cargado
        # Esta variable debe llenarse cuando llamas a self.cargar_datos_rma(rma_id)
        datos = self.datos_rma_maestro 
        
        # 🚨 Verificación: Asegúrate de que los datos clave existen
        codigo_rma = datos.get('codigo_rma')
        nombre_cliente = datos.get('cliente')
        
        if not codigo_rma or not nombre_cliente:
             messagebox.showerror("Error", "Los datos del RMA no están cargados. Intente recargar el expediente.")
             return
        
        # 1.5. Obtener datos de artículos desde rma_detalles
        conn, cursor = self.master.conectar_db()
        if not conn:
            messagebox.showerror("Error", "No se pudo conectar a la base de datos.")
            return
        
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT referencia_articulo, cantidad_entregada 
                FROM rma_detalles 
                WHERE rma_id = ?
                ORDER BY id
            """, (self.current_rma_id,))
            articulos_data = cursor.fetchall()
            conn.close()
        except Exception as e:
            conn.close()
            messagebox.showerror("Error", f"Error al obtener datos de artículos: {e}")
            return

        # 2. Rutas - Plantilla sigue siendo local
        plantilla_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas", "Plantilla_RMA.docx")
        
        # Nombre del archivo final: Ej. RMA2024-001_Informe_20240920.docx
        fecha_str = datetime.datetime.now().strftime("%Y%m%d")
        nombre_archivo_final = f"{codigo_rma}_Informe_{fecha_str}.docx"

        try:
            # 3. Cargar la plantilla y preparar datos para mapeo
            document = docx.Document(plantilla_path)
            
            # 3.1. Preparar lista de artículos para los marcadores en formato estructurado
            if articulos_data:
                # Crear formato de lista estructurada para artículos
                lista_articulos = []
                for i, (ref_articulo, cantidad) in enumerate(articulos_data, 1):
                    ref_str = str(ref_articulo) if ref_articulo else 'N/A'
                    cant_str = str(cantidad) if cantidad else '0'
                    lista_articulos.append(f"{i}. {ref_str} - Cantidad: {cant_str}")
                
                # Unir con saltos de línea para formato de lista
                articulos_formateados = '\n'.join(lista_articulos)
                
                # También crear versiones simples por compatibilidad
                referencias_articulos = [str(ref) if ref else 'N/A' for ref, _ in articulos_data]
                cantidades_articulos = [str(cant) if cant else '0' for _, cant in articulos_data]
            else:
                # Valores por defecto si no hay artículos
                articulos_formateados = "No se han registrado artículos para este RMA."
                referencias_articulos = ['N/A']
                cantidades_articulos = ['N/A']
            
            # 3.2. Mapeo expandido: [Marcador en Word]: [Valor a insertar]
            mapeo = {
                # Campos existentes
                '[[CODIGO_RMA]]': codigo_rma,
                '[[CLIENTE]]': nombre_cliente,
                '[[FECHA_EMISION]]': datos.get('fecha_emision', 'N/A'),
                '[[FECHA_RECEPCION]]': datos.get('fecha_recepcion', 'N/A'),
                '[[ESTADO_ACTUAL]]': datos.get('estado', 'N/A'),
                '[[USUARIO_CREADOR]]': datos.get('creado_por', self.username),
                
                # Campos nuevos de rma_maestro
                '[[NUMERO_DOC]]': datos.get('numero_documento_cliente', 'N/A'),
                '[[MOTIVO]]': datos.get('motivo', 'N/A'),
                '[[NUMERO_ALBARAN]]': datos.get('numero_albaran', 'N/A'),
                
                # Campos de artículos mejorados
                '[[LISTA_ARTICULOS]]': articulos_formateados,  # Lista estructurada completa
                '[[REF_ARTICULO]]': ', '.join(referencias_articulos[:3]),  # Mantener compatibilidad
                '[[CANTIDAD]]': ', '.join(cantidades_articulos[:3]),  # Mantener compatibilidad
                
                # Campos adicionales para flexibilidad
                '[[TOTAL_ARTICULOS]]': str(len(articulos_data)) if articulos_data else '0'
            }
            
            # 4. Reemplazar marcadores preservando el formato de la plantilla
            reemplazar_texto_preservando_formato = self._reemplazar_marcadores_preservando_formato

            # 4.2. Función auxiliar para insertar Obs_Tecnica enriquecida (texto + imágenes)
            def insertar_obs_tecnica_enriquecida(document, marcador='[[OBS_TECNICA]]'):
                """
                Localiza el párrafo que contiene el marcador [[OBS_TECNICA]], lo elimina
                e inserta en su lugar el contenido del editor enriquecido (texto con formato
                e imágenes) leyendo directamente del widget self.entry_Obs_Tecnica.
                Si el editor no existe o está vacío, elimina el marcador sin dejar rastro.
                """
                import json as _json
                import base64 as _base64
                import io as _io
                from docx.shared import Pt, RGBColor
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                from PIL import Image as _PILImage
                from lib.rich_text_editor import DEFAULT_FAMILY as _RTE_DEFAULT_FAMILY, DEFAULT_SIZE as _RTE_DEFAULT_SIZE

                # Obtener JSON del editor
                obs_json = ""
                if hasattr(self, 'entry_Obs_Tecnica') and hasattr(self.entry_Obs_Tecnica, 'get_content'):
                    obs_json = self.entry_Obs_Tecnica.get_content() or ""
                import json as _jd
                try:
                    for _s in _jd.loads(obs_json).get("segments",[]):
                        print(f"[DBG] type={_s.get('type')} bold={_s.get('bold')} size={_s.get('size')} family={_s.get('family')} color={_s.get('color')} nombre={_s.get('nombre')} content={str(_s.get('content',''))[:40]}")
                except: pass

                # Localizar el párrafo con el marcador
                parrafo_marcador = None
                idx_marcador = None
                for i, p in enumerate(document.paragraphs):
                    if marcador in p.text:
                        parrafo_marcador = p
                        idx_marcador = i
                        break

                if parrafo_marcador is None:
                    return  # Marcador no encontrado en la plantilla, no hacer nada

                # Si no hay contenido en el editor, eliminar el párrafo marcador y salir
                if not obs_json:
                    parrafo_marcador._element.getparent().remove(parrafo_marcador._element)
                    return

                # Intentar parsear el JSON
                try:
                    data = _json.loads(obs_json)
                    segmentos = data.get("segments", [])
                except Exception:
                    # Si falla el parse, insertar como texto plano
                    segmentos = [{"type": "text", "content": obs_json,
                                  "bold": False, "italic": False,
                                  "underline": False, "size": 11, "color": None}]

                if not segmentos:
                    parrafo_marcador._element.getparent().remove(parrafo_marcador._element)
                    return

                # Referencia al elemento padre XML y posición del marcador
                parent_elem = parrafo_marcador._element.getparent()
                pos_insercion = list(parent_elem).index(parrafo_marcador._element)

                # Insertar los segmentos en orden inverso para mantener la posición correcta
                # (cada nuevo elemento se inserta en pos_insercion, empujando los anteriores)
                elementos_nuevos = []

                # Agrupar segmentos consecutivos de texto en un mismo párrafo,
                # cada imagen va en su propio párrafo
                grupos = []       # lista de listas: cada sublista es un párrafo
                grupo_actual = []
                for seg in segmentos:
                    if seg.get("type") in ("image", "image_ref"):
                        if grupo_actual:
                            grupos.append(("texto", grupo_actual))
                            grupo_actual = []
                        grupos.append(("imagen", seg))
                    else:
                        # Dividir por saltos de línea dentro del mismo segmento de texto
                        contenido = seg.get("content", "")
                        lineas = contenido.split("\n")
                        for j, linea in enumerate(lineas):
                            seg_linea = dict(seg)
                            seg_linea["content"] = linea
                            grupo_actual.append(seg_linea)
                            if j < len(lineas) - 1:
                                # Salto de línea → nuevo párrafo
                                grupos.append(("texto", grupo_actual))
                                grupo_actual = []
                if grupo_actual:
                    grupos.append(("texto", grupo_actual))

                def _nuevo_parrafo_docx():
                    """Crea un elemento <w:p> vacío."""
                    return OxmlElement('w:p')

                from docx.text.paragraph import Paragraph as _Paragraph
                from docx.oxml.ns import qn as _qn
                from docx.oxml import OxmlElement as _OE

                def _hex_to_rgb_str(hex_color):
                    """Convierte #RRGGBB a 'RRGGBB' en mayúsculas."""
                    return hex_color.lstrip('#').upper()

                def _make_parrafo_xml(grupo_segs):
                    """
                    Construye un elemento <w:p> completo en XML puro para un grupo
                    de segmentos de texto. Esto garantiza que el formato se aplica
                    correctamente aunque el párrafo se inserte manualmente en el body.
                    """
                    p = _OE('w:p')

                    # ── Propiedades de párrafo (pPr) ──────────────────────────
                    seg0 = grupo_segs[0] if grupo_segs else {}
                    align  = seg0.get("align")
                    indent_px = seg0.get("indent", 0)

                    if align or indent_px:
                        pPr = _OE('w:pPr')
                        if align:
                            jc = _OE('w:jc')
                            jc.set(_qn('w:val'), align)
                            pPr.append(jc)
                        if indent_px:
                            # 1 pt = 20 twips; 1 px @ 96dpi ≈ 0.75 pt → 15 twips
                            twips = str(int(indent_px * 15))
                            ind = _OE('w:ind')
                            ind.set(_qn('w:left'), twips)
                            pPr.append(ind)
                        p.append(pPr)

                    # ── Runs ──────────────────────────────────────────────────
                    for seg in grupo_segs:
                        texto = seg.get("content", "")
                        r = _OE('w:r')

                        # rPr — propiedades del run
                        rPr = _OE('w:rPr')
                        tiene_formato = False

                        # Familia de fuente — siempre explícita (si el usuario no la
                        # cambió, se usa la misma por defecto que muestra el editor).
                        # Si no se escribe, Word aplica la fuente por defecto de la
                        # plantilla/estilo, que puede no coincidir con el editor.
                        family = seg.get("family") or _RTE_DEFAULT_FAMILY
                        rFonts = _OE('w:rFonts')
                        rFonts.set(_qn('w:ascii'),    family)
                        rFonts.set(_qn('w:hAnsi'),    family)
                        rFonts.set(_qn('w:eastAsia'), family)
                        rPr.append(rFonts)
                        tiene_formato = True

                        # Negrita
                        if seg.get("bold"):
                            rPr.append(_OE('w:b'))
                            rPr.append(_OE('w:bCs'))
                            tiene_formato = True

                        # Cursiva
                        if seg.get("italic"):
                            rPr.append(_OE('w:i'))
                            rPr.append(_OE('w:iCs'))
                            tiene_formato = True

                        # Subrayado
                        if seg.get("underline"):
                            u = _OE('w:u')
                            u.set(_qn('w:val'), 'single')
                            rPr.append(u)
                            tiene_formato = True

                        # Tachado
                        if seg.get("strikethrough"):
                            rPr.append(_OE('w:strike'))
                            tiene_formato = True

                        # Tamaño (en half-points: pt * 2) — siempre explícito, igual
                        # que con la familia, para que coincida con el tamaño por
                        # defecto que se ve en el editor en lugar de heredar el de
                        # la plantilla.
                        size = seg.get("size") or _RTE_DEFAULT_SIZE
                        sz = _OE('w:sz')
                        sz.set(_qn('w:val'), str(int(size * 2)))
                        szCs = _OE('w:szCs')
                        szCs.set(_qn('w:val'), str(int(size * 2)))
                        rPr.append(sz)
                        rPr.append(szCs)
                        tiene_formato = True

                        # Color de fuente
                        color = seg.get("color")
                        if color and color.startswith("#") and len(color) == 7:
                            clr = _OE('w:color')
                            clr.set(_qn('w:val'), _hex_to_rgb_str(color))
                            rPr.append(clr)
                            tiene_formato = True

                        # Color de fondo / resaltado (shd)
                        bgcolor = seg.get("bgcolor")
                        if bgcolor and bgcolor.startswith("#") and len(bgcolor) == 7:
                            shd = _OE('w:shd')
                            shd.set(_qn('w:val'),   'clear')
                            shd.set(_qn('w:color'), 'auto')
                            shd.set(_qn('w:fill'),  _hex_to_rgb_str(bgcolor))
                            rPr.append(shd)
                            tiene_formato = True

                        if tiene_formato:
                            r.append(rPr)

                        # Texto — preservar espacios con xml:space="preserve"
                        t = _OE('w:t')
                        t.text = texto
                        if texto and (texto[0] == ' ' or texto[-1] == ' '):
                            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        r.append(t)
                        p.append(r)

                    return p

                for tipo, contenido in grupos:
                    if tipo == "texto":
                        p_elem = _make_parrafo_xml(contenido)
                        elementos_nuevos.append(p_elem)

                    elif tipo == "imagen":
                        try:
                            tipo_seg = contenido.get("type", "image")
                            img_bytes = None

                            if tipo_seg == "image_ref":
                                # Descargar desde B2 usando el nombre del adjunto
                                nombre_adj = contenido.get("nombre", "")
                                if nombre_adj and usar_b2():
                                    try:
                                        b2_api, bucket = get_b2_client()
                                        if b2_api and bucket:
                                            # Buscar ruta relativa del adjunto por nombre
                                            conn2, cur2 = self.master.conectar_db()
                                            cur2.execute(
                                                "SELECT ruta_relativa FROM rma_adjuntos "
                                                "WHERE rma_id = ? AND nombre_archivo = ?",
                                                (self.current_rma_id, nombre_adj))
                                            fila = cur2.fetchone()
                                            conn2.close()
                                            if fila:
                                                ruta_b2 = normalizar_ruta_b2(
                                                    f"{B2_ROOT_FOLDER}/{fila[0]}")
                                                tmp_fd, tmp_path = tempfile.mkstemp(
                                                    suffix=os.path.splitext(nombre_adj)[1])
                                                os.close(tmp_fd)
                                                bucket.download_file_by_name(ruta_b2).save_to(tmp_path)
                                                with open(tmp_path, "rb") as f:
                                                    img_bytes = f.read()
                                                os.unlink(tmp_path)
                                    except Exception as e_b2:
                                        print(f"[Informe] Error descargando imagen B2: {e_b2}")
                            elif tipo_seg == "image":
                                b64 = contenido.get("b64", "")
                                if b64:
                                    img_bytes = _base64.b64decode(b64)

                            if not img_bytes:
                                continue

                            img_stream = _io.BytesIO(img_bytes)
                            from docx.shared import Cm
                            pil_img = _PILImage.open(_io.BytesIO(img_bytes))
                            w_px, h_px = pil_img.size
                            max_cm = 14.0
                            w_cm = w_px / 96 * 2.54
                            h_cm = h_px / 96 * 2.54
                            if w_cm > max_cm:
                                factor = max_cm / w_cm
                                w_cm = max_cm
                                h_cm = h_cm * factor

                            p_elem = OxmlElement('w:p')
                            parrafo_docx = _Paragraph(p_elem, document)
                            run = parrafo_docx.add_run()
                            run.add_picture(img_stream, width=Cm(w_cm), height=Cm(h_cm))
                            elementos_nuevos.append(p_elem)
                        except Exception as img_err:
                            p_elem = OxmlElement('w:p')
                            parrafo_docx = _Paragraph(p_elem, document)
                            parrafo_docx.add_run(f"[Error al insertar imagen: {img_err}]")
                            elementos_nuevos.append(p_elem)

                # Eliminar el párrafo marcador
                parent_elem.remove(parrafo_marcador._element)

                # Insertar los nuevos párrafos en la posición correcta
                for offset, elem in enumerate(elementos_nuevos):
                    parent_elem.insert(pos_insercion + offset, elem)

            # 4.1. Aplicar reemplazos de texto plano preservando formato
            reemplazar_texto_preservando_formato(document, mapeo)

            # 4.3. Insertar contenido enriquecido de Obs_Tecnica en su marcador
            insertar_obs_tecnica_enriquecida(document)

            # 5. Guardar temporalmente para subirlo a Backblaze B2
            temp_dir = tempfile.mkdtemp(prefix="informe_rma_")
            temp_file_path = os.path.join(temp_dir, nombre_archivo_final)
            document.save(temp_file_path)
            
            # 6. Decidir dónde guardar (Dropbox o local)
            if usar_b2():
                # Subir a Backblaze B2
                exito, ruta_relativa = self._subir_archivo_b2(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'backblaze'
                ubicacion_desc = "Dropbox"
            else:
                # Guardar localmente (fallback)
                exito, ruta_relativa = self._subir_archivo_local(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'local'
                ubicacion_desc = "local"
            
            # 7. Limpiar archivo temporal
            try:
                os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            if not exito:
                messagebox.showerror("Error", f"No se pudo guardar el informe en {ubicacion_desc}.")
                return
            
            # 8. Registrar en la Base de Datos
            conn, cursor = self.master.conectar_db()
            try:
                # Verificar esquema de BD antes de insertar
                self._verificar_columna_tipo_almacenamiento(cursor)
                
                # Preparar inserción con o sin tipo_almacenamiento según el esquema
                if getattr(self, '_usar_tipo_almacenamiento', False):
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento) 
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username,
                        tipo_almacenamiento
                    ))
                else:
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username
                    ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().isoformat(),
                    f"Generado documento de Informe: {nombre_archivo_final} ({'☁️ Backblaze B2' if usar_b2() else '💾 Local'})", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                
                try:
                    # Actualizar historial si está visible
                    if hasattr(self, 'historial_tab'):
                        self.mostrar_historial(self.historial_tab)
                except AttributeError:
                    # Si la pestaña historial_tab no está definida, ignoramos.
                    pass
                
                # Mensaje personalizado según donde se guardó
                if usar_b2():
                    messagebox.showinfo("Éxito", f"✅ Informe '{nombre_archivo_final}' generado y subido a Backblaze B2 correctamente.\n\n📁 Ubicación: {ruta_relativa}")
                else:
                    messagebox.showinfo("Éxito", f"✅ Informe '{nombre_archivo_final}' generado y guardado localmente.")
                
            except Exception as db_e:
                if hasattr(conn, 'rollback'):
                    conn.rollback()
                messagebox.showerror("Error DB", f"Informe generado, pero error al registrar en DB.\nError: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el informe dinámico. Asegúrese de que la plantilla existe y python-docx está instalado.\nError: {e}")

    def generar_reposicion_devolucion(self):
        """
        Genera el documento de Reposición/Devolución usando la plantilla
        "Reposicion_RMA.docx", lo guarda en Backblaze B2 y lo registra como adjunto.
        """
        # 1. Validaciones y Obtención de Datos
        if not self.current_rma_id:
            messagebox.showerror("Error", "Debe cargar un RMA guardado para generar el documento de Reposición/Devolución.")
            return

        # Asumimos que self.datos_rma_maestro contiene los datos del RMA cargado
        datos = self.datos_rma_maestro 
        
        # Obtener datos clave para el archivo y la ruta
        codigo_rma = datos.get('codigo_rma')
        nombre_cliente = datos.get('cliente')
        
        if not codigo_rma or not nombre_cliente:
             messagebox.showerror("Error", "No se pudieron cargar los datos clave del RMA. Intente recargar el expediente.")
             return

        # 2. Rutas y Nombres de Archivo - Plantilla sigue siendo local
        nombre_plantilla = "Reposicion_RMA.docx" 
        plantilla_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas", nombre_plantilla)
        
        # Nombre del archivo final: Ej. RMA2024-001_Reposicion_20251016.docx
        fecha_str = datetime.datetime.now().strftime("%Y%m%d")
        nombre_archivo_final = f"{codigo_rma}_Reposicion_{fecha_str}.docx"

        # 3. Verificar la Plantilla
        if not os.path.exists(plantilla_path):
            messagebox.showerror("Error", f"No se encontró la plantilla requerida en:\n{plantilla_path}")
            return
            
        # 1.5. Obtener artículos del expediente (rma_detalles) para la tabla de material
        conn_art, cursor_art = self.master.conectar_db()
        articulos_reposicion = []
        if conn_art:
            try:
                cursor_art.execute("""
                    SELECT referencia_articulo, cantidad_entregada
                    FROM rma_detalles
                    WHERE rma_id = ?
                    ORDER BY id
                """, (self.current_rma_id,))
                articulos_reposicion = cursor_art.fetchall()
            except Exception as e:
                print(f"Error al obtener artículos para el documento de Reposición: {e}")
            finally:
                conn_art.close()

        try:
            # 4. Cargar la plantilla y definir mapeo de marcadores
            document = docx.Document(plantilla_path)

            # Mapeo de los marcadores de la plantilla actual
            mapeo = {
                '[[CODIGO_RMA]]':  codigo_rma,
                '[[CLIENTE]]':     nombre_cliente,
                '[[FECHA]]':       datetime.datetime.now().strftime("%Y-%m-%d"),
                '[[DOC_CLIENTE]]': datos.get('numero_documento_cliente', 'N/A'),
                '[[NOM_CLIENTE]]': datos.get('persona_de_contacto', 'N/A'),
                '[[ALBARAN]]':     datos.get('numero_albaran', 'N/A'),
            }

            # 5. Reemplazar marcadores preservando el formato de la plantilla
            # (antes se hacía con `p.text = p.text.replace(...)`, que en
            # python-docx borra todos los runs del párrafo y los sustituye por
            # uno solo con formato por defecto, perdiendo el de la plantilla)
            self._reemplazar_marcadores_preservando_formato(document, mapeo)

            # 5.1. Rellenar la tabla de material (UNIDADES / REFERENCIA) con los
            # artículos del expediente, una fila por artículo. La plantilla trae
            # una fila de datos en blanco ya formateada (fuente/tamaño de la
            # plantilla); se reutiliza para el primer artículo y se clonan filas
            # adicionales para el resto. MED. CORTE y DESCRIPCION quedan en
            # blanco: no hay un dato equivalente en el expediente.
            if document.tables and articulos_reposicion:
                tabla_material = document.tables[0]
                fila_plantilla = tabla_material.rows[-1]

                def _rellenar_fila(fila, referencia, cantidad):
                    unidades_cell, referencia_cell = fila.cells[0], fila.cells[1]
                    unidades_cell.paragraphs[0].add_run(str(cantidad) if cantidad else '')
                    referencia_cell.paragraphs[0].add_run(str(referencia) if referencia else 'N/A')

                for i, (ref_articulo, cantidad) in enumerate(articulos_reposicion):
                    if i == 0:
                        fila_destino = fila_plantilla
                    else:
                        fila_destino = tabla_material.add_row()
                    _rellenar_fila(fila_destino, ref_articulo, cantidad)

            # 6. Guardar temporalmente para subirlo a Backblaze B2
            temp_dir = tempfile.mkdtemp(prefix="reposicion_rma_")
            temp_file_path = os.path.join(temp_dir, nombre_archivo_final)
            document.save(temp_file_path)
            
            # 7. Decidir dónde guardar (Dropbox o local)
            if usar_b2():
                # Subir a Backblaze B2
                exito, ruta_relativa = self._subir_archivo_b2(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'backblaze'
                ubicacion_desc = "Dropbox"
            else:
                # Guardar localmente (fallback)
                exito, ruta_relativa = self._subir_archivo_local(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'local'
                ubicacion_desc = "local"
            
            # 8. Limpiar archivo temporal
            try:
                os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            if not exito:
                messagebox.showerror("Error", f"No se pudo guardar el documento de reposición en {ubicacion_desc}.")
                return
            
            # 9. Registrar en la Base de Datos
            conn, cursor = self.master.conectar_db()
            try:
                # Verificar esquema de BD antes de insertar
                self._verificar_columna_tipo_almacenamiento(cursor)
                
                # Preparar inserción con o sin tipo_almacenamiento según el esquema
                if getattr(self, '_usar_tipo_almacenamiento', False):
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento) 
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username,
                        tipo_almacenamiento
                    ))
                else:
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username
                    ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().isoformat(),
                    f"Generado documento de Reposición/Devolución: {nombre_archivo_final} ({'☁️ Backblaze B2' if usar_b2() else '💾 Local'})", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                
                try:
                    # Actualizar historial si está visible
                    if hasattr(self, 'historial_tab'):
                        self.mostrar_historial(self.historial_tab)
                except AttributeError:
                    # Si la pestaña historial_tab no está definida, ignoramos.
                    pass
                
                # Mensaje personalizado según donde se guardó
                if usar_b2():
                    messagebox.showinfo("Éxito", f"✅ Documento de Reposición/Devolución '{nombre_archivo_final}' generado y subido a Backblaze B2 correctamente.\n\n📁 Ubicación: {ruta_relativa}")
                else:
                    messagebox.showinfo("Éxito", f"✅ Documento de Reposición/Devolución '{nombre_archivo_final}' generado y guardado localmente.")
                
            except Exception as db_e:
                if hasattr(conn, 'rollback'):
                    conn.rollback()
                messagebox.showerror("Error DB", f"Documento generado, pero error al registrar en DB/Historial. Error: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el documento. Asegúrese de que la plantilla existe y es un archivo .docx válido.\nError: {e}")
            conn, cursor = self.master.conectar_db()
            try:
                # Registro en rma_adjuntos
                cursor.execute("""
                    INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    nombre_archivo_final, 
                    ruta_relativa, 
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                    self.username
                ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                    f"Generado documento de Reposición/Devolución: {nombre_archivo_final}", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                try:
                    # Llamamos al método que recarga el contenido de la pestaña
                    self.mostrar_historial(self.historial_tab) 
                except AttributeError:
                    # Si la pestaña historial_tab no está definida (ej. en modo "nuevo"), ignoramos.
                    pass
                
                messagebox.showinfo("Éxito", f"Documento de Reposición/Devolución '{nombre_archivo_final}' generado y adjuntado correctamente.")
                
            except Exception as db_e:
                if hasattr(conn, 'rollback'):
                    conn.rollback()
                messagebox.showerror("Error DB", f"Documento generado, pero error al registrar en DB/Historial. Error: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el documento. Asegúrese de que la plantilla existe y es un archivo .docx válido.\nError: {e}")

    def abrir_plantilla_informe_manual_sinusoactualmente(self):
        """
        Abre la plantilla de informe de Word y la carpeta de destino de los adjuntos.
        """
        # Verificación inicial: ¿Estamos editando un RMA?
        if not self.rma_actual_id:
            messagebox.showerror("Error", "Primero debe guardar el expediente o cargar uno existente para generar un informe.")
            return

        # 1. Definir rutas
        # Obtiene la ruta absoluta de la carpeta 'plantillas'
        plantilla_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas")
        plantilla_path = os.path.join(plantilla_dir, "Plantilla_RMA.docx") # Asegúrate de que el nombre coincide
        
        # 2. Verificar que la plantilla existe
        if not os.path.exists(plantilla_path):
            messagebox.showerror("Error", f"No se encontró la plantilla de informe en:\n{plantilla_path}")
            return
            
        # 3. Abrir la plantilla con el programa asociado (Word)
        try:
            # os.startfile es la forma más limpia en Windows para abrir archivos
            if os.name == 'nt': 
                 os.startfile(plantilla_path)
            else:
                 # Común para macOS o Linux (puede requerir ajustes según la distribución)
                 subprocess.Popen(['xdg-open', plantilla_path]) 

            messagebox.showinfo("Instrucción", 
                "Se ha abierto la plantilla de Word.\n"
                "A continuación se abrirá la carpeta de destino. Por favor, "
                "**GUARDE el documento de Word FINAL ahí**.")
            
            # 4. Abrir la carpeta de destino para que el usuario guarde el resultado
            self.abrir_dialogo_adjunto(modo_abrir_carpeta=True)

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la plantilla.\nError: {e}")

    def exportar_a_excel(self, datos, columnas, filename):
        """Exporta los datos de la estadística actual a un archivo Excel."""
        if not datos:
            messagebox.showwarning("Exportar", "No hay datos para exportar.")
            return

        # 1. Crear el DataFrame de Pandas
        df = pd.DataFrame(datos, columns=columnas)
        
        # 2. Abrir diálogo para guardar archivo
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx", 
            filetypes=[("Archivos Excel", "*.xlsx")],
            initialfile=filename
        )
        
        if path:
            try:
                # 3. Guardar en Excel (openpyxl se usa internamente por pandas)
                df.to_excel(path, index=False) # index=False evita añadir la columna de índice de pandas
                messagebox.showinfo("Exportación Exitosa", f"Datos exportados a:\n{path}")
            except Exception as e:
                messagebox.showerror("Error de Exportación", f"No se pudo guardar el archivo Excel.\nError: {e}")

    def mostrar_tabla_estadistica(self, datos, columnas, export_filename, frame, formato_moneda=False):
        """
        Dibuja un listado genérico de resultados usando un CTkScrollableFrame 
        y añade el botón de exportación.
        """
        
        if not datos:
            ctk.CTkLabel(frame, text="No se encontraron datos para esta estadística.", text_color="gray").pack(pady=20)
            return

        # Función para abrir el expediente al hacer clic en el código RMA
        def abrir_expediente(codigo_rma):
            """Abre el expediente en una nueva ventana independiente para consulta/edición.

            La ventana muestra un resumen y permite abrir el expediente en el panel principal si
            el usuario prefiere editar en la interfaz habitual.
            """
            conn, cursor = self.master.conectar_db()
            if not conn:
                return

            try:
                # Obtener datos maestro
                cursor.execute("SELECT id, codigo_rma, cliente, fecha_gestion, motivo FROM rma_maestro WHERE codigo_rma = ?", (codigo_rma,))
                maestro = cursor.fetchone()
                if not maestro:
                    messagebox.showerror("Error", f"No se encontró el expediente con código {codigo_rma}")
                    return

                rma_id = maestro[0]
                codigo = maestro[1]
                cliente = maestro[2]
                fecha = maestro[3]
                motivo = maestro[4]

                # Crear ventana independiente
                vent = ctk.CTkToplevel(self)
                vent.title(f"Expediente {codigo}")
                vent.geometry("900x700")
                vent.minsize(700, 500)

                # Contenedor principal
                cont = ctk.CTkFrame(vent)
                cont.pack(fill="both", expand=True, padx=12, pady=12)

                # Cabecera - campos editables
                header_frame = ctk.CTkFrame(cont)
                header_frame.pack(fill="x", pady=(0,6))

                ctk.CTkLabel(header_frame, text=f"Nº EXPEDIENTE: {codigo}", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0,6))

                ctk.CTkLabel(header_frame, text="Cliente:").grid(row=1, column=0, sticky="w")
                entry_cliente = ctk.CTkEntry(header_frame)
                entry_cliente.grid(row=1, column=1, sticky="ew", padx=(8,0))
                entry_cliente.insert(0, str(cliente) if cliente is not None else "")

                ctk.CTkLabel(header_frame, text="Nº Documento:").grid(row=2, column=0, sticky="w")
                entry_num_doc = ctk.CTkEntry(header_frame)
                entry_num_doc.grid(row=2, column=1, sticky="ew", padx=(8,0))
                # intentar rellenar si existe
                try:
                    entry_num_doc.insert(0, str(maestro[3]) if maestro[3] is not None else "")
                except Exception:
                    pass

                ctk.CTkLabel(header_frame, text="Fecha Gestión:").grid(row=3, column=0, sticky="w")
                entry_fecha = ctk.CTkEntry(header_frame)
                entry_fecha.grid(row=3, column=1, sticky="ew", padx=(8,0))
                entry_fecha.insert(0, str(fecha) if fecha is not None else "")

                ctk.CTkLabel(header_frame, text="Email contacto:").grid(row=4, column=0, sticky="w")
                entry_email = ctk.CTkEntry(header_frame)
                entry_email.grid(row=4, column=1, sticky="ew", padx=(8,0))
                # maestro[6] será email si está presente
                try:
                    entry_email.insert(0, str(maestro[6]) if len(maestro) > 6 and maestro[6] is not None else "")
                except Exception:
                    pass

                header_frame.grid_columnconfigure(1, weight=1)

                ctk.CTkLabel(cont, text="Motivo:").pack(anchor="w")
                txt_motivo = ctk.CTkTextbox(cont, height=80)
                txt_motivo.pack(fill="x", pady=(0,8))
                try:
                    txt_motivo.insert("1.0", str(motivo) if motivo is not None else "")
                except Exception:
                    pass

                # Mostrar artículos relacionados (lista simple)
                ctk.CTkLabel(cont, text="Artículos asociados:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(6,4))
                art_frame = ctk.CTkScrollableFrame(cont)
                art_frame.pack(fill="both", expand=True, pady=(0,8))

                # Obtener detalles incluyendo su id para permitir edición
                cursor.execute("SELECT id, referencia_articulo, cantidad_segun_documento, cantidad_entregada, estado_producto FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                detalles = cursor.fetchall()
                if detalles:
                    for d in detalles:
                        det_id = d[0]
                        ref = d[1]
                        cant_doc = d[2]
                        cant_ent = d[3]
                        estado = d[4]

                        row_fr = ctk.CTkFrame(art_frame)
                        row_fr.pack(fill="x", padx=4, pady=2)

                        lbl = ctk.CTkLabel(row_fr, text=f"{ref} — Doc: {cant_doc} — Ent: {cant_ent} — Estado: {estado}")
                        lbl.pack(side="left", anchor="w")

                        def make_editar(did, lbl_widget):
                            def editar():
                                ed_win = ctk.CTkToplevel(vent)
                                ed_win.title(f"Editar artículo {ref}")
                                ed_win.geometry("420x200")

                                ctk.CTkLabel(ed_win, text=f"Referencia: {ref}").pack(anchor="w", padx=8, pady=(8,4))
                                ctk.CTkLabel(ed_win, text="Cantidad Entregada:").pack(anchor="w", padx=8)
                                ent_cant = ctk.CTkEntry(ed_win)
                                ent_cant.pack(fill="x", padx=8, pady=(0,6))
                                ent_cant.insert(0, str(cant_ent) if cant_ent is not None else "")

                                ctk.CTkLabel(ed_win, text="Estado Producto:").pack(anchor="w", padx=8)
                                ent_estado = ctk.CTkEntry(ed_win)
                                ent_estado.pack(fill="x", padx=8, pady=(0,6))
                                ent_estado.insert(0, str(estado) if estado is not None else "")

                                def guardar_detalle():
                                    new_cant = ent_cant.get().strip()
                                    new_estado = ent_estado.get().strip()
                                    try:
                                        conn2 = connect_db()
                                        cur2 = conn2.cursor()
                                        cur2.execute("UPDATE rma_detalles SET cantidad_entregada = ?, estado_producto = ? WHERE id = ?", (new_cant, new_estado, did))
                                        conn2.commit()
                                        conn2.close()
                                        # Actualizar etiqueta
                                        lbl_widget.configure(text=f"{ref} — Doc: {cant_doc} — Ent: {new_cant} — Estado: {new_estado}")
                                        ed_win.destroy()
                                        # Refrescar los datos de la tabla
                                        refrescar_estadisticas()
                                    except sqlite3.Error as e:
                                        messagebox.showerror("Error BD", f"No se pudo actualizar el detalle: {e}")

                                ctk.CTkButton(ed_win, text="Guardar", command=guardar_detalle).pack(side="left", padx=12, pady=12)
                                ctk.CTkButton(ed_win, text="Cerrar", command=ed_win.destroy).pack(side="right", padx=12, pady=12)

                            return editar

                        btn_ed = ctk.CTkButton(row_fr, text="Editar", command=make_editar(det_id, lbl))
                        btn_ed.pack(side="right")
                else:
                    ctk.CTkLabel(art_frame, text="No hay artículos registrados.").pack(pady=8)

                # Botones en el footer de la ventana
                footer = ctk.CTkFrame(cont)
                footer.pack(fill="x", pady=(8,0))

                # Botones en el footer de la ventana
                # Función para refrescar la tabla de estadísticas
                def refrescar_estadisticas():
                    """Refresca los datos de la tabla de estadísticas."""
                    # Limpiar y recargar los datos
                    try:
                        # Re-ejecutar la carga de datos actual
                        self._cargar_datos_articulos_incidencia()
                    except Exception as e:
                        print(f"Error al refrescar estadísticas: {e}")

                def guardar_maestro():
                    new_cliente = entry_cliente.get().strip()
                    new_num_doc = entry_num_doc.get().strip()
                    new_fecha = entry_fecha.get().strip()
                    new_email = entry_email.get().strip().lower()  # Convertir a minúsculas
                    new_motivo = txt_motivo.get("1.0", "end-1c").strip()

                    try:
                        conn2 = connect_db()
                        cur2 = conn2.cursor()
                        cur2.execute(
                            "UPDATE rma_maestro SET cliente = ?, numero_documento_cliente = ?, fecha_gestion = ?, motivo = ?, email_de_contacto = ? WHERE id = ?",
                            (new_cliente, new_num_doc, new_fecha, new_motivo, new_email, rma_id)
                        )
                        conn2.commit()
                        conn2.close()
                        messagebox.showinfo("Guardado", "Expediente actualizado correctamente.")
                        # Refrescar datos de la tabla
                        refrescar_estadisticas()
                    except sqlite3.Error as e:
                        messagebox.showerror("Error BD", f"No se pudo actualizar el expediente: {e}")

                def abrir_en_panel():
                    # Usar RmaEditorWindow (no self.mostrar_nuevo_rma directo): esa función
                    # dibuja la ficha sobre el content_frame que se le pase, y su botón
                    # "Cerrar" asume que está en una ventana emergente propia. Llamarla
                    # directamente aquí dibujaba la ficha sobre la ventana principal pero con
                    # ese botón "Cerrar", que terminaba destruyendo la ventana principal entera.
                    try:
                        from lib.rma_editor_window import RmaEditorWindow
                        RmaEditorWindow(self, rma_id)
                        vent.destroy()
                    except Exception as e:
                        messagebox.showerror("Error", f"No se pudo abrir el expediente: {e}")

                ctk.CTkButton(footer, text="Guardar cambios", command=guardar_maestro).pack(side="left")
                ctk.CTkButton(footer, text="✏️ Abrir expediente", command=abrir_en_panel).pack(side="left", padx=8)
                ctk.CTkButton(footer, text="Cerrar", command=vent.destroy).pack(side="right")

            except sqlite3.Error as e:
                messagebox.showerror("Error de Base de Datos", f"Error al consultar el expediente: {e}")
            finally:
                conn.close()

        # Mensaje informativo para el usuario
        ctk.CTkLabel(frame, text="Pulse el Nº EXPEDIENTE (primera columna) para abrirlo en una nueva ventana.", text_color="gray").pack(pady=(6,4))

        # 1. Botón de Exportar
        btn_export = ctk.CTkButton(
            frame, 
            text="💾 Exportar a Excel", 
            command=lambda: self.exportar_a_excel(datos, columnas, export_filename) # Llama al método de exportación
        )
        btn_export.pack(pady=(5, 8))

        # 2. Contenedor de la Tabla
        tabla_scroll_frame = ctk.CTkScrollableFrame(frame, label_text="Resultados")
        tabla_scroll_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        
        # La columna 0 (típicamente el nombre del cliente/artículo) se expande
        for col_idx in range(len(columnas)):
            # Damos un peso 3 a la primera columna (Referencia Artículo)
            # y peso 1 al resto para asegurar que la Referencia se estire más.
            if col_idx == 0:
                tabla_scroll_frame.grid_columnconfigure(col_idx, weight=3) 
            else:
                tabla_scroll_frame.grid_columnconfigure(col_idx, weight=1) 
        
        # 3. Encabezados
        header_font = ctk.CTkFont(weight="bold")
        for col_idx, col_name in enumerate(columnas):
            # Alineación del encabezado
            sticky = "w" if col_idx == 0 else "e"
            ctk.CTkLabel(tabla_scroll_frame, text=col_name, font=header_font).grid(row=0, column=col_idx, padx=10, pady=5, sticky=sticky)
            
        # 4. Datos
        for row_idx, row_data in enumerate(datos):
            for col_idx, cell_value in enumerate(row_data):
                
                text_to_display = str(cell_value if cell_value is not None else "N/A")
                
                # 🚨 Aplicar formato de moneda si se especifica y estamos en la columna de valor (la última)
                if formato_moneda and col_idx == len(row_data) - 1 and cell_value is not None:
                    try:
                        # Intenta usar locale para el formato de moneda (€)
                        text_to_display = locale.currency(float(cell_value), grouping=True, symbol=True)
                    except (ValueError, TypeError, locale.Error):
                        # Fallback si no es un número o si falla el locale
                        text_to_display = f"{float(cell_value):,.2f} €"

                # Alineación del contenido
                sticky = "w" if col_idx == 0 else "e"
                
                # Crear el label como widget para poder bindear eventos
                label_widget = ctk.CTkLabel(
                    tabla_scroll_frame,
                    text=text_to_display,
                    justify="left"
                )
                label_widget.grid(row=row_idx + 1, column=col_idx, padx=10, pady=2, sticky=sticky)

                # Si es la primera columna (código RMA), hacerlo clicable
                if col_idx == 0:
                    # Cambiar cursor a mano y subrayar/colorear al hover
                    try:
                        original_color = label_widget.cget("text_color")
                    except Exception:
                        original_color = None

                    # Set cursor to hand2 if supported
                    try:
                        label_widget.configure(cursor="hand2")
                    except Exception:
                        pass

                    def on_enter(e):
                        try:
                            label_widget.configure(text_color="#1a73e8")
                        except Exception:
                            pass

                    def on_leave(e):
                        try:
                            if original_color is not None:
                                label_widget.configure(text_color=original_color)
                        except Exception:
                            pass

                    # Bind click to abrir_expediente using the displayed text (codigo)
                    label_widget.bind("<Button-1>", lambda e, code=text_to_display: abrir_expediente(code))
                    label_widget.bind("<Enter>", on_enter)
                    label_widget.bind("<Leave>", on_leave)

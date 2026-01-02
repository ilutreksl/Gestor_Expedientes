import customtkinter as ctk
import tkinter.messagebox as messagebox
from tkinter import Toplevel
import sqlite3
import os
try:
    # Carga variables de entorno desde .env si existe (opcional)
    from dotenv import load_dotenv  # type: ignore
    load_dotenv()
except Exception:
    pass
import bcrypt
import sys
import os 
import datetime
import webbrowser
import docx
import requests
from dotenv import load_dotenv
import re

# Cargar variables de entorno
load_dotenv()
from PIL import Image, ImageTk
#from tkcalendar import Calendar
from CTkDatePicker import CTkDatePicker
from lib.changelog_window import mostrar_ventana_cambios
from lib.rma_utils import obtener_ultima_actividad, calcular_tiempos_expediente, obtener_color_tiempo, obtener_promedio_cliente
from lib.video_utils import comprimir_video_inteligente
from lib.avisos_manager import AvisosManager
from lib.backup_manager import BackupManagerB2
from lib.estados_manager import EstadosArticuloManager
from lib.personas_manager import PersonasManager
from lib.resultado_expediente_manager import ResultadoExpedienteManager

import tkinter as tk
from tkinter import ttk

# Monkey patch para prevenir errores en ventanas destruidas
# Guardamos referencias a los métodos originales
_original_tk_after = tk.Misc.after
_original_toplevel_destroy = Toplevel.destroy

# Diccionario global para rastrear IDs de callbacks por ventana
_window_after_ids = {}

def _safe_after(self, ms, func=None, *args):
    """Wrapper seguro para after() que registra los IDs"""
    # Si func es None, es un sleep, no lo trackeamos
    if func is None:
        return _original_tk_after(self, ms)
    
    # Función wrapper que verifica si la ventana existe antes de ejecutar
    def safe_func(*func_args):
        try:
            # Verificar si el widget todavía existe
            if self.winfo_exists():
                return func(*func_args)
        except:
            pass  # Ventana ya destruida, ignorar silenciosamente
    
    # Programar el callback con la función segura
    after_id = _original_tk_after(self, ms, safe_func, *args)
    
    # Registrar el ID si es una ventana Toplevel
    if isinstance(self, (Toplevel, ctk.CTkToplevel)):
        window_id = id(self)
        if window_id not in _window_after_ids:
            _window_after_ids[window_id] = []
        _window_after_ids[window_id].append(after_id)
    
    return after_id

def _safe_destroy(self):
    """Wrapper seguro para destroy() que cancela todos los callbacks"""
    window_id = id(self)
    
    # Cancelar todos los callbacks programados para esta ventana
    if window_id in _window_after_ids:
        for after_id in _window_after_ids[window_id][:]:
            try:
                self.after_cancel(after_id)
            except:
                pass
        del _window_after_ids[window_id]
    
    # Liberar grab si existe
    try:
        self.grab_release()
    except:
        pass
    
    # Llamar al destroy original
    try:
        _original_toplevel_destroy(self)
    except:
        pass

# Aplicar los monkey patches
tk.Misc.after = _safe_after
Toplevel.destroy = _safe_destroy
ctk.CTkToplevel.destroy = _safe_destroy
import threading
import tempfile

# Importaciones para Dropbox
import dropbox
from dropbox.exceptions import ApiError, AuthError
try:
    from dropbox_config import (
        DROPBOX_ACCESS_TOKEN, DROPBOX_ROOT_FOLDER, 
        DROPBOX_APP_KEY, DROPBOX_APP_SECRET
    )
    # Intentar importar refresh token si existe
    try:
        from dropbox_config import DROPBOX_REFRESH_TOKEN
    except ImportError:
        DROPBOX_REFRESH_TOKEN = None
except ImportError:
    print("ADVERTENCIA: No se pudo cargar dropbox_config.py. El sistema de adjuntos usará almacenamiento local.")
    DROPBOX_ACCESS_TOKEN = None
    DROPBOX_REFRESH_TOKEN = None
    DROPBOX_APP_KEY = None
    DROPBOX_APP_SECRET = None
    DROPBOX_ROOT_FOLDER = "/Adjuntos_RMA"

# Variables globales para cache de cliente de Dropbox
_dropbox_client_cache = None
_last_token_check = 0

# ================================
# FUNCIONES DE COMPRESIÓN DE IMÁGENES
# ================================

def es_imagen(filepath):
    """Detecta si un archivo es una imagen basándose en la extensión."""
    extensiones_imagen = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif', '.webp', '.heic', '.heif'}
    ext = os.path.splitext(filepath)[1].lower()
    return ext in extensiones_imagen

def es_video(filepath):
    """Detecta si un archivo es un video basándose en la extensión."""
    extensiones_video = {'.mp4', '.mov', '.avi', '.mkv', '.wmv', '.flv', '.webm', '.m4v', '.mpg', '.mpeg', '.3gp'}
    ext = os.path.splitext(filepath)[1].lower()
    return ext in extensiones_video

def comprimir_imagen_inteligente(filepath_original, callback_progreso=None):
    """
    Comprime una imagen siguiendo la estrategia inteligente (Opción 1):
    - Si > 2MB: Redimensionar a 1920x1080 máx, calidad 85%
    - Si > 500KB: Solo recomprimir con calidad 90%
    - Si < 500KB: No modificar
    
    Returns: (filepath_comprimido, tamaño_original_mb, tamaño_final_mb) o (None, 0, 0) si hay error
    """
    try:
        from PIL import Image
        import tempfile
        
        if callback_progreso:
            callback_progreso("🔍 Analizando imagen...")
        
        # Obtener tamaño original
        tamaño_original = os.path.getsize(filepath_original)
        tamaño_original_mb = tamaño_original / (1024 * 1024)
        
        # Si es muy pequeña, no comprimir
        if tamaño_original < 500 * 1024:  # 500KB
            if callback_progreso:
                callback_progreso(f"✅ Imagen pequeña ({tamaño_original_mb:.1f}MB), no necesita compresión")
            return filepath_original, tamaño_original_mb, tamaño_original_mb
        
        if callback_progreso:
            callback_progreso(f"📏 Imagen {tamaño_original_mb:.1f}MB - iniciando compresión...")
        
        # Abrir imagen
        with Image.open(filepath_original) as img:
            # Convertir HEIC/HEIF a RGB si es necesario
            if img.format in ['HEIC', 'HEIF'] or img.mode in ['RGBA', 'LA']:
                if callback_progreso:
                    callback_progreso("🔄 Convirtiendo formato...")
                # Crear fondo blanco para transparencias
                if img.mode in ['RGBA', 'LA']:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                else:
                    img = img.convert('RGB')
            
            # Determinar parámetros de compresión según tamaño
            if tamaño_original > 2 * 1024 * 1024:  # > 2MB
                # Redimensionar y comprimir agresivamente
                max_width, max_height = 1920, 1080
                calidad = 85
                if callback_progreso:
                    callback_progreso("🎯 Redimensionando a Full HD y comprimiendo...")
            else:
                # Solo recomprimir
                max_width, max_height = img.size
                calidad = 90
                if callback_progreso:
                    callback_progreso("🔧 Recomprimiendo con calidad optimizada...")
            
            # Redimensionar manteniendo aspecto si es necesario
            if img.width > max_width or img.height > max_height:
                img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                if callback_progreso:
                    callback_progreso(f"📐 Redimensionado a {img.width}x{img.height}")
            
            # Crear archivo temporal comprimido
            temp_fd, temp_path = tempfile.mkstemp(suffix='.jpg')
            os.close(temp_fd)
            
            # Guardar imagen comprimida
            if callback_progreso:
                callback_progreso("💾 Guardando imagen optimizada...")
                
            img.save(temp_path, 'JPEG', quality=calidad, optimize=True, progressive=True)
        
        # Verificar resultado
        tamaño_final = os.path.getsize(temp_path)
        tamaño_final_mb = tamaño_final / (1024 * 1024)
        reduccion = ((tamaño_original - tamaño_final) / tamaño_original) * 100
        
        if callback_progreso:
            callback_progreso(f"✅ Compresión completada: {tamaño_original_mb:.1f}MB → {tamaño_final_mb:.1f}MB ({reduccion:.1f}% reducción)")
        
        return temp_path, tamaño_original_mb, tamaño_final_mb
        
    except Exception as e:
        if callback_progreso:
            callback_progreso(f"❌ Error en compresión: {e}")
        print(f"Error comprimiendo imagen {filepath_original}: {e}")
        return None, 0, 0


class Tooltip:
    """Simple tooltip for tkinter widgets. Shows a small Toplevel on hover."""
    def __init__(self, widget, text, delay=400):
        self.widget = widget
        self.text = text
        self.delay = delay
        self._id = None
        self._tipwindow = None
        widget.bind("<Enter>", self._schedule)
        widget.bind("<Leave>", self._hide)
        widget.bind("<ButtonPress>", self._hide)

    def _schedule(self, event=None):
        self._unschedule()
        try:
            if not USER_SETTINGS.get("show_tooltips", True):
                return
        except Exception:
            pass
        self._id = self.widget.after(self.delay, self._show)

    def _unschedule(self):
        if self._id:
            try:
                self.widget.after_cancel(self._id)
            except Exception:
                pass
            self._id = None

    def _show(self):
        if self._tipwindow or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
        self._tipwindow = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1,
                         font=("Segoe UI", 9))
        label.pack(ipadx=4, ipady=2)

    def _hide(self, event=None):
        self._unschedule()
        tw = self._tipwindow
        self._tipwindow = None
        if tw:
            try:
                tw.destroy()
            except Exception:
                pass

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    print("⚠️ Pandas no está disponible. La exportación a Excel estará deshabilitada.")

import locale
import shutil # Para copiar archivos
import tkinter.filedialog as filedialog # Para el diálogo de selección de archivos
import subprocess # Para abrir archivos en diferentes OS
# Notificaciones nativas (Windows)
try:
    from win10toast import ToastNotifier
except Exception:
    ToastNotifier = None

# Módulo para rellenado de PDFs
try:
    from lib.pdf_fill import fill_pdf_for_rma, get_pdf_field_names, fill_pdf
except Exception:
    # Si no está instalado/ disponible aún, seguiremos sin la funcionalidad
    fill_pdf_for_rma = None
    get_pdf_field_names = None
    fill_pdf = None

# Definición de las variables globales de la base de datos
DB_NAME = "rma_app.db"
# Mensaje de advertencia sobre la limitación de SQLite en red compartida
ADVERTENCIA_MULTIUSUARIO = "⚠️ ADVERTENCIA: Esta app usa SQLite, NO es segura para múltiples usuarios escribiendo a la vez en red compartida. ¡Riesgo de corrupción de datos si escriben a la vez!"

APP_VERSION = "v1.0.2"
DB_FILENAME = "rma_app.db"

# Session global para Turso (reutiliza conexiones HTTP)
_turso_session = None

# Caché simple para consultas frecuentes (estados, etc.)
_query_cache = {}
_cache_timeout = 300  # 5 minutos

def _get_turso_session():
    """Obtiene o crea una sesión HTTP persistente para Turso"""
    global _turso_session
    if _turso_session is None:
        import requests
        _turso_session = requests.Session()
        _turso_session.headers.update({
            "Content-Type": "application/json"
        })
    return _turso_session

def _get_cached_query(cache_key, query_func, ttl=None):
    """Sistema de caché para queries frecuentes (estados, usuarios, etc.)"""
    import time
    global _query_cache
    
    if ttl is None:
        ttl = _cache_timeout
    
    now = time.time()
    
    # Verificar si existe en caché y no ha expirado
    if cache_key in _query_cache:
        cached_data, timestamp = _query_cache[cache_key]
        if now - timestamp < ttl:
            return cached_data
    
    # Si no está en caché o expiró, ejecutar query y guardar
    result = query_func()
    _query_cache[cache_key] = (result, now)
    return result

def invalidate_cache(pattern=None):
    """Invalida caché completo o por patrón"""
    global _query_cache
    if pattern is None:
        _query_cache.clear()
    else:
        keys_to_delete = [k for k in _query_cache.keys() if pattern in k]
        for key in keys_to_delete:
            del _query_cache[key]


def parse_date_to_iso(value: str) -> str:
    """Intenta parsear una cadena de fecha en varios formatos comunes y devuelve
    la fecha normalizada en formato ISO YYYY-MM-DD.

    Lanza ValueError si no puede parsearse.
    """
    if value is None:
        raise ValueError("None provided")
    v = str(value).strip()
    if v == "":
        raise ValueError("Empty date")

    # Intentar formatos más comunes
    formatos = ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d", "%m/%d/%Y"]
    for fmt in formatos:
        try:
            dt = datetime.datetime.strptime(v, fmt)
            return dt.strftime("%Y-%m-%d")
        except Exception:
            continue

    # Si ninguno encaja, intentar parseo flexible usando dateutil si está disponible
    try:
        from dateutil import parser as _parser  # type: ignore
        dt = _parser.parse(v, dayfirst=True)
        return dt.strftime("%Y-%m-%d")
    except Exception:
        raise ValueError(f"Formato de fecha no reconocido: {value}")

# --- Conector unificado: Turso (libSQL) si hay credenciales, si no SQLite local ---
def connect_db(timeout: float | None = None):
    """Devuelve una conexión a la BD.
    - Si existen TURSO_DATABASE_URL y TURSO_AUTH_TOKEN en el entorno, usa Turso (libSQL).
    - En caso contrario, usa SQLite local (rma_app.db).
    El API devuelto implementa DB-API 2.0 (cursor/execute/commit).
    """
    turso_url = os.getenv("TURSO_DATABASE_URL")
    turso_token = os.getenv("TURSO_AUTH_TOKEN")
    if turso_url and turso_token:
        try:
            import requests
            import json

            class TursoCursor:
                def __init__(self, url: str, token: str, connection=None):
                    self._url = url
                    self._token = token
                    self._connection = connection
                    self._result = None
                    self.lastrowid = None
                    self.rowcount = -1
                    self._fetch_idx = 0
                    self.description = None  # DB-API 2.0: lista de tuplas con info de columnas

                def execute(self, sql: str, params: tuple | list | None = None):
                    # Convertir placeholders ? a args posicionales para Turso
                    if params:
                        # Turso espera args como lista de objetos con {type, value}
                        args = [{"type": "text", "value": str(v)} for v in params]
                    else:
                        args = []
                    
                    # Usar session persistente para mejor rendimiento
                    session = _get_turso_session()
                    
                    # Hacer request a la API REST de Turso (v2/pipeline format)
                    request_payload = {"requests": [{"type": "execute", "stmt": {"sql": sql, "args": args}}]}
                    
                    response = session.post(
                        self._url,
                        headers={
                            "Authorization": f"Bearer {self._token}"
                        },
                        json=request_payload,
                        timeout=10  # Timeout de 10 segundos
                    )
                    
                    if response.status_code != 200:
                        print(f"Error - Response status: {response.status_code}, text: {response.text}")
                        raise Exception(f"Turso API error: {response.status_code} - {response.text}")
                    
                    data = response.json()
                    results = data.get("results", [])
                    
                    if results and len(results) > 0:
                        # Verificar si hay error en la respuesta
                        result_item = results[0]
                        if result_item.get("type") == "error":
                            error_msg = result_item.get("error", {}).get("message", "Unknown error")
                            # No imprimir errores de verificación de esquema que son normales
                            if not ("tipo_almacenamiento" in sql and "SELECT" in sql):
                                print(f"SQL Error: {error_msg}")
                                print(f"Query: {sql}")
                                print(f"Params: {params}")
                            # Lanzar excepción para que pueda ser manejada por el código que llama
                            raise sqlite3.OperationalError(f"SQLite error: {error_msg}")
                        else:
                            result = result_item.get("response", {}).get("result", {})
                            self._result = result
                    else:
                        self._result = {"rows": []}
                    
                    # Extraer metadatos
                    self.lastrowid = self._result.get("last_insert_rowid")
                    self.rowcount = self._result.get("rows_affected", -1)
                    
                    # Construir description según DB-API 2.0
                    # Formato: (name, type_code, display_size, internal_size, precision, scale, null_ok)
                    cols = self._result.get("cols", [])
                    if cols:
                        self.description = [
                            (col.get("name"), None, None, None, None, None, None)
                            for col in cols
                        ]
                    else:
                        self.description = None
                    
                    self._fetch_idx = 0
                    return self
                
                def executemany(self, sql: str, params_list: list):
                    """Ejecuta múltiples queries en un solo batch (mucho más rápido)"""
                    if not params_list:
                        return self
                    
                    # Construir batch de requests para pipeline
                    requests_batch = []
                    for params in params_list:
                        if params:
                            args = [{"type": "text", "value": str(v)} for v in params]
                        else:
                            args = []
                        requests_batch.append({
                            "type": "execute",
                            "stmt": {"sql": sql, "args": args}
                        })
                    
                    session = _get_turso_session()
                    response = session.post(
                        self._url,
                        headers={
                            "Authorization": f"Bearer {self._token}"
                        },
                        json={"requests": requests_batch},
                        timeout=30  # Timeout mayor para batch
                    )
                    
                    if response.status_code != 200:
                        raise Exception(f"Turso API error: {response.status_code} - {response.text}")
                    
                    # Para executemany, solo guardamos el último resultado
                    data = response.json()
                    results = data.get("results", [])
                    if results and len(results) > 0:
                        last_result = results[-1].get("response", {}).get("result", {})
                        self._result = last_result
                        self.lastrowid = last_result.get("last_insert_rowid")
                        self.rowcount = sum(
                            r.get("response", {}).get("result", {}).get("rows_affected", 0)
                            for r in results
                        )
                    
                    return self

                def fetchall(self):
                    if not self._result:
                        return []
                    
                    rows = self._result.get("rows", [])
                    
                    # Si no hay filas directamente, intentar extraer de otros posibles formatos
                    if not rows and "values" in self._result:
                        rows = self._result.get("values", [])
                    
                    if not rows and "data" in self._result:
                        rows = self._result.get("data", [])
                    
                    # Turso puede devolver diferentes formatos
                    result_rows = []
                    for row in rows:
                        if isinstance(row, list):
                            # Caso 1: Lista de valores directos
                            if row and not isinstance(row[0], dict):
                                result_rows.append(tuple(row))
                            else:
                                # Caso 2: Lista de objetos {"type": ..., "value": ...}
                                values = []
                                for cell in row:
                                    if isinstance(cell, dict):
                                        value = cell.get("value")
                                        if value is None:
                                            value = cell.get("v")  # Otra posible key
                                        values.append(value)
                                    else:
                                        values.append(cell)
                                result_rows.append(tuple(values))
                        else:
                            # Caso 3: Fila como tupla directa
                            result_rows.append(tuple(row) if not isinstance(row, tuple) else row)
                    
                    return result_rows

                def fetchone(self):
                    rows = self.fetchall()
                    if self._fetch_idx < len(rows):
                        row = rows[self._fetch_idx]
                        self._fetch_idx += 1
                        return row
                    return None

            class TursoConnection:
                def __init__(self, url: str, token: str):
                    self._url = url
                    self._token = token

                def cursor(self):
                    return TursoCursor(self._url, self._token, connection=self)

                def commit(self):
                    # No-op: Turso auto-commits
                    pass

                def close(self):
                    # No cerramos la session para que se reutilice entre conexiones
                    # La session se mantendrá viva durante toda la ejecución de la app
                    pass

            # Convertir URL: libsql://xxx o wss://xxx -> https://xxx
            api_url = turso_url.replace("libsql://", "https://").replace("wss://", "https://")
            if not api_url.endswith("/"):
                api_url += "/"
            api_url += "v2/pipeline"
            
            return TursoConnection(api_url, turso_token)
        except Exception as e:
            # Si falla, caer a SQLite local con log sencillo
            print("[WARN] No se pudo conectar a Turso (libSQL HTTP). Usando SQLite local. Causa:", e)
            return sqlite3.connect(DB_NAME, timeout=timeout if timeout is not None else 5)
    # Fallback por defecto: SQLite local
    return sqlite3.connect(DB_NAME, timeout=timeout if timeout is not None else 5)

def optimize_database():
    """Crea índices en la base de datos para mejorar el rendimiento de consultas"""
    try:
        conn = connect_db()
        cursor = conn.cursor()
        
        # Índices para búsquedas frecuentes en rma_maestro
        indices = [
            "CREATE INDEX IF NOT EXISTS idx_rma_codigo ON rma_maestro(codigo_rma)",
            "CREATE INDEX IF NOT EXISTS idx_rma_cliente ON rma_maestro(cliente)",
            "CREATE INDEX IF NOT EXISTS idx_rma_estado ON rma_maestro(estado)",
            "CREATE INDEX IF NOT EXISTS idx_rma_doc_cliente ON rma_maestro(numero_documento_cliente)",
            "CREATE INDEX IF NOT EXISTS idx_rma_fecha ON rma_maestro(fecha_emision)",
            # Índices para rma_detalles
            "CREATE INDEX IF NOT EXISTS idx_detalle_rma_id ON rma_detalles(rma_id)",
            # Índices para rma_historial
            "CREATE INDEX IF NOT EXISTS idx_historial_rma_id ON rma_historial(rma_id)",
            # Índices para rma_adjuntos
            "CREATE INDEX IF NOT EXISTS idx_adjuntos_rma_id ON rma_adjuntos(rma_id)",
        ]
        
        for idx_sql in indices:
            try:
                cursor.execute(idx_sql)
            except Exception as e:
                # Algunos índices pueden fallar en Turso, continuar
                print(f"Info: {e}")
        
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error al crear índices: {e}")

# --- NUEVA VARIABLE GLOBAL ---
ADJUNTOS_ROOT_DIR = "Adjuntos_RMA" # Carpeta principal para guardar todos los archivos adjuntos
# -----------------------------

# --- CONFIGURACIÓN Y FUNCIONES PARA DROPBOX ---
def get_dropbox_client():
    """
    Crea y retorna un cliente de Dropbox con manejo automático de tokens expirados.
    Intenta usar refresh token si está disponible, sino usa access token temporal.
    Retorna None si no está configurado o hay error irrecuperable.
    """
    global _dropbox_client_cache, _last_token_check
    import time
    
    # Cache del cliente para evitar múltiples verificaciones
    current_time = time.time()
    if _dropbox_client_cache and (current_time - _last_token_check < 300):  # Cache por 5 minutos
        return _dropbox_client_cache
    
    # Verificar si tenemos las credenciales básicas
    if not DROPBOX_APP_KEY or not DROPBOX_APP_SECRET:
        print("DROPBOX: App key y app secret requeridos")
        return None
    
    # Método 1: Intentar usar refresh token (permanente)
    if DROPBOX_REFRESH_TOKEN:
        try:
            dbx = dropbox.Dropbox(
                oauth2_refresh_token=DROPBOX_REFRESH_TOKEN,
                app_key=DROPBOX_APP_KEY,
                app_secret=DROPBOX_APP_SECRET
            )
            # Verificar que funciona
            dbx.users_get_current_account()
            _dropbox_client_cache = dbx
            _last_token_check = current_time
            print("DROPBOX: Conectado usando refresh token")
            return dbx
        except (AuthError, ApiError) as e:
            print(f"DROPBOX: Error con refresh token: {e}")
        except Exception as e:
            print(f"DROPBOX: Error inesperado con refresh token: {e}")
    
    # Método 2: Intentar usar access token temporal (4 horas)
    if DROPBOX_ACCESS_TOKEN and DROPBOX_ACCESS_TOKEN != "tu_access_token_aqui":
        try:
            dbx = dropbox.Dropbox(DROPBOX_ACCESS_TOKEN)
            # Verificar que funciona
            dbx.users_get_current_account()
            _dropbox_client_cache = dbx
            _last_token_check = current_time
            print("DROPBOX: Conectado usando access token temporal")
            return dbx
        except AuthError as e:
            if 'expired_access_token' in str(e):
                print("DROPBOX: Token de acceso expirado. Necesitas generar uno nuevo.")
                print("DROPBOX: Ve a https://www.dropbox.com/developers/apps > Settings > OAuth 2 > Generate access token")
            else:
                print(f"DROPBOX: Error de autenticación: {e}")
        except (ApiError) as e:
            print(f"DROPBOX: Error de API: {e}")
        except Exception as e:
            print(f"DROPBOX: Error inesperado: {e}")
    
    # Si llegamos aquí, no pudimos conectar
    print("DROPBOX: No se pudo establecer conexión. Sistema funcionará en modo local.")
    _dropbox_client_cache = None
    _last_token_check = current_time
    return None

def usar_dropbox():
    """
    Determina si se debe usar Dropbox o almacenamiento local.
    """
    return get_dropbox_client() is not None

def normalizar_ruta_dropbox(ruta):
    """
    Normaliza una ruta para Dropbox (debe empezar con /).
    """
    if not ruta.startswith('/'):
        ruta = '/' + ruta
    return ruta.replace('\\', '/')

# ------------------------------------------------

# --- User settings persistence (simple JSON file) ---
import json

USER_SETTINGS: dict = {}

def _get_user_settings_path() -> str:
    # Guardar en la carpeta del proyecto por simplicidad
    return os.path.join(os.getcwd(), "user_settings.json")

def load_user_settings(username: str = None) -> dict:
    defaults = {
        "date_format": "YYYY-MM-DD",
        "show_tooltips": True,
        "compact_mode": True,
        "icon_size": 24,
        "theme": "themes/BH_rime.json",
        "appearance_mode": "light"
    }
    path = _get_user_settings_path()
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as fh:
                data = json.load(fh)
                if not isinstance(data, dict):
                    return defaults
                # Backwards-compatible file structure:
                # { "global": {...}, "users": { "username": {...} } }
                merged = defaults.copy()
                # If old flat format (keys at root), merge them as global
                root_global = {}
                if any(k in data for k in defaults.keys()):
                    root_global.update(data)

                # Merge global overrides
                if isinstance(data.get("global"), dict):
                    root_global.update(data.get("global"))

                merged.update(root_global)

                # Merge per-user overrides if username provided
                if username and isinstance(data.get("users"), dict):
                    user_section = data.get("users", {}).get(username)
                    if isinstance(user_section, dict):
                        merged.update(user_section)

                return merged
    except Exception as e:
        print(f"Warning: no se pudieron cargar user_settings.json: {e}")
    return defaults

def save_user_settings(settings: dict, username: str = None) -> bool:
    # Do not persist attachments_dir - it's fixed by the app
    settings_to_save = settings.copy()
    settings_to_save.pop("attachments_dir", None)
    path = _get_user_settings_path()
    try:
        # Load existing file if present to preserve other users/global settings
        root = {}
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as fh:
                    root = json.load(fh) or {}
            except Exception:
                root = {}

        if not isinstance(root, dict):
            root = {}

        # Ensure structure
        if "global" not in root or not isinstance(root.get("global"), dict):
            # If file was flat (legacy), migrate existing keys into global
            flat_keys = {k: v for k, v in root.items() if k not in ("users", "global")}
            root = {"global": flat_keys, "users": {}}

        if username:
            users = root.setdefault("users", {})
            users[username] = settings_to_save
        else:
            root["global"] = settings_to_save

        with open(path, "w", encoding="utf-8") as fh:
            json.dump(root, fh, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error al guardar user_settings.json: {e}")
        try:
            # Mostrar cuadro de diálogo de error al usuario cuando falle el guardado
            messagebox.showerror("Error al guardar ajustes", f"No se pudieron guardar los ajustes: {e}")
        except Exception:
            pass
        return False

#try:
#    locale.setlocale(locale.LC_ALL, 'es_ES.UTF-8')
#except locale.Error:
#    try:
#        locale.setlocale(locale.LC_ALL, 'es_ES')
#    except locale.Error:
#        print("Advertencia: No se pudo configurar el locale para moneda.")

# ----------------------------------------------------------------------
# 1. CLASE DE LA VENTANA DE LOGIN
# ----------------------------------------------------------------------

class LoginApp(ctk.CTk):
    """Clase principal de la aplicación, encargada del login."""
    def __init__(self):
        super().__init__()
        
        # Configuraciones básicas de la ventana
        self.title("Gestión RMA - Login")
        self.geometry("400x300")
        self.resizable(False, False)
        
        # Agregar icono personalizado de ILUTREK
        try:
            self.iconbitmap("Icono_Ilutrek.ico")
        except Exception:
            pass  # Si no se puede cargar el icono, continuar sin él
        
        ctk.set_appearance_mode("light") 
        ctk.set_default_color_theme("themes/BH_rime.json")

        self.crear_widgets_login()

    def crear_widgets_login(self):
        """Diseña y coloca los elementos de la ventana de login."""
        
        login_frame = ctk.CTkFrame(self, 
                                   width=400, 
                                   height=300, 
                                   corner_radius=10,
                                   #fg_color=("gray95", "gray10")
                                   )
        login_frame.pack(pady=20, padx=40, fill="both", expand=True)

        ctk.CTkLabel(login_frame, text="Iniciar Sesión", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(20, 10))

        self.username_entry = ctk.CTkEntry(login_frame, placeholder_text="Nombre de Usuario")
        self.username_entry.pack(pady=12, padx=10)
        # Permitir login con ENTER desde el campo de usuario
        self.username_entry.bind("<Return>", lambda event: self.verificar_login())

        self.password_entry = ctk.CTkEntry(login_frame, placeholder_text="Contraseña", show="*")
        self.password_entry.pack(pady=12, padx=10)
        # Permitir login con ENTER desde el campo de contraseña
        self.password_entry.bind("<Return>", lambda event: self.verificar_login())

        login_button = ctk.CTkButton(login_frame, 
                                  text="Iniciar Sesión", 
                                  command=self.verificar_login,
                                  # AJUSTE CLAVE: Color de fondo y hover a tonos de gris
                                  #fg_color="gray50",     # Fondo del botón: Gris medio
                                  #hover_color="gray40",   # Color al pasar el ratón: Gris oscuro
                                  #text_color="white"
                                  )     # Color del texto (blanco para contraste)
        login_button.pack(pady=12, padx=10)
        
        self.error_label = ctk.CTkLabel(login_frame, text="", text_color="red")
        self.error_label.pack(pady=5)

    def conectar_db(self):
        """Intenta conectar a la base de datos."""
        try:
            # Añadimos un timeout de 5 segundos. Si hay variables de entorno de Turso,
            # se conectará a Turso; en caso contrario, a SQLite local.
            conn = connect_db(timeout=5)
            cursor = conn.cursor()
            return conn, cursor
        except sqlite3.Error as e:
            self.error_label.configure(text=f"Error de DB: {e}", text_color="red")
            print(f"Error de base de datos: {e}")
            return None

    def cargar_tema_usuario(self, username):
        """Carga y aplica el tema personalizado del usuario antes de crear la ventana principal."""
        try:
            user_settings = load_user_settings(username)
            
            # Aplicar tema personalizado si está configurado
            tema_usuario = user_settings.get("theme", "themes/BH_rime.json")
            
            # Validar que el tema existe y es válido
            if tema_usuario and tema_usuario != "System":
                if os.path.exists(tema_usuario):
                    try:
                        ctk.set_default_color_theme(tema_usuario)
                    except Exception as e:
                        # Si falla, usar tema por defecto
                        ctk.set_default_color_theme("themes/BH_rime.json")
                else:
                    ctk.set_default_color_theme("themes/BH_rime.json")
            else:
                ctk.set_default_color_theme("themes/BH_rime.json")
            
            # Aplicar modo de apariencia personalizado
            modo_usuario = user_settings.get("appearance_mode", "light")
            # BH_rime siempre debe usar modo claro
            if tema_usuario == "themes/BH_rime.json":
                modo_usuario = "light"
            try:
                ctk.set_appearance_mode(modo_usuario)
            except Exception as e:
                ctk.set_appearance_mode("light")
                
        except Exception as e:
            # Aplicar valores por defecto
            ctk.set_default_color_theme("themes/BH_rime.json")
            ctk.set_appearance_mode("light")

    def verificar_login(self):
        """Comprueba las credenciales del usuario."""
        username = self.username_entry.get()
        password = self.password_entry.get()
        
        if not username or not password:
            self.error_label.configure(text="Rellena todos los campos.")
            return

        conn, cursor = self.conectar_db()
    
        # Manejar el caso de que la conexión haya fallado
        if conn is None:
            messagebox.showerror("Error de Conexión", "No se pudo conectar a la base de datos.")
            return
        # -----------------------
        if not conn:
            return

        cursor = conn.cursor()
        
        cursor.execute("SELECT password_hash, rol FROM usuarios WHERE nombre_usuario = ?", (username,))
        resultado = cursor.fetchone()
        conn.close()

        if resultado:
            password_hash, rol = resultado
            try:
                if bcrypt.checkpw(password.encode('utf-8'), password_hash.encode('utf-8')):
                    self.error_label.configure(text=f"✅ Acceso concedido", text_color="green")
                    self.abrir_ventana_principal(username, rol)
                else:
                    self.error_label.configure(text="Usuario o contraseña incorrectos.")
            except Exception as e:
                 self.error_label.configure(text="Error de seguridad en credenciales.")
                 print(f"Error bcrypt: {e}")
        else:
            self.error_label.configure(text="Usuario o contraseña incorrectos.")

    def abrir_ventana_principal(self, username, rol):
        """Abre la ventana principal de la aplicación."""
        
        # IMPORTANTE: Cargar y aplicar tema del usuario ANTES de crear la ventana principal
        self.cargar_tema_usuario(username)
        
        self.withdraw() # Ocultamos la ventana de login
        # Mostrar un splash/transición entre login y la ventana principal
        try:
            splash = tk.Toplevel(self)
            splash.overrideredirect(True)
            splash.configure(bg="white")
            sw = splash.winfo_screenwidth()
            sh = splash.winfo_screenheight()
            w, h = 380, 120
            x = (sw - w) // 2
            y = (sh - h) // 2
            splash.geometry(f"{w}x{h}+{x}+{y}")

            frame = tk.Frame(splash, bg="white")
            frame.pack(fill="both", expand=True)
            tk.Label(frame, text="Cargando Gestor RMA...", font=("Segoe UI", 11, "bold"), bg="white").pack(pady=(16, 6))
            status_label = tk.Label(frame, text="Preparando la interfaz...", font=("Segoe UI", 9), bg="white")
            status_label.pack(pady=(0, 8))
            try:
                progress = ttk.Progressbar(frame, mode="indeterminate", length=320)
                progress.pack(pady=(0, 6))
                progress.start(10)
            except Exception:
                progress = None

            # Forzar dibujado del splash antes de crear la ventana principal
            try:
                splash.update()
            except Exception:
                pass
        except Exception:
            splash = None
            status_label = None
            progress = None

        try:
            if not hasattr(self, 'ventana_principal') or not self.ventana_principal.winfo_exists():
                # Instanciar la ventana principal (esto puede tardar si carga muchos datos)
                # Actualizar label de estado antes de la carga (si es posible)
                try:
                    if status_label:
                        status_label.config(text="Cargando datos y recursos...")
                        splash.update()
                except Exception:
                    pass

                self.ventana_principal = VentanaPrincipal(self, username, rol)

        finally:
            # Cerrar el splash y detener la barra si existen
            try:
                if progress:
                    progress.stop()
            except Exception:
                pass
            try:
                if splash:
                    splash.destroy()
            except Exception:
                pass

        print(ADVERTENCIA_MULTIUSUARIO)

# ----------------------------------------------------------------------
# 2. CLASE DE LA VENTANA PRINCIPAL DE LA APLICACIÓN
# ----------------------------------------------------------------------

class VentanaPrincipal(ctk.CTkToplevel):
    """Ventana principal que gestiona el listado, creación y edición de RMAs."""
    
    # Opciones predefinidas para desplegables
    # Inicializar el gestor de estados
    estados_manager = EstadosArticuloManager()
    # Inicializar el gestor de personas
    personas_manager = PersonasManager()
    
    # Inicializar el gestor de resultado expediente
    resultado_expediente_manager = ResultadoExpedienteManager()
    
    OPCIONES = {
        "Autorizacion": ["SI", "NO"],
        "Autorizado_Por": personas_manager.cargar_personas(),  # Cargar desde JSON
        "Gestionado_Por": personas_manager.cargar_personas(),  # Cargar desde JSON
        "Resultado_Expediente": resultado_expediente_manager.cargar_resultados(),  # Cargar desde JSON
        "Estado_Producto": estados_manager.cargar_estados()  # Cargar desde JSON
    }
    
    def get_color_por_estado(self, estado):
        """Obtiene el color del dashboard para un estado específico, manteniendo coherencia visual"""
        # Mapeo de colores del dashboard para coherencia visual
        colores_dashboard = {
            'Completado': "#27ae60",
            'Pendiente de Autorización': "#e74c3c", 
            'Pendiente de Autorizacion': "#e74c3c",  # Variación sin tilde
            'Recibido': "#3498db",
            'En Trámite': "#f39c12",
            'En Tramite': "#f39c12",  # Variación sin tilde
            'En Proceso': "#f39c12",  # Variación alternativa
            'Procesando': "#f39c12",  # Variación alternativa
            'Autorizado': "#9b59b6"
        }
        return colores_dashboard.get(estado, "#7f8c8d")  # Color gris por defecto
    
    def __init__(self, master, username, rol):
        super().__init__(master)
        self.master = master
        self.username = username
        self.rol = rol
        try:
            self.toaster = ToastNotifier() if ToastNotifier else None
        except Exception:
            self.toaster = None

        # Configuración para compatibilidad de esquema de BD
        self._usar_tipo_almacenamiento = True  # Por defecto asumimos que funciona

        # Inicializar referencias de iconos por defecto para evitar AttributeError
        # Icon placeholders (se asignarán en crear_diseno)
        self.icon_list = None
        self.icon_articles = None
        self.icon_stats = None
        self.icon_tasks = None
        self.icon_reports = None
        self.icon_info = None
        self.icon_edit = None
        self.icon_user = None
        self.icon_papel = None
        self.icon_mas = None
        # Cargar ajustes de usuario (ya aplicados en LoginApp, solo necesitamos cargarlos aquí)
        try:
            self.user_settings = load_user_settings(self.username)
        except Exception:
            self.user_settings = {}
            
        # Exponer a nivel de módulo para que Tooltip y otros lean la preferencia
        try:
            global USER_SETTINGS
            USER_SETTINGS = self.user_settings
        except Exception:
            pass
        
        # ----------------------------------------------------
        # 🛠️ AJUSTE DE PESO PARA EXPANDIR EL ÁREA DE TRABAJO 🛠️
        # ----------------------------------------------------
        # Configurar la expansión horizontal y vertical de la ventana principal
        
        # Columna 0 (Sidebar): Peso 0 (No se expande)
        self.master.grid_columnconfigure(0, weight=0) 
        
        # Columna 1 (Content Frame): Peso 1 (Ocupa todo el ancho restante)
        self.master.grid_columnconfigure(1, weight=1) 
        
        # Fila 0: Peso 1 (Ocupa toda la altura)
        self.master.grid_rowconfigure(0, weight=1)
        
        self.protocol("WM_DELETE_WINDOW", self.cerrar_app)
        # Título de la ventana principal
        try:
            self.title("Gestión RMA - Expedientes")
            # Agregar icono personalizado de ILUTREK
            self.iconbitmap("Icono_Ilutrek.ico")
        except Exception:
            # CTkToplevel puede usar wm_title alternativamente
            try:
                self.wm_title("Gestión RMA - Expedientes")
                self.iconbitmap("Icono_Ilutrek.ico")
            except Exception:
                pass  # Si no se puede cargar, continuar sin icono
            except Exception:
                pass
        
        self.crear_diseno()
        
        # Establecer tamaño inicial y centrar la ventana en la pantalla
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        # Tamaño de la ventana (80% de la pantalla o 1400x700, lo que sea mayor)
        window_width = max(1400, int(screen_width * 0.8))
        window_height = max(700, int(screen_height * 0.8))
        
        # Calcular posición central
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        # Aplicar geometría completa (tamaño + posición)
        self.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # Establecer tamaño mínimo para acomodar el dashboard
        self.minsize(1400, 700)
        
        # Configurar atajos de teclado globales
        self.bind_all("<Control-f>", lambda e: self.mostrar_busqueda_global())
        self.bind_all("<Control-F>", lambda e: self.mostrar_busqueda_global())
        self.bind_all("<Control-n>", lambda e: self.mostrar_nuevo_rma())
        self.bind_all("<Control-N>", lambda e: self.mostrar_nuevo_rma())
        # Refrescar listado con F5 (como en muchas aplicaciones)
        try:
            self.bind_all("<F5>", lambda e: self.aplicar_filtros_rma())
        except Exception:
            # Failsafe: no bloquear si bind falla en algún entorno
            pass
        
        # Iniciar comprobación periódica de tareas (notificaciones para el creador)
        try:
            self.programar_chequeo_tareas()
        except Exception:
            pass
        
        # Inicializar gestor de avisos
        try:
            root_path = os.path.dirname(os.path.abspath(__file__))
            self.avisos_manager = AvisosManager(root_path)
            # Mostrar aviso si está activo (con un pequeño delay para que la ventana principal se renderice primero)
            self.after(500, lambda: self.avisos_manager.mostrar_aviso_popup(self))
        except Exception as e:
            print(f"Error al inicializar avisos: {e}")

    def verificar_columna_motivo(self):
        """
        Verifica si la columna 'motivo' existe en rma_maestro y la añade si no.
        Esto es una migración simple para SQLite.
        NOTA: Solo funciona con SQLite local. Turso debe tener la estructura completa desde el dump.
        """
        # Si estamos usando Turso, saltar esta verificación (ALTER TABLE no funciona bien)
        if os.getenv("TURSO_DATABASE_URL"):
            print("ℹ️ Usando Turso - verificación de columnas omitida (usar dump SQL completo)")
            return
        
        conn, cursor = self.master.conectar_db()
        if not conn: return

        try:
            # Obtenemos las columnas actuales de la tabla rma_maestro
            cursor.execute("PRAGMA table_info('rma_maestro')")
            cols = [row[1] for row in cursor.fetchall()]

            # Columnas que queremos asegurar que existen (nombre en DB, tipo y default opcional)
            columnas_necesarias = {
                'motivo': "TEXT",
                'rma_proveedor': "TEXT DEFAULT ''",
                'modelo': "TEXT DEFAULT ''",
                'n_serie': "TEXT DEFAULT ''",
                'ref_proveedor': "TEXT DEFAULT ''",
                'obs_tecnica': "TEXT DEFAULT ''"
            }

            for col_name, col_def in columnas_necesarias.items():
                if col_name not in cols:
                    try:
                        cursor.execute(f"ALTER TABLE rma_maestro ADD COLUMN {col_name} {col_def}")
                    except Exception as e:
                        print(f"Error al añadir columna '{col_name}': {e}")

            conn.commit()

        except Exception as e:
            print(f"Error comprobando/añadiendo columnas en rma_maestro: {e}")
        finally:
            conn.close()
    
    
    def cerrar_app(self):
        """Maneja el cierre de la ventana principal y de toda la app."""
        self.master.destroy()

    def limpiar_contenido(self):
        """Limpia todos los widgets del marco de contenido principal."""
        try:
            if self.content_frame and self.content_frame.winfo_exists():
                for widget in self.content_frame.winfo_children():
                    try:
                        widget.destroy()
                    except:
                        pass
        except Exception as e:
            print(f"Error al limpiar contenido: {e}")

    def crear_diseno(self):
        """Define la estructura principal con un panel lateral y un área de contenido."""
        
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # --- Panel Lateral de Navegación (Columna 0) ---
        self.sidebar_frame = ctk.CTkFrame(self,
                                          width=100,
                                          corner_radius=0,
                                          #fg_color="gray90"
                                          )
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        
        ruta_logo = os.path.join(os.path.dirname(__file__), "Icono_Ilutrek.png")
        self.logo_image = ctk.CTkImage(light_image=Image.open(ruta_logo),
                                           dark_image=Image.open(ruta_logo),
                                           size=(100, 100)) # Ajusta este tamaño (width, height)
            
        # 💡 PASO 2: COLOCAR LA IMAGEN EN LA FILA 0
        # Usamos un CTkLabel para contener la imagen.
        self.logo_label = ctk.CTkLabel(self.sidebar_frame, 
                                       text="", 
                                       image=self.logo_image)
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))
        # Cargar iconos (outline 24x24) para botones icon-only
        ruta_icons = os.path.join(os.path.dirname(__file__), "icons")
        # Mantener referencias para evitar GC (ImageTk objects must be referenced)
        self._icon_refs = {}

        def _ensure_icon_png(fname, shape="rect", fg=(43,108,176,255)):
            """Asegura que exista un PNG válido en icons/fname. Si no existe o está corrupto,
            genera uno simple de 24x24 y lo guarda en disco.
            """
            path = os.path.join(ruta_icons, fname)
            try:
                if os.path.exists(path):
                    # Intentar abrir con PIL para verificar integridad
                    Image.open(path).verify()
                    return True
            except Exception:
                # Archivo corrupto o no imagen válida -> reescribir
                pass

            # Crear carpeta si no existe
            try:
                os.makedirs(ruta_icons, exist_ok=True)
            except Exception:
                pass

            # Generar imagen sencilla
            try:
                img = Image.new("RGBA", (24, 24), (0, 0, 0, 0))
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                if shape == "list":
                    draw.line([4, 7, 20, 7], fill=fg, width=2)
                    draw.line([4, 12, 20, 12], fill=fg, width=2)
                    draw.line([4, 17, 20, 17], fill=fg, width=2)
                elif shape == "dot":
                    draw.ellipse([6, 6, 18, 18], outline=fg, width=2)
                elif shape == "pencil":
                    draw.line([6, 17, 17, 6], fill=fg, width=2)
                    draw.polygon([(17,6),(19,8),(15,10)], fill=fg)
                else:
                    draw.rectangle([4, 4, 20, 20], outline=fg, width=2)

                img.save(path, format="PNG")
                return True
            except Exception:
                return False

        def _load_icon(fname):
            path = os.path.join(ruta_icons, fname)
            if not os.path.exists(path):
                return None
            try:
                # Prefer CTkImage (works well with customtkinter)
                img = ctk.CTkImage(light_image=Image.open(path), dark_image=Image.open(path), size=(24, 24))
                self._icon_refs[fname] = img
                return img
            except Exception:
                try:
                    # Fallback to ImageTk.PhotoImage if CTkImage fails
                    pil = Image.open(path).resize((24, 24), Image.LANCZOS)
                    tkimg = ImageTk.PhotoImage(pil)
                    self._icon_refs[fname] = tkimg
                    return tkimg
                except Exception:
                    return None

        def _make_placeholder_icon(key=None, shape="rect", fg="#2b6cb0", bg=None):
            """Genera un icono simple de 24x24 con PIL y lo convierte a PhotoImage/CTkImage.
            shape: 'rect', 'dot', 'pencil' (simple), 'list'
            """
            try:
                img = Image.new("RGBA", (24, 24), bg if bg is not None else (0, 0, 0, 0))
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                if shape == "rect":
                    draw.rectangle([4, 4, 20, 20], outline=fg, width=2)
                elif shape == "dot":
                    draw.ellipse([6, 6, 18, 18], outline=fg, width=2)
                elif shape == "list":
                    # three horizontal lines
                    draw.line([5, 7, 19, 7], fill=fg, width=2)
                    draw.line([5, 12, 19, 12], fill=fg, width=2)
                    draw.line([5, 17, 19, 17], fill=fg, width=2)
                elif shape == "pencil":
                    # simple pencil: diagonal line + tip
                    draw.line([6, 17, 17, 6], fill=fg, width=2)
                    draw.polygon([(17,6),(19,8),(15,10)], fill=fg)
                else:
                    draw.rectangle([4, 4, 20, 20], outline=fg, width=2)

                # Try CTkImage first
                try:
                    ctki = ctk.CTkImage(light_image=img, dark_image=img, size=(24, 24))
                    keyname = f"gen-{key or shape}"
                    self._icon_refs[keyname] = ctki
                    return ctki
                except Exception:
                    tkimg = ImageTk.PhotoImage(img)
                    keyname = f"gen-{key or shape}"
                    self._icon_refs[keyname] = tkimg
                    return tkimg
            except Exception:
                return None

        # Asegurar que existan ficheros PNG válidos en disco (sobrescribe corruptos)
        _ensure_icon_png("ic_list_outline_24.png", shape="list")
        _ensure_icon_png("ic_articles_outline_24.png", shape="rect")
        _ensure_icon_png("ic_stats_outline_24.png", shape="dot")
        _ensure_icon_png("ic_tasks_outline_24.png", shape="rect")
        _ensure_icon_png("ic_reports_outline_24.png", shape="rect")
        _ensure_icon_png("boton-de-informacion.png", shape="dot")
        _ensure_icon_png("ic_edit_outline_24.png", shape="pencil")
        # Additional user-provided icons
        _ensure_icon_png("user.png", shape="rect")
        _ensure_icon_png("papel.png", shape="rect")
        _ensure_icon_png("mas.png", shape="rect")
        _ensure_icon_png("settings.png", shape="rect")

        # Cargar iconos (preferir CTkImage, fallback a ImageTk)
        self.icon_list = _load_icon("ic_list_outline_24.png") or _make_placeholder_icon("list", shape="list")
        self.icon_articles = _load_icon("ic_articles_outline_24.png") or _make_placeholder_icon("articles", shape="rect")
        self.icon_stats = _load_icon("ic_stats_outline_24.png") or _make_placeholder_icon("stats", shape="dot")
        self.icon_tasks = _load_icon("ic_tasks_outline_24.png") or _make_placeholder_icon("tasks", shape="rect")
        self.icon_reports = _load_icon("ic_reports_outline_24.png") or _make_placeholder_icon("reports", shape="rect")
        self.icon_info = _load_icon("boton-de-informacion.png") or _make_placeholder_icon("info", shape="dot")
        self.icon_edit = _load_icon("ic_edit_outline_24.png") or _make_placeholder_icon("edit", shape="pencil")
        # Load additional icons (user, papel, mas)
        self.icon_user = _load_icon("user.png")
        self.icon_papel = _load_icon("papel.png")
        self.icon_mas = _load_icon("mas.png")
        self.icon_settings = _load_icon("settings.png")
        self.icon_busqueda = _load_icon("busqueda.png")
        self.icon_bd = _load_icon("bd.png")
        self.icon_report = _load_icon("report.png")
        
        # ... (Botones btn_lista, btn_buscar, btn_reportar en filas 1, 2, 3) ...

        # Fila alta: Espacio vacío para empujar los elementos inferiores.
        # Reservamos una fila alta para 'push' y colocamos la info de usuario en filas finales.
        self.sidebar_frame.grid_rowconfigure(20, weight=1) # <--- Espaciador principal.

        # Fila alta: Espacio vacío para empujar los elementos inferiores.
        self.sidebar_frame.grid_rowconfigure(50, weight=1)  # Uso fila 50 para tener espacio para más botones

        # Información de Usuario y Rol (al final)
        self.lbl_usuario_rol = ctk.CTkLabel(self.sidebar_frame, text=f"Usuario: {self.username} ({self.rol})", font=ctk.CTkFont(size=12))
        self.lbl_usuario_rol.grid(row=51, column=0, padx=20, pady=(10, 5), sticky="s")

        # Versión de la App
        ctk.CTkLabel(self.sidebar_frame, text=f"Versión: {APP_VERSION}", font=ctk.CTkFont(size=10, slant="italic")).grid(row=52, column=0, padx=20, pady=(0, 2), sticky="s")

        # Copyright
        año_actual = datetime.datetime.now().year
        ctk.CTkLabel(self.sidebar_frame, text=f"© {año_actual} ILUTREK, S.L.", font=ctk.CTkFont(size=10, weight="bold")).grid(row=53, column=0, padx=20, pady=(2, 10), sticky="s")
        
        ctk.CTkLabel(self.sidebar_frame, text="MENÚ", font=ctk.CTkFont(family="Verdana", size=20, weight="bold")).grid(row=1, column=0, padx=20, pady=(20, 10))

        # Usamos índices de fila secuenciales para evitar solapamientos.
        fila = 2

        # Botón Búsqueda Global - busca en todos los campos
        sidebar_bg = self.sidebar_frame.cget("fg_color") if hasattr(self.sidebar_frame, 'cget') else None
        self.btn_busqueda_global = ctk.CTkButton(self.sidebar_frame,
                                               text="",
                                               image=self.icon_busqueda,
                                               width=44,
                                               height=44,
                                               fg_color=sidebar_bg,
                                               hover_color=sidebar_bg,
                                               command=self.mostrar_busqueda_global)
        self.btn_busqueda_global.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_busqueda_global, "Búsqueda Global - Buscar en todos los campos")
        fila += 1

        # Botón de gestión de usuarios (solo visible para administradores)
        if str(self.rol).strip().lower() in ("admin", "administrador"):
            sidebar_bg = self.sidebar_frame.cget("fg_color") if hasattr(self.sidebar_frame, 'cget') else None
            self.btn_usuarios = ctk.CTkButton(self.sidebar_frame,
                                             text="",
                                             image=(self.icon_user or self.icon_tasks),
                                             width=44,
                                             height=44,
                                             fg_color=sidebar_bg if sidebar_bg is not None else None,
                                             hover_color=sidebar_bg if sidebar_bg is not None else None,
                                             command=self.mostrar_gestion_usuarios)
            self.btn_usuarios.grid(row=fila, column=0, padx=12, pady=6)
            Tooltip(self.btn_usuarios, "Gestión de usuarios")
            fila += 1
            
        sidebar_bg = self.sidebar_frame.cget("fg_color") if hasattr(self.sidebar_frame, 'cget') else None

        self.btn_lista = ctk.CTkButton(self.sidebar_frame,
                                       text="",
                                       image=self.icon_list,
                                       width=44,
                                       height=44,
                                       fg_color=sidebar_bg,
                                       hover_color=sidebar_bg,
                                       command=self.mostrar_lista_rma)
        self.btn_lista.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_lista, "Listado de expedientes")
        fila += 1

        # Botón Artículos: abre ventana con listado de artículos y conteo de expedientes asociados
        self.btn_articulos = ctk.CTkButton(self.sidebar_frame,
                                           text="",
                                           image=self.icon_articles,
                                           width=44,
                                           height=44,
                                           fg_color=sidebar_bg,
                                           hover_color=sidebar_bg,
                                           command=self.mostrar_articulos_window)
        self.btn_articulos.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_articulos, "Artículos")
        fila += 1

        self.btn_estadisticas = ctk.CTkButton(self.sidebar_frame,
                                              text="",
                                              image=self.icon_stats,
                                              width=44,
                                              height=44,
                                              fg_color=sidebar_bg,
                                              hover_color=sidebar_bg,
                                              command=self.mostrar_ventana_estadisticas)
        self.btn_estadisticas.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_estadisticas, "Filtrar / Estadísticas")
        fila += 1

        # Botón de Tareas (lista y creación de tareas por expediente) con badge
        # Crear frame contenedor para el botón + badge
        self.frame_tareas = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.frame_tareas.grid(row=fila, column=0, padx=20, pady=6)
        
        self.btn_tareas = ctk.CTkButton(self.frame_tareas,
                                        text="",
                                        image=self.icon_tasks,
                                        width=44,
                                        height=44,
                                        fg_color=sidebar_bg,
                                        hover_color=sidebar_bg,
                                        command=self.mostrar_gestion_tareas)
        self.btn_tareas.grid(row=0, column=0)
        
        # Badge para contador de tareas pendientes
        self.badge_tareas = ctk.CTkLabel(self.frame_tareas,
                                       text="0",
                                       width=18,
                                       height=18,
                                       corner_radius=9,
                                       fg_color="#e74c3c",  # Rojo más intenso
                                       text_color="white",
                                       font=ctk.CTkFont(size=9, weight="bold"))
        self.badge_tareas.grid(row=0, column=0, sticky="ne", padx=(32, 0), pady=(2, 0))
        
        # Inicializar badge (oculto inicialmente)
        self.badge_tareas.grid_remove()
        
        Tooltip(self.btn_tareas, "Tareas")
        fila += 1

        # Botón de Gestión de Clientes
        self.btn_clientes = ctk.CTkButton(self.sidebar_frame,
                                        text="",
                                        image=self.icon_user,  # Usar icono de usuario para clientes
                                        width=44,
                                        height=44,
                                        fg_color=sidebar_bg,
                                        hover_color=sidebar_bg,
                                        command=self.mostrar_clientes)
        self.btn_clientes.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_clientes, "Gestión de Clientes")
        fila += 1

        # Botón Gestión RMP (Proveedores -> Expedientes)
        self.btn_gestion_rmp = ctk.CTkButton(self.sidebar_frame,
                                             text="",
                                             image=(self.icon_papel or self.icon_reports),
                                             width=44,
                                             height=44,
                                             fg_color=sidebar_bg,
                                             hover_color=sidebar_bg,
                                             command=self.mostrar_gestion_rmp)
        self.btn_gestion_rmp.grid(row=fila, column=0, padx=12, pady=6)
        Tooltip(self.btn_gestion_rmp, "Gestión RMP")
        fila += 1

        # Mostrar botón de Backup solo para administradores y Dpto. Tecnico
        rol_norm = str(self.rol).strip().lower()
        if rol_norm in ("admin", "administrador", "dpto. tecnico", "dpto tecnico", "dpto técnico"):
            self.btn_buscar = ctk.CTkButton(self.sidebar_frame,
                                           text="",
                                           image=self.icon_bd,
                                           width=44,
                                           height=44,
                                           fg_color=sidebar_bg,
                                           hover_color=sidebar_bg,
                                           command=self.mostrar_gestor_backups)
            self.btn_buscar.grid(row=fila, column=0, padx=20, pady=6)
            Tooltip(self.btn_buscar, "Gestor de Backups en Backblaze B2")
            fila += 1

        self.btn_reportar = ctk.CTkButton(self.sidebar_frame,
                                          text="",
                                          image=self.icon_report,
                                          width=44,
                                          height=44,
                                          fg_color=sidebar_bg,
                                          hover_color=sidebar_bg,
                                          command=self.mostrar_formulario_github)
        self.btn_reportar.grid(row=fila, column=0, padx=20, pady=6)
        Tooltip(self.btn_reportar, "Reportar un problema / abrir GitHub")
        
        # Botón Ajustes - abre diálogo de preferencias del usuario
        try:
            fila += 1
            self.btn_ajustes = ctk.CTkButton(self.sidebar_frame,
                                             text="",
                                             image=(self.icon_settings or self.icon_info),
                                             width=44,
                                             height=44,
                                             fg_color=sidebar_bg if sidebar_bg is not None else None,
                                             hover_color=sidebar_bg if sidebar_bg is not None else None,
                                             command=self.mostrar_ajustes)
            self.btn_ajustes.grid(row=fila, column=0, padx=20, pady=6)
            Tooltip(self.btn_ajustes, "Ajustes")
        except Exception:
            pass
        
        # Botón de menú Admin (solo para administradores) - AL FINAL
        if str(self.rol).strip().lower() in ("admin", "administrador"):
            fila += 1
            self.btn_admin_menu = ctk.CTkButton(self.sidebar_frame,
                                               text="⚙️ Admin",
                                               width=100,
                                               height=35,
                                               font=ctk.CTkFont(size=12, weight="bold"),
                                               fg_color="#ff6b35",
                                               hover_color="#e55a2b",
                                               command=self.mostrar_menu_admin)
            self.btn_admin_menu.grid(row=fila, column=0, padx=12, pady=8)
            Tooltip(self.btn_admin_menu, "Funciones de Administración")
        
        # --- Contenido Principal (Columna 1) ---
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent") # 'transparent' para que herede el fondo 'Light' (blanco)
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        self.content_frame.grid_columnconfigure(0, weight=1)
        
        # Establecer tamaño mínimo de ventana para acomodar el dashboard
        self.minsize(1400, 700)
        
        self.mostrar_lista_rma()
        
        # Actualizar badge de tareas al cargar la ventana
        self.actualizar_badge_tareas()
    
    def contar_tareas_pendientes(self):
        """Cuenta las tareas que no están completadas (Pendiente + En progreso)."""
        try:
            conn, cursor = self.conectar_db()
            if not conn:
                return 0
            
            cursor.execute("""
                SELECT COUNT(*) FROM tareas 
                WHERE estado NOT IN ('Completado', 'Completada', 'Finalizada')
                AND estado IS NOT NULL
            """)
            count = int(cursor.fetchone()[0])  # Convertir a entero
            conn.close()
            return count
        except Exception as e:
            print(f"Error al contar tareas pendientes: {e}")
            return 0

    def verificar_tareas_pendientes_expediente(self, rma_id):
        """Verifica si un expediente específico tiene tareas pendientes."""
        try:
            conn, cursor = self.conectar_db()
            if not conn:
                return 0, []
            
            # Primero obtenemos el código RMA del expediente
            cursor.execute("SELECT codigo_rma FROM rma_maestro WHERE id = ?", (rma_id,))
            resultado_rma = cursor.fetchone()
            if not resultado_rma:
                conn.close()
                return 0, []
            
            codigo_rma = resultado_rma[0]
            
            # Ahora buscamos tareas pendientes usando el código RMA
            cursor.execute("""
                SELECT COUNT(*), GROUP_CONCAT(titulo || ' (' || estado || ')') as tareas_pendientes
                FROM tareas 
                WHERE codigo_rma = ? 
                AND estado NOT IN ('Completado', 'Completada', 'Finalizada')
                AND estado IS NOT NULL
            """, (codigo_rma,))
            
            resultado = cursor.fetchone()
            count = int(resultado[0]) if resultado[0] else 0
            tareas_pendientes = resultado[1].split(',') if resultado[1] else []
            
            conn.close()
            return count, tareas_pendientes
        except Exception as e:
            print(f"Error al verificar tareas pendientes del expediente: {e}")
            return 0, []

    def actualizar_badge_tareas(self):
        """Actualiza el badge visual del botón de tareas."""
        if not hasattr(self, 'badge_tareas'):
            return
        
        count = self.contar_tareas_pendientes()
        if count > 0:
            # Mostrar badge con el número
            self.badge_tareas.configure(text=str(count))
            self.badge_tareas.grid()  # Hacer visible
            # Actualizar tooltip con información de tareas pendientes
            Tooltip(self.btn_tareas, f"Tareas ({count} pendientes)")
        else:
            # Ocultar badge si no hay tareas pendientes
            self.badge_tareas.grid_remove()
            # Restaurar tooltip normal
            Tooltip(self.btn_tareas, "Tareas")

    def conectar_db(self):
        """Intenta conectar a la base de datos (método heredado de master)."""
        return self.master.conectar_db()

    def cargar_historial_busquedas(self):
        """Carga el historial de búsquedas del usuario desde configuración."""
        if not hasattr(self, 'historial_busquedas'):
            self.historial_busquedas = self.user_settings.get("historial_busquedas", [])
        return self.historial_busquedas[:10]  # Máximo 10 búsquedas recientes

    def guardar_busqueda_en_historial(self, termino, filtros=None):
        """Guarda una nueva búsqueda en el historial."""
        if not termino.strip():
            return
        
        # Crear entrada del historial
        entrada = {
            "termino": termino.strip(),
            "filtros": filtros or {},
            "fecha": datetime.datetime.now().isoformat()
        }
        
        # Cargar historial actual
        if not hasattr(self, 'historial_busquedas'):
            self.historial_busquedas = self.user_settings.get("historial_busquedas", [])
        
        # Remover búsqueda duplicada si existe
        self.historial_busquedas = [h for h in self.historial_busquedas 
                                   if h.get("termino") != termino.strip()]
        
        # Añadir al principio
        self.historial_busquedas.insert(0, entrada)
        
        # Mantener solo las últimas 10
        self.historial_busquedas = self.historial_busquedas[:10]
        
        # Guardar en configuración
        self.user_settings["historial_busquedas"] = self.historial_busquedas
        save_user_settings(self.user_settings, self.username)

    def limpiar_historial_busquedas(self):
        """Limpia todo el historial de búsquedas."""
        self.historial_busquedas = []
        self.user_settings["historial_busquedas"] = []
        save_user_settings(self.user_settings, self.username)
    
    def actualizar_dashboard(self, año=None):
        """Actualiza las estadísticas del dashboard para el año seleccionado."""
        if año is None:
            año = self.combo_año_dashboard.get()
        
        # Limpiar estadísticas actuales de forma más exhaustiva
        try:
            for widget in self.stats_frame.winfo_children():
                widget.destroy()
            # Forzar actualización del frame
            self.stats_frame.update_idletasks()
        except Exception as e:
            print(f"Error limpiando dashboard: {e}")
        
        # Obtener estadísticas de la base de datos
        stats = self.obtener_estadisticas_expedientes(año)
        
        # Obtener artículos problemáticos
        periodo = self.combo_periodo.get()
        articulos_problematicos = self.obtener_articulos_problematicos(año, periodo)
        
        # Crear interfaz de estadísticas
        self.crear_interfaz_estadisticas(stats, articulos_problematicos)
    
    def obtener_estadisticas_expedientes(self, año):
        """Obtiene las estadísticas de expedientes para el año especificado."""
        try:
            conn, cursor = self.conectar_db()
            if not conn:
                return {}
            
            # Definir los estados a consultar (incluyendo variaciones posibles)
            estados_mapeados = {
                'Completado': ['Completado', 'Finalizado', 'Cerrado'],
                'Pendiente de Autorización': ['Pendiente de Autorización', 'Pendiente de Autorizacion', 'Pendiente Autorización', 'Pendiente', 'Sin Autorizar'],
                'Recibido': ['Recibido', 'Recepcionado'],
                'En Trámite': ['En Trámite', 'En Tramite', 'En Proceso', 'Procesando'],
                'Autorizado': ['Autorizado', 'Aprobado']
            }
            
            estadisticas = {}
            
            for estado_display, variaciones in estados_mapeados.items():
                count = 0
                for variacion in variaciones:
                    cursor.execute("""
                        SELECT COUNT(*) 
                        FROM rma_maestro 
                        WHERE estado = ? 
                        AND strftime('%Y', fecha_emision) = ?
                    """, (variacion, año))
                    
                    result = cursor.fetchone()[0]
                    count += int(result) if result is not None else 0
                
                estadisticas[estado_display] = count
            
            # Obtener total del año
            cursor.execute("""
                SELECT COUNT(*) 
                FROM rma_maestro 
                WHERE strftime('%Y', fecha_emision) = ?
            """, (año,))
            
            total_result = cursor.fetchone()[0]
            estadisticas['Total'] = int(total_result) if total_result is not None else 0
            
            conn.close()
            return estadisticas
            
        except Exception as e:
            print(f"Error obteniendo estadísticas: {e}")
            return {}
    
    def obtener_articulos_problematicos(self, año, periodo):
        """Obtiene los 10 artículos con más problemas según el período especificado."""
        try:
            conn, cursor = self.conectar_db()
            if not conn:
                return []
            
            # Estados problemáticos a considerar (los que realmente aparecerán)
            estados_problematicos = [
                "NO FUNCIONA, ABONAR",
                "NO FUNCIONA ; NO ABONAR",
                "REPOSICION FALLO PRODUCTO", 
                "REPOSICION ; ABONAR",
                "FALLO SOLDADURA ; ABONAR",
                "FALLO SOLDADURA ; NO ABONAR",
                "FALLO MODULO ; ABONAR"
            ]
            
            # Crear placeholders para la consulta SQL
            placeholders = ','.join(['?' for _ in estados_problematicos])
            
            # Determinar condición de fecha según el período
            fecha_condicion = ""
            if periodo == "Trimestral":
                # Trimestre actual basado en el mes actual
                mes_actual = datetime.datetime.now().month
                if mes_actual <= 3:
                    trimestre = 1
                elif mes_actual <= 6:
                    trimestre = 2
                elif mes_actual <= 9:
                    trimestre = 3
                else:
                    trimestre = 4
                
                mes_inicio = (trimestre - 1) * 3 + 1
                mes_fin = trimestre * 3
                fecha_condicion = f"AND strftime('%Y', rm.fecha_emision) = '{año}' AND CAST(strftime('%m', rm.fecha_emision) AS INTEGER) BETWEEN {mes_inicio} AND {mes_fin}"
                
            elif periodo == "Semestral":
                # Primer semestre (1-6) o segundo semestre (7-12) según el mes actual
                mes_actual = datetime.datetime.now().month
                if mes_actual <= 6:
                    semestre = 1
                    fecha_condicion = f"AND strftime('%Y', rm.fecha_emision) = '{año}' AND CAST(strftime('%m', rm.fecha_emision) AS INTEGER) BETWEEN 1 AND 6"
                else:
                    semestre = 2
                    fecha_condicion = f"AND strftime('%Y', rm.fecha_emision) = '{año}' AND CAST(strftime('%m', rm.fecha_emision) AS INTEGER) BETWEEN 7 AND 12"
                    
            else:  # Anual
                fecha_condicion = f"AND strftime('%Y', rm.fecha_emision) = '{año}'"
            
            # Consulta para obtener artículos problemáticos
            query = f"""
                SELECT rd.referencia_articulo, COUNT(*) as problemas
                FROM rma_detalles rd
                JOIN rma_maestro rm ON rd.rma_id = rm.id
                WHERE rd.estado_producto IN ({placeholders})
                {fecha_condicion}
                GROUP BY rd.referencia_articulo
                ORDER BY problemas DESC
                LIMIT 10
            """
            
            cursor.execute(query, estados_problematicos)
            resultados = cursor.fetchall()
            
            # Convertir a lista de diccionarios para facilitar el manejo
            articulos_problematicos = [
                {
                    'referencia_articulo': row[0],
                    'problemas': int(row[1])
                }
                for row in resultados
            ]
            
            conn.close()
            return articulos_problematicos
            
        except Exception as e:
            print(f"Error obteniendo artículos problemáticos: {e}")
            return []
    
    def crear_interfaz_estadisticas(self, stats, articulos_problematicos):
        """Crea la interfaz visual para mostrar las estadísticas de forma simple y rápida."""
        # Verificar que el frame exista y esté limpio
        if not hasattr(self, 'stats_frame') or not self.stats_frame.winfo_exists():
            return
            
        # Doble verificación: limpiar cualquier widget remanente
        for widget in self.stats_frame.winfo_children():
            widget.destroy()
        
        # Título del año
        año_label = ctk.CTkLabel(self.stats_frame, 
                                text=f"Expedientes {self.combo_año_dashboard.get()}", 
                                font=ctk.CTkFont(size=12, weight="bold"))
        año_label.pack(pady=(8, 12))
        
        # Lista simple de estadísticas sin colores complejos
        estadisticas_texto = [
            f"📋 Total: {stats.get('Total', 0)}",
            f"✅ Completado: {stats.get('Completado', 0)}",
            f"⏳ Pend. Autor.: {stats.get('Pendiente de Autorización', 0)}",
            f"📥 Recibido: {stats.get('Recibido', 0)}",
            f"🔄 En Trámite: {stats.get('En Trámite', 0)}",
            f"✔️ Autorizado: {stats.get('Autorizado', 0)}"
        ]
        
        # Mostrar cada estadística en una fila simple
        for texto in estadisticas_texto:
            fila_label = ctk.CTkLabel(self.stats_frame, 
                                    text=texto, 
                                    font=ctk.CTkFont(size=10),
                                    anchor="w")
            fila_label.pack(fill="x", padx=10, pady=1)
        
        # Separador simple
        separador = ctk.CTkLabel(self.stats_frame, text="─" * 25, 
                               font=ctk.CTkFont(size=8))
        separador.pack(pady=(8, 4))
        
        # Título de artículos problemáticos simplificado
        periodo_texto = self.combo_periodo.get()
        titulo_problematicos = ctk.CTkLabel(self.stats_frame, 
                                          text=f"🔴 Problemáticos ({periodo_texto})", 
                                          font=ctk.CTkFont(size=10, weight="bold"))
        titulo_problematicos.pack(pady=(0, 4))
        
        # Lista simple de artículos problemáticos (solo top 5 para mejor rendimiento)
        if articulos_problematicos:
            for i, articulo in enumerate(articulos_problematicos[:5], 1):
                # Texto simple sin marcos complejos
                problema_texto = f"{i}. {articulo['referencia_articulo'][:15]}... ({articulo['problemas']})"
                problema_label = ctk.CTkLabel(self.stats_frame, 
                                            text=problema_texto, 
                                            font=ctk.CTkFont(size=9),
                                            anchor="w")
                problema_label.pack(fill="x", padx=10, pady=1)
        else:
            no_datos_label = ctk.CTkLabel(self.stats_frame, 
                                        text="Sin datos", 
                                        font=ctk.CTkFont(size=9))
            no_datos_label.pack(pady=2)
        
        # Botón de actualización simple
        btn_actualizar = ctk.CTkButton(self.stats_frame, 
                                     text="� Actualizar",
                                     command=self.actualizar_dashboard,
                                     width=80, height=25,
                                     font=ctk.CTkFont(size=10))
        btn_actualizar.pack(pady=(10, 8))
    
    
    # ----------------------------------------------------------------------
    # 3. MÉTODOS AUXILIARES Y GENERACIÓN DE CÓDIGO RMA
    # ----------------------------------------------------------------------

    def obtener_temas_disponibles(self):
        """Obtiene la lista de temas disponibles desde la carpeta themes/"""
        import glob
        import os
        try:
            temas = []
            # Buscar archivos .json en la carpeta themes
            archivos_tema = glob.glob("themes/*.json")
            for archivo in archivos_tema:
                # Normalizar la ruta y extraer solo el nombre del archivo sin la extensión
                nombre_archivo = os.path.basename(archivo)
                nombre_tema = os.path.splitext(nombre_archivo)[0]
                # Convertir nombres a formato más amigable
                if nombre_tema.lower() == "bh_rime":
                    display = "BH Rime (Predeterminado)"
                else:
                    # Reemplazar guiones/underscores por espacios y title case
                    display = nombre_tema.replace('_', ' ').replace('-', ' ').title()
                temas.append(display)

            # Asegurar que BH Rime esté primero
            temas_ordenados = []
            if any(t.startswith("BH Rime") for t in temas):
                # encontrar la primera que contenga 'BH Rime'
                for t in temas:
                    if t.startswith("BH Rime"):
                        temas_ordenados.append(t)
                        temas.remove(t)
                        break
            temas_ordenados.extend(sorted(temas))

            return temas_ordenados if temas_ordenados else ["BH Rime (Predeterminado)"]
        except Exception:
            return ["BH Rime (Predeterminado)", "Rime", "Metal", "Pink", "Red"]

    def tema_display_a_archivo(self, tema_display):
        """Convierte el nombre mostrado del tema al nombre del archivo"""
        # Limpiar el nombre de entrada por si tiene rutas o caracteres extraños
        tema_limpio = tema_display.replace("Themes\\", "").replace("Themes/", "").replace("themes\\", "").replace("themes/", "").strip()

        # Mapeo rápido para nombres comunes
        mapping = {
            "BH Rime (Predeterminado)": "BH_rime.json",
            "Rime": "rime.json",
            "Metal": "metal.json",
            "Pink": "pink.json",
            "Red": "red.json"
        }
        if tema_limpio in mapping:
            return mapping[tema_limpio]

        # Intentar buscar un archivo en themes/ que coincida con el display
        import glob, os
        posibles = glob.glob("themes/*.json")
        # Normalizar display para comparar
        display_norm = tema_limpio.lower().replace(' ', '').replace('_', '').replace('-', '')
        for p in posibles:
            base = os.path.splitext(os.path.basename(p))[0]
            if base.lower() == tema_limpio.lower().replace(' ', '_') or base.lower() == tema_limpio.lower():
                return os.path.basename(p)
            if base.lower().replace('_', '').replace('-', '') == display_norm:
                return os.path.basename(p)

        # Intentar construir nombres probables
        cand1 = f"{tema_limpio}.json"
        cand2 = f"{tema_limpio.replace(' ', '_')}.json"
        if os.path.exists(os.path.join('themes', cand1)):
            return cand1
        if os.path.exists(os.path.join('themes', cand2)):
            return cand2

        # Fallback: BH_rime
        return 'BH_rime.json'

    def archivo_a_tema_display(self, archivo_tema):
        """Convierte el nombre del archivo al nombre mostrado"""
        import os
        if not archivo_tema:
            return "BH Rime (Predeterminado)"
        # Si viene con path, extraer basename
        archivo = os.path.basename(archivo_tema)
        name = os.path.splitext(archivo)[0]
        if name.lower() == 'bh_rime':
            return 'BH Rime (Predeterminado)'
        # Normalizar: underscore/dash -> space, title case
        display = name.replace('_', ' ').replace('-', ' ').title()
        return display

    def mostrar_ajustes(self):
        """Abre un diálogo modal para que el usuario modifique sus preferencias."""
        try:
            dlg = ctk.CTkToplevel(self)
            dlg.transient(self)
            dlg.grab_set()
            dlg.title("Ajustes")
            
            # Agregar icono personalizado
            try:
                dlg.iconbitmap("Icono_Ilutrek.ico")
            except Exception:
                pass

            frm = ctk.CTkFrame(dlg, fg_color="transparent")
            frm.grid(row=0, column=0, padx=12, pady=12, sticky="nsew")
            frm.grid_columnconfigure(1, weight=1)

            # Date format
            ctk.CTkLabel(frm, text="Formato de fecha:").grid(row=0, column=0, sticky="w", padx=6, pady=6)
            date_values = ["YYYY-MM-DD", "DD/MM/YYYY", "MM/DD/YYYY"]
            date_menu = ctk.CTkOptionMenu(frm, values=date_values, width=180)
            date_menu.set(self.user_settings.get("date_format", "YYYY-MM-DD"))
            date_menu.grid(row=0, column=1, sticky="e", padx=6, pady=6)

            # Tooltips
            ctk.CTkLabel(frm, text="Mostrar tooltips:").grid(row=1, column=0, sticky="w", padx=6, pady=6)
            var_tooltips = tk.BooleanVar(value=self.user_settings.get("show_tooltips", True))
            switch_tooltips = ctk.CTkSwitch(frm, text="", variable=var_tooltips, width=40)
            switch_tooltips.grid(row=1, column=1, sticky="w", padx=6, pady=6)

            # Compact mode
            ctk.CTkLabel(frm, text="Modo compacto:").grid(row=2, column=0, sticky="w", padx=6, pady=6)
            var_compact = tk.BooleanVar(value=self.user_settings.get("compact_mode", True))
            switch_compact = ctk.CTkSwitch(frm, text="", variable=var_compact, width=40)
            switch_compact.grid(row=2, column=1, sticky="w", padx=6, pady=6)

            # Tema de la aplicación
            ctk.CTkLabel(frm, text="Tema:").grid(row=3, column=0, sticky="w", padx=6, pady=6)
            temas_disponibles = self.obtener_temas_disponibles()
            tema_menu = ctk.CTkOptionMenu(frm, values=temas_disponibles, width=250)
            # Obtener tema actual del usuario
            tema_actual = self.user_settings.get("theme", "themes/BH_rime.json")
            tema_display_actual = self.archivo_a_tema_display(tema_actual.replace("themes/", ""))
            if tema_display_actual in temas_disponibles:
                tema_menu.set(tema_display_actual)
            else:
                tema_menu.set("BH Rime (Predeterminado)")
            tema_menu.grid(row=3, column=1, sticky="ew", padx=6, pady=6)

            # Modo claro/oscuro (solo para temas que no sean BH_rime)
            ctk.CTkLabel(frm, text="Modo:").grid(row=4, column=0, sticky="w", padx=6, pady=6)
            modo_values = ["Claro", "Oscuro"]
            modo_menu = ctk.CTkOptionMenu(frm, values=modo_values, width=150)
            modo_actual = self.user_settings.get("appearance_mode", "light")
            modo_menu.set("Claro" if modo_actual == "light" else "Oscuro")
            modo_menu.grid(row=4, column=1, sticky="w", padx=6, pady=6)

            # Función para habilitar/deshabilitar el selector de modo según el tema
            def actualizar_modo_disponible(*args):
                tema_seleccionado = tema_menu.get()
                if tema_seleccionado == "BH Rime (Predeterminado)":
                    # BH_rime solo funciona en modo claro
                    modo_menu.set("Claro")
                    modo_menu.configure(state="disabled")
                else:
                    modo_menu.configure(state="normal")
            
            # Conectar el evento de cambio de tema
            tema_menu.configure(command=actualizar_modo_disponible)
            # Aplicar estado inicial
            actualizar_modo_disponible()

            # Email
            ctk.CTkLabel(frm, text="Email:").grid(row=5, column=0, sticky="w", padx=6, pady=6)
            entry_email = ctk.CTkEntry(frm, width=300)
            # Try to prefill from DB if exists
            try:
                conn, cursor = self.master.conectar_db()
                if conn and cursor:
                    cursor.execute("PRAGMA table_info('usuarios')")
                    cols = [r[1] for r in cursor.fetchall()]
                    if 'email' in cols:
                        cursor.execute("SELECT email FROM usuarios WHERE nombre_usuario = ?", (self.username,))
                        row = cursor.fetchone()
                        if row and row[0]:
                            entry_email.insert(0, row[0])
                    # close connection
                    try:
                        conn.close()
                    except Exception:
                        pass
            except Exception:
                pass
            entry_email.grid(row=5, column=1, sticky="ew", padx=6, pady=6)

            # New password
            ctk.CTkLabel(frm, text="Nueva contraseña:").grid(row=6, column=0, sticky="w", padx=6, pady=6)
            entry_password = ctk.CTkEntry(frm, width=300, show="*")
            entry_password.grid(row=6, column=1, sticky="ew", padx=6, pady=6)

            ctk.CTkLabel(frm, text="Confirmar contraseña:").grid(row=7, column=0, sticky="w", padx=6, pady=6)
            entry_password2 = ctk.CTkEntry(frm, width=300, show="*")
            entry_password2.grid(row=7, column=1, sticky="ew", padx=6, pady=6)

            # Buttons
            btn_frame = ctk.CTkFrame(dlg, fg_color="transparent")
            btn_frame.grid(row=99, column=0, sticky="ew", padx=12, pady=(6,12))
            btn_frame.grid_columnconfigure(0, weight=1)

            def guardar():
                # Obtener archivo del tema seleccionado
                tema_seleccionado = tema_menu.get()
                archivo_tema = self.tema_display_a_archivo(tema_seleccionado)
                
                # Obtener modo seleccionado
                modo_seleccionado = modo_menu.get()
                appearance_mode = "light" if modo_seleccionado == "Claro" else "dark"
                
                # Guardar valores actuales para comparar
                tema_anterior = self.user_settings.get("theme", "")
                modo_anterior = self.user_settings.get("appearance_mode", "")
                
                new = {
                    "date_format": date_menu.get(),
                    "show_tooltips": bool(var_tooltips.get()),
                    "compact_mode": bool(var_compact.get()),
                    "theme": f"themes/{archivo_tema}",
                    "appearance_mode": appearance_mode
                }
                
                # Actualizar configuraciones
                self.user_settings.update(new)
                ok = save_user_settings(self.user_settings, self.username)
                
                # Exponer globalmente y persistir valores en ejecución
                try:
                    global USER_SETTINGS
                    USER_SETTINGS = self.user_settings
                except Exception:
                    pass
                
                # Mostrar mensaje informativo si se cambió el tema o modo
                nuevo_tema = new.get("theme", "")
                nuevo_modo = new.get("appearance_mode", "")
                
                if nuevo_tema != tema_anterior or nuevo_modo != modo_anterior:
                    try:
                        # Asegurar que el messagebox aparezca en primer plano
                        dlg.focus_force()
                        messagebox.showinfo("Ajustes Guardados", 
                                          "Los cambios de tema y modo se aplicarán al reiniciar la aplicación.",
                                          parent=dlg)
                    except Exception:
                        # Fallback sin parent
                        try:
                            messagebox.showinfo("Ajustes Guardados", 
                                              "Los cambios de tema y modo se aplicarán al reiniciar la aplicación.")
                        except Exception:
                            pass
                
                # Redibujar listado para aplicar compact mode
                try:
                    self.mostrar_lista_rma()
                except Exception:
                    pass
                # Update email/password in DB if provided
                try:
                    # Update email column if present (and value provided)
                    email_val = entry_email.get().strip()
                    pw = entry_password.get()
                    pw2 = entry_password2.get()
                    conn_cursor = self.master.conectar_db()
                    if conn_cursor:
                        conn, cursor = conn_cursor
                        try:
                            # Ensure email column exists
                            cursor.execute("PRAGMA table_info('usuarios')")
                            cols = [r[1] for r in cursor.fetchall()]
                            if 'email' not in cols:
                                try:
                                    cursor.execute("ALTER TABLE usuarios ADD COLUMN email TEXT")
                                except Exception:
                                    pass
                            if email_val:
                                try:
                                    cursor.execute("UPDATE usuarios SET email = ? WHERE nombre_usuario = ?", (email_val, self.username))
                                except Exception as e:
                                    print(f"Error actualizando email: {e}")
                            # If password fields provided and match, update hash
                            if pw:
                                if pw != pw2:
                                    messagebox.showerror("Error", "Las contraseñas no coinciden.")
                                else:
                                    try:
                                        hashed = bcrypt.hashpw(pw.encode('utf-8'), bcrypt.gensalt())
                                        cursor.execute("UPDATE usuarios SET password_hash = ? WHERE nombre_usuario = ?", (hashed.decode('utf-8'), self.username))
                                    except Exception as e:
                                        print(f"Error actualizando contraseña: {e}")
                            try:
                                conn.commit()
                            except Exception:
                                pass
                            try:
                                conn.close()
                            except Exception:
                                pass
                        except Exception as e:
                            print(f"Error actualizando credenciales en DB: {e}")
                except Exception:
                    pass
                dlg.destroy()

            def cancelar():
                dlg.destroy()

            ctk.CTkButton(btn_frame, text="Cancelar", command=cancelar).grid(row=0, column=0, padx=6)
            ctk.CTkButton(btn_frame, text="Ver cambios", command=lambda: mostrar_ventana_cambios(dlg)).grid(row=0, column=1, padx=6)
            ctk.CTkButton(btn_frame, text="Guardar", command=guardar).grid(row=0, column=2, padx=6)

        except Exception as e:
            print(f"Error abriendo diálogo de ajustes: {e}")
    
    def obtener_quincenas_futuras(self):
        """Genera las quincenas (Q1/Q2) para los próximos 12 meses."""
        hoy = datetime.date.today()
        quincenas = []
        for i in range(12):
            mes = hoy.month + i
            anio = hoy.year + (mes - 1) // 12
            mes = (mes - 1) % 12 + 1
            
            anio_str = str(anio)[2:]
            mes_str = str(mes).zfill(2)
            
            quincenas.append(f"Q1-{mes_str}-{anio_str}")
            quincenas.append(f"Q2-{mes_str}-{anio_str}")
        return quincenas

    def obtener_siguiente_rma(self):
        """Calcula el siguiente código RMA para mostrar como número temporal (Ej: RMA25001)."""
        conn, cursor = self.master.conectar_db()
        if not conn: return "ERROR-DB"

        cursor = conn.cursor()
        
        anio_actual_str = str(datetime.datetime.now().year)[2:]
        prefijo_busqueda = f"RMA{anio_actual_str}%" 
        
        cursor.execute("""
            SELECT codigo_rma FROM rma_maestro 
            WHERE codigo_rma LIKE ? 
            ORDER BY id DESC 
            LIMIT 1
        """, (prefijo_busqueda,))
        
        ultimo_rma = cursor.fetchone()
        siguiente_numero = 1
        
        if ultimo_rma:
            numero_str = ultimo_rma[0].replace(f"RMA{anio_actual_str}", "")
            try:
                siguiente_numero = int(numero_str) + 1
            except ValueError:
                siguiente_numero = 1
        
        codigo_numerico = str(siguiente_numero).zfill(3)
        conn.close()
        
        return f"RMA{anio_actual_str}{codigo_numerico}"
    
    def generar_codigo_rma_final(self, cursor):
        """
        Genera y asigna el código RMA definitivo dentro de una transacción.
        Esta función debe ejecutarse dentro de una transacción activa.
        """
        anio_actual_str = str(datetime.datetime.now().year)[2:]
        prefijo_busqueda = f"RMA{anio_actual_str}%" 
        
        # Buscar el último número asignado en la misma transacción
        cursor.execute("""
            SELECT codigo_rma FROM rma_maestro 
            WHERE codigo_rma LIKE ? 
            ORDER BY id DESC 
            LIMIT 1
        """, (prefijo_busqueda,))
        
        ultimo_rma = cursor.fetchone()
        siguiente_numero = 1
        
        if ultimo_rma:
            numero_str = ultimo_rma[0].replace(f"RMA{anio_actual_str}", "")
            try:
                siguiente_numero = int(numero_str) + 1
            except ValueError:
                siguiente_numero = 1
        
        codigo_numerico = str(siguiente_numero).zfill(3)
        return f"RMA{anio_actual_str}{codigo_numerico}"

    def crear_campo(self, parent, fila, label_text, campo_bd, valor_defecto="", deshabilitado=False, tipo="entry", opciones=None):
        """Función auxiliar para crear etiquetas y campos de entrada/desplegables en el formulario."""
        ctk.CTkLabel(parent, text=label_text).grid(row=fila, column=0, padx=10, pady=5, sticky="w")

        
        widget = None
        
        if tipo == "entry":
            # 1. Crear el Entry en estado normal (por defecto)
            widget = ctk.CTkEntry(parent, width=300, state="normal") 
            
            # 2. Insertar el valor (solo se puede en estado normal)
            if valor_defecto:
                widget.insert(0, valor_defecto)
                
            # 3. Deshabilitar si se indica
            if deshabilitado:
                widget.configure(state="disabled")

        elif tipo == "optionmenu":
            widget = ctk.CTkOptionMenu(parent, 
                                       values=opciones, 
                                       width=300,
                                       #fg_color="gray80",        # Color del botón principal
                                       #button_color="gray70",    # Color del botón de flecha
                                       #button_hover_color="gray60", # Color al pasar el ratón por el botón de flecha
                                       #text_color="black"
                                       )
            if valor_defecto in opciones:
                widget.set(valor_defecto)
            elif opciones:
                widget.set(opciones[0])
            # Nota: CTkOptionMenu no tiene estado 'disabled' de forma nativa como el Entry, 
            # pero no es necesario para los campos que estamos usando ahora.
        elif tipo == "date":
            # 📅 Campo de selección de fecha
            widget = CTkDatePicker(parent, width=300)
            # Aplicar el formato de fecha según la preferencia del usuario (solo visual)
            try:
                pref = getattr(self, 'user_settings', {}).get('date_format', 'YYYY-MM-DD')
                fmt_map = {
                    'YYYY-MM-DD': '%Y-%m-%d',
                    'DD/MM/YYYY': '%d/%m/%Y',
                    'MM/DD/YYYY': '%m/%d/%Y'
                }
                widget_fmt = fmt_map.get(pref, '%Y-%m-%d')
                widget.set_date_format(widget_fmt)
            except Exception:
                # Fallback seguro
                try:
                    widget.set_date_format('%Y-%m-%d')
                except Exception:
                    pass

            # Si hay valor por defecto, establecerlo
            if valor_defecto:
                try:
                    from datetime import datetime
                    fecha = datetime.strptime(valor_defecto, "%Y-%m-%d")
                    widget.set_date(fecha)
                except ValueError:
                    pass  # Ignorar si el formato no es válido

            if deshabilitado:
                widget.configure(state="disabled")

            
        if widget:
            widget.grid(row=fila, column=1, padx=10, pady=5, sticky="ew")
            # Guardamos la referencia para acceder a los datos después
            setattr(self, f"entry_{campo_bd}", widget)

    def _prompt_select_autorizado_por(self, opciones):
        """Muestra un diálogo modal simple para seleccionar quien autoriza.
        Devuelve la opción seleccionada o None si se cancela.
        """
        if not opciones:
            return None

        dlg = Toplevel(self)
        dlg.title("Seleccionar Autorizado Por")
        dlg.transient(self)
        dlg.grab_set()
        # Centrar el diálogo respecto a la ventana principal
        try:
            width = 360
            height = 120
            # Forzar cálculo de medidas
            dlg.update_idletasks()
            px = self.winfo_rootx()
            py = self.winfo_rooty()
            pw = self.winfo_width() or self.winfo_screenwidth()
            ph = self.winfo_height() or self.winfo_screenheight()
            x = px + max(0, (pw - width) // 2)
            y = py + max(0, (ph - height) // 2)
            dlg.geometry(f"{width}x{height}+{x}+{y}")
        except Exception:
            try:
                dlg.geometry("360x120")
            except Exception:
                pass

        ctk.CTkLabel(dlg, text="Selecciona quien autoriza:", anchor="w").pack(fill='x', padx=12, pady=(12,6))
        var = tk.StringVar(value=opciones[0])
        opt = ctk.CTkOptionMenu(dlg, values=opciones)
        opt.set(opciones[0])
        opt.pack(fill='x', padx=12, pady=(0,8))

        result = {'value': None}

        def _ok():
            try:
                result['value'] = opt.get()
            except Exception:
                result['value'] = None
            dlg.destroy()

        def _cancel():
            dlg.destroy()

        btnf = ctk.CTkFrame(dlg)
        btnf.pack(fill='x', pady=(6,10), padx=12)
        ctk.CTkButton(btnf, text='OK', width=80, command=_ok).pack(side='right', padx=(6,0))
        ctk.CTkButton(btnf, text='Cancelar', width=80, command=_cancel).pack(side='right')

        self.wait_window(dlg)
        return result['value']
    

    # ----------------------------------------------------------------------
    # 4. LÓGICA DE LISTADO DE RMAS
    # ----------------------------------------------------------------------

    def mostrar_lista_rma(self):
        """Muestra el listado completo de RMAs, filtros y el dashboard de estadísticas."""
        self.limpiar_contenido()
        
        # Configurar layout principal con dos columnas
        self.content_frame.grid_columnconfigure(0, weight=3, minsize=800)  # Lista principal
        self.content_frame.grid_columnconfigure(1, weight=0, minsize=200)  # Dashboard (ancho fijo)
        self.content_frame.grid_rowconfigure(0, weight=1)
        
        # === COLUMNA IZQUIERDA: LISTA Y FILTROS ===
        lista_column = ctk.CTkFrame(self.content_frame)
        lista_column.grid(row=0, column=0, sticky="nsew", padx=(10, 5), pady=10)
        
        # Configurar expansión para la columna de lista
        lista_column.grid_rowconfigure(0, weight=0)  # Título
        lista_column.grid_rowconfigure(1, weight=0)  # Filtros  
        lista_column.grid_rowconfigure(2, weight=1)  # Listado
        lista_column.grid_columnconfigure(0, weight=1)

        # 1. Título y Botón Crear
        title_frame = ctk.CTkFrame(lista_column, fg_color="transparent")
        title_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 5))
        title_frame.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(title_frame, text="LISTADO DE EXPEDIENTES", 
                    font=ctk.CTkFont(family="Verdana", size=20, weight="bold")).grid(row=0, column=0, sticky="w")
        
        # Botón Crear Nuevo RMA
        try:
            btn_bg = None
            if hasattr(self, 'sidebar_frame') and hasattr(self.sidebar_frame, 'cget'):
                btn_bg = self.sidebar_frame.cget("fg_color")

            ctk.CTkButton(title_frame,
                          text="",
                          image=(self.icon_mas or self.icon_mas),
                          width=36,
                          height=36,
                          fg_color=(btn_bg if btn_bg is not None else None),
                          hover_color=(btn_bg if btn_bg is not None else None),
                          command=lambda: self._abrir_editor_rma(rma_id=None)).grid(row=0, column=1, padx=(20, 0), sticky="e")
            try:
                Tooltip(title_frame.winfo_children()[-1], "Crear nuevo RMA")
            except Exception:
                pass
        except Exception:
            ctk.CTkButton(title_frame,
                          text="➕ Crear Nuevo RMA",
                          command=lambda: self._abrir_editor_rma(rma_id=None)).grid(row=0, column=1, padx=(20, 0), sticky="e")

        # 2. Panel de Búsqueda y Filtros
        filtro_frame = ctk.CTkFrame(lista_column, fg_color="transparent")
        filtro_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=10)
        
        # Búsqueda por texto
        ctk.CTkLabel(filtro_frame, text="Buscar:").grid(row=0, column=0, padx=(0, 5), pady=5, sticky="w")
        self.entry_busqueda = ctk.CTkEntry(filtro_frame, placeholder_text="Código RMA, Cliente o Doc.", width=250)
        self.entry_busqueda.grid(row=0, column=1, padx=10, pady=5, sticky="w")
        # Bind para ejecutar búsqueda con Enter
        self.entry_busqueda.bind("<Return>", lambda e: self.aplicar_filtros_rma())
        
        # Filtro por Estado
        estados_posibles = self.OPCIONES.get("Estado", ["Todos"])
        if "Todos" not in estados_posibles:
            estados_posibles.insert(0, "Todos")
        if 'Exportado' not in estados_posibles:
            estados_posibles.append('Exportado')
            
        ctk.CTkLabel(filtro_frame, text="Estado:").grid(row=0, column=2, padx=(20, 5), pady=5, sticky="w")
        self.filtro_estado = ctk.CTkOptionMenu(filtro_frame, 
                                               values=estados_posibles, 
                                               width=120)
        self.filtro_estado.set("Todos")
        self.filtro_estado.grid(row=0, column=3, padx=10, pady=5, sticky="w")
        
        # Filtro por Año
        ctk.CTkLabel(filtro_frame, text="Año:").grid(row=0, column=4, padx=(20, 5), pady=5, sticky="w")
        # Inicializar con el año actual, se actualizará dinámicamente en cargar_lista_rma
        año_actual = str(datetime.datetime.now().year)
        self.filtro_año = ctk.CTkOptionMenu(filtro_frame, 
                                            values=[año_actual], 
                                            width=80)
        self.filtro_año.set(año_actual)
        self.filtro_año.grid(row=0, column=5, padx=10, pady=5, sticky="w")
        
        # Botón de Aplicar Filtro
        btn_aplicar_filtro = ctk.CTkButton(filtro_frame,
                                           text="🔍 Aplicar Filtros", 
                                           command=self.aplicar_filtros_rma)
        btn_aplicar_filtro.grid(row=0, column=6, padx=(20, 0), pady=5, sticky="w")
        
        filtro_frame.grid_columnconfigure(1, weight=1)

        # 3. Listado de RMAs
        self.lista_rma_frame = ctk.CTkScrollableFrame(lista_column, 
                                                     label_text="Pulse F5 para actualizar el listado ; Pinche dos veces sobre el expediente para abrirlo.")
        self.lista_rma_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=10)
        self.lista_rma_frame.grid_columnconfigure(0, weight=1)
        
        # === COLUMNA DERECHA: DASHBOARD ===
        dashboard_column = ctk.CTkFrame(self.content_frame, width=200, fg_color=("#f0f0f0", "#2b2b2b"))
        dashboard_column.grid(row=0, column=1, sticky="nsew", padx=(5, 10), pady=10)
        dashboard_column.grid_propagate(False)  # Mantener el ancho fijo
        
        # Header del dashboard
        dashboard_header = ctk.CTkFrame(dashboard_column, fg_color="transparent")
        dashboard_header.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(dashboard_header, text="📊 Estadísticas", 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(side="left")
        
        # Selector de año
        año_frame = ctk.CTkFrame(dashboard_column, fg_color="transparent")
        año_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(año_frame, text="Año:", font=ctk.CTkFont(size=11)).pack(side="left")
        
        años_disponibles = [str(año) for año in range(2020, datetime.datetime.now().year + 2)]
        self.combo_año_dashboard = ctk.CTkOptionMenu(año_frame, values=años_disponibles, 
                                                    command=self.actualizar_dashboard,
                                                    width=60)
        self.combo_año_dashboard.set(str(datetime.datetime.now().year))
        self.combo_año_dashboard.pack(side="right")
        
        # Selector de período para artículos problemáticos
        periodo_frame = ctk.CTkFrame(dashboard_column, fg_color="transparent")
        periodo_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(periodo_frame, text="Período:", font=ctk.CTkFont(size=11)).pack(side="left")
        
        self.combo_periodo = ctk.CTkOptionMenu(periodo_frame, 
                                             values=["Anual", "Semestral", "Trimestral"],
                                             command=self.actualizar_dashboard,
                                             width=80)
        self.combo_periodo.set("Anual")
        self.combo_periodo.pack(side="right")
        
        # Frame para las estadísticas
        self.stats_frame = ctk.CTkFrame(dashboard_column)
        self.stats_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Cargar datos iniciales - con filtro del año actual por defecto para optimizar rendimiento
        año_actual = str(datetime.datetime.now().year)
        self.cargar_lista_rma("", "Todos", año_actual)  # Cargar solo expedientes del año actual
        self.actualizar_dashboard()  # Cargar estadísticas del dashboard


    def cargar_lista_rma(self, texto_busqueda="", estado_filtro="Todos", año_filtro=None):
        """
        Carga los estados únicos de la DB para el filtro, y luego carga los RMA 
        desde la DB aplicando los filtros (texto, estado, año).
        """
        
        # Si no se especifica año, usar el año actual por defecto
        if año_filtro is None:
            año_filtro = str(datetime.datetime.now().year)
        
        # Limpiar el frame (siempre)
        for widget in self.lista_rma_frame.winfo_children():
            widget.destroy()

        conn, cursor = self.master.conectar_db()
        if not conn: 
            ctk.CTkLabel(self.lista_rma_frame, text="Error de conexión a la base de datos.").grid(row=0, column=0, padx=10, pady=10)
            return
        cursor = conn.cursor()
        
        try:
            # 1. OBTENER ESTADOS ÚNICOS PARA EL FILTRO - CON CACHÉ
            # Solo hacemos esto si el filtro_estado ya existe (es decir, en mostrar_lista_rma ya se creó la interfaz)
            if hasattr(self, 'filtro_estado'):
                try:
                    # Usar caché para estados (se actualiza cada 5 minutos o al invalidar)
                    def query_estados():
                        cursor.execute("SELECT DISTINCT estado FROM rma_maestro WHERE estado IS NOT NULL AND estado != '' ORDER BY estado ASC")
                        return [fila[0] for fila in cursor.fetchall()]
                    
                    estados_db = _get_cached_query('estados_rma', query_estados)
                    
                    # Crear la lista final de opciones: "Todos" + estados únicos de la DB
                    estados_posibles = ["Todos"] + estados_db
                    
                    # Actualizar el OptionMenu (sin cambiar la selección actual si es válida)
                    seleccion_actual = self.filtro_estado.get()
                    self.filtro_estado.configure(values=estados_posibles)
                    
                    # Mantener la selección actual si todavía existe, si no, poner "Todos"
                    if seleccion_actual in estados_posibles:
                        self.filtro_estado.set(seleccion_actual)
                    else:
                        self.filtro_estado.set("Todos")
                except Exception as e:
                    print(f"Error al cargar estados para filtro: {e}")
                    # Continuar con valores por defecto
                    if hasattr(self, 'filtro_estado'):
                        self.filtro_estado.configure(values=["Todos"])
                        self.filtro_estado.set("Todos")
            
            # 1.5. OBTENER AÑOS DISPONIBLES PARA EL FILTRO - CON CACHÉ
            if hasattr(self, 'filtro_año'):
                try:
                    # Usar caché para años (se actualiza cada 5 minutos o al invalidar)
                    def query_años():
                        cursor.execute("""
                            SELECT DISTINCT SUBSTR(codigo_rma, 4, 2) as año_corto 
                            FROM rma_maestro 
                            WHERE codigo_rma LIKE 'RMA%' 
                            ORDER BY año_corto DESC
                        """)
                        años_cortos = [fila[0] for fila in cursor.fetchall()]
                        # Convertir años cortos (25, 24, 23...) a años completos (2025, 2024, 2023...)
                        años_completos = []
                        for año_corto in años_cortos:
                            try:
                                año_int = int(año_corto)
                                # Asumimos que años 00-30 son 2000-2030, y 31-99 son 1931-1999
                                if año_int <= 30:
                                    año_completo = 2000 + año_int
                                else:
                                    año_completo = 1900 + año_int
                                años_completos.append(str(año_completo))
                            except ValueError:
                                continue
                        return años_completos if años_completos else [str(datetime.datetime.now().year)]
                    
                    años_db = _get_cached_query('años_rma', query_años)
                    
                    # Actualizar el OptionMenu de años (sin cambiar la selección actual si es válida)
                    seleccion_actual_año = self.filtro_año.get()
                    self.filtro_año.configure(values=años_db)
                    
                    # Mantener la selección actual si todavía existe, si no, poner el año actual
                    if seleccion_actual_año in años_db:
                        self.filtro_año.set(seleccion_actual_año)
                    else:
                        # Si no existe la selección actual, usar el año actual o el más reciente
                        año_actual = str(datetime.datetime.now().year)
                        if año_actual in años_db:
                            self.filtro_año.set(año_actual)
                        elif años_db:
                            self.filtro_año.set(años_db[0])  # El más reciente
                        else:
                            self.filtro_año.set(año_actual)
                            
                except Exception as e:
                    print(f"Error al cargar años para filtro: {e}")
                    # Continuar con valores por defecto
                    año_actual = str(datetime.datetime.now().year)
                    if hasattr(self, 'filtro_año'):
                        self.filtro_año.configure(values=[año_actual])
                        self.filtro_año.set(año_actual)
                    
                # 2. CARGAR LOS REGISTROS APLICANDO LOS FILTROS
                # (Aquí mantenemos tu lógica SQL que ya estaba funcionando)
            
                sql = "SELECT id, codigo_rma, cliente, numero_documento_cliente, fecha_emision, estado, fecha_autorizacion, fecha_recepcion, fecha_proceso, fecha_gestion FROM rma_maestro WHERE 1=1"
                params = []
            
                # Aplicar filtro de ESTADO
                estado_filtro_actual = self.filtro_estado.get() # Usamos el valor que se ha configurado
                if estado_filtro_actual and estado_filtro_actual != "Todos":
                    sql += " AND estado = ?"
                    params.append(estado_filtro_actual)
            
                # Aplicar filtro de AÑO
                año_filtro_actual = self.filtro_año.get() if hasattr(self, 'filtro_año') else año_filtro
                if año_filtro_actual:
                    # Extraer los últimos 2 dígitos del año para el formato RMA (ej: 2025 -> "25")
                    try:
                        año_corto = str(año_filtro_actual)[-2:]
                        sql += " AND codigo_rma LIKE ?"
                        params.append(f"RMA{año_corto}%")
                    except Exception as e:
                        print(f"Error aplicando filtro de año: {e}")
                
                # Aplicar filtro de BÚSQUEDA por texto
                if texto_busqueda:
                    busqueda_like = f"%{texto_busqueda}%"
                    sql += " AND (codigo_rma LIKE ? OR cliente LIKE ? OR numero_documento_cliente LIKE ?)"
                    params.append(busqueda_like)
                    params.append(busqueda_like)
                    params.append(busqueda_like) 

                # Ordenar y Ejecutar
                sql += " ORDER BY id DESC"
                try:
                    cursor.execute(sql, tuple(params))
                    registros = cursor.fetchall()
                except Exception as e:
                    print(f"Error ejecutando query principal: {e}")
                    print(f"SQL: {sql}")
                    print(f"Params: {params}")
                    raise
            
                conn.close()

                # 3. Dibujar la tabla de resultados (Encabezados y Registros)
            
                # ... (código para dibujar encabezados y la tabla de registros, sigue igual) ...
            
                # Encabezados
                header_font = ctk.CTkFont(weight="bold")
                # Configurar anchos de columnas para mejor distribución
                self.lista_rma_frame.grid_columnconfigure(0, weight=0, minsize=100)  # CÓDIGO RMA
                self.lista_rma_frame.grid_columnconfigure(1, weight=2, minsize=200)  # CLIENTE (reducido)
                self.lista_rma_frame.grid_columnconfigure(2, weight=1, minsize=150)  # DOCUMENTO
                self.lista_rma_frame.grid_columnconfigure(3, weight=0, minsize=130)  # ÚLTIMA ACTIVIDAD
                self.lista_rma_frame.grid_columnconfigure(4, weight=1, minsize=180)  # ESTADO (ampliado)
                self.lista_rma_frame.grid_columnconfigure(5, weight=0, minsize=110)  # FECHA EMISIÓN
            
                ctk.CTkLabel(self.lista_rma_frame, text="CÓDIGO RMA", font=header_font).grid(row=0, column=0, padx=5, pady=5, sticky="w")
                ctk.CTkLabel(self.lista_rma_frame, text="CLIENTE", font=header_font).grid(row=0, column=1, padx=5, pady=5, sticky="w")
                ctk.CTkLabel(self.lista_rma_frame, text="DOCUMENTO DE CLIENTE", font=header_font).grid(row=0, column=2, padx=5, pady=5, sticky="w")
                ctk.CTkLabel(self.lista_rma_frame, text="ÚLTIMA ACTIVIDAD", font=header_font).grid(row=0, column=3, padx=5, pady=5, sticky="w")
                ctk.CTkLabel(self.lista_rma_frame, text="ESTADO", font=header_font).grid(row=0, column=4, padx=5, pady=5, sticky="w")
                ctk.CTkLabel(self.lista_rma_frame, text="FECHA EMISIÓN", font=header_font).grid(row=0, column=5, padx=5, pady=5, sticky="w")
                # Eliminada columna 'ACCIONES' para usar doble clic en la fila
                if not registros:
                    ctk.CTkLabel(self.lista_rma_frame, text="No se encontraron expedientes con los filtros aplicados.", text_color="gray").grid(row=1, column=0, columnspan=6, padx=10, pady=20)
                    return

                # Registros (filas) - usar fondo por defecto del tema (sin cebra)
                colors = ("#FFFFFF", "#F3F4F6")  # variable no usada para cebra pero mantenida por compatibilidad
                # Determinar color de fondo para botones (para hacer 'transparent-like')
                btn_bg = None
                if hasattr(self, 'sidebar_frame') and hasattr(self.sidebar_frame, 'cget'):
                    btn_bg = self.sidebar_frame.cget("fg_color")

                # Altura de fila según compact_mode
                row_height = 22 if getattr(self, 'user_settings', {}).get('compact_mode', True) else 32

                for i, reg in enumerate(registros):
                    rma_id, codigo_rma, cliente, numero_documento_cliente, fecha_emision, estado, fecha_autorizacion, fecha_recepcion, fecha_proceso, fecha_gestion = reg
                    row = i + 1

                    # Calcular la última actividad usando la función auxiliar
                    ultima_actividad = obtener_ultima_actividad(
                        fecha_emision, 
                        fecha_autorizacion, 
                        fecha_recepcion, 
                        fecha_proceso, 
                        fecha_gestion
                    )

                    # Mapeo de color según estado (coherente con dashboard)
                    color = self.get_color_por_estado(estado)

                    # Usar fondo por defecto del tema en lugar de cebra
                    bg = "transparent"
                    # Crear un frame por columna para alinear exactamente con los encabezados
                    # Reducimos la altura de cada fila usando height pequeño para que sean más finas.
                    # Height reducido para filas compactas
                    # Hacer que la columna de 'ACCIONES' tenga un fondo fijo (sin cebra)
                    actions_bg = None
                    try:
                        if hasattr(self, 'lista_rma_frame') and hasattr(self.lista_rma_frame, 'cget'):
                            actions_bg = self.lista_rma_frame.cget('fg_color')
                    except Exception:
                        actions_bg = None
                    if actions_bg is None:
                        actions_bg = colors[0]

                    f0 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)
                    f1 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)
                    f2 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)
                    f3 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)
                    f4 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)
                    f5 = ctk.CTkFrame(self.lista_rma_frame, fg_color="transparent", height=row_height)

                    # Colocar cada columna en la grilla principal para que se alinee con encabezados
                    f0.grid(row=row, column=0, sticky="nsew", padx=0, pady=0)
                    f1.grid(row=row, column=1, sticky="nsew", padx=0, pady=0)
                    f2.grid(row=row, column=2, sticky="nsew", padx=0, pady=0)
                    f3.grid(row=row, column=3, sticky="nsew", padx=0, pady=0)
                    f4.grid(row=row, column=4, sticky="nsew", padx=0, pady=0)
                    f5.grid(row=row, column=5, sticky="nsew", padx=0, pady=0)

                    # Contenido de cada columna con padding muy reducido para filas más finas
                    # Crear labels y mantener referencias para enlazar eventos (doble click)
                    lbl0 = ctk.CTkLabel(f0, text=codigo_rma)
                    lbl0.pack(anchor="w", padx=4, pady=0)
                    lbl1 = ctk.CTkLabel(f1, text=cliente)
                    lbl1.pack(anchor="w", padx=4, pady=0)
                    lbl2 = ctk.CTkLabel(f2, text=numero_documento_cliente)
                    lbl2.pack(anchor="w", padx=4, pady=0)
                    lbl3 = ctk.CTkLabel(f3, text=ultima_actividad)
                    lbl3.pack(anchor="w", padx=4, pady=0)
                    lbl4 = ctk.CTkLabel(f4, text=estado, text_color=color)
                    lbl4.pack(anchor="w", padx=4, pady=0)
                    lbl5 = ctk.CTkLabel(f5, text=fecha_emision)
                    lbl5.pack(anchor="w", padx=4, pady=0)
                    # En lugar de botón de editar, abrimos el editor con doble clic en cualquier columna de la fila

                    # Hover efectos para toda la fila: aplicar a cada columna
                    cols = [f0, f1, f2, f3, f4, f5]
                    # Guardar el color original por columna para restaurarlo correctamente
                    originals = ["transparent", "transparent", "transparent", "transparent", "transparent", "transparent"]

                    def _on_enter(e, cols=cols):
                        for rf in cols:
                            try:
                                rf.configure(fg_color="transparent")
                            except Exception:
                                pass

                    def _on_leave(e, cols=cols, originals=originals):
                        for idx, rf in enumerate(cols):
                            try:
                                rf.configure(fg_color=originals[idx])
                            except Exception:
                                pass

                    # También enlazamos los labels internos para asegurar que capturan el doble click
                    inner_widgets = [lbl0, lbl1, lbl2, lbl3, lbl4, lbl5]

                    for rf in cols:
                        rf.bind("<Enter>", _on_enter)
                        rf.bind("<Leave>", _on_leave)
                        # Mostrar cursor 'hand2' para indicar que la fila es clicable
                        try:
                            rf.configure(cursor="hand2")
                        except Exception:
                            pass

                    # Enlazar eventos a labels (algunas platforms capturan el evento en el label)
                    for w in inner_widgets:
                        try:
                            w.configure(cursor="hand2")
                        except Exception:
                            pass
                        try:
                            w.bind("<Double-Button-1>", lambda e=None, r=rma_id: self._abrir_editor_rma(rma_id=r))
                        except Exception:
                            pass
            
        except Exception as e:
            print(f"Error al cargar lista de RMA: {e}")
            if conn: conn.close()
            ctk.CTkLabel(self.lista_rma_frame, text=f"Error al cargar la lista: {e}").grid(row=1, column=0, columnspan=6, padx=10, pady=20)

    def crear_copia_seguridad_db(self):
        """
        Crea una copia de seguridad completa de la base de datos principal
        usando el diálogo 'Guardar como' para que el usuario elija la ubicación.
        La base de datos se asume en la carpeta actual: rma_app.db.
        """
        
        # 1. Determinar la ruta de origen de la BD
        try:
            # Usamos os.getcwd() para obtener la ruta del directorio de trabajo actual
            # y os.path.join para construir la ruta completa de la base de datos.
            db_path_origen = os.path.join(os.getcwd(), DB_FILENAME)

            # Si se han configurado variables de TURSO, podemos intentar volcar desde Turso
            turso_url_check = os.getenv("TURSO_DATABASE_URL")
            turso_token_check = os.getenv("TURSO_AUTH_TOKEN")

            # Opcional: una verificación rápida para ver si el archivo existe.
            # Pero si hay credenciales de Turso, NO abortamos aquí: intentaremos volcar Turso.
            if not os.path.exists(db_path_origen) and not (turso_url_check and turso_token_check):
                messagebox.showerror(
                    "Error de Archivo",
                    f"No se encontró la base de datos local en la ruta esperada: {db_path_origen}\n"
                    f"Asegúrate de que el archivo '{DB_FILENAME}' está en la misma carpeta, o configura TURSO_DATABASE_URL/TURSO_AUTH_TOKEN para volcar desde Turso."
                )
                return

        except Exception as e:
            messagebox.showerror("Error de Ruta", f"No se pudo determinar la ruta de la base de datos: {e}")
            return

        # 2. Generar un nombre de archivo sugerido con fecha y hora
        fecha_actual = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_sugerido = f"backup_rma_app_{fecha_actual}.db"

        # 3. Abrir el diálogo para que el usuario seleccione la ruta de destino
        path_destino = filedialog.asksaveasfilename(
            defaultextension=".db", 
            filetypes=[
                ("Base de Datos SQLite", "*.db"), 
                ("Base de Datos SQLite", "*.sqlite"),
                ("Todos los archivos", "*.*")
            ],
            initialfile=nombre_sugerido,
            title="Guardar Copia de Seguridad de la Base de Datos"
        )

        if not path_destino:
            # El usuario canceló el diálogo
            return

        # 4. Realizar la copia de seguridad
        conn_origen = None
        conn_destino = None
        # Registro de operaciones para mostrar al final
        operations_log: list[str] = []

        def show_log_window(lines: list[str], title: str = "Registro de Copia de Seguridad"):
            """Muestra una ventana Toplevel con el log de operaciones."""
            try:
                win = ctk.CTkToplevel(self)
                win.title(title)
                win.geometry("700x400")
                win.grab_set()

                frame = ctk.CTkFrame(win)
                frame.pack(fill="both", expand=True, padx=10, pady=10)

                txt = ctk.CTkTextbox(frame, wrap="word")
                txt.pack(fill="both", expand=True)
                txt.configure(state="normal")
                txt.delete("1.0", "end")
                txt.insert("1.0", "\n".join(lines))
                txt.configure(state="disabled")

                btn_close = ctk.CTkButton(win, text="Cerrar", command=win.destroy)
                btn_close.pack(pady=8)
            except Exception:
                # Fallback simple si CustomTk no funciona por alguna razón
                try:
                    messagebox.showinfo(title, "\n".join(lines))
                except Exception:
                    print("Log:\n" + "\n".join(lines))

        # Si están definidas las variables de Turso, intentamos volcar la BD remota
        turso_url = os.getenv("TURSO_DATABASE_URL")
        turso_token = os.getenv("TURSO_AUTH_TOKEN")

        if turso_url and turso_token:
            try:
                operations_log.append(f"Intentando volcar Turso: {turso_url}")
                # Aviso al usuario: volcando desde Turso
                try:
                    messagebox.showinfo("Volcando desde Turso", "Se va a intentar volcar la base de datos desde Turso. Esta operación puede tardar varios minutos dependiendo del tamaño de las tablas.")
                except Exception:
                    # No bloquear si messagebox falla
                    pass
                # Conectamos a la BD (esto usará la implementación Turso si las env vars están)
                remote_conn = connect_db(timeout=30)
                remote_cursor = remote_conn.cursor()

                # Obtenemos los CREATE TABLE de la BD remota (evitar tablas internas sqlite_)
                remote_cursor.execute("SELECT name, sql FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
                tablas = remote_cursor.fetchall()

                # Crear la BD destino local y recrear tablas
                conn_destino = sqlite3.connect(path_destino)
                dest_cur = conn_destino.cursor()

                for tabla in tablas:
                    if len(tabla) < 2:
                        continue
                    nombre, create_sql = tabla[0], tabla[1]
                    if not create_sql:
                        continue
                    try:
                        dest_cur.execute(create_sql)
                        operations_log.append(f"CREATE TABLE ejecutado para: {nombre}")
                    except Exception:
                        # Si falla crear la tabla (por compatibilidad), intentar skip
                        operations_log.append(f"Advertencia: no se pudo crear tabla {nombre} (se omite)")
                        pass

                    # Copiar los datos de la tabla en bloques para no consumir mucha memoria
                    try:
                        chunk_size = 1000

                        # Intentar obtener el total de filas para poder mostrar progreso
                        total_rows = None
                        try:
                            remote_cursor.execute(f"SELECT COUNT(*) FROM \"{nombre}\"")
                            cnt = remote_cursor.fetchone()
                            total_rows = int(cnt[0]) if cnt and cnt[0] is not None else 0
                        except Exception:
                            # Si COUNT falla por permisos o particularidades, seguiremos sin total
                            total_rows = None

                        if total_rows == 0:
                            operations_log.append(f"Tabla {nombre}: 0 filas (omitida)")
                            continue

                        offset = 0
                        copied = 0
                        insert_sql = None

                        while True:
                            # Leer por bloques usando LIMIT/OFFSET
                            remote_cursor.execute(f"SELECT * FROM \"{nombre}\" LIMIT {chunk_size} OFFSET {offset}")
                            rows = remote_cursor.fetchall()
                            if not rows:
                                break

                            # Preparar placeholders la primera vez que tengamos descripción
                            if insert_sql is None:
                                colcount = len(remote_cursor.description) if remote_cursor.description else 0
                                placeholders = ",".join(["?" for _ in range(colcount)])
                                insert_sql = f"INSERT INTO \"{nombre}\" VALUES ({placeholders})"

                            # Insertar bloque en la DB destino
                            try:
                                dest_cur.executemany(insert_sql, rows)
                                conn_destino.commit()
                                copied += len(rows)
                                offset += len(rows)
                                if total_rows is None:
                                    operations_log.append(f"Tabla {nombre}: {copied} filas copiadas (progreso por bloques)")
                                else:
                                    operations_log.append(f"Tabla {nombre}: {copied}/{total_rows} filas copiadas")
                            except Exception:
                                conn_destino.rollback()
                                operations_log.append(f"Error insertando bloque en tabla {nombre} (omitida)")
                                # Saltar a la siguiente tabla
                                break

                        if copied > 0:
                            operations_log.append(f"Tabla {nombre}: copia finalizada, {copied} filas copiadas")
                        else:
                            operations_log.append(f"Tabla {nombre}: ninguna fila copiada")
                    except Exception as e:
                        # Ignorar errores de copia de datos de tablas concretas
                        try:
                            conn_destino.rollback()
                        except Exception:
                            pass
                        operations_log.append(f"Error copiando datos de tabla {nombre} (omitida): {e}")
                        continue

                # Cerrar conexión remota si existe
                try:
                    remote_conn.close()
                except Exception:
                    pass

                operations_log.append(f"Copia de Turso completada correctamente en: {path_destino}")
                # Intentar guardar el log en un archivo junto al .db de destino
                try:
                    # Guardar logs en carpeta centralizada: <project>/logs/backups/
                    logs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs', 'backups')
                    os.makedirs(logs_dir, exist_ok=True)
                    log_name = os.path.splitext(os.path.basename(path_destino))[0] + '.log'
                    log_path = os.path.join(logs_dir, log_name)
                    with open(log_path, 'w', encoding='utf-8') as lf:
                        lf.write('\n'.join(operations_log))
                    operations_log.append(f"Registro guardado en: {log_path}")
                except Exception as e:
                    operations_log.append(f"No se pudo guardar el registro en disco: {e}")

                # Mostrar log detallado al usuario
                show_log_window(operations_log, title="Registro de copia (Turso)")
                return

            except Exception as e:
                # Si falla el volcado remoto, registramos y caeremos al backup local
                operations_log.append(f"Info: no se pudo volcar Turso directamente: {e}")
                try:
                    remote_conn.close()
                except Exception:
                    pass

        # Si no hay Turso o el volcado falló, intentamos el backup local tradicional
        try:
            # 4.1. Conectar a la base de datos de origen (solo lectura)
            if not os.path.exists(db_path_origen):
                raise FileNotFoundError(f"Archivo de BD local no encontrado: {db_path_origen}")

            conn_origen = sqlite3.connect(db_path_origen)
            # 4.2. Conectar a la base de datos de destino (se crea si no existe)
            conn_destino = sqlite3.connect(path_destino)

            # 4.3. Usar el método de backup integrado de SQLite
            with conn_destino:
                conn_origen.backup(conn_destino)

            operations_log.append(f"Copia local realizada correctamente en: {path_destino}")
            # Intentar guardar el log en un archivo junto al .db de destino
            try:
                # Guardar logs en carpeta centralizada: <project>/logs/backups/
                logs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs', 'backups')
                os.makedirs(logs_dir, exist_ok=True)
                log_name = os.path.splitext(os.path.basename(path_destino))[0] + '.log'
                log_path = os.path.join(logs_dir, log_name)
                with open(log_path, 'w', encoding='utf-8') as lf:
                    lf.write('\n'.join(operations_log))
                operations_log.append(f"Registro guardado en: {log_path}")
            except Exception as e:
                operations_log.append(f"No se pudo guardar el registro en disco: {e}")

            show_log_window(operations_log, title="Registro de copia (local)")

        except sqlite3.Error as e:
            messagebox.showerror("Error de Base de Datos", f"Ocurrió un error al crear la copia de seguridad: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error inesperado: {e}")
        finally:
            if conn_origen:
                conn_origen.close()
            if conn_destino:
                conn_destino.close()

    # ----------------------------------------------------------------------
    # 4.5. BÚSQUEDA GLOBAL AVANZADA
    # ----------------------------------------------------------------------

    def mostrar_busqueda_global(self):
        """Muestra la interfaz de búsqueda global avanzada con filtros múltiples e historial."""
        self.limpiar_contenido()
        
        # Header
        header_frame = ctk.CTkFrame(self.content_frame)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(header_frame, 
                    text="🔍 BÚSQUEDA GLOBAL AVANZADA", 
                    font=ctk.CTkFont(size=20, weight="bold")).pack(pady=10)
        
        # Botón para volver
        btn_volver = ctk.CTkButton(header_frame, text="← Volver a Lista", 
                                  command=self.mostrar_lista_rma,
                                  width=150, height=30)
        btn_volver.pack(anchor="e", padx=10, pady=(0,5))
        
        # Frame principal con dos columnas
        main_frame = ctk.CTkFrame(self.content_frame)
        main_frame.pack(fill="both", expand=True, padx=10, pady=5)
        main_frame.grid_columnconfigure(0, weight=2, minsize=400)  # Búsqueda principal
        main_frame.grid_columnconfigure(1, weight=1, minsize=250)  # Historial
        
        # === COLUMNA IZQUIERDA: BÚSQUEDA Y FILTROS ===
        search_column = ctk.CTkFrame(main_frame)
        search_column.grid(row=0, column=0, sticky="nsew", padx=(10,5), pady=10)
        
        # Campo de búsqueda principal
        ctk.CTkLabel(search_column, text="Buscar en todos los campos:", 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=10, pady=(10,5))
        
        entrada_frame = ctk.CTkFrame(search_column, fg_color="transparent")
        entrada_frame.pack(fill="x", padx=10, pady=5)
        
        self.entry_busqueda_global = ctk.CTkEntry(entrada_frame, 
                                                 placeholder_text="Escribe cualquier texto para buscar...",
                                                 height=35,
                                                 font=ctk.CTkFont(size=12))
        self.entry_busqueda_global.pack(fill="x", padx=(0,10))
        
        # Botones de búsqueda
        botones_frame = ctk.CTkFrame(search_column, fg_color="transparent")
        botones_frame.pack(fill="x", padx=10, pady=5)
        
        btn_buscar = ctk.CTkButton(botones_frame, text="🔍 Buscar", 
                                  command=self.ejecutar_busqueda_global,
                                  width=100, height=35)
        btn_buscar.pack(side="left", padx=(0,5))
        
        btn_limpiar = ctk.CTkButton(botones_frame, text="🗑️ Limpiar", 
                                   command=self.limpiar_busqueda_global,
                                   width=100, height=35)
        btn_limpiar.pack(side="left", padx=5)
        
        # === FILTROS AVANZADOS ===
        filtros_frame = ctk.CTkFrame(search_column)
        filtros_frame.pack(fill="x", padx=10, pady=10)
        
        # Header de filtros con botón expandir/contraer
        filtros_header = ctk.CTkFrame(filtros_frame, fg_color="transparent")
        filtros_header.pack(fill="x", padx=10, pady=5)
        
        self.filtros_expandido = ctk.BooleanVar(value=False)
        self.btn_toggle_filtros = ctk.CTkButton(filtros_header, 
                                              text="🔽 Mostrar Filtros Avanzados",
                                              command=self.toggle_filtros_avanzados,
                                              width=200, height=30)
        self.btn_toggle_filtros.pack(side="left")
        
        # Frame de filtros (inicialmente oculto)
        self.filtros_content = ctk.CTkFrame(filtros_frame)
        # No pack inicialmente - se mostrará al expandir
        
        # Crear controles de filtros
        self.crear_controles_filtros()
        
        # Bind Enter key
        self.entry_busqueda_global.bind("<Return>", lambda e: self.ejecutar_busqueda_global())
        
        # === COLUMNA DERECHA: HISTORIAL ===
        historial_column = ctk.CTkFrame(main_frame)
        historial_column.grid(row=0, column=1, sticky="nsew", padx=(5,10), pady=10)
        
        # Header del historial
        historial_header = ctk.CTkFrame(historial_column, fg_color="transparent")
        historial_header.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(historial_header, text="� Historial de Búsquedas", 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(side="left")
        
        btn_limpiar_historial = ctk.CTkButton(historial_header, text="🗑️", 
                                            command=self.limpiar_historial_busquedas_ui,
                                            width=30, height=25)
        btn_limpiar_historial.pack(side="right")
        
        # Lista del historial
        self.historial_frame = ctk.CTkScrollableFrame(historial_column, height=200)
        self.historial_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.actualizar_historial_ui()
        
        # Información de ayuda
        ayuda_frame = ctk.CTkFrame(search_column, fg_color="transparent")
        ayuda_frame.pack(fill="x", padx=10, pady=5)
        
        ayuda_text = """💡 Búsqueda en: Expedientes (Código, Cliente, Estado, etc.) y Productos (Referencia, Serie, etc.)
⌨️ Atajo: Ctrl+F | 🔽 Usa filtros para búsquedas más específicas"""
        
        ctk.CTkLabel(ayuda_frame, text=ayuda_text, 
                    font=ctk.CTkFont(size=11), 
                    text_color="gray",
                    justify="left").pack(anchor="w")
        
        # Frame para resultados (span ambas columnas)
        self.resultados_frame = ctk.CTkScrollableFrame(main_frame, height=300)
        self.resultados_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)
        main_frame.grid_rowconfigure(1, weight=1)

    def crear_controles_filtros(self):
        """Crea los controles de filtros avanzados."""
        # Configurar grid
        self.filtros_content.grid_columnconfigure(0, weight=1)
        self.filtros_content.grid_columnconfigure(1, weight=1)
        
        row = 0
        
        # === FILTROS PARA EXPEDIENTES ===
        exp_label = ctk.CTkLabel(self.filtros_content, text="📋 Filtros de Expedientes", 
                               font=ctk.CTkFont(size=13, weight="bold"))
        exp_label.grid(row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10,5))
        row += 1
        
        # Estado del expediente
        ctk.CTkLabel(self.filtros_content, text="Estado:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_estado = ctk.CTkOptionMenu(self.filtros_content, 
                                             values=["Todos", "Pendiente", "Autorizado", "Recibido", "Completado"])
        self.filtro_estado.set("Todos")
        self.filtro_estado.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        # Rango de fechas
        ctk.CTkLabel(self.filtros_content, text="Fecha desde:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_fecha_desde = ctk.CTkEntry(self.filtros_content, placeholder_text="YYYY-MM-DD")
        self.filtro_fecha_desde.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        ctk.CTkLabel(self.filtros_content, text="Fecha hasta:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_fecha_hasta = ctk.CTkEntry(self.filtros_content, placeholder_text="YYYY-MM-DD")
        self.filtro_fecha_hasta.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        # Cliente específico
        ctk.CTkLabel(self.filtros_content, text="Cliente:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_cliente = ctk.CTkEntry(self.filtros_content, placeholder_text="Nombre del cliente")
        self.filtro_cliente.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        # === FILTROS PARA PRODUCTOS ===
        prod_label = ctk.CTkLabel(self.filtros_content, text="📦 Filtros de Productos", 
                               font=ctk.CTkFont(size=13, weight="bold"))
        prod_label.grid(row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(15,5))
        row += 1
        
        # Estado del producto
        ctk.CTkLabel(self.filtros_content, text="Estado Producto:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_estado_producto = ctk.CTkEntry(self.filtros_content, placeholder_text="Estado del producto")
        self.filtro_estado_producto.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        # Referencia específica
        ctk.CTkLabel(self.filtros_content, text="Referencia:").grid(row=row, column=0, sticky="w", padx=10, pady=2)
        self.filtro_referencia = ctk.CTkEntry(self.filtros_content, placeholder_text="Referencia del artículo")
        self.filtro_referencia.grid(row=row, column=1, sticky="ew", padx=10, pady=2)
        row += 1
        
        # Botones de filtros
        botones_filtros = ctk.CTkFrame(self.filtros_content, fg_color="transparent")
        botones_filtros.grid(row=row, column=0, columnspan=2, sticky="ew", padx=10, pady=10)
        
        btn_aplicar_filtros = ctk.CTkButton(botones_filtros, text="🔍 Buscar con Filtros", 
                                          command=self.ejecutar_busqueda_con_filtros,
                                          height=30)
        btn_aplicar_filtros.pack(side="left", padx=(0,5))
        
        btn_limpiar_filtros = ctk.CTkButton(botones_filtros, text="🧹 Limpiar Filtros", 
                                          command=self.limpiar_filtros,
                                          height=30)
        btn_limpiar_filtros.pack(side="left")

    def toggle_filtros_avanzados(self):
        """Muestra u oculta los filtros avanzados."""
        if self.filtros_expandido.get():
            # Contraer
            self.filtros_content.pack_forget()
            self.btn_toggle_filtros.configure(text="🔽 Mostrar Filtros Avanzados")
            self.filtros_expandido.set(False)
        else:
            # Expandir
            self.filtros_content.pack(fill="x", padx=10, pady=(0,10))
            self.btn_toggle_filtros.configure(text="🔼 Ocultar Filtros Avanzados")
            self.filtros_expandido.set(True)

    def limpiar_filtros(self):
        """Limpia todos los filtros avanzados."""
        self.filtro_estado.set("Todos")
        self.filtro_fecha_desde.delete(0, 'end')
        self.filtro_fecha_hasta.delete(0, 'end')
        self.filtro_cliente.delete(0, 'end')
        self.filtro_estado_producto.delete(0, 'end')
        self.filtro_referencia.delete(0, 'end')

    def actualizar_historial_ui(self):
        """Actualiza la interfaz del historial de búsquedas."""
        # Limpiar historial actual
        for widget in self.historial_frame.winfo_children():
            widget.destroy()
        
        historial = self.cargar_historial_busquedas()
        
        if not historial:
            ctk.CTkLabel(self.historial_frame, text="Sin búsquedas recientes", 
                        text_color="gray").pack(pady=10)
            return
        
        for i, entrada in enumerate(historial):
            self.crear_entrada_historial(entrada, i)

    def crear_entrada_historial(self, entrada, index):
        """Crea una entrada visual en el historial."""
        frame = ctk.CTkFrame(self.historial_frame)
        frame.pack(fill="x", padx=5, pady=2)
        
        # Texto de la búsqueda
        texto = entrada.get("termino", "")
        if len(texto) > 25:
            texto = texto[:22] + "..."
        
        btn_entrada = ctk.CTkButton(frame, text=texto, 
                                  command=lambda: self.usar_busqueda_historial(entrada),
                                  height=25, font=ctk.CTkFont(size=11))
        btn_entrada.pack(side="left", fill="x", expand=True, padx=5, pady=2)
        
        # Fecha (opcional)
        try:
            fecha = entrada.get("fecha", "")
            if fecha:
                fecha_obj = datetime.datetime.fromisoformat(fecha.replace('Z', '+00:00'))
                fecha_str = fecha_obj.strftime("%d/%m")
                ctk.CTkLabel(frame, text=fecha_str, text_color="gray", 
                           font=ctk.CTkFont(size=9)).pack(side="right", padx=5)
        except:
            pass

    def usar_busqueda_historial(self, entrada):
        """Aplica una búsqueda del historial."""
        # Cargar término de búsqueda
        self.entry_busqueda_global.delete(0, 'end')
        self.entry_busqueda_global.insert(0, entrada.get("termino", ""))
        
        # Cargar filtros si existen
        filtros = entrada.get("filtros", {})
        if filtros and hasattr(self, 'filtro_estado'):
            # Expandir filtros si hay filtros guardados
            if not self.filtros_expandido.get():
                self.toggle_filtros_avanzados()
            
            # Aplicar filtros guardados
            self.filtro_estado.set(filtros.get("estado", "Todos"))
            self.filtro_fecha_desde.delete(0, 'end')
            self.filtro_fecha_desde.insert(0, filtros.get("fecha_desde", ""))
            self.filtro_fecha_hasta.delete(0, 'end')
            self.filtro_fecha_hasta.insert(0, filtros.get("fecha_hasta", ""))
            self.filtro_cliente.delete(0, 'end')
            self.filtro_cliente.insert(0, filtros.get("cliente", ""))
            self.filtro_estado_producto.delete(0, 'end')
            self.filtro_estado_producto.insert(0, filtros.get("estado_producto", ""))
            self.filtro_referencia.delete(0, 'end')
            self.filtro_referencia.insert(0, filtros.get("referencia", ""))
        
        # Ejecutar búsqueda
        if filtros:
            self.ejecutar_busqueda_con_filtros()
        else:
            self.ejecutar_busqueda_global()

    def limpiar_historial_busquedas_ui(self):
        """Limpia el historial desde la interfaz."""
        self.limpiar_historial_busquedas()
        self.actualizar_historial_ui()

    def ejecutar_busqueda_con_filtros(self):
        """Ejecuta búsqueda combinada con filtros avanzados."""
        termino = self.entry_busqueda_global.get().strip()
        
        # Recopilar filtros
        filtros = {
            "estado": self.filtro_estado.get() if self.filtro_estado.get() != "Todos" else "",
            "fecha_desde": self.filtro_fecha_desde.get().strip(),
            "fecha_hasta": self.filtro_fecha_hasta.get().strip(),
            "cliente": self.filtro_cliente.get().strip(),
            "estado_producto": self.filtro_estado_producto.get().strip(),
            "referencia": self.filtro_referencia.get().strip()
        }
        
        # Validar fechas
        if filtros["fecha_desde"] and not self.validar_fecha(filtros["fecha_desde"]):
            messagebox.showerror("Error", "Formato de fecha inválido (YYYY-MM-DD)")
            return
        if filtros["fecha_hasta"] and not self.validar_fecha(filtros["fecha_hasta"]):
            messagebox.showerror("Error", "Formato de fecha inválido (YYYY-MM-DD)")
            return
        
        # Guardar en historial si hay criterios de búsqueda
        if termino or any(v for v in filtros.values() if v):
            self.guardar_busqueda_en_historial(termino, filtros)
            self.actualizar_historial_ui()
        
        # Limpiar resultados previos
        for widget in self.resultados_frame.winfo_children():
            widget.destroy()
        
        try:
            resultados_maestro = self.buscar_en_maestro_con_filtros(termino, filtros)
            resultados_detalles = self.buscar_en_detalles_con_filtros(termino, filtros)
            
            self.mostrar_resultados_busqueda_avanzada(resultados_maestro, resultados_detalles, termino, filtros)
            
        except Exception as e:
            messagebox.showerror("Error en la búsqueda", f"Error al ejecutar búsqueda con filtros: {str(e)}")
            print(f"Error en búsqueda con filtros: {e}")

    def buscar_en_maestro_con_filtros(self, termino, filtros):
        """Busca en rma_maestro aplicando filtros."""
        query = "SELECT * FROM rma_maestro WHERE 1=1"
        params = []
        
        # Filtro de texto general
        if termino:
            query += """ AND (
                numero_rma LIKE ? OR cliente LIKE ? OR nombre_contacto LIKE ? OR 
                email_contacto LIKE ? OR direccion_cliente LIKE ? OR 
                telefono_contacto LIKE ? OR motivo_devolucion LIKE ? OR 
                numero_seguimiento LIKE ? OR observaciones LIKE ?
            )"""
            termino_param = f"%{termino}%"
            params.extend([termino_param] * 9)
        
        # Filtros específicos
        if filtros["estado"]:
            query += " AND estado_expediente = ?"
            params.append(filtros["estado"])
        
        if filtros["fecha_desde"]:
            query += " AND fecha_creacion >= ?"
            params.append(filtros["fecha_desde"])
        
        if filtros["fecha_hasta"]:
            query += " AND fecha_creacion <= ?"
            params.append(filtros["fecha_hasta"])
        
        if filtros["cliente"]:
            query += " AND cliente LIKE ?"
            params.append(f"%{filtros['cliente']}%")
        
        query += " ORDER BY fecha_creacion DESC LIMIT 50"
        
        conn, cursor = self.master.conectar_db()
        if not conn:
            return []
        
        try:
            cursor.execute(query, params)
            result = cursor.fetchall()
            conn.close()
            return result
        except Exception as e:
            conn.close()
            raise e

    def buscar_en_detalles_con_filtros(self, termino, filtros):
        """Busca en rma_detalles aplicando filtros."""
        query = "SELECT * FROM rma_detalles WHERE 1=1"
        params = []
        
        # Filtro de texto general
        if termino:
            query += """ AND (
                numero_rma LIKE ? OR numero_serie LIKE ? OR referencia_articulo LIKE ? OR 
                descripcion_articulo LIKE ? OR estado_producto LIKE ? OR observaciones LIKE ?
            )"""
            termino_param = f"%{termino}%"
            params.extend([termino_param] * 6)
        
        # Filtros específicos para productos
        if filtros["estado_producto"]:
            query += " AND estado_producto LIKE ?"
            params.append(f"%{filtros['estado_producto']}%")
        
        if filtros["referencia"]:
            query += " AND referencia_articulo LIKE ?"
            params.append(f"%{filtros['referencia']}%")
        
        # Filtros relacionados con el expediente padre
        if filtros["estado"] or filtros["fecha_desde"] or filtros["fecha_hasta"] or filtros["cliente"]:
            query += """ AND numero_rma IN (
                SELECT numero_rma FROM rma_maestro WHERE 1=1"""
            
            if filtros["estado"]:
                query += " AND estado_expediente = ?"
                params.append(filtros["estado"])
            
            if filtros["fecha_desde"]:
                query += " AND fecha_creacion >= ?"
                params.append(filtros["fecha_desde"])
            
            if filtros["fecha_hasta"]:
                query += " AND fecha_creacion <= ?"
                params.append(filtros["fecha_hasta"])
            
            if filtros["cliente"]:
                query += " AND cliente LIKE ?"
                params.append(f"%{filtros['cliente']}%")
            
            query += ")"
        
        query += " ORDER BY numero_rma DESC LIMIT 50"
        
        conn, cursor = self.master.conectar_db()
        if not conn:
            return []
        
        try:
            cursor.execute(query, params)
            result = cursor.fetchall()
            conn.close()
            return result
        except Exception as e:
            conn.close()
            raise e

    def mostrar_resultados_busqueda_avanzada(self, resultados_maestro, resultados_detalles, termino, filtros):
        """Muestra los resultados de la búsqueda avanzada."""
        total_resultados = len(resultados_maestro) + len(resultados_detalles)
        
        # Título de resultados
        titulo_frame = ctk.CTkFrame(self.resultados_frame, fg_color="transparent")
        titulo_frame.pack(fill="x", padx=10, pady=(10,5))
        
        filtros_activos = [k for k, v in filtros.items() if v]
        if termino or filtros_activos:
            criterios = []
            if termino:
                criterios.append(f"Texto: '{termino}'")
            if filtros_activos:
                criterios.append(f"Filtros: {', '.join(filtros_activos)}")
            criterios_text = " | ".join(criterios)
        else:
            criterios_text = "Todos los registros"
        
        ctk.CTkLabel(titulo_frame, 
                    text=f"🔍 Búsqueda Avanzada: {criterios_text}",
                    font=ctk.CTkFont(size=14, weight="bold")).pack(side="left")
        
        ctk.CTkLabel(titulo_frame, 
                    text=f"({total_resultados} resultados)",
                    font=ctk.CTkFont(size=12), 
                    text_color="gray").pack(side="right")
        
        if total_resultados == 0:
            ctk.CTkLabel(self.resultados_frame, 
                        text="❌ No se encontraron resultados con los criterios especificados",
                        font=ctk.CTkFont(size=13),
                        text_color="orange").pack(pady=20)
            return
        
        # Resultados de expedientes
        if resultados_maestro:
            exp_frame = ctk.CTkFrame(self.resultados_frame)
            exp_frame.pack(fill="x", padx=10, pady=5)
            
            ctk.CTkLabel(exp_frame, 
                        text=f"📋 Expedientes ({len(resultados_maestro)} encontrados)",
                        font=ctk.CTkFont(size=13, weight="bold")).pack(anchor="w", padx=10, pady=5)
            
            for expediente in resultados_maestro:
                self.crear_resultado_expediente_avanzado(exp_frame, expediente, termino)
        
        # Resultados de productos
        if resultados_detalles:
            prod_frame = ctk.CTkFrame(self.resultados_frame)
            prod_frame.pack(fill="x", padx=10, pady=5)
            
            ctk.CTkLabel(prod_frame, 
                        text=f"📦 Productos ({len(resultados_detalles)} encontrados)",
                        font=ctk.CTkFont(size=13, weight="bold")).pack(anchor="w", padx=10, pady=5)
            
            for producto in resultados_detalles:
                self.crear_resultado_producto_avanzado(prod_frame, producto, termino)

    def crear_resultado_expediente_avanzado(self, parent, expediente, termino_resaltado):
        """Crea un resultado visual avanzado para expediente."""
        resultado_frame = ctk.CTkFrame(parent)
        resultado_frame.pack(fill="x", padx=10, pady=2)
        
        # Info principal
        info_frame = ctk.CTkFrame(resultado_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=10, pady=5)
        
        # RMA y estado
        header_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        header_frame.pack(fill="x")
        
        rma_texto = f"RMA: {expediente[1]}"  # numero_rma
        ctk.CTkLabel(header_frame, text=rma_texto, 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(side="left")
        
        estado = expediente[8] if expediente[8] else "Sin estado"  # estado_expediente
        color_estado = self.get_color_por_estado(estado)
        ctk.CTkLabel(header_frame, text=f"🏷️ {estado}", 
                    text_color=color_estado, font=ctk.CTkFont(size=11)).pack(side="right")
        
        # Cliente y fecha
        cliente_fecha = f"👤 {expediente[2]} | 📅 {expediente[9]}"  # cliente, fecha_creacion
        ctk.CTkLabel(info_frame, text=cliente_fecha, 
                    font=ctk.CTkFont(size=11), text_color="gray").pack(anchor="w")
        
        # Botón para abrir
        btn_abrir = ctk.CTkButton(resultado_frame, text="📂 Abrir Expediente", 
                                 command=lambda: self.abrir_expediente_desde_busqueda(expediente[1]),
                                 height=25, width=120)
        btn_abrir.pack(side="right", padx=10, pady=5)

    def crear_resultado_producto_avanzado(self, parent, producto, termino_resaltado):
        """Crea un resultado visual avanzado para producto."""
        resultado_frame = ctk.CTkFrame(parent)
        resultado_frame.pack(fill="x", padx=10, pady=2)
        
        # Info principal
        info_frame = ctk.CTkFrame(resultado_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=10, pady=5)
        
        # Referencia y RMA
        header_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        header_frame.pack(fill="x")
        
        ref_texto = f"📦 {producto[3]}"  # referencia_articulo
        ctk.CTkLabel(header_frame, text=ref_texto, 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(side="left")
        
        rma_texto = f"RMA: {producto[1]}"  # numero_rma
        ctk.CTkLabel(header_frame, text=rma_texto, 
                    font=ctk.CTkFont(size=11), text_color="blue").pack(side="right")
        
        # Descripción y estado
        descripcion = producto[4] if producto[4] else "Sin descripción"  # descripcion_articulo
        if len(descripcion) > 60:
            descripcion = descripcion[:57] + "..."
        ctk.CTkLabel(info_frame, text=f"📝 {descripcion}", 
                    font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w")
        
        estado_prod = producto[6] if producto[6] else "Sin estado"  # estado_producto
        ctk.CTkLabel(info_frame, text=f"🔧 Estado: {estado_prod}", 
                    font=ctk.CTkFont(size=10), text_color="orange").pack(anchor="w")
        
        # Botón para abrir
        btn_abrir = ctk.CTkButton(resultado_frame, text="📂 Ver en Expediente", 
                                 command=lambda: self.abrir_expediente_desde_busqueda(producto[1]),
                                 height=25, width=120)
        btn_abrir.pack(side="right", padx=10, pady=5)

    def validar_fecha(self, fecha_str):
        """Valida formato de fecha YYYY-MM-DD."""
        try:
            datetime.datetime.strptime(fecha_str, "%Y-%m-%d")
            return True
        except ValueError:
            return False

    def abrir_expediente_desde_busqueda(self, numero_rma):
        """Abre un expediente específico desde los resultados de búsqueda."""
        try:
            # Cerrar ventana de búsqueda
            if hasattr(self, 'ventana_busqueda_global') and self.ventana_busqueda_global:
                self.ventana_busqueda_global.destroy()
                self.ventana_busqueda_global = None
            
            # Cambiar al marco de expedientes
            self.mostrar_expedientes()
            
            # Buscar y mostrar el expediente específico
            self.buscar_expediente_especifico(numero_rma)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al abrir expediente: {str(e)}")

    def buscar_expediente_especifico(self, numero_rma):
        """Busca y muestra un expediente específico."""
        try:
            # Limpiar lista actual
            for widget in self.expedientes_frame.winfo_children():
                widget.destroy()
            
            # Buscar expediente
            conn, cursor = self.master.conectar_db()
            if not conn:
                return
            
            try:
                cursor.execute("SELECT * FROM rma_maestro WHERE numero_rma = ?", (numero_rma,))
                expediente = cursor.fetchone()
                
                if expediente:
                    # Mostrar como resultado único
                    resultado_frame = ctk.CTkFrame(self.expedientes_frame)
                    resultado_frame.pack(fill="x", padx=10, pady=5)
                    
                    # Crear botón de expediente con información destacada
                    self.crear_boton_expediente(resultado_frame, expediente, destacado=True)
                    
                    # Mensaje de búsqueda exitosa
                    mensaje_frame = ctk.CTkFrame(self.expedientes_frame)
                    mensaje_frame.pack(fill="x", padx=10, pady=5)
                    ctk.CTkLabel(mensaje_frame, 
                                text=f"✅ Expediente {numero_rma} encontrado desde búsqueda",
                                font=ctk.CTkFont(size=12), 
                                text_color="green").pack(pady=10)
                else:
                    ctk.CTkLabel(self.expedientes_frame, 
                                text=f"❌ No se encontró el expediente {numero_rma}",
                                font=ctk.CTkFont(size=12), 
                                text_color="red").pack(pady=20)
                
                conn.close()
                
            except Exception as e:
                conn.close()
                raise e
                    
        except Exception as e:
            messagebox.showerror("Error", f"Error al buscar expediente específico: {str(e)}")
            print(f"Error en buscar_expediente_especifico: {e}")
        
        # Mensaje inicial
        ctk.CTkLabel(self.resultados_frame, 
                    text="👆 Introduce un término de búsqueda arriba para comenzar",
                    font=ctk.CTkFont(size=14),
                    text_color="gray").pack(pady=50)
        
        # Focus en el campo de búsqueda
        self.entry_busqueda_global.focus()

    def ejecutar_busqueda_global(self):
        """Ejecuta la búsqueda global en todos los campos."""
        termino = self.entry_busqueda_global.get().strip()
        
        if not termino:
            self.mostrar_mensaje_busqueda("⚠️ Por favor, introduce un término de búsqueda")
            return
        
        if len(termino) < 2:
            self.mostrar_mensaje_busqueda("⚠️ El término de búsqueda debe tener al menos 2 caracteres")
            return
        
        # Limpiar resultados anteriores
        for widget in self.resultados_frame.winfo_children():
            widget.destroy()
        
        # Mostrar indicador de carga
        loading_label = ctk.CTkLabel(self.resultados_frame, 
                                   text="🔄 Buscando...", 
                                   font=ctk.CTkFont(size=14))
        loading_label.pack(pady=20)
        self.update()
        
        try:
            # Guardar búsqueda en historial
            self.guardar_busqueda_en_historial(termino)
            self.actualizar_historial_ui()
            
            # Realizar búsqueda
            resultados_expedientes, resultados_productos, resultados_historial, resultados_tareas = self.buscar_en_todos_los_campos(termino)
            
            # Limpiar indicador de carga
            loading_label.destroy()
            
            # Mostrar resultados
            self.mostrar_resultados_busqueda(resultados_expedientes, resultados_productos, resultados_historial, resultados_tareas, termino)
            
        except Exception as e:
            loading_label.destroy()
            self.mostrar_mensaje_busqueda(f"❌ Error en la búsqueda: {str(e)}")

    def buscar_en_todos_los_campos(self, termino):
        """Busca el término en todos los campos de las tablas principales."""
        conn, cursor = self.master.conectar_db()
        if not conn:
            return [], []
        
        cursor = conn.cursor()
        termino_like = f"%{termino}%"
        
        try:
            # Búsqueda en rma_maestro (expedientes)
            # Obtenemos primero todas las columnas disponibles
            cursor.execute("PRAGMA table_info(rma_maestro)")
            columnas_maestro = [col[1] for col in cursor.fetchall()]
            
            # Construir query dinámicamente para todos los campos de texto
            campos_busqueda_maestro = []
            params_maestro = []
            
            for col in columnas_maestro:
                if col.lower() not in ['id']:  # Excluir campos numéricos ID
                    campos_busqueda_maestro.append(f"{col} LIKE ?")
                    params_maestro.append(termino_like)
            
            if campos_busqueda_maestro:
                sql_maestro = f"""
                SELECT id, codigo_rma, cliente, numero_documento_cliente, estado, fecha_emision,
                       rma_proveedor, motivo, fecha_gestion
                FROM rma_maestro 
                WHERE {' OR '.join(campos_busqueda_maestro)}
                ORDER BY fecha_emision DESC
                LIMIT 100
                """
                cursor.execute(sql_maestro, params_maestro)
                resultados_expedientes = cursor.fetchall()
            else:
                resultados_expedientes = []
            
            # Búsqueda en rma_detalles (productos)
            cursor.execute("PRAGMA table_info(rma_detalles)")
            columnas_detalles = [col[1] for col in cursor.fetchall()]
            
            campos_busqueda_detalles = []
            params_detalles = []
            
            for col in columnas_detalles:
                if col.lower() not in ['id', 'rma_id', 'cantidad_segun_documento', 'cantidad_entregada', 'precio_unitario']:
                    campos_busqueda_detalles.append(f"d.{col} LIKE ?")
                    params_detalles.append(termino_like)
            
            if campos_busqueda_detalles:
                sql_detalles = f"""
                SELECT DISTINCT d.id, d.rma_id, d.referencia_articulo, d.cantidad_segun_documento,
                       d.cantidad_entregada, d.estado_producto, d.precio_unitario, m.codigo_rma, m.cliente
                FROM rma_detalles d
                JOIN rma_maestro m ON d.rma_id = m.id
                WHERE {' OR '.join(campos_busqueda_detalles)}
                ORDER BY m.fecha_emision DESC
                LIMIT 100
                """
                cursor.execute(sql_detalles, params_detalles)
                resultados_productos = cursor.fetchall()
            else:
                resultados_productos = []
            
            # Búsqueda en historial (opcional si la tabla existe)
            resultados_historial = []
            try:
                sql_historial = """
                SELECT DISTINCT h.rma_id, h.fecha_cambio, h.descripcion_cambio, h.usuario, 
                       m.codigo_rma, m.cliente
                FROM rma_historial h
                JOIN rma_maestro m ON h.rma_id = m.id
                WHERE h.descripcion_cambio LIKE ? OR h.usuario LIKE ?
                ORDER BY h.fecha_cambio DESC
                LIMIT 100
                """
                cursor.execute(sql_historial, (termino_like, termino_like))
                resultados_historial = cursor.fetchall()
            except Exception:
                pass  # La tabla rma_historial puede no existir
            
            # Búsqueda en tareas
            resultados_tareas = []
            try:
                sql_tareas = """
                SELECT DISTINCT t.id, t.codigo_rma, t.titulo, t.descripcion, t.estado,
                       t.fecha_vencimiento, t.creado_por, m.id as rma_id, m.cliente
                FROM tareas t
                LEFT JOIN rma_maestro m ON t.codigo_rma = m.codigo_rma
                WHERE t.titulo LIKE ? OR t.descripcion LIKE ? OR t.estado LIKE ? OR t.creado_por LIKE ?
                ORDER BY t.fecha_vencimiento IS NULL, t.fecha_vencimiento DESC
                LIMIT 100
                """
                cursor.execute(sql_tareas, (termino_like, termino_like, termino_like, termino_like))
                resultados_tareas = cursor.fetchall()
            except Exception:
                pass  # Error en búsqueda de tareas
            
            conn.close()
            return resultados_expedientes, resultados_productos, resultados_historial, resultados_tareas
            
        except Exception as e:
            conn.close()
            raise e

    def mostrar_resultados_busqueda(self, expedientes, productos, historial, tareas, termino):
        """Muestra los resultados de la búsqueda de forma organizada."""
        total_resultados = len(expedientes) + len(productos) + len(historial) + len(tareas)
        
        if total_resultados == 0:
            self.mostrar_mensaje_busqueda(f"🔍 No se encontraron resultados para '{termino}'")
            return
        
        # Header de resultados
        header_resultados = ctk.CTkLabel(self.resultados_frame,
                                       text=f"📊 Se encontraron {total_resultados} resultados para '{termino}'",
                                       font=ctk.CTkFont(size=16, weight="bold"))
        header_resultados.pack(pady=10, anchor="w")
        
        # Mostrar expedientes
        if expedientes:
            self.mostrar_seccion_expedientes(expedientes)
        
        # Mostrar productos
        if productos:
            self.mostrar_seccion_productos(productos)
        
        # Mostrar historial
        if historial:
            self.mostrar_seccion_historial(historial)
        
        # Mostrar tareas
        if tareas:
            self.mostrar_seccion_tareas(tareas)

    def mostrar_seccion_expedientes(self, expedientes):
        """Muestra la sección de expedientes encontrados."""
        seccion_frame = ctk.CTkFrame(self.resultados_frame)
        seccion_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(seccion_frame, 
                    text=f"📋 EXPEDIENTES ENCONTRADOS ({len(expedientes)})",
                    font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        for exp in expedientes:
            exp_id, codigo_rma, cliente, num_doc, estado, fecha, rma_prov, motivo, fecha_gestion = exp
            
            item_frame = ctk.CTkFrame(seccion_frame)
            item_frame.pack(fill="x", padx=10, pady=2)
            
            # Información principal
            info_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            info_frame.pack(fill="x", padx=10, pady=5)
            
            # Línea 1: Código RMA y fecha
            linea1 = ctk.CTkFrame(info_frame, fg_color="transparent")
            linea1.pack(fill="x")
            
            ctk.CTkLabel(linea1, text=f"📋 {codigo_rma or 'Sin código'}", 
                        font=ctk.CTkFont(weight="bold")).pack(side="left")
            ctk.CTkLabel(linea1, text=f"📅 {fecha or 'Sin fecha'}", 
                        text_color="gray").pack(side="right")
            
            # Línea 2: Cliente y documento
            if cliente or num_doc:
                linea2 = ctk.CTkLabel(info_frame, 
                                    text=f"👤 {cliente or 'Sin cliente'} | 📄 {num_doc or 'Sin doc.'}",
                                    text_color="gray")
                linea2.pack(anchor="w")
            
            # Línea 3: Estado
            if estado:
                color_estado = self.get_color_por_estado(estado)
                ctk.CTkLabel(info_frame, text=f"🏷️ {estado}", 
                           text_color=color_estado).pack(anchor="w")
            
            # Botón para abrir expediente
            btn_abrir = ctk.CTkButton(item_frame, text="📖 Abrir Expediente", 
                                     command=lambda eid=exp_id: self._abrir_editor_rma(rma_id=eid),
                                     width=150, height=30)
            btn_abrir.pack(side="right", padx=10, pady=5)

    def mostrar_seccion_productos(self, productos):
        """Muestra la sección de productos encontrados."""
        seccion_frame = ctk.CTkFrame(self.resultados_frame)
        seccion_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(seccion_frame, 
                    text=f"📦 PRODUCTOS ENCONTRADOS ({len(productos)})",
                    font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        for prod in productos:
            prod_id, rma_id, ref_articulo, cantidad_doc, cantidad_entregada, estado_prod, precio_unit, codigo_rma, cliente = prod
            
            item_frame = ctk.CTkFrame(seccion_frame)
            item_frame.pack(fill="x", padx=10, pady=2)
            
            # Información principal
            info_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            info_frame.pack(fill="x", padx=10, pady=5)
            
            # Línea 1: Referencia y precio
            linea1 = ctk.CTkFrame(info_frame, fg_color="transparent")
            linea1.pack(fill="x")
            
            ctk.CTkLabel(linea1, text=f"📦 {ref_articulo or 'Sin referencia'}", 
                        font=ctk.CTkFont(weight="bold")).pack(side="left")
            if precio_unit:
                ctk.CTkLabel(linea1, text=f"💰 {precio_unit}€", 
                           text_color="blue").pack(side="right")
            
            # Línea 2: Cantidades
            if cantidad_doc or cantidad_entregada:
                ctk.CTkLabel(info_frame, text=f"� Doc: {cantidad_doc or 0} | Entregada: {cantidad_entregada or 0}", 
                           text_color="gray").pack(anchor="w")
            
            # Línea 3: Expediente asociado
            ctk.CTkLabel(info_frame, text=f"📋 Expediente: {codigo_rma} | 👤 {cliente or 'Sin cliente'}", 
                        text_color="gray").pack(anchor="w")
            
            # Línea 4: Estado del producto
            if estado_prod:
                color_estado = {"Nuevo": "green", "Usado": "orange", "Defectuoso": "red"}.get(estado_prod, "gray")
                ctk.CTkLabel(info_frame, text=f"⚪ {estado_prod}", 
                           text_color=color_estado).pack(anchor="w")
            
            # Botón para abrir expediente
            btn_abrir = ctk.CTkButton(item_frame, text="📖 Ver en Expediente", 
                                     command=lambda rid=rma_id: self._abrir_editor_rma(rma_id=rid),
                                     width=150, height=30)
            btn_abrir.pack(side="right", padx=10, pady=5)

    def mostrar_seccion_historial(self, historial):
        """Muestra la sección de historial encontrado."""
        seccion_frame = ctk.CTkFrame(self.resultados_frame)
        seccion_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(seccion_frame, 
                    text=f"📜 HISTORIAL ENCONTRADO ({len(historial)})",
                    font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        for hist in historial:
            # Columnas: rma_id, fecha_cambio, descripcion_cambio, usuario, codigo_rma, cliente
            rma_id, fecha_cambio, descripcion_cambio, usuario, codigo_rma, cliente = hist
            
            item_frame = ctk.CTkFrame(seccion_frame)
            item_frame.pack(fill="x", padx=10, pady=2)
            
            # Información principal
            info_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            info_frame.pack(fill="x", padx=10, pady=5)
            
            # Línea 1: Fecha y usuario
            linea1 = ctk.CTkFrame(info_frame, fg_color="transparent")
            linea1.pack(fill="x")
            
            ctk.CTkLabel(linea1, text=f"📅 {fecha_cambio or 'Sin fecha'}", 
                        font=ctk.CTkFont(weight="bold")).pack(side="left")
            ctk.CTkLabel(linea1, text=f"👨 {usuario or 'Sin usuario'}", 
                        text_color="blue").pack(side="right")
            
            # Línea 2: Descripción del cambio
            desc_text = (descripcion_cambio[:100] + "...") if descripcion_cambio and len(descripcion_cambio) > 100 else (descripcion_cambio or "Sin descripción")
            ctk.CTkLabel(info_frame, text=desc_text, text_color="gray").pack(anchor="w")
            
            # Línea 3: Expediente
            ctk.CTkLabel(info_frame, 
                        text=f"📋 {codigo_rma or 'Sin código'} | 👤 Cliente: {cliente or 'Sin cliente'}",
                        text_color="gray").pack(anchor="w")
            
            # Botón para abrir expediente
            btn_abrir = ctk.CTkButton(item_frame, text="📖 Ver Expediente", 
                                     command=lambda rid=rma_id: self._abrir_editor_rma(rma_id=rid),
                                     width=150, height=30)
            btn_abrir.pack(side="right", padx=10, pady=5)

    def mostrar_seccion_tareas(self, tareas):
        """Muestra la sección de tareas encontradas."""
        seccion_frame = ctk.CTkFrame(self.resultados_frame)
        seccion_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(seccion_frame, 
                    text=f"✅ TAREAS ENCONTRADAS ({len(tareas)})",
                    font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        for tarea in tareas:
            # Columnas: id, codigo_rma, titulo, descripcion, estado, fecha_vencimiento, creado_por, rma_id, cliente
            tarea_id, codigo_rma_tarea, titulo, descripcion, estado, fecha_vencimiento, creado_por, rma_id, cliente = tarea
            
            item_frame = ctk.CTkFrame(seccion_frame)
            item_frame.pack(fill="x", padx=10, pady=2)
            
            # Información principal
            info_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
            info_frame.pack(fill="x", padx=10, pady=5)
            
            # Línea 1: Título de la tarea
            linea1 = ctk.CTkFrame(info_frame, fg_color="transparent")
            linea1.pack(fill="x")
            
            ctk.CTkLabel(linea1, text=f"📝 {titulo or 'Sin título'}"[:60] + ("..." if titulo and len(titulo) > 60 else ""),
                        font=ctk.CTkFont(weight="bold")).pack(side="left")
            
            # Línea 2: Descripción
            if descripcion:
                desc_text = (descripcion[:80] + "...") if len(descripcion) > 80 else descripcion
                ctk.CTkLabel(info_frame, text=desc_text, text_color="gray").pack(anchor="w")
            
            # Línea 3: Estado y fecha de vencimiento
            color_estado = {"Pendiente": "orange", "En Progreso": "blue", "Completada": "green", "Cancelada": "red"}.get(estado or "", "gray")
            estado_text = f"🏷️ Estado: {estado or 'Sin estado'}"
            if fecha_vencimiento:
                estado_text += f" | ⏰ Vence: {fecha_vencimiento}"
            ctk.CTkLabel(info_frame, text=estado_text, text_color=color_estado).pack(anchor="w")
            
            # Línea 4: Creado por
            if creado_por:
                ctk.CTkLabel(info_frame, text=f"👨 Creado por: {creado_por}", 
                            text_color="gray").pack(anchor="w")
            
            # Línea 5: Expediente asociado
            if codigo_rma_tarea:
                ctk.CTkLabel(info_frame, text=f"📋 Expediente: {codigo_rma_tarea} | 👤 {cliente or 'Sin cliente'}", 
                            text_color="gray").pack(anchor="w")
            
            # Botón para abrir expediente
            if rma_id:
                btn_abrir = ctk.CTkButton(item_frame, text="📖 Ver Expediente", 
                                         command=lambda rid=rma_id: self._abrir_editor_rma(rma_id=rid),
                                         width=150, height=30)
                btn_abrir.pack(side="right", padx=10, pady=5)

    def mostrar_mensaje_busqueda(self, mensaje):
        """Muestra un mensaje en el área de resultados."""
        for widget in self.resultados_frame.winfo_children():
            widget.destroy()
        
        ctk.CTkLabel(self.resultados_frame, text=mensaje, 
                    font=ctk.CTkFont(size=14),
                    text_color="gray").pack(pady=50)

    def limpiar_busqueda_global(self):
        """Limpia el campo de búsqueda y resultados."""
        self.entry_busqueda_global.delete(0, 'end')
        self.mostrar_mensaje_busqueda("👆 Introduce un término de búsqueda arriba para comenzar")
        self.entry_busqueda_global.focus()

    def test_conexion_busqueda(self):
        """Prueba rápida de la conexión para búsqueda."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn:
                return False
            
            cursor = conn.cursor()
            # Prueba simple: obtener count de expedientes
            cursor.execute("SELECT COUNT(*) FROM rma_maestro")
            count = cursor.fetchone()[0]
            conn.close()
            
            return True
            
        except Exception as e:
            print(f"❌ Error de conexión en búsqueda: {e}")
            return False

    # ----------------------------------------------------------------------
    # 5. LÓGICA DE FORMULARIO (CREAR/EDITAR)
    # ----------------------------------------------------------------------

    def _abrir_editor_rma(self, rma_id=None):
        """Abre el editor de RMA en una ventana separada."""
        from lib.rma_editor_window import RmaEditorWindow
        
        try:
            # Restaurar el content_frame original si fue modificado por una ventana modal
            if hasattr(self, '_original_content_frame') and self._original_content_frame:
                try:
                    if self._original_content_frame.winfo_exists():
                        self.content_frame = self._original_content_frame
                except:
                    pass
            
            # Abrir directamente la nueva ventana
            # Si hay otras ventanas abiertas, no las cerramos automáticamente
            # El usuario las puede cerrar manualmente si lo desea
            RmaEditorWindow(self, rma_id)
                
        except Exception as e:
            messagebox.showerror("Error", f"Error al abrir editor RMA: {str(e)}")
    
    def _mostrar_widget_tiempos(self, parent_frame, rma_id):
        """Muestra el widget de tiempos de tramitación en la ficha del expediente."""
        try:
            # Obtener datos del expediente
            conn = connect_db()
            if not conn:
                return
            
            cursor = conn.cursor()
            cursor.execute("""
                SELECT fecha_emision, fecha_autorizacion, fecha_recepcion, 
                       fecha_proceso, fecha_gestion, cliente
                FROM rma_maestro
                WHERE id = ?
            """, (rma_id,))
            
            row = cursor.fetchone()
            if not row:
                conn.close()
                return
            
            fecha_emision, fecha_autorizacion, fecha_recepcion, fecha_proceso, fecha_gestion, cliente = row
            
            # Calcular tiempos del expediente
            tiempos = calcular_tiempos_expediente(
                fecha_emision, fecha_autorizacion, fecha_recepcion, 
                fecha_proceso, fecha_gestion
            )
            
            # Obtener promedio del cliente si está cerrado
            promedio_cliente = None
            if tiempos['cerrado'] and cliente:
                promedio_info = obtener_promedio_cliente(cliente, conn)
                promedio_cliente = promedio_info['promedio_total']
            
            conn.close()
            
            # Si no hay datos de tiempo, no mostrar nada
            if tiempos['dias_total'] is None:
                return
            
            # Frame para el widget de tiempos (row 2 en la columna 0)
            tiempos_frame = ctk.CTkFrame(parent_frame, fg_color="#f0f0f0", corner_radius=8)
            tiempos_frame.grid(row=2, column=0, padx=10, pady=(5, 10), sticky="ew")
            
            # Título del widget
            ctk.CTkLabel(tiempos_frame, text="📊 TIEMPOS DE TRAMITACIÓN", 
                        font=ctk.CTkFont(size=12, weight="bold"),
                        text_color="#2c3e50").pack(anchor="w", padx=10, pady=(8, 5))
            
            # Frame horizontal para los indicadores
            indicadores_frame = ctk.CTkFrame(tiempos_frame, fg_color="transparent")
            indicadores_frame.pack(fill="x", padx=10, pady=(0, 8))
            
            # Días totales
            dias_total = tiempos['dias_total']
            color_total = obtener_color_tiempo(dias_total)
            estado_texto = " (Cerrado)" if tiempos['cerrado'] else " (En curso)"
            
            total_frame = ctk.CTkFrame(indicadores_frame, fg_color="white", corner_radius=6)
            total_frame.pack(side="left", padx=(0, 8), pady=2, fill="x", expand=True)
            
            ctk.CTkLabel(total_frame, text="Total:", font=ctk.CTkFont(size=10)).pack(anchor="w", padx=8, pady=(4, 0))
            ctk.CTkLabel(total_frame, text=f"{dias_total} días{estado_texto}", 
                        font=ctk.CTkFont(size=13, weight="bold"),
                        text_color=color_total).pack(anchor="w", padx=8, pady=(0, 4))
            
            # Promedio del cliente (si está disponible)
            if promedio_cliente is not None:
                promedio_frame = ctk.CTkFrame(indicadores_frame, fg_color="white", corner_radius=6)
                promedio_frame.pack(side="left", padx=(0, 8), pady=2, fill="x", expand=True)
                
                dias_prom = int(promedio_cliente)
                color_prom = obtener_color_tiempo(dias_prom)
                
                ctk.CTkLabel(promedio_frame, text=f"Promedio {cliente}:", 
                           font=ctk.CTkFont(size=10)).pack(anchor="w", padx=8, pady=(4, 0))
                ctk.CTkLabel(promedio_frame, text=f"{dias_prom} días", 
                           font=ctk.CTkFont(size=13, weight="bold"),
                           text_color=color_prom).pack(anchor="w", padx=8, pady=(0, 4))
            
            # Tiempos entre fases (en una línea horizontal)
            fases_frame = ctk.CTkFrame(indicadores_frame, fg_color="white", corner_radius=6)
            fases_frame.pack(side="left", padx=0, pady=2, fill="x", expand=True)
            
            ctk.CTkLabel(fases_frame, text="Fases:", font=ctk.CTkFont(size=10)).pack(anchor="w", padx=8, pady=(4, 0))
            
            # Frame horizontal para las fases
            fases_detalle_frame = ctk.CTkFrame(fases_frame, fg_color="transparent")
            fases_detalle_frame.pack(fill="x", padx=8, pady=(0, 4))
            
            fases = [
                ('E→A', tiempos['dias_e_a']),
                ('A→R', tiempos['dias_a_r']),
                ('R→P', tiempos['dias_r_p']),
                ('P→C', tiempos['dias_p_c'])
            ]
            
            for nombre, dias in fases:
                if dias is not None:
                    color = obtener_color_tiempo(dias)
                    lbl = ctk.CTkLabel(fases_detalle_frame, 
                                      text=f"{nombre}: {dias}d", 
                                      font=ctk.CTkFont(size=11),
                                      text_color=color)
                    lbl.pack(side="left", padx=4)
            
        except Exception as e:
            print(f"Error al mostrar widget de tiempos: {e}")

    def mostrar_nuevo_rma(self, rma_id=None):
        """Muestra el formulario para crear (rma_id=None) o editar (rma_id=ID) un RMA."""
        self.limpiar_contenido()
        self.rma_actual_id = rma_id
        self.articulos_data = [] # Lista temporal que contendrá los artículos
        
        self.datos_rma_maestro = {}
        
        if rma_id is not None:
            # Modo Edición
            self.mode = 'editar'
            # Usaremos 'current_rma_id' para ser consistentes con la función de guardar
            self.current_rma_id = rma_id
        else:
            # Modo Nuevo
            self.mode = 'nuevo'
            self.current_rma_id = None
        
        es_edicion = rma_id is not None

        # Garantizar que exista el atributo antes de llamar (evita AttributeError si la función se define más abajo)
        if not hasattr(self, 'cargar_lista_tareas_rma'):
            self.cargar_lista_tareas_rma = lambda: None
        
        # --- Cabecera ---
        titulo_texto = "EDITAR EXPEDIENTE" if es_edicion else "CREAR NUEVO EXPEDIENTE"
        header_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 5))
        header_frame.grid_columnconfigure(0, weight=1)
        
        # 🛠️ 1. AJUSTE DE PESO: Fila 0 (Cabecera Principal)
        self.content_frame.grid_rowconfigure(0, weight=0) # No se expande
        # --------------------------------------------------------------------------

        ctk.CTkLabel(header_frame, text=titulo_texto, font=ctk.CTkFont(size=24, weight="bold")).grid(row=0, column=0, sticky="w")
        
        # Detectar si estamos en una ventana modal (RmaEditorWindow)
        # Si es así, el botón debe cerrar la ventana modal
        ventana_actual = self.content_frame.winfo_toplevel()
        es_ventana_modal = ventana_actual != self.master
        
        if es_ventana_modal:
            # Estamos en RmaEditorWindow - cerrar la ventana
            btn_volver = ctk.CTkButton(header_frame, 
                                       text="✖️ Cerrar", 
                                       command=ventana_actual.destroy)
        else:
            # Estamos en la ventana principal - volver a la lista
            btn_volver = ctk.CTkButton(header_frame, 
                                       text="⬅️ Volver", 
                                       command=self.mostrar_lista_rma)
        
        btn_volver.grid(row=0, column=1, padx=(20, 0), sticky="e")
        
        # --------------------------------------------------------------------------
        # 2. Fila Principal (Código RMA + Comentarios Fijos) (Fila 1)
        # --------------------------------------------------------------------------
        # Creamos un frame de control para la Fila 1
        fila1_control_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        fila1_control_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 10))
        
        # Configuramos las columnas dentro de este frame: Col 0 es fija, Col 1 se expande
        fila1_control_frame.grid_columnconfigure(0, weight=0) # Fija para el código
        fila1_control_frame.grid_columnconfigure(1, weight=1) # Expansiva para los comentarios

        # A) NÚMERO DE EXPEDIENTE (Columna 0)
        if es_edicion:
            # Consultar el código RMA real desde la base de datos
            conn = connect_db()
            cur = conn.cursor()
            cur.execute("SELECT codigo_rma FROM rma_maestro WHERE id = ?", (rma_id,))
            row = cur.fetchone()
            conn.close()
            codigo_rma_mostrar = row[0] if row else "(Desconocido)"
            texto_expediente = f"Nº EXPEDIENTE: {codigo_rma_mostrar}"
            color_texto = "grey30"
        else:
            codigo_rma_temporal = self.obtener_siguiente_rma()
            texto_expediente = f"Nº EXPEDIENTE: {codigo_rma_temporal} (Temporal)"
            color_texto = "orange"  # Color naranja para indicar que es temporal
        
        self.lbl_codigo_rma = ctk.CTkLabel(fila1_control_frame, text=texto_expediente, 
                     font=ctk.CTkFont(size=18, weight="bold"), 
                     text_color=color_texto)
        self.lbl_codigo_rma.grid(row=0, column=0, padx=10, pady=5, sticky="w") 
        
        # Añadir texto explicativo para números temporales
        if not es_edicion:
            self.lbl_explicacion = ctk.CTkLabel(fila1_control_frame, 
                         text="⚠️ Este número es temporal. El número definitivo se asignará al guardar.", 
                         font=ctk.CTkFont(size=10),
                         text_color="gray50")
            self.lbl_explicacion.grid(row=1, column=0, padx=10, pady=(0, 5), sticky="w")
        
        # Widget de tiempos de tramitación (solo en modo edición)
        if es_edicion:
            self._mostrar_widget_tiempos(fila1_control_frame, rma_id)
        
        # Guardar si es modo edición para usar en el guardado
        self.es_modo_edicion = es_edicion 
        
        
        # B) CAJA DE COMENTARIOS (Columna 1)
        comentarios_frame = ctk.CTkFrame(fila1_control_frame, fg_color="transparent") 
        # Cambiamos el contenedor a fila1_control_frame y lo ponemos en column=1
        comentarios_frame.grid(row=0, column=1, sticky="ew", padx=10, pady=5)
        comentarios_frame.grid_columnconfigure(0, weight=1) 

        # Etiqueta
        ctk.CTkLabel(comentarios_frame, text="Comentarios (Guarde al momento con el botón ➕):", 
                     text_color="black",
                     font=ctk.CTkFont(size=11, weight="bold")).grid(row=0, column=0, padx=5, pady=(5, 0), sticky="nw")
        
        # PRECIO TOTAL EXPEDIENTE (visible siempre)
        precio_total_frame = ctk.CTkFrame(comentarios_frame, fg_color="#e8f5e9", corner_radius=8)
        precio_total_frame.grid(row=3, column=0, padx=5, pady=(10, 5), sticky="ew")
        
        ctk.CTkLabel(precio_total_frame, text="💰 PRECIO TOTAL EXPEDIENTE:", 
                     font=ctk.CTkFont(size=12, weight="bold"),
                     text_color="#2e7d32").pack(side="left", padx=10, pady=8)
        
        self.lbl_precio_total = ctk.CTkLabel(precio_total_frame, text="0.00 €", 
                                             font=ctk.CTkFont(size=16, weight="bold"), 
                                             text_color="#1b5e20")
        self.lbl_precio_total.pack(side="left", padx=5, pady=8)
        
        comentario_input_frame = ctk.CTkFrame(comentarios_frame, fg_color="transparent")
        comentario_input_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=(5, 5))
        comentario_input_frame.grid_columnconfigure(0, weight=1) # El textbox se expande
        
        # 1. Textbox (con altura reducida) - va en la Columna 0 del nuevo frame
        self.textbox_comentarios = ctk.CTkTextbox(comentario_input_frame, 
                                                  height=40, # Altura compacta
                                                  #fg_color="gray95", 
                                                  #text_color="black",
                                                  wrap="word")
        self.textbox_comentarios.grid(row=0, column=0, sticky="ew")

        # 2. Botón de Guardar Comentario - va en la Columna 1 del nuevo frame
        ctk.CTkButton(comentario_input_frame, 
                      text="➕", 
                      width=40, 
                      #fg_color="gray70", 
                      #hover_color="gray60", 
                      #text_color="black", 
                      command=self.guardar_comentario_historial
                      ).grid(row=0, column=1, padx=(5, 0), sticky="e")
        
        # 🛠️ 2. AJUSTE DE PESO: Fila 1 (Código RMA y Status)
        self.content_frame.grid_rowconfigure(1, weight=0) # No se expande
        # --------------------------------------------------------------------------
        
       
        
        # 3. Vista con pestañas (Tabview) para el formulario y el historial
        self.tabview = ctk.CTkTabview(self.content_frame)
        self.tabview.grid(row=2, column=0, sticky="nsew", padx=10, pady=10)
        self.content_frame.grid_rowconfigure(2, weight=1)

        # -----------------------------------------------------------
        # -- 1. PESTAÑAS PRINCIPALES DEL FORMULARIO (NUEVAS) --
        # -----------------------------------------------------------
        general_tab = self.tabview.add("📝 General")
        estados_fechas_tab = self.tabview.add("⏱️ Estados y Fechas")
        articulos_tab = self.tabview.add("📦 Artículos")
        contabilidad_tab = self.tabview.add("💰 Contabilidad")
        # Nueva pestaña para información técnica — por si en el futuro añadimos más campos técnicos
        info_tecnica_tab = self.tabview.add("🔧 Información Técnica")
        # Determinar título de la pestaña según el modo de almacenamiento
        if usar_dropbox():
            adjuntos_tab = self.tabview.add("📎 Adjuntos (Dropbox)")
        else:
            adjuntos_tab = self.tabview.add("📎 Adjuntos (Local)")
        historial_tab = self.tabview.add("📜 Historial de Cambios")
        # Pestaña de Tareas por RMA (creación/edición desde la ficha del expediente)
        tareas_tab = self.tabview.add("🗒️ Tareas")
        self.historial_tab = historial_tab

        # Configurar todas las pestañas con un marco scrollable (excepto Adjuntos/Historial, si es necesario)
        # Hacemos los marcos scrollable y los frames internos transparentes
        general_scroll = ctk.CTkScrollableFrame(general_tab, label_text="Datos de Identificación")
        general_scroll.pack(fill="both", expand=True, padx=10, pady=10)
        general_frame = ctk.CTkFrame(general_scroll, fg_color="transparent")
        general_frame.pack(fill="x", padx=10, pady=10)
        general_frame.grid_columnconfigure(1, weight=1)

        estados_fechas_scroll = ctk.CTkScrollableFrame(estados_fechas_tab, label_text="Trazabilidad y Proceso")
        estados_fechas_scroll.pack(fill="both", expand=True, padx=10, pady=10)
        estados_fechas_frame = ctk.CTkFrame(estados_fechas_scroll, fg_color="transparent")
        estados_fechas_frame.pack(fill="x", padx=10, pady=10)
        estados_fechas_frame.grid_columnconfigure(1, weight=1)
        
        contabilidad_scroll = ctk.CTkScrollableFrame(contabilidad_tab, label_text="Cierre del Expediente")
        contabilidad_scroll.pack(fill="both", expand=True, padx=10, pady=10)
        contabilidad_frame = ctk.CTkFrame(contabilidad_scroll, fg_color="transparent")
        contabilidad_frame.pack(fill="x", padx=10, pady=10)
        contabilidad_frame.grid_columnconfigure(1, weight=1)

        # Configurar la nueva pestaña Información Técnica
        info_tecnica_scroll = ctk.CTkScrollableFrame(info_tecnica_tab, label_text="Información Técnica")
        info_tecnica_scroll.pack(fill="both", expand=True, padx=10, pady=10)
        info_tecnica_frame = ctk.CTkFrame(info_tecnica_scroll, fg_color="transparent")
        info_tecnica_frame.pack(fill="x", padx=10, pady=10)
        info_tecnica_frame.grid_columnconfigure(1, weight=1)

        # El historial solo se muestra si estamos editando
        if es_edicion:
            self.mostrar_historial(historial_tab)
            
        # Configurar la pestaña de Tareas - siempre visible pero muestra mensaje diferente si es nuevo
        tareas_scroll = ctk.CTkScrollableFrame(tareas_tab, label_text="Tareas asociadas")
        tareas_scroll.pack(fill="both", expand=True, padx=10, pady=10)
        tareas_frame = ctk.CTkFrame(tareas_scroll, fg_color="transparent")
        tareas_frame.pack(fill="x", padx=10, pady=10)

        # Lista de tareas para este RMA
        # Crear solo una vez el frame y el título
        self.tareas_list_frame = ctk.CTkFrame(tareas_frame)
        self.tareas_list_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        tareas_frame.grid_columnconfigure(0, weight=1)
        self.tareas_title_label = ctk.CTkLabel(self.tareas_list_frame, text="Tareas asociadas", font=("Arial", 14, "bold"))
        self.tareas_title_label.pack(pady=(10, 5))
        # Si no es edición, mostramos un mensaje informativo
        if not es_edicion:
            ctk.CTkLabel(self.tareas_list_frame, text="Guarde el expediente para poder crear tareas.").pack(pady=10)
        # Siempre refrescar la lista de tareas al abrir la ficha
        self.cargar_lista_tareas_rma()

        # Función para refrescar historial tras operaciones
        def refrescar_historial():
            self.mostrar_historial(historial_tab)

        # Exponer la función para uso en otras partes
        self.refrescar_historial = refrescar_historial

        # Botón para crear nueva tarea (auto-llena código RMA y creador)
        def crear_tarea_en_rma():
            # Solo permitir si el RMA fue guardado
            if self.current_rma_id is None:
                messagebox.showwarning("Guardar primero", "Guarde el expediente antes de crear tareas.")
                return

            # Abrir un pequeño diálogo para título, descripción y fecha
            dlg = ctk.CTkToplevel(self)
            dlg.title("Crear tarea")
            dlg.geometry("400x300")
            dlg.grab_set()

            ctk.CTkLabel(dlg, text=f"RMA: {self.lbl_codigo_rma.cget('text').split(': ')[1]}").pack(pady=5)
            ctk.CTkLabel(dlg, text=f"Creador: {self.username}").pack(pady=5)

            ctk.CTkLabel(dlg, text="Título:").pack(pady=(10,0))
            titulo_entry = ctk.CTkEntry(dlg)
            titulo_entry.pack(padx=10, pady=5, fill='x')

            ctk.CTkLabel(dlg, text="Descripción:").pack(pady=(10,0))
            desc_text = tk.Text(dlg, height=5)
            desc_text.pack(padx=10, pady=5, fill='both', expand=True)

            ctk.CTkLabel(dlg, text="Fecha Vencimiento (YYYY-MM-DD):").pack(pady=(5,0))
            fecha_entry = ctk.CTkEntry(dlg)
            fecha_entry.pack(padx=10, pady=5, fill='x')

            def confirmar_crear():
                titulo = titulo_entry.get().strip()
                descripcion = desc_text.get("1.0", "end").strip()
                fecha_v = fecha_entry.get().strip() or None
                codigo_rma = self.lbl_codigo_rma.cget('text').split(': ')[1]
                
                if not titulo:
                    messagebox.showerror("Error", "El título es obligatorio.")
                    return
                    
                # Validar formato de fecha
                if fecha_v:
                    try:
                        fecha_obj = datetime.datetime.strptime(fecha_v, "%Y-%m-%d")
                        if fecha_obj.date() < datetime.date.today():
                            if not messagebox.askyesno("Advertencia", 
                                "La fecha de vencimiento es anterior a hoy. ¿Desea continuar?"):
                                return
                    except ValueError:
                        messagebox.showerror("Error", "Formato de fecha inválido. Use YYYY-MM-DD")
                        return
                try:
                    conn = connect_db()
                    cur = conn.cursor()
                    cur.execute(
                        "INSERT INTO tareas (codigo_rma, titulo, descripcion, fecha_vencimiento, estado, creado_por, creado_en, notificado) VALUES (?, ?, ?, ?, ?, ?, ?, 0)",
                        (codigo_rma, titulo, descripcion, fecha_v, 'Pendiente', self.username, datetime.datetime.now().isoformat())
                    )
                    conn.commit()
                    
                    # Registrar en historial del RMA
                    cur.execute("SELECT id FROM rma_maestro WHERE codigo_rma = ?", (codigo_rma,))
                    rma_row = cur.fetchone()
                    if rma_row:
                        rma_id = rma_row[0]
                        cur.execute("""
                            INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                            VALUES (?, ?, ?, ?)
                        """, (rma_id, datetime.datetime.now().isoformat(), self.username,
                             f"Nueva tarea creada: {titulo}"))
                        conn.commit()
                        conn.close()
                    
                    dlg.destroy()
                    # Recargar la lista e historial si corresponde
                    if hasattr(self, 'cargar_lista_tareas_rma'):
                        self.cargar_lista_tareas_rma()
                    if hasattr(self, 'refrescar_historial'):
                        self.refrescar_historial()
                    # Actualizar badge de tareas
                    self.actualizar_badge_tareas()
                    messagebox.showinfo("Éxito", "✅ Tarea creada correctamente")
                except sqlite3.Error as e:
                    messagebox.showerror("Error BD", f"No se pudo crear la tarea: {e}")
                    if 'conn' in locals():
                        conn.close()

            ctk.CTkButton(dlg, text="Crear", command=confirmar_crear).pack(pady=10)

        # Mostrar el botón de crear solo en modo edición
        if es_edicion:
            ctk.CTkButton(tareas_frame, text="➕ Crear Tarea", command=crear_tarea_en_rma).grid(row=1, column=0, sticky="w", padx=5, pady=(5,15))

        def editar_tarea_dialog(task):
                # task is a dict with task fields
                dlg = ctk.CTkToplevel(self)
                dlg.title("Editar tarea")
                dlg.geometry("420x360")
                dlg.grab_set()

                ctk.CTkLabel(dlg, text=f"ID: {task['id']} - RMA: {task['codigo_rma']}").pack(pady=5)
                ctk.CTkLabel(dlg, text="Título:").pack(pady=(10,0))
                titulo_entry = ctk.CTkEntry(dlg)
                titulo_entry.insert(0, task['titulo'])
                titulo_entry.pack(padx=10, pady=5, fill='x')

                ctk.CTkLabel(dlg, text="Descripción:").pack(pady=(10,0))
                desc_text = tk.Text(dlg, height=6)
                desc_text.insert('1.0', task['descripcion'] or '')
                desc_text.pack(padx=10, pady=5, fill='both', expand=True)

                ctk.CTkLabel(dlg, text="Fecha Vencimiento (YYYY-MM-DD):").pack(pady=(5,0))
                fecha_entry = ctk.CTkEntry(dlg)
                fecha_entry.insert(0, task.get('fecha_vencimiento') or '')
                fecha_entry.pack(padx=10, pady=5, fill='x')

                estado_var = ctk.StringVar(value=task.get('estado', 'Pendiente'))
                estado_opt = ctk.CTkOptionMenu(dlg, values=["Pendiente", "En Progreso", "Completado"], variable=estado_var)
                estado_opt.pack(pady=5)

                def guardar_edicion():
                    nuevo_titulo = titulo_entry.get().strip()
                    nueva_desc = desc_text.get('1.0', 'end').strip()
                    nueva_fecha = fecha_entry.get().strip() or None
                    nuevo_estado = estado_var.get()
                    try:
                        conn = connect_db()
                        cur = conn.cursor()
                        cur.execute("UPDATE tareas SET titulo = ?, descripcion = ?, fecha_vencimiento = ?, estado = ? WHERE id = ?",
                                    (nuevo_titulo, nueva_desc, nueva_fecha, nuevo_estado, task['id']))
                        conn.commit()
                        # Registrar en historial del RMA
                        try:
                            # Obtener el ID del RMA para el historial
                            cur.execute("SELECT id FROM rma_maestro WHERE codigo_rma = ?", (task['codigo_rma'],))
                            rma_row = cur.fetchone()
                            if rma_row:
                                rma_id = rma_row[0]
                                # Registrar el cambio en el historial
                                cur.execute("""
                                    INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                                    VALUES (?, ?, ?, ?)
                                """, (rma_id, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                                     self.username, 
                                     f"Tarea ID {task['id']} editada - {task['titulo']} -> {nuevo_titulo} (Estado: {nuevo_estado})")
                                )
                                conn.commit()
                        except sqlite3.Error as e:
                            print(f"Error al registrar historial de tarea: {e}")
                        conn.close()
                        dlg.destroy()
                        self.cargar_lista_tareas_rma()
                        # Actualizar badge de tareas
                        self.actualizar_badge_tareas()
                        messagebox.showinfo("Éxito", "✅ Tarea actualizada correctamente")
                    except sqlite3.Error as e:
                        messagebox.showerror("Error BD", f"No se pudo actualizar la tarea: {e}")

                ctk.CTkButton(dlg, text="Guardar", command=guardar_edicion).pack(pady=10)

        def eliminar_tarea_rma(task_id, codigo_rma=None, titulo=None):
            if not messagebox.askyesno("Confirmar", "¿Eliminar esta tarea? Esta acción no se puede deshacer."):
                return
            try:
                conn = connect_db()
                cur = conn.cursor()
                cur.execute("DELETE FROM tareas WHERE id = ?", (task_id,))
                # Registrar en el historial la eliminación
                try:
                    cur.execute("SELECT id FROM rma_maestro WHERE codigo_rma = ?", (codigo_rma,))
                    rma_row = cur.fetchone()
                    if rma_row and titulo:
                        rma_id = rma_row[0]
                        cur.execute("""
                            INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                            VALUES (?, ?, ?, ?)
                        """, (rma_id, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                             self.username,
                             f"Tarea eliminada: {titulo}")
                        )
                except Exception as e:
                    print(f"Error al registrar eliminación en historial: {e}")
                conn.commit()
                conn.close()
                if hasattr(self, 'cargar_lista_tareas_rma'):
                    self.cargar_lista_tareas_rma()
                messagebox.showinfo("Eliminada", "❌ Tarea eliminada correctamente")
            except sqlite3.Error as e:
                messagebox.showerror("Error BD", f"No se pudo eliminar la tarea: {e}")

        def mostrar_tarea_row(task):
            row = ctk.CTkFrame(self.tareas_list_frame)
            row.pack(fill="x", padx=5, pady=3)
            # Determinar color según vencimiento y estado
            fecha_v = task.get('fecha_vencimiento')
            estado = task.get('estado')
            color_texto = "black"
            if estado == "Completado":
                color_texto = "green"
            elif fecha_v:
                try:
                    fecha_venc = datetime.datetime.strptime(fecha_v, "%Y-%m-%d").date()
                    hoy = datetime.date.today()
                    dias_restantes = (fecha_venc - hoy).days
                    if dias_restantes < 0:
                        color_texto = "red"  # Vencida
                    elif dias_restantes <= 3:
                        color_texto = "orange"  # Próxima a vencer
                except ValueError:
                    pass

            ctk.CTkLabel(row, 
                        text=f"{task['titulo']} - Vence: {fecha_v or 'Sin fecha'} - Estado: {estado}",
                        text_color=color_texto).pack(side='left', padx=5)
            ctk.CTkButton(row, text="Editar", width=60, command=lambda t=task: editar_tarea_dialog(t)).pack(side='right', padx=5)
            ctk.CTkButton(row, text="Eliminar", width=60, command=lambda t=task: eliminar_tarea_rma(t['id'], t['codigo_rma'], t['titulo'])).pack(side='right', padx=5)

        def cargar_lista_tareas_rma():
            # Limpiar solo las filas de tareas, no el título
            for w in self.tareas_list_frame.winfo_children():
                if w != self.tareas_title_label:
                    w.destroy()

            if self.current_rma_id is None:
                ctk.CTkLabel(self.tareas_list_frame, text="Guarde el expediente para ver las tareas.").pack(pady=10)
                return
            try:
                # Usar el código RMA actual de la ficha
                codigo = self.lbl_codigo_rma.cget('text').split(': ')[1].strip()
                conn = connect_db()
                cur = conn.cursor()
                cur.execute("SELECT id, codigo_rma, titulo, descripcion, fecha_vencimiento, estado, creado_por FROM tareas WHERE codigo_rma = ? ORDER BY fecha_vencimiento IS NULL, fecha_vencimiento ASC", (codigo,))
                filas = cur.fetchall()
                conn.close()

                if not filas:
                    ctk.CTkLabel(self.tareas_list_frame, text="No hay tareas asociadas a este RMA.", text_color="gray").pack(pady=10)
                    return
                for tid, codigo_rma, titulo, desc, fecha_v, estado, creador in filas:
                    task = {'id': tid, 'codigo_rma': codigo_rma, 'titulo': titulo, 'descripcion': desc, 'fecha_vencimiento': fecha_v, 'estado': estado, 'creado_por': creador}
                    mostrar_tarea_row(task)
            except sqlite3.Error as e:
                messagebox.showerror("Error BD", f"Error cargando tareas: {e}")

        # Exponer la función para recarga externa
        self.cargar_lista_tareas_rma = cargar_lista_tareas_rma
        # Cargar las tareas si estamos en edición
        if es_edicion:
            self.cargar_lista_tareas_rma()
    # Nota: No recrear tareas_scroll aquí para evitar duplicados en la pestaña de Tareas.
            
        # -----------------------------------------------------------
        # -- 2. MOVER LLAMADAS A crear_campo A SUS NUEVOS FRAMES --
        # -----------------------------------------------------------
        
        # V A L O R E S  A U T O M Á T I C O S
        fecha_emision_valor = datetime.datetime.now().strftime("%Y-%m-%d")
        usuario_actual = self.username
        
        # A) PESTAÑA GENERAL
        # Campos de Cliente y Contacto - Cliente y Número de Documento son de solo lectura en modo edición (excepto admin)
        es_admin = self.username.lower() == "admin"
        deshabilitar_campos = es_edicion and not es_admin
        
        self.crear_campo(general_frame, 0, "Cliente:", "Cliente", deshabilitado=deshabilitar_campos)
        self.crear_campo(general_frame, 1, "Núm. Doc. Cliente:", "Numero_Documento_Cliente", deshabilitado=deshabilitar_campos)
        self.crear_campo(general_frame, 2, "Persona de Contacto:", "Persona_de_Contacto")
        self.crear_campo(general_frame, 3, "Email de Contacto:", "Email_de_Contacto")
        self.crear_campo(general_frame, 4, "Autorización:", "Autorizacion", tipo="optionmenu", opciones=self.OPCIONES["Autorizacion"], valor_defecto="NO")
        self.crear_campo(general_frame, 5, "Motivo Devolucion:", "motivo")
        
        # Fechas y Creador (Solo lectura excepto admin)
        self.crear_campo(general_frame, 6, "Fecha Emisión:", "Fecha_Emision", 
                         valor_defecto=fecha_emision_valor, deshabilitado=not es_admin)
        self.crear_campo(general_frame, 7, "Creado Por:", "Creado_Por", 
                         valor_defecto=usuario_actual, deshabilitado=not es_admin)


        # B) PESTAÑA ESTADOS Y FECHAS
        # Fechas de Autorización, Recepción, Proceso y Gestión
        fila_estados = 0
        self.crear_campo(estados_fechas_frame, fila_estados, "Fecha Autorización:", "Fecha_Autorizacion", tipo="date"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Autorizado Por:", "Autorizado_Por", tipo="optionmenu", opciones=self.OPCIONES["Autorizado_Por"], valor_defecto=self.OPCIONES["Autorizado_Por"][0]); fila_estados += 1
        
        ctk.CTkLabel(estados_fechas_frame, text="--- RECEPCIÓN ---", font=ctk.CTkFont(weight="bold")).grid(row=fila_estados, column=0, columnspan=2, pady=(10, 5), sticky="w"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Fecha Recepción:", "Fecha_Recepcion", tipo="date"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Recepcionado Por:", "Recepcionado_Por"); fila_estados += 1
        
        ctk.CTkLabel(estados_fechas_frame, text="--- PROCESO ---", font=ctk.CTkFont(weight="bold")).grid(row=fila_estados, column=0, columnspan=2, pady=(10, 5), sticky="w"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Fecha Proceso:", "Fecha_Proceso", tipo="date"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Procesado Por:", "Procesado_Por"); fila_estados += 1
        
        ctk.CTkLabel(estados_fechas_frame, text="--- CIERRE/GESTIÓN ---", font=ctk.CTkFont(weight="bold")).grid(row=fila_estados, column=0, columnspan=2, pady=(10, 5), sticky="w"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Fecha Gestión:", "Fecha_Gestion", tipo="date"); fila_estados += 1
        self.crear_campo(estados_fechas_frame, fila_estados, "Gestionado Por:", "Gestionado_Por", tipo="optionmenu", opciones=self.OPCIONES["Gestionado_Por"], valor_defecto=self.OPCIONES["Gestionado_Por"][0]); fila_estados += 1
        opciones_quincenas = ["Seleccionar..."] + self.obtener_quincenas_futuras()
        self.crear_campo(estados_fechas_frame, fila_estados, "Fecha para Factura:", "Fecha_para_factura", tipo="optionmenu", opciones=opciones_quincenas, valor_defecto="Seleccionar..."); fila_estados += 1

        
        # C) PESTAÑA INFORMACIÓN TÉCNICA (campos técnicos)

        # --- Sincronización Autorización / Fecha_Autorizacion / Autorizado_Por ---
        try:
            auth_widget = getattr(self, 'entry_Autorizacion', None)
            if auth_widget is not None:
                def _on_autorizacion_change(choice=None):
                    try:
                        sel = choice if choice is not None else None
                        # CTkOptionMenu may call the function with the selected value
                        if sel is None:
                            try:
                                sel = auth_widget.get()
                            except Exception:
                                sel = None

                        if sel == 'SI':
                            # Poner fecha de autorización hoy
                            hoy = datetime.datetime.now()
                            if hasattr(self, 'entry_Fecha_Autorizacion'):
                                try:
                                    if hasattr(self.entry_Fecha_Autorizacion, 'set_date'):
                                        self.entry_Fecha_Autorizacion.set_date(hoy)
                                    else:
                                        # Tratar como Entry
                                        self.entry_Fecha_Autorizacion.configure(state='normal')
                                        try:
                                            self.entry_Fecha_Autorizacion.delete(0, ctk.END)
                                        except Exception:
                                            pass
                                        try:
                                            self.entry_Fecha_Autorizacion.insert(0, hoy.strftime('%Y-%m-%d'))
                                        except Exception:
                                            pass
                                except Exception:
                                    pass

                            # Pedir autor (selector)
                            try:
                                if hasattr(self, 'entry_Autorizado_Por'):
                                    opcion = self._prompt_select_autorizado_por(self.OPCIONES.get('Autorizado_Por', []))
                                    if opcion:
                                        try:
                                            self.entry_Autorizado_Por.set(opcion)
                                        except Exception:
                                            try:
                                                self.entry_Autorizado_Por.configure(state='normal')
                                                self.entry_Autorizado_Por.delete(0, ctk.END)
                                                self.entry_Autorizado_Por.insert(0, opcion)
                                            except Exception:
                                                pass
                            except Exception:
                                pass

                        else:
                            # Si se pone 'NO', borrar fecha y dejar autor por defecto
                            if hasattr(self, 'entry_Fecha_Autorizacion'):
                                try:
                                    if hasattr(self.entry_Fecha_Autorizacion, 'set_date'):
                                        try:
                                            self.entry_Fecha_Autorizacion.set_date('')
                                        except Exception:
                                            # Intentar borrar texto
                                            try:
                                                self.entry_Fecha_Autorizacion.delete(0, ctk.END)
                                            except Exception:
                                                pass
                                    else:
                                        try:
                                            self.entry_Fecha_Autorizacion.configure(state='normal')
                                            self.entry_Fecha_Autorizacion.delete(0, ctk.END)
                                        except Exception:
                                            pass
                                except Exception:
                                    pass

                            if hasattr(self, 'entry_Autorizado_Por'):
                                try:
                                    default = self.OPCIONES.get('Autorizado_Por', [None])[0]
                                    if default:
                                        try:
                                            self.entry_Autorizado_Por.set(default)
                                        except Exception:
                                            try:
                                                self.entry_Autorizado_Por.configure(state='normal')
                                                self.entry_Autorizado_Por.delete(0, ctk.END)
                                            except Exception:
                                                pass
                                except Exception:
                                    pass
                    except Exception as e:
                        print('Error manejando cambio Autorizacion:', e)

                try:
                    # Intentar configurar el callback
                    auth_widget.configure(command=_on_autorizacion_change)
                except Exception:
                    # Fallback: si no acepta configure, intentar setear manualmente (no garantizado)
                    try:
                        auth_widget.set_callback = _on_autorizacion_change
                    except Exception:
                        pass
        except Exception:
            pass
        
        # --- Sincronización INVERSA: Fecha_Autorizacion -> Autorizacion ---
        try:
            fecha_auth_widget = getattr(self, 'entry_Fecha_Autorizacion', None)
            auth_widget = getattr(self, 'entry_Autorizacion', None)
            
            if fecha_auth_widget is not None and auth_widget is not None:
                # Variable para almacenar el último valor conocido
                self._last_fecha_autorizacion = None
                
                def _verificar_fecha_autorizacion():
                    """Verifica periódicamente si cambió Fecha_Autorizacion y actualiza Autorizacion"""
                    try:
                        # Obtener la fecha del widget
                        fecha_valor = None
                        
                        # Intentar obtener el valor según el tipo de widget
                        if hasattr(fecha_auth_widget, 'get_date'):
                            try:
                                fecha_valor = fecha_auth_widget.get_date()
                            except:
                                pass
                        
                        if fecha_valor is None and hasattr(fecha_auth_widget, 'get'):
                            try:
                                fecha_valor = fecha_auth_widget.get()
                            except:
                                pass
                        
                        # Convertir a string para comparación
                        fecha_str = str(fecha_valor).strip() if fecha_valor else ""
                        
                        # Solo actualizar si cambió
                        if fecha_str != self._last_fecha_autorizacion:
                            self._last_fecha_autorizacion = fecha_str
                            
                            # Verificar si hay fecha válida
                            tiene_fecha = bool(fecha_str) and fecha_str not in ('None', 'null', '')
                            
                            # Actualizar el campo Autorizacion
                            if tiene_fecha:
                                # Hay fecha -> Autorización debe ser SI
                                try:
                                    current_auth = auth_widget.get()
                                    if current_auth != 'SI':
                                        auth_widget.set('SI')
                                except:
                                    pass
                            else:
                                # No hay fecha -> Autorización debe ser NO
                                try:
                                    current_auth = auth_widget.get()
                                    if current_auth != 'NO':
                                        auth_widget.set('NO')
                                except:
                                    pass
                        
                        # Programar siguiente verificación (cada 500ms)
                        if hasattr(self, 'content_frame') and self.content_frame.winfo_exists():
                            self.content_frame.after(500, _verificar_fecha_autorizacion)
                    except Exception as e:
                        # Si hay error, intentar reprogramar de todos modos
                        try:
                            if hasattr(self, 'content_frame') and self.content_frame.winfo_exists():
                                self.content_frame.after(500, _verificar_fecha_autorizacion)
                        except:
                            pass
                
                # Iniciar la verificación periódica
                try:
                    self.content_frame.after(500, _verificar_fecha_autorizacion)
                except Exception as e:
                    print(f'Error iniciando verificación periódica: {e}')
        except Exception as e:
            print(f'Error configurando sincronización inversa: {e}')
        # Fila 0: RMA Proveedor (label cambiado a 'RMA Proveedor')
        self.crear_campo(info_tecnica_frame, 0, "RMA Proveedor:", "Rma_Proveedor")
        # Fila 1: Modelo
        self.crear_campo(info_tecnica_frame, 1, "Modelo:", "Modelo")
        # Fila 2: N. Serie
        self.crear_campo(info_tecnica_frame, 2, "N. Serie:", "N_Serie")
        # Fila 3: Ref. Proveedor
        self.crear_campo(info_tecnica_frame, 3, "Ref. Proveedor:", "Ref_Proveedor")
        # Fila 4: Observaciones Técnicas (caja de texto mayor)
        ctk.CTkLabel(info_tecnica_frame, text="Observaciones Técnicas:").grid(row=4, column=0, padx=10, pady=5, sticky="w")
        self.entry_Obs_Tecnica = ctk.CTkTextbox(info_tecnica_frame, height=120, wrap="word")
        self.entry_Obs_Tecnica.grid(row=4, column=1, padx=10, pady=5, sticky="ew")

        # Botón para generar la Solicitud de RMA desde la plantilla PDF
        ctk.CTkButton(
            info_tecnica_frame,
            text="Generar Solicitud de RMA",
            command=self.autorrellena_pdf
        ).grid(row=5, column=1, padx=10, pady=(8, 12), sticky="e")

        # C) PESTAÑA ARTÍCULOS (Mantener la lógica de listado y añadir artículo)
        articulos_tab.grid_columnconfigure(0, weight=1)
        articulos_frame = ctk.CTkFrame(articulos_tab) # Este marco no necesita scroll, la lista interna sí
        articulos_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        articulos_frame.grid_columnconfigure(0, weight=1)
        articulos_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(articulos_frame, text="**DETALLE DE ARTÍCULOS**", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, columnspan=6, pady=(10, 5), sticky="w")
        
        # 4. Entradas para añadir un nuevo artículo (Input Article Frame)
        input_articulo_frame = ctk.CTkFrame(articulos_frame)
        input_articulo_frame.grid(row=1, column=0, columnspan=6, sticky="ew", pady=(5, 10))
        # ... (Mantener la definición de las entradas self.art_ref, self.art_cant_doc, etc.)
        # ... (NO TOCAR ESTA SECCIÓN, está bien como está, solo moverla)

        # Etiquetas
        ctk.CTkLabel(input_articulo_frame, text="Ref. Artículo", font=ctk.CTkFont(size=11)).grid(row=0, column=0, padx=5)
        ctk.CTkLabel(input_articulo_frame, text="Cant. Doc.", font=ctk.CTkFont(size=11)).grid(row=0, column=1, padx=5)
        ctk.CTkLabel(input_articulo_frame, text="Cant. Entregada", font=ctk.CTkFont(size=11)).grid(row=0, column=2, padx=5)
        ctk.CTkLabel(input_articulo_frame, text="Estado", font=ctk.CTkFont(size=11)).grid(row=0, column=3, padx=5)
        ctk.CTkLabel(input_articulo_frame, text="Precio Unitario", font=ctk.CTkFont(size=11)).grid(row=0, column=4, padx=5)
        
        # Entradas
        self.art_ref = ctk.CTkEntry(input_articulo_frame, width=150)
        self.art_ref.grid(row=1, column=0, padx=5, pady=2, sticky="ew")
        self.art_cant_doc = ctk.CTkEntry(input_articulo_frame, width=80)
        self.art_cant_doc.grid(row=1, column=1, padx=5, pady=2, sticky="ew")
        self.art_cant_entregada = ctk.CTkEntry(input_articulo_frame, width=80)
        self.art_cant_entregada.grid(row=1, column=2, padx=5, pady=2, sticky="ew")
        self.art_estado = ctk.CTkOptionMenu(input_articulo_frame, values=self.OPCIONES["Estado_Producto"], width=150)
        self.art_estado.grid(row=1, column=3, padx=5, pady=2, sticky="ew")
        self.art_precio = ctk.CTkEntry(input_articulo_frame, width=100)
        self.art_precio.grid(row=1, column=4, padx=5, pady=2, sticky="ew")

        # Botones de Acción de Artículos
        # Guardar referencia al botón para permitir cambiar su comportamiento durante la edición
        self.btn_anadir_articulo = ctk.CTkButton(input_articulo_frame,
                                                 text="➕",
                                                 width=30,
                                                 command=self.anadir_articulo)
        self.btn_anadir_articulo.grid(row=1, column=5, padx=5, pady=2)

        # Permitir pulsar ENTER en los campos para añadir/actualizar artículo
        try:
            self.art_ref.bind("<Return>", self._enter_articulo)
            self.art_cant_doc.bind("<Return>", self._enter_articulo)
            self.art_cant_entregada.bind("<Return>", self._enter_articulo)
            self.art_precio.bind("<Return>", self._enter_articulo)
        except Exception:
            pass
        
        # 5. Listado de Artículos ya añadidos
        self.articulos_list_frame = ctk.CTkFrame(articulos_frame)
        self.articulos_list_frame.grid(row=2, column=0, columnspan=6, sticky="ew", padx=10, pady=10)
        self.actualizar_listado_articulos()

        
        # D) PESTAÑA CONTABILIDAD
        fila_cont = 0
        self.crear_campo(contabilidad_frame, fila_cont, "Resultado Expediente:", "Resultado_Expediente", tipo="optionmenu", opciones=self.OPCIONES["Resultado_Expediente"], valor_defecto=self.OPCIONES["Resultado_Expediente"][0]); fila_cont += 1
        self.crear_campo(contabilidad_frame, fila_cont, "Número Albarán:", "Numero_Albaran"); fila_cont += 1
        self.crear_campo(contabilidad_frame, fila_cont, "Fecha Doc. Cliente:", "Fecha_Doc_Cliente"); fila_cont += 1
        
        # Nuevos campos de reposición y abono
        self.crear_campo(contabilidad_frame, fila_cont, "Nº Albarán Reposición:", "numero_albaran_reposicion"); fila_cont += 1
        self.crear_campo(contabilidad_frame, fila_cont, "Fecha Albarán Reposición:", "fecha_albaran_reposicion", tipo="date"); fila_cont += 1
        self.crear_campo(contabilidad_frame, fila_cont, "Nº Factura Abono:", "numero_factura_abono"); fila_cont += 1
        self.crear_campo(contabilidad_frame, fila_cont, "Fecha Factura Abono:", "fecha_factura_abono", tipo="date"); fila_cont += 1
        
        # E) PESTAÑA ADJUNTOS
        # 1. Botón para Añadir Adjunto
        self.btn_subir_adjunto = ctk.CTkButton(
            adjuntos_tab, 
            text="➕ Subir Archivo", 
            #fg_color="gray80",        # Fondo del botón: Gris claro
            #hover_color="gray70",     # Efecto hover: Ligeramente más oscuro
            #text_color="black",
            font=ctk.CTkFont(size=14, weight="bold"),
            # CRÍTICO: Usar lambda para forzar el argumento a False
            command=lambda: self.abrir_dialogo_adjunto(modo_abrir_carpeta=False) 
        )
        self.btn_subir_adjunto.pack(pady=(10, 5), padx=10, fill='x')

        # 2. Frame para el Listado de Adjuntos
        # Usamos un ScrollableFrame para que la lista crezca
        self.adjuntos_list_frame = ctk.CTkScrollableFrame(adjuntos_tab, label_text="Archivos Adjuntos")
        self.adjuntos_list_frame.pack(pady=5, padx=10, fill="both", expand=True)

        # Cargar los adjuntos si estamos en modo edición (la función la crearemos en el paso 4)
        if self.mode == 'editar':
            self.cargar_lista_adjuntos(self.current_rma_id)
        else:
            # En modo 'nuevo', mostramos un mensaje hasta que el RMA se guarde
            ctk.CTkLabel(self.adjuntos_list_frame, 
                         text="Guarde el expediente primero para poder adjuntar archivos.")\
                .pack(pady=20)
        

        # 4. Botón de Guardar Definitivo
        
        # 💡 CREAR UN FRAME PARA AGRUPAR LOS BOTONES DE ACCIÓN (Fila 3)
        btn_action_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        btn_action_frame.grid(row=3, column=0, padx=20, pady=20, sticky="ew")
        
        # Botón de cliente (lado derecho, siempre visible)
        def abrir_ficha_cliente_desde_expediente():
            # Obtener nombre del cliente desde el campo
            if hasattr(self, 'entry_Cliente'):
                nombre_cliente = self.entry_Cliente.get().strip()
                if nombre_cliente:
                    # Buscar cliente_id por nombre
                    try:
                        conn, cursor = self.master.conectar_db()
                        if conn:
                            cursor.execute("SELECT cliente_id FROM clientes WHERE nombre = ?", (nombre_cliente,))
                            resultado = cursor.fetchone()
                            conn.close()
                            
                            if resultado:
                                self.abrir_ficha_cliente(resultado[0])
                            else:
                                messagebox.showinfo("Cliente no encontrado", 
                                                  f"El cliente '{nombre_cliente}' no está registrado en el sistema.\n\n"
                                                  "Puede usar el botón 'Migrar desde RMAs' en la ventana de clientes.")
                    except Exception as e:
                        print(f"Error buscando cliente: {e}")
                        messagebox.showerror("Error", "No se pudo buscar el cliente")
                else:
                    messagebox.showwarning("Sin cliente", "Primero debe especificar un cliente en el expediente")
        
        self.btn_ver_cliente = ctk.CTkButton(
            btn_action_frame,
            text="👤",
            command=abrir_ficha_cliente_desde_expediente,
            width=35,
            height=35,
            font=ctk.CTkFont(size=16)
        )
        self.btn_ver_cliente.pack(side="right", padx=(10, 0))
        Tooltip(self.btn_ver_cliente, "Ver ficha del cliente")
        
        # Botón de artículos (lado derecho, junto al de cliente)
        def abrir_ficha_articulo_desde_expediente():
            # Si no hay artículos en memoria, mostrar mensaje
            if not self.articulos_data:
                messagebox.showinfo("Sin artículos", 
                                  "Este expediente no tiene artículos asociados aún.\n\n"
                                  "Agregue artículos en la sección 'ARTÍCULOS' de este formulario.")
                return
            
            # Si solo hay un artículo, abrir su ficha directamente
            if len(self.articulos_data) == 1:
                referencia = self.articulos_data[0].get('referencia_articulo', '')
                if referencia:
                    self.mostrar_estados_por_articulo(referencia)
                return
            
            # Si hay múltiples artículos, mostrar selector
            from lib.articulo_utils import mostrar_selector_referencias
            
            def callback_seleccion(referencia):
                self.mostrar_estados_por_articulo(referencia)
            
            mostrar_selector_referencias(self.articulos_data, self.content_frame.winfo_toplevel(), callback_seleccion)
        
        self.btn_ver_articulo = ctk.CTkButton(
            btn_action_frame,
            text="📦",
            command=abrir_ficha_articulo_desde_expediente,
            width=35,
            height=35,
            font=ctk.CTkFont(size=16)
        )
        self.btn_ver_articulo.pack(side="right", padx=(5, 0))
        Tooltip(self.btn_ver_articulo, "Ver ficha de artículo")
        
        # 1. Botón de Generar Informe (Solo en modo edición)
        if es_edicion:
            self.btn_generar_informe = ctk.CTkButton(
                btn_action_frame, 
                text="📄 Informe (Word)", 
                command=self.generar_informe_dinamico, 
                #text_color="black",
                #fg_color="grey80", 
                #hover_color="grey70",
                font=ctk.CTkFont(size=14, weight="bold")
            )
            self.btn_generar_informe.pack(side="left", padx=(5, 15)) # 15px de margen derecho
            self.btn_generar_reposicion = ctk.CTkButton(
                btn_action_frame, 
                text="🔄 Reposicion/Devolucion", # Texto largo para mayor claridad
                command=self.generar_reposicion_devolucion, 
                #text_color="black",
                #fg_color="grey80",
                #hover_color="grey70",
                font=ctk.CTkFont(size=14, weight="bold")
            )
            self.btn_generar_reposicion.pack(side="left", padx=(5, 15)) # 5px izquierda, 15px derecha antes de Actualizar
            self.btn_enviar_email = ctk.CTkButton(btn_action_frame, 
                text="📧 Enviar Email",
                command=self.enviar_email_contacto,
                #text_color="black",
                #fg_color="grey80",
                #hover_color="grey70",
                font=ctk.CTkFont(size=14, weight="bold"))               
            self.btn_enviar_email.pack(side="left", padx=(5, 15))

        # 2. Botón de Guardar
        if es_edicion:
            guardar_texto = "💾 ACTUALIZAR"
            color_boton = None  # Color por defecto
        else:
            guardar_texto = "💾 GUARDAR Y ASIGNAR NÚMERO"
            color_boton = "orange"  # Color naranja para destacar que asignará el número final
        
        self.btn_guardar_rma = ctk.CTkButton(
            btn_action_frame, 
            text=guardar_texto, 
            fg_color=color_boton,        # Color especial para nuevos expedientes
            #hover_color="gray70",     # Efecto hover: Ligeramente más oscuro
            #text_color="black",
            font=ctk.CTkFont(size=14, weight="bold"),
            command=self.guardar_rma_placeholder
        )
        self.btn_guardar_rma.pack(side="left", padx=(0, 5))
        
        # Botón de eliminar expediente (solo para admin en modo edición)
        if es_edicion and self.username.lower() == "admin":
            self.btn_eliminar_rma = ctk.CTkButton(
                btn_action_frame,
                text="🗑️ ELIMINAR EXPEDIENTE",
                fg_color="red",
                hover_color="darkred",
                font=ctk.CTkFont(size=14, weight="bold"),
                command=lambda: self.eliminar_expediente(rma_id)
            )
            self.btn_eliminar_rma.pack(side="left", padx=(5, 0))
        
        # Lógica de Edición
        if es_edicion:
            # self.lbl_codigo_rma.configure(text="Cargando datos...")
            self.cargar_datos_rma(rma_id) # Se implementará después

    def guardar_comentario_historial(self):
        """
        Guarda el contenido del Textbox de comentarios como una entrada del historial, 
        usando la estructura simple de la tabla rma_historial.
        """
        
        # 1. Validar ID del Expediente
        rma_id = self.rma_actual_id
        if rma_id is None:
            messagebox.showwarning("Aviso", "Debes guardar el expediente principal primero (botón GUARDAR RMA) para poder añadir comentarios.")
            return

        # 2. Obtener y validar el texto
        nuevo_comentario = self.textbox_comentarios.get("1.0", "end-1c").strip()
        
        if not nuevo_comentario:
            messagebox.showwarning("Aviso", "No hay texto en la caja de comentarios para guardar.")
            return
            
        try:
            conn = connect_db()
            cursor = conn.cursor()
            
            # 3. Preparar los datos
            fecha_cambio = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            usuario = self.username
            
            # 4. Crear la descripción completa del cambio
            # Para un historial simple, el comentario es la descripción completa
            descripcion_cambio = f"COMENTARIO MANUAL: {nuevo_comentario}"
            
            # 5. Ejecutar la inserción SQL
            cursor.execute("""
                INSERT INTO rma_historial 
                (rma_id, fecha_cambio, usuario, descripcion_cambio)
                VALUES (?, ?, ?, ?)
            """, (rma_id, fecha_cambio, usuario, descripcion_cambio))

            # 6. Commit, Limpieza y Feedback
            conn.commit()
            conn.close()
            
            self.textbox_comentarios.delete("1.0", "end")
            
            # 7. Recargar historial (si la pestaña está visible)
            if hasattr(self, 'historial_tab') and self.historial_tab:
                # Llama a la función que recarga el listado de la pestaña historial
                self.mostrar_historial(self.historial_tab)
            
            try:
                # 1. Obtener el texto completo (ej: "Nº EXPEDIENTE: RMA25001")
                texto_label = self.lbl_codigo_rma.cget("text")
                # 2. Extraer solo el código (ej: "RMA25001")
                codigo_rma = texto_label.split(': ')[1]
            except Exception:
                # Fallback seguro en caso de que el label tenga un formato inesperado
                codigo_rma = f"ID {rma_id}"
                
            # 8. Mostrar mensaje de éxito usando la variable local 'codigo_rma'
            messagebox.showinfo("Comentario Guardado", f"El comentario ha sido guardado en el historial del RMA {codigo_rma}.")
            
        except sqlite3.Error as e:
            messagebox.showerror("Error de Base de Datos", f"Error al guardar el comentario: {e}")
    
    def anadir_articulo(self):
        """Añade una fila de artículo a la lista temporal."""
        try:
            referencia = self.art_ref.get()
            # Permitir decimales en las cantidades
            cant_doc = float(self.art_cant_doc.get().replace(',', '.') or 0.0)
            cant_entregada = float(self.art_cant_entregada.get().replace(',', '.') or 0.0)
            estado = self.art_estado.get()
            # Reemplazar comas por puntos para que float funcione
            precio_unitario = float(self.art_precio.get().replace(',', '.') or 0.0) 
        except ValueError:
            print("Error: Cantidad y Precio deben ser números.")
            return

        if not referencia:
            print("Error: La Referencia es obligatoria.")
            return

        nuevo_articulo = {
            "referencia_articulo": referencia,
            "cantidad_segun_documento": cant_doc,
            "cantidad_entregada": cant_entregada,
            "estado_producto": estado,
            "precio_unitario": precio_unitario
        }
        
        self.articulos_data.append(nuevo_articulo)
        self.actualizar_listado_articulos()
        self.limpiar_articulo()

    def limpiar_articulo(self):
        """Limpia los campos de entrada de un solo artículo."""
        self.art_ref.delete(0, ctk.END)
        self.art_cant_doc.delete(0, ctk.END)
        self.art_cant_entregada.delete(0, ctk.END)
        self.art_precio.delete(0, ctk.END)
        self.art_estado.set(self.OPCIONES["Estado_Producto"][0])
        # Si estábamos en modo edición, salir de él
        try:
            if hasattr(self, 'editing_articulo_index') and self.editing_articulo_index is not None:
                self.editing_articulo_index = None
                # Restaurar el botón de añadir a su comportamiento original
                try:
                    self.btn_anadir_articulo.configure(text="➕", command=self.anadir_articulo)
                except Exception:
                    pass
        except Exception:
            pass

    def eliminar_articulo(self, index):
        """Elimina un artículo de la lista temporal y actualiza la vista."""
        if 0 <= index < len(self.articulos_data):
            self.articulos_data.pop(index)
            self.actualizar_listado_articulos()

    def editar_articulo(self, index):
        """Carga un artículo en los campos de entrada para su edición."""
        if 0 <= index < len(self.articulos_data):
            art = self.articulos_data[index]
            try:
                self.art_ref.delete(0, ctk.END); self.art_ref.insert(0, str(art.get('referencia_articulo', '')))
                self.art_cant_doc.delete(0, ctk.END); self.art_cant_doc.insert(0, str(art.get('cantidad_segun_documento', '')))
                self.art_cant_entregada.delete(0, ctk.END); self.art_cant_entregada.insert(0, str(art.get('cantidad_entregada', '')))
                self.art_precio.delete(0, ctk.END); self.art_precio.insert(0, str(art.get('precio_unitario', '')))
                try:
                    self.art_estado.set(art.get('estado_producto', self.OPCIONES['Estado_Producto'][0]))
                except Exception:
                    pass
            except Exception:
                pass
            # Marcar índice en edición y cambiar botón
            self.editing_articulo_index = index
            try:
                self.btn_anadir_articulo.configure(text="✔️", command=self.actualizar_articulo)
            except Exception:
                pass

    def actualizar_articulo(self):
        """Actualiza el artículo seleccionado con los valores actuales de los campos."""
        if not hasattr(self, 'editing_articulo_index') or self.editing_articulo_index is None:
            return
        idx = self.editing_articulo_index
        try:
            referencia = self.art_ref.get()
            cant_doc = float(self.art_cant_doc.get().replace(',', '.') or 0.0)
            cant_entregada = float(self.art_cant_entregada.get().replace(',', '.') or 0.0)
            estado = self.art_estado.get()
            precio_unitario = float(self.art_precio.get().replace(',', '.') or 0.0)
        except ValueError:
            messagebox.showwarning("Error", "Cantidad y Precio deben ser números válidos.")
            return

        if not referencia:
            messagebox.showwarning("Error", "La referencia es obligatoria.")
            return

        nuevo_articulo = {
            "referencia_articulo": referencia,
            "cantidad_segun_documento": cant_doc,
            "cantidad_entregada": cant_entregada,
            "estado_producto": estado,
            "precio_unitario": precio_unitario
        }
        # Reemplazar en la lista
        try:
            self.articulos_data[idx] = nuevo_articulo
        except Exception:
            return

        # Finalizar edición
        self.editing_articulo_index = None
        try:
            self.btn_anadir_articulo.configure(text="➕", command=self.anadir_articulo)
        except Exception:
            pass

        self.actualizar_listado_articulos()
        self.limpiar_articulo()

    def _enter_articulo(self, event=None):
        """Handler para la tecla ENTER en los campos de artículo: añade o actualiza según el modo."""
        try:
            if hasattr(self, 'editing_articulo_index') and self.editing_articulo_index is not None:
                self.actualizar_articulo()
            else:
                self.anadir_articulo()
        except Exception:
            # No propagamos excepciones por un ENTER accidental
            pass

    def actualizar_listado_articulos(self):
        """Redibuja la tabla con los artículos de la lista temporal."""
        for widget in self.articulos_list_frame.winfo_children():
            widget.destroy()
            
        if not self.articulos_data:
            ctk.CTkLabel(self.articulos_list_frame, text="No hay artículos asociados a este RMA.", text_color="gray").pack(pady=10)
            return
            
        # Dibujar encabezados y filas usando grid directamente en el contenedor principal
        cols = ["Ref. Artículo", "Cant. Doc.", "Cant. Entregada", "Estado", "Precio Unitario", "Acción", ""]
        weights = [2, 1, 1, 2, 1, 0, 0]
        header_font = ctk.CTkFont(weight="bold", size=12)

        # Configurar columnas del contenedor para que se alineen entre filas
        for i, w in enumerate(weights):
            try:
                self.articulos_list_frame.grid_columnconfigure(i, weight=w)
            except Exception:
                pass

        # Encabezados en la fila 0
        for i, col in enumerate(cols):
            ctk.CTkLabel(self.articulos_list_frame, text=col, font=header_font).grid(row=0, column=i, padx=5, pady=5, sticky="w")

        # Filas: colocar directamente labels en la grilla del contenedor para mantener columnas alineadas
        for i, item in enumerate(self.articulos_data):
            row = i + 1
            try:
                ctk.CTkLabel(self.articulos_list_frame, text=item["referencia_articulo"]).grid(row=row, column=0, padx=5, pady=2, sticky="w")
                ctk.CTkLabel(self.articulos_list_frame, text=item["cantidad_segun_documento"]).grid(row=row, column=1, padx=5, pady=2, sticky="w")
                ctk.CTkLabel(self.articulos_list_frame, text=item["cantidad_entregada"]).grid(row=row, column=2, padx=5, pady=2, sticky="w")
                ctk.CTkLabel(self.articulos_list_frame, text=item["estado_producto"]).grid(row=row, column=3, padx=5, pady=2, sticky="w")
                ctk.CTkLabel(self.articulos_list_frame, text=f"{item['precio_unitario']:.2f} €").grid(row=row, column=4, padx=5, pady=2, sticky="w")

                # Acciones: Eliminar y Editar
                ctk.CTkButton(self.articulos_list_frame, text="X", width=30, fg_color="red", hover_color="darkred",
                              command=lambda idx=i: self.eliminar_articulo(idx)).grid(row=row, column=5, padx=5, pady=2, sticky="w")
                try:
                    ctk.CTkButton(self.articulos_list_frame, text="✏️", width=30,
                                  command=lambda idx=i: self.editar_articulo(idx)).grid(row=row, column=6, padx=2, pady=2, sticky="w")
                except Exception:
                    pass
            except Exception:
                # Silenciar errores de renderizado individual para no romper toda la lista
                pass
            
        # --- NUEVO: Calcular y actualizar el Precio Total en la etiqueta de Contabilidad ---
        precio_total = 0.0
        
        for item in self.articulos_data:
            try:
                # Convertir cantidad_entregada a float de forma segura
                cantidad = item.get('cantidad_entregada', 0)
                if isinstance(cantidad, str):
                    # Limpiar el string y convertir a float
                    cantidad = float(cantidad.replace(',', '.')) if cantidad.strip() else 0.0
                elif cantidad is None:
                    cantidad = 0.0
                else:
                    cantidad = float(cantidad)
                
                # Convertir precio_unitario a float de forma segura
                precio = item.get('precio_unitario', 0.0)
                if isinstance(precio, str):
                    # Limpiar el string y convertir a float
                    precio = float(precio.replace(',', '.')) if precio.strip() else 0.0
                elif precio is None:
                    precio = 0.0
                else:
                    precio = float(precio)
                
                # Multiplicar de forma segura
                precio_total += cantidad * precio
                
            except (ValueError, TypeError) as e:
                # Si hay un error en la conversión, ignorar este artículo y continuar
                print(f"Error al calcular precio para artículo {item.get('referencia_articulo', 'N/A')}: {e}")
                continue
        
        # Esto es seguro porque lbl_precio_total se crea en mostrar_nuevo_rma
        if hasattr(self, 'lbl_precio_total'):
            self.lbl_precio_total.configure(text=f"{precio_total:.2f} €")

        # Guardar el precio total en la BD si estamos editando un expediente ya existente
        try:
            self.guardar_precio_total_expediente()
        except Exception:
            # No bloquear la UI si la actualización en BD falla aquí; el guardado final también actualizará
            pass

    def guardar_precio_total_expediente(self):
        """Recalcula el precio total desde `self.articulos_data` y lo persiste
        en `rma_maestro.precio_total_expediente` si el RMA está abierto (tiene id)."""
        # Calcular total localmente
        precio_total = 0.0
        for item in getattr(self, 'articulos_data', []):
            try:
                cantidad = item.get('cantidad_entregada', 0) or 0
                if isinstance(cantidad, str):
                    cantidad = float(cantidad.replace(',', '.')) if cantidad.strip() else 0.0
                else:
                    cantidad = float(cantidad)

                precio = item.get('precio_unitario', 0.0) or 0.0
                if isinstance(precio, str):
                    precio = float(precio.replace(',', '.')) if precio.strip() else 0.0
                else:
                    precio = float(precio)

                precio_total += cantidad * precio
            except Exception:
                continue

        # Actualizar etiqueta si existe
        try:
            if hasattr(self, 'lbl_precio_total'):
                self.lbl_precio_total.configure(text=f"{precio_total:.2f} €")
        except Exception:
            pass

        # Si estamos editando un expediente ya guardado, escribir en BD
        try:
            if getattr(self, 'rma_actual_id', None) is not None:
                conn, cursor = self.master.conectar_db()
                if conn:
                    try:
                        cursor.execute("UPDATE rma_maestro SET precio_total_expediente = ? WHERE id = ?", (precio_total, self.rma_actual_id))
                        conn.commit()
                    finally:
                        conn.close()
        except Exception as e:
            print(f"Error guardando precio_total_expediente: {e}")


    def guardar_rma_placeholder(self):
        """Punto de entrada para guardar/actualizar."""
        if self.rma_actual_id is None:
            self.guardar_nuevo_rma()
        else:
            self.actualizar_rma()

    def eliminar_expediente(self, rma_id):
        """
        Elimina un expediente completo incluyendo todos sus datos relacionados.
        Solo disponible para usuario admin.
        """
        # Verificar permisos de admin
        if self.username.lower() != "admin":
            messagebox.showerror("Acceso Denegado", "Solo el administrador puede eliminar expedientes.")
            return
        
        try:
            conn = self.master.conectar_db()
            cursor = conn.cursor()
            
            # Obtener información del expediente para mostrar en la advertencia
            cursor.execute("""
                SELECT codigo_rma, cliente, numero_documento_cliente, resultado_expediente
                FROM rma_maestro 
                WHERE id = ?
            """, (rma_id,))
            info_exp = cursor.fetchone()
            
            if not info_exp:
                messagebox.showerror("Error", "No se encontró el expediente.")
                return
            
            codigo_rma, cliente, num_doc, resultado = info_exp
            
            # Contar datos relacionados
            cursor.execute("SELECT COUNT(*) FROM rma_detalles WHERE rma_id = ?", (rma_id,))
            num_articulos = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM rma_historial WHERE rma_id = ?", (rma_id,))
            num_historial = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM tareas WHERE expediente_codigo = ?", (codigo_rma,))
            num_tareas = cursor.fetchone()[0]
            
            conn.close()
            
            # Crear ventana de advertencia personalizada
            ventana_advertencia = ctk.CTkToplevel(self)
            ventana_advertencia.title("⚠️ ADVERTENCIA - Eliminar Expediente")
            ventana_advertencia.geometry("600x450")
            ventana_advertencia.transient(self)
            ventana_advertencia.grab_set()
            
            # Frame principal
            main_frame = ctk.CTkFrame(ventana_advertencia)
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Título de advertencia
            lbl_titulo = ctk.CTkLabel(
                main_frame,
                text="⚠️ ADVERTENCIA: ELIMINACIÓN PERMANENTE",
                font=ctk.CTkFont(size=18, weight="bold"),
                text_color="red"
            )
            lbl_titulo.pack(pady=(0, 20))
            
            # Información del expediente
            info_frame = ctk.CTkFrame(main_frame)
            info_frame.pack(fill="x", pady=10)
            
            info_texto = f"""
EXPEDIENTE A ELIMINAR:

• Código RMA: {codigo_rma}
• Cliente: {cliente}
• Número Documento: {num_doc}
• Resultado: {resultado or 'Sin resultado'}

DATOS RELACIONADOS QUE SE ELIMINARÁN:

• Artículos en RMA: {num_articulos}
• Registros de historial: {num_historial}
• Tareas asociadas: {num_tareas}

⚠️ ESTA ACCIÓN NO SE PUEDE DESHACER ⚠️
            """
            
            lbl_info = ctk.CTkLabel(
                info_frame,
                text=info_texto,
                font=ctk.CTkFont(size=13),
                justify="left"
            )
            lbl_info.pack(padx=10, pady=10)
            
            # Campo de confirmación
            lbl_confirmar = ctk.CTkLabel(
                main_frame,
                text='Para confirmar, escriba "ELIMINAR" en el campo siguiente:',
                font=ctk.CTkFont(size=12, weight="bold")
            )
            lbl_confirmar.pack(pady=(10, 5))
            
            entry_confirmar = ctk.CTkEntry(
                main_frame,
                width=300,
                font=ctk.CTkFont(size=14)
            )
            entry_confirmar.pack(pady=5)
            
            # Frame de botones
            btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            btn_frame.pack(pady=20)
            
            def ejecutar_eliminacion():
                """Ejecuta la eliminación si la confirmación es correcta"""
                confirmacion = entry_confirmar.get().strip().upper()
                
                if confirmacion != "ELIMINAR":
                    messagebox.showwarning("Confirmación Incorrecta", 
                                         'Debe escribir exactamente "ELIMINAR" para confirmar.')
                    return
                
                try:
                    conn = self.master.conectar_db()
                    cursor = conn.cursor()
                    
                    # Eliminar en orden de dependencias
                    # 1. Detalles de artículos
                    cursor.execute("DELETE FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                    
                    # 2. Historial
                    cursor.execute("DELETE FROM rma_historial WHERE rma_id = ?", (rma_id,))
                    
                    # 3. Tareas asociadas
                    cursor.execute("DELETE FROM tareas WHERE expediente_codigo = ?", (codigo_rma,))
                    
                    # 4. Finalmente el registro maestro
                    cursor.execute("DELETE FROM rma_maestro WHERE id = ?", (rma_id,))
                    
                    conn.commit()
                    conn.close()
                    
                    messagebox.showinfo("Expediente Eliminado", 
                                      f"El expediente {codigo_rma} y todos sus datos relacionados han sido eliminados correctamente.")
                    
                    ventana_advertencia.destroy()
                    
                    # Cerrar la ventana del expediente actual
                    if hasattr(self, 'destroy'):
                        self.destroy()
                    
                    # Refrescar la tabla principal si existe
                    if hasattr(self.master, 'actualizar_tabla_rmas'):
                        self.master.actualizar_tabla_rmas()
                        
                except Exception as e:
                    messagebox.showerror("Error", f"Error al eliminar el expediente:\n{str(e)}")
            
            btn_eliminar = ctk.CTkButton(
                btn_frame,
                text="🗑️ ELIMINAR EXPEDIENTE",
                fg_color="red",
                hover_color="darkred",
                font=ctk.CTkFont(size=14, weight="bold"),
                command=ejecutar_eliminacion,
                width=200
            )
            btn_eliminar.pack(side="left", padx=5)
            
            btn_cancelar = ctk.CTkButton(
                btn_frame,
                text="✖ Cancelar",
                fg_color="gray",
                hover_color="darkgray",
                font=ctk.CTkFont(size=14),
                command=ventana_advertencia.destroy,
                width=150
            )
            btn_cancelar.pack(side="left", padx=5)
            
            # Focus en el campo de confirmación
            entry_confirmar.focus()
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al procesar la eliminación:\n{str(e)}")

    def guardar_nuevo_rma(self):
        """Valida los campos y realiza la inserción en rma_maestro y rma_detalles."""
        
        # 1. Recolección y Validación de campos obligatorios
        datos_maestro = {}
        campos_a_insertar = [
            'Cliente', 'Numero_Documento_Cliente', 'Persona_de_Contacto', 'Email_de_Contacto',
            'Autorizacion', 'Autorizado_Por', 'Fecha_Autorizacion', 'Fecha_Recepcion',
            'Recepcionado_Por', 'Fecha_Gestion', 'Gestionado_Por', 'Fecha_Proceso', 'Procesado_Por',
            'Fecha_para_factura', 'Numero_Albaran', 'Fecha_Doc_Cliente', 'Resultado_Expediente', 'motivo', 'Rma_Proveedor',
            'Modelo', 'N_Serie', 'Ref_Proveedor', 'Obs_Tecnica',
            'numero_albaran_reposicion', 'fecha_albaran_reposicion', 'numero_factura_abono', 'fecha_factura_abono'
        ]
        
        for campo in campos_a_insertar:
            entry = getattr(self, f"entry_{campo}")
            # Obtener el valor según el tipo de widget; soportar CTkDatePicker con get_date()/get()
            try:
                if isinstance(entry, ctk.CTkTextbox):
                    valor = entry.get("1.0", "end-1c").strip()
                elif hasattr(entry, 'get_date'):
                    # CTkDatePicker exposes get_date() and also get() alias
                    valor = entry.get_date()
                elif hasattr(entry, 'get'):
                    valor = entry.get()
                else:
                    valor = entry.cget("text")
            except Exception:
                # Fallback seguro
                try:
                    if hasattr(entry, 'get_date'):
                        valor = entry.get_date()
                    else:
                        valor = entry.get()
                except Exception:
                    valor = ''
            
            # Validación de obligatorios
            if campo in ["Cliente", "Numero_Documento_Cliente", "Persona_de_Contacto", "Email_de_Contacto", "motivo"] and not valor:
                print(f"Error: El campo {campo.replace('_', ' ')} es obligatorio.")
                messagebox.showinfo("Advertencia", f"Error: El campo {campo.replace('_', ' ')} es obligatorio.")
                # Aquí deberías mostrar un mensaje de error en la interfaz
                return
            
            # Para los campos de fecha, validar y normalizar a ISO YYYY-MM-DD
            DATE_FIELDS = {'Fecha_Autorizacion', 'Fecha_Recepcion', 'Fecha_Proceso', 'Fecha_Gestion', 'Fecha_Emision', 'Fecha_Doc_Cliente', 'fecha_albaran_reposicion', 'fecha_factura_abono'}
            if campo in DATE_FIELDS:
                # Permitir valor vacío
                if valor is None or str(valor).strip() == "":
                    datos_maestro[campo.lower()] = ''
                else:
                    try:
                        datos_maestro[campo.lower()] = parse_date_to_iso(valor)
                    except ValueError as e:
                        messagebox.showerror("Fecha inválida", f"El campo {campo} debe ser una fecha válida. Valor: {valor}")
                        return
            else:
                # Campos no fecha: conversiones especiales
                if campo == 'Autorizacion':
                    try:
                        v = str(valor).strip().upper() if valor is not None else ''
                        datos_maestro[campo.lower()] = 1 if v in ('1', 'SI', 'S', 'TRUE', 'T') else 0
                    except Exception:
                        datos_maestro[campo.lower()] = 0
                elif campo == 'Email_de_Contacto':
                    datos_maestro[campo.lower()] = valor.lower() if valor else ''
                else:
                    datos_maestro[campo.lower()] = valor

        # Campos automáticos/calculados
        # NOTA: El código_rma se generará al momento del guardado para evitar condiciones de carrera
        if hasattr(self, 'es_modo_edicion') and self.es_modo_edicion:
            # En modo edición, conservamos el código existente
            datos_maestro['codigo_rma'] = self.lbl_codigo_rma.cget("text").split(": ")[1]
        else:
            # En modo nuevo, lo generaremos dentro de la transacción
            # No asignamos código_rma aquí
            pass
            
        datos_maestro['fecha_emision'] = self.entry_Fecha_Emision.get()
        datos_maestro['creado_por'] = self.entry_Creado_Por.get()
        
        # Definir estado inicial basado en Autorización
        # 1. INTEGRACIÓN DE LA TRAZABILIDAD
        datos_maestro['estado'] = self.determinar_estado_rma(datos_maestro)
        
        # 2. Calcular Precio Total y validar Artículos
        precio_total = 0.0
        
        for item in self.articulos_data:
            try:
                # Convertir cantidad_entregada a float de forma segura
                cantidad = item.get('cantidad_entregada', 0)
                if isinstance(cantidad, str):
                    cantidad = float(cantidad.replace(',', '.')) if cantidad.strip() else 0.0
                elif cantidad is None:
                    cantidad = 0.0
                else:
                    cantidad = float(cantidad)
                
                # Convertir precio_unitario a float de forma segura
                precio = item.get('precio_unitario', 0.0)
                if isinstance(precio, str):
                    precio = float(precio.replace(',', '.')) if precio.strip() else 0.0
                elif precio is None:
                    precio = 0.0
                else:
                    precio = float(precio)
                
                precio_total += cantidad * precio
                
            except (ValueError, TypeError) as e:
                print(f"Error al calcular precio para artículo {item.get('referencia_articulo', 'N/A')}: {e}")
                continue
        
        datos_maestro['precio_total_expediente'] = precio_total

        # 2.5 Validación: Numero de documento del cliente no debe repetirse
        # Se permiten repeticiones para valores provisionales que contengan 'email' o 'telefonic'
        numero_doc = str(datos_maestro.get('numero_documento_cliente', '')).strip()
        if numero_doc:
            numero_doc_norm = numero_doc.lower()
            # Si el valor contiene las palabras provisionales (email/e-mail/telefonica/telefonico), no realizar la comprobación de duplicados
            if not re.search(r"\b(e-?mail|telefonica|telefonico)\b", numero_doc_norm, re.IGNORECASE):
                # Conectar a la DB para comprobar duplicados
                conn_check, cursor_check = self.master.conectar_db()
                if not conn_check:
                    return
                try:
                    cursor_check = conn_check.cursor()
                    cursor_check.execute("SELECT COUNT(*) FROM rma_maestro WHERE lower(numero_documento_cliente) = ?", (numero_doc_norm,))
                    count = cursor_check.fetchone()[0]
                    # Convertir a entero para evitar errores de comparación
                    count = int(count) if count is not None else 0
                    if count > 0:
                        conn_check.close()
                        messagebox.showwarning("Valor duplicado", f"El número de documento '{numero_doc}' ya existe en otro expediente.")
                        return
                except sqlite3.Error as e:
                    print(f"Error comprobando duplicados de numero_documento_cliente: {e}")
                    conn_check.close()
                    return
                # Cerramos la conexión de comprobación; la conexión principal se abrirá más abajo para la inserción
                conn_check.close()

        # if not self.articulos_data:
        #     print("Error: Debe añadir al menos un artículo.")
        #     return

        # Aviso: Si el número de documento contiene palabras provisionales, pedir confirmación al usuario
        try:
            if numero_doc and re.search(r"\b(e-?mail|telefonica|telefonico)\b", numero_doc, re.IGNORECASE):
                respuesta_prov = messagebox.askyesno("Valor provisional detectado",
                                                     "El campo 'Núm. Doc. Cliente' contiene un valor provisional (p.ej. 'Email' o 'Telefonica').\n¿Deseas continuar guardando el expediente con este valor?")
                if not respuesta_prov:
                    return
        except Exception:
            # Si falla mostrar el diálogo, continuar sin bloquear el guardado
            pass

        # 3. Inserción en la Base de Datos
        conn, cursor = self.master.conectar_db()
        if not conn: return
        cursor = conn.cursor()
        
        try:
            # Para nuevos RMA, generar el código final dentro de la transacción
            if not hasattr(self, 'es_modo_edicion') or not self.es_modo_edicion:
                codigo_rma_final = self.generar_codigo_rma_final(cursor)
                datos_maestro['codigo_rma'] = codigo_rma_final
            
            # 3a. Inserción en rma_maestro
            columnas_maestro = ', '.join(datos_maestro.keys())
            placeholders_maestro = ', '.join('?' * len(datos_maestro))
            valores_maestro = tuple(datos_maestro.values())
            
            cursor.execute(f"""
                INSERT INTO rma_maestro ({columnas_maestro}, estado) 
                VALUES ({placeholders_maestro}, ?)
            """, valores_maestro + (datos_maestro['estado'],)) # El estado se añade al final

            # Obtener el ID del RMA recién creado
            rma_id_generado = cursor.lastrowid
            
            # 3b. Inserción en rma_detalles - OPTIMIZADO con executemany
            if self.articulos_data:
                # Preparar todos los artículos para inserción batch
                primer_articulo = self.articulos_data[0].copy()
                primer_articulo['rma_id'] = rma_id_generado
                
                columnas_detalle = ', '.join(primer_articulo.keys())
                placeholders_detalle = ', '.join('?' * len(primer_articulo))
                
                # Preparar lista de valores para executemany
                valores_batch = []
                for articulo in self.articulos_data:
                    articulo_copia = articulo.copy()
                    articulo_copia['rma_id'] = rma_id_generado
                    valores_batch.append(tuple(articulo_copia.values()))
                
                # Insertar todos los artículos en una sola llamada (mucho más rápido con Turso)
                cursor.executemany(f"""
                    INSERT INTO rma_detalles ({columnas_detalle}) 
                    VALUES ({placeholders_detalle})
                """, valores_batch)

            # 3c. Inserción en rma_historial (modificar descripción si no hay artículos)
            num_articulos = len(self.articulos_data)
            descripcion = f"RMA creado. Cliente: {datos_maestro['cliente']}. Artículos: {num_articulos}. Total: {precio_total:.2f} €"
            cursor.execute("""
                INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                VALUES (?, ?, ?, ?)
            """, (rma_id_generado, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), self.username, descripcion))

            conn.commit()
            
            # Invalidar caché de estados (puede que se haya creado un nuevo estado)
            invalidate_cache('estados_rma')
            
            # Mostrar mensaje de confirmación con número RMA final
            if not hasattr(self, 'es_modo_edicion') or not self.es_modo_edicion:
                # Es un nuevo expediente - mostrar el número final asignado
                codigo_final = datos_maestro['codigo_rma']
                messagebox.showinfo("¡Expediente Creado!", 
                                    f"✅ Expediente creado exitosamente\n\n"
                                    f"📋 Número final asignado: {codigo_final}\n"
                                    f"👤 Cliente: {datos_maestro['cliente']}\n"
                                    f"💰 Total: {precio_total:.2f} €\n\n"
                                    f"Ya puede adjuntar archivos y gestionar el expediente.")
                
                # Actualizar el label con el código final
                self.lbl_codigo_rma.configure(
                    text=f"Nº EXPEDIENTE: {codigo_final}",
                    text_color="green"  # Verde para indicar que ya está guardado
                )
                
                # Actualizar el botón de guardar a verde para indicar que ya está guardado
                if hasattr(self, 'btn_guardar_rma'):
                    self.btn_guardar_rma.configure(
                        text="💾 ACTUALIZAR",
                        fg_color="green"  # Verde para indicar que ya está guardado
                    )
                
                # Ocultar el texto explicativo ya que ahora es definitivo
                if hasattr(self, 'lbl_explicacion'):
                    self.lbl_explicacion.configure(text="✅ Número definitivo asignado")
                    self.lbl_explicacion.configure(text_color="green")
            else:
                messagebox.showinfo("Expediente Guardado", "El expediente se ha actualizado correctamente.")
            
            # IMPORTANTE: Actualizar ambos IDs para que el siguiente guardado use actualizar_rma()
            self.current_rma_id = rma_id_generado  # Asigna el ID al atributo de instancia
            self.rma_actual_id = rma_id_generado   # CRÍTICO: Actualizar para que guardar_rma_placeholder() use actualizar_rma()
            self.mode = 'editar'                    # Cambia la ventana a modo edición
            self.es_modo_edicion = True            # Actualizar el indicador
            
            # Eliminamos la llamada a self.mostrar_lista_rma() para mantener la vista abierta.
            
            # Actualizar el título de la pestaña si fuera necesario
            self.tabview.set("📝 General")
            
            # Ya no necesitamos este mensaje duplicado - se maneja arriba
            # messagebox.showinfo("Éxito", f"RMA {datos_maestro['codigo_rma']} creado y guardado. Ahora puede adjuntar archivos.")
            
            # NO volver al listado - mantener el expediente abierto para que puedan trabajar con él
            # self.mostrar_lista_rma()

        except sqlite3.IntegrityError as e:
            conn.rollback()
            print(f"Error al guardar (Integridad): {e}. Es posible que el código RMA ya exista.")
        except sqlite3.Error as e:
            conn.rollback()
            print(f"Error general de DB al guardar: {e}")
        finally:
            conn.close()

    # ----------------------------------------------------------------------
    # 6. LÓGICA DE EDICIÓN Y ACTUALIZACIÓN (CARGA Y GUARDADO DE CAMBIOS)
    # ----------------------------------------------------------------------

    def obtener_datos_actuales_maestro(self):
        """Recupera los datos del formulario MAESTRO actuales."""
        datos_maestro = {}
        # Lista de campos de la tabla rma_maestro que corresponden a entries
        campos_a_recuperar = [
            'Cliente', 'Numero_Documento_Cliente', 'Persona_de_Contacto', 'Email_de_Contacto',
            'Autorizacion', 'Autorizado_Por', 'Fecha_Autorizacion', 'Fecha_Recepcion',
            'Recepcionado_Por', 'Fecha_Gestion', 'Gestionado_Por', 'Fecha_Proceso', 'Procesado_Por',
            'Fecha_para_factura', 'Numero_Albaran', 'Fecha_Doc_Cliente', 'Resultado_Expediente',
            'Fecha_Emision', 'Creado_Por', 'motivo', 'Rma_Proveedor', 'Modelo', 'N_Serie', 'Ref_Proveedor', 'Obs_Tecnica',
            'numero_albaran_reposicion', 'fecha_albaran_reposicion', 'numero_factura_abono', 'fecha_factura_abono'
        ]

        for campo in campos_a_recuperar:
            entry_name = f"entry_{campo}"
            if hasattr(self, entry_name):
                entry = getattr(self, entry_name)
                # Intenta obtener el valor de la entrada o del optionmenu
                # Soportar CTkTextbox (requiere índices para get)
                try:
                    if isinstance(entry, ctk.CTkTextbox):
                        valor = entry.get("1.0", "end-1c").strip()
                    elif hasattr(entry, 'get_date'):
                        valor = entry.get_date()
                    elif hasattr(entry, 'get'):
                        valor = entry.get()
                    else:
                        valor = entry.cget("text")
                except Exception:
                    # Fallback
                    try:
                        if hasattr(entry, 'get_date'):
                            valor = entry.get_date()
                        else:
                            valor = entry.get()
                    except Exception:
                        valor = ''

                # Normalizar fechas si corresponde
                DATE_FIELDS = {'Fecha_Autorizacion', 'Fecha_Recepcion', 'Fecha_Proceso', 'Fecha_Gestion', 'Fecha_Emision', 'Fecha_Doc_Cliente', 'fecha_albaran_reposicion', 'fecha_factura_abono'}
                if campo in DATE_FIELDS:
                    if valor is None or str(valor).strip() == "":
                        datos_maestro[campo.lower()] = ''
                    else:
                        try:
                            datos_maestro[campo.lower()] = parse_date_to_iso(valor)
                        except ValueError:
                            # Mostrar error y abortar la recolección para forzar corrección
                            messagebox.showerror("Fecha inválida", f"El campo {campo} debe contener una fecha válida. Valor: {valor}")
                            return None
                else:
                    # Conversión especial para Autorizacion (SI/NO a 1/0)
                    if campo == 'Autorizacion':
                        try:
                            v = str(valor).strip().upper() if valor is not None else ''
                            datos_maestro['autorizacion'] = 1 if v in ('1', 'SI', 'S', 'TRUE', 'T') else 0
                        except Exception:
                            datos_maestro['autorizacion'] = 0
                    # Conversión especial para Email_de_Contacto (siempre en minúsculas)
                    elif campo == 'Email_de_Contacto':
                        datos_maestro[campo.lower()] = valor.lower() if valor else ''
                    else:
                        datos_maestro[campo.lower()] = valor
        
        datos_maestro['codigo_rma'] = self.lbl_codigo_rma.cget("text").split(": ")[1]
        
        return datos_maestro

    def autorrellena_pdf(self):
        """Autorrellena la plantilla PDF con los datos del RMA actual y la guarda en Dropbox.

        Busca 'Plantilla_SOLICITUD RMA.pdf' en la carpeta plantillas/. Si no existe, abre
        un diálogo para seleccionar la plantilla. Luego llama a la función de librería
        para rellenar el PDF y lo sube a Dropbox como adjunto.
        """
        # Validaciones
        if not hasattr(self, 'current_rma_id') or not self.current_rma_id:
            messagebox.showwarning("Guardar primero", "Guarde el expediente antes de generar el PDF.")
            return

        # Obtener código RMA del label
        try:
            codigo_rma = self.lbl_codigo_rma.cget("text").split(": ")[1]
        except Exception:
            messagebox.showerror("Error", "No se pudo determinar el código RMA.")
            return

        # Determinar plantilla por defecto
        base_dir = os.path.dirname(os.path.abspath(__file__))
        plantilla_def = os.path.join(base_dir, 'plantillas', 'Plantilla_SOLICITUD RMA.pdf')
        if os.path.exists(plantilla_def):
            plantilla_path = plantilla_def
        else:
            # Pedir al usuario que seleccione la plantilla
            plantilla_path = filedialog.askopenfilename(
                title='Seleccionar plantilla PDF', 
                initialdir=os.path.join(base_dir, 'plantillas'), 
                filetypes=[('PDF files', '*.pdf')]
            )
            if not plantilla_path:
                return

        # Comprobar que la función de relleno esté disponible
        if fill_pdf is None:
            messagebox.showerror("Dependencia falta", "La funcionalidad de rellenado PDF no está disponible. Instala la dependencia o revisa el módulo lib.pdf_fill.")
            return

        # Obtener datos del RMA desde Turso usando la conexión de la aplicación
        try:
            conn, cursor = self.master.conectar_db()
            cursor.execute("SELECT codigo_rma, modelo, n_serie, obs_tecnica FROM rma_maestro WHERE codigo_rma = ?", (codigo_rma,))
            datos_rma = cursor.fetchone()
            conn.close()
            
            if not datos_rma:
                messagebox.showerror("Error", f"No se encontraron datos para el RMA {codigo_rma} en la base de datos.")
                return
            
            codigo_bd, modelo, n_serie, obs_tecnica = datos_rma
            
            # Debug: Mostrar datos obtenidos de Turso
            print(f"DEBUG PDF - Datos desde TURSO para {codigo_bd}:")
            print(f"  modelo: '{modelo}'" + (" ✓" if modelo else " [VACÍO]"))
            print(f"  n_serie: '{n_serie}'" + (" ✓" if n_serie else " [VACÍO]"))
            print(f"  obs_tecnica: '{obs_tecnica}'" + (" ✓" if obs_tecnica else " [VACÍO]"))
            
        except Exception as e:
            messagebox.showerror("Error BD", f"Error obteniendo datos de la base de datos: {e}")
            return

        # Preparar nombres de archivo
        nombre_base = f"{codigo_rma}_SOLICITUD RMA.pdf"
        
        # Si el archivo ya existe en Dropbox, añadimos timestamp para evitar sobrescribir
        nombre_salida = nombre_base
        if usar_dropbox():
            # Verificar si ya existe en Dropbox
            dbx = get_dropbox_client()
            if dbx:
                try:
                    ruta_check = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{codigo_rma}/{nombre_base}")
                    dbx.files_get_metadata(ruta_check)
                    # Si llegamos aquí, el archivo existe, añadir timestamp
                    fecha_str = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    nombre_salida = f"{codigo_rma}_SOLICITUD RMA_{fecha_str}.pdf"
                except:
                    # El archivo no existe, podemos usar el nombre base
                    pass

        # Crear archivo temporal para el PDF generado
        temp_dir = tempfile.mkdtemp(prefix="solicitud_rma_")
        temp_pdf_path = os.path.join(temp_dir, nombre_salida)

        # Rellenar PDF con datos obtenidos de Turso
        try:
            # Campos que queremos que queden vacíos y editables
            force_empty = [
                'Cantidad afectada', 'N Pedido', 'Nº de pedido  Albarán', 'Nº de pedido Albarán',
                'N Pedido Albarán', 'Nº de pedido Albaran', 'Nº RMA'
            ]

            # Preparar valores para rellenar el PDF directamente con datos de Turso
            valores_pdf = {}
            
            # Aplicar mapping: De campos PDF a valores de BD obtenidos de Turso
            if modelo:
                valores_pdf['Modelo'] = modelo
            if n_serie: 
                valores_pdf['Número de serie'] = n_serie
            if obs_tecnica:
                valores_pdf['Ubicación de las fuentes en la instalación'] = obs_tecnica
            if codigo_rma:
                valores_pdf['Referencia para devolución'] = codigo_rma
            
            # Campos que queremos que queden vacíos y editables
            for campo_vacio in force_empty:
                valores_pdf[campo_vacio] = ''
            
            print(f"DEBUG PDF - Valores que se van a escribir en PDF:")
            for campo, valor in valores_pdf.items():
                print(f"  '{campo}' = '{valor}'")
            
            # DEBUG: Verificar campos disponibles en el PDF antes de rellenar
            try:
                from lib.pdf_fill import get_pdf_field_names
                if get_pdf_field_names:
                    campos_pdf = get_pdf_field_names(plantilla_path)
                    print(f"\nDEBUG PDF - Campos disponibles en la plantilla ({len(campos_pdf)} total):")
                    for i, campo in enumerate(campos_pdf, 1):
                        estado = "✓ MATCH" if campo in valores_pdf else "○ Sin mapear"
                        print(f"  {i:2d}. '{campo}' {estado}")
                    
                    # Verificar matching exacto
                    print(f"\nDEBUG PDF - Verificación de coincidencias:")
                    for campo_valor, valor in valores_pdf.items():
                        if campo_valor in campos_pdf:
                            print(f"  ✅ '{campo_valor}' -> EXISTE en PDF")
                        else:
                            print(f"  ❌ '{campo_valor}' -> NO EXISTE en PDF")
                            # Buscar campos similares
                            similares = [c for c in campos_pdf if campo_valor.lower() in c.lower() or c.lower() in campo_valor.lower()]
                            if similares:
                                print(f"     Similares: {similares}")
            except Exception as debug_e:
                print(f"DEBUG PDF - Error obteniendo campos PDF: {debug_e}")
            
            # Rellenar PDF usando función directa (sin consulta a BD adicional)
            print(f"\nDEBUG PDF - Ejecutando fill_pdf...")
            print(f"  Plantilla: {plantilla_path}")
            print(f"  Salida: {temp_pdf_path}")
            
            # Verificar que la plantilla existe y es accesible
            if os.path.exists(plantilla_path):
                size_plantilla = os.path.getsize(plantilla_path)
                print(f"  Plantilla existe: {size_plantilla} bytes")
            else:
                print(f"  ❌ Plantilla NO existe")
                
            fill_pdf(plantilla_path, temp_pdf_path, valores_pdf)
            
            # Verificar que el archivo de salida se generó
            if os.path.exists(temp_pdf_path):
                size_salida = os.path.getsize(temp_pdf_path)
                print(f"  ✅ PDF generado: {size_salida} bytes")
                
                # Verificar si el tamaño cambió (indicaría que se modificó)
                if abs(size_salida - size_plantilla) > 100:  # Al menos 100 bytes de diferencia
                    print(f"  ✅ PDF modificado (diferencia: {size_salida - size_plantilla} bytes)")
                else:
                    print(f"  ⚠️  PDF no modificado o cambio mínimo (diferencia: {size_salida - size_plantilla} bytes)")
            else:
                print(f"  ❌ PDF NO se generó")
                
            print(f"DEBUG PDF - fill_pdf completado")
        except Exception as e:
            # Limpiar archivo temporal
            try:
                os.remove(temp_pdf_path)
                os.rmdir(temp_dir)
            except:
                pass
            messagebox.showerror("Error Rellenado", f"Error al rellenar la plantilla: {e}")
            return

        # Decidir dónde guardar (Dropbox o local)
        if usar_dropbox():
            # Subir a Dropbox
            exito, ruta_relativa = self._subir_archivo_dropbox(temp_pdf_path, codigo_rma, nombre_salida)
            tipo_almacenamiento = 'dropbox'
            ubicacion_desc = "Dropbox"
        else:
            # Guardar localmente (fallback)
            exito, ruta_relativa = self._subir_archivo_local(temp_pdf_path, codigo_rma, nombre_salida)
            tipo_almacenamiento = 'local'
            ubicacion_desc = "local"
        
        # Limpiar archivo temporal
        try:
            os.remove(temp_pdf_path)
            os.rmdir(temp_dir)
        except:
            pass
        
        if not exito:
            messagebox.showerror("Error", f"No se pudo guardar la solicitud PDF en {ubicacion_desc}.")
            return

        # Registrar en la base de datos como adjunto
        try:
            conn, cursor = self.master.conectar_db()
            
            # Verificar esquema de BD antes de insertar
            self._verificar_columna_tipo_almacenamiento(cursor)
            
            # Preparar inserción con o sin tipo_almacenamiento según el esquema
            if getattr(self, '_usar_tipo_almacenamiento', False):
                cursor.execute("""
                    INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    self.current_rma_id,
                    nombre_salida,
                    ruta_relativa,
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    self.username,
                    tipo_almacenamiento
                ))
            else:
                cursor.execute("""
                    INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    self.current_rma_id,
                    nombre_salida,
                    ruta_relativa,
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    self.username
                ))
            
            # También registrar entrada en rma_historial con información de ubicación
            cursor.execute("""
                INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                VALUES (?, ?, ?, ?)
            """, (
                self.current_rma_id,
                datetime.datetime.now().isoformat(),
                self.username,
                f"Generada Solicitud RMA: {nombre_salida} ({'☁️ Dropbox' if usar_dropbox() else '💾 Local'})"
            ))

            conn.commit()
            conn.close()
            
            # Refrescar lista de adjuntos
            try:
                self.cargar_lista_adjuntos(self.current_rma_id)
            except Exception:
                pass
            
            # Actualizar historial si está visible
            try:
                if hasattr(self, 'historial_tab'):
                    self.mostrar_historial(self.historial_tab)
            except AttributeError:
                pass
                
            # Feedback al usuario personalizado según ubicación
            if usar_dropbox():
                messagebox.showinfo("Éxito", f"✅ Solicitud PDF '{nombre_salida}' generada y subida a Dropbox correctamente.\n\n📁 Ubicación: {ruta_relativa}")
            else:
                messagebox.showinfo("Éxito", f"✅ Solicitud PDF '{nombre_salida}' generada y guardada localmente.")
                
        except Exception as e:
            # Limpiar archivo temporal si hubo error
            try:
                os.remove(temp_pdf_path)
                os.rmdir(temp_dir)
            except:
                pass
            messagebox.showwarning("Aviso", f"PDF generado, pero no se pudo registrar en la BD: {e}")
            return


    def cargar_datos_rma(self, rma_id):
        """Carga el RMA maestro y sus detalles en el formulario."""
        conn, cursor = self.master.conectar_db()
        if not conn: return
        cursor = conn.cursor()
        
        try:
            # 1. Cargar RMA Maestro
            cursor.execute("SELECT * FROM rma_maestro WHERE id = ?", (rma_id,))
            columnas_maestro = [col[0] for col in cursor.description]
            datos_maestro = dict(zip(columnas_maestro, cursor.fetchone()))
            
            self.datos_rma_maestro = datos_maestro
            
            # 2. Cargar RMA Detalles (Artículos)
            cursor.execute("SELECT referencia_articulo, cantidad_segun_documento, cantidad_entregada, estado_producto, precio_unitario FROM rma_detalles WHERE rma_id = ?", (rma_id,))
            columnas_detalle = [col[0] for col in cursor.description]
            articulos_db = [dict(zip(columnas_detalle, fila)) for fila in cursor.fetchall()]

            conn.close()

            # 3. Rellenar Formulario Maestro
            self.lbl_codigo_rma.configure(text=f"Nº EXPEDIENTE: {datos_maestro['codigo_rma']}")
            
            # >>> VERIFICACIÓN CRÍTICA DEL WIDGET MOTIVO <<<
            # Si esta línea falla, el nombre de la variable está mal.
            entry_motivo = None
            
            # Prueba 1: entry_Motivo (La que falló)
            if 'entry_Motivo' in self.__dict__:
                entry_motivo = self.entry_Motivo
            # Prueba 2: entry_motivo (Todo minúscula, la más probable)
            elif 'entry_motivo' in self.__dict__:
                entry_motivo = self.entry_motivo
            # Prueba 3: entry_motivo_ (Por si hay un guion bajo final)
            elif 'entry_motivo_' in self.__dict__:
                entry_motivo = getattr(self, "entry_motivo_", None)

            # Si encontramos el widget con el nombre correcto:
            if entry_motivo and 'motivo' in datos_maestro:
                valor_motivo = datos_maestro['motivo']
                
                if isinstance(entry_motivo, ctk.CTkEntry):
                    # 1. Borrar el placeholder
                    entry_motivo.configure(state="normal")
                    entry_motivo.delete(0, ctk.END)
                    
                    # 2. Insertar el valor
                    valor_a_insertar = str(valor_motivo) if valor_motivo is not None and str(valor_motivo).strip() != "" else ""
                    entry_motivo.insert(0, valor_a_insertar)
                    
                    # 3. Restaurar estado (si no era un campo deshabilitado, se queda en normal)
                    # NOTA: En este formulario es un campo editable, así que se queda en 'normal'
            # -----------------------------------------------
            
            # --- TRATAMIENTO ESPECIAL PARA FECHA_PARA_FACTURA ---
            # Este campo usa quincenas futuras, necesita manejo especial
            
            if hasattr(self, 'entry_Fecha_para_factura'):
                widget_fecha = self.entry_Fecha_para_factura
                
                if isinstance(widget_fecha, ctk.CTkOptionMenu):
                    # Probar con ambos nombres posibles
                    valor_fecha = datos_maestro.get('fecha_para_factura') or datos_maestro.get('Fecha_para_factura') or datos_maestro.get('Fecha_Para_factura')
                    
                    # Normalizar: si es None o la cadena 'None', tratarlo como vacío
                    if valor_fecha is None or str(valor_fecha).strip().lower() == 'none' or str(valor_fecha).strip() == '':
                        valor_fecha = None
                    
                    if valor_fecha:
                        # Hay un valor guardado
                        valores_actuales = list(widget_fecha.cget("values"))
                        valor_str = str(valor_fecha).strip()
                        
                        # Si el valor no está en las opciones (fecha antigua), añadirlo después de "Seleccionar..."
                        if valor_str and valor_str not in valores_actuales:
                            valores_actuales.insert(1, valor_str)  # Después de "Seleccionar..."
                            widget_fecha.configure(values=valores_actuales)
                        
                        # Establecer el valor guardado
                        widget_fecha.set(valor_str)
                    else:
                        # No hay valor guardado, establecer "Seleccionar..."
                        widget_fecha.set("Seleccionar...")
            
            # --- Mapeo de Columna DB a Variable de Formulario ---
            for columna, valor in datos_maestro.items():
                
                # Excluir 'id', 'precio_total_expediente', 'estado' y 'fecha_para_factura' (ya procesado arriba)
                if columna in ['id', 'precio_total_expediente', 'estado', 'fecha_para_factura']:
                    continue

                # ---------------------------------------------------------------------------------
                # 1. DETERMINAR EL NOMBRE CORRECTO DE LA VARIABLE DE LA INTERFAZ (entry_name)
                # ---------------------------------------------------------------------------------
                entry_name = None
                
                # Primero intentar con el nombre exacto del campo (como se crea en crear_campo)
                nombre_exacto = f"entry_{columna}"
                if hasattr(self, nombre_exacto):
                    entry_name = nombre_exacto
                # Caso A: Campos simples (Sin guiones bajos, como 'motivo', 'cliente', 'creado_por')
                # ESTA ES LA RUTA CRÍTICA QUE 'motivo' DEBE SEGUIR
                elif '_' not in columna:
                    # Convierte a título para hacer coincidir la convención (ej: motivo -> Motivo)
                    entry_name = f"entry_{columna.title()}" 
                
                # Caso B: Campos con guiones bajos (como 'persona_de_contacto', 'numero_documento_cliente')
                else:
                    # Aplica tu lógica compleja de transformación:
                    formato_titulo = columna.replace('_', ' ').title()
                    entry_key_name = formato_titulo.replace(' De ', ' de ')
                    entry_key_name = entry_key_name.replace(' ', '_')
                    entry_name = f"entry_{entry_key_name}"


                # ---------------------------------------------------------------------------------
                # 2. TRATAMIENTO Y RELLENO DEL WIDGET EN BASE AL NOMBRE
                # ---------------------------------------------------------------------------------
                
                # Si encontramos un nombre de variable, intentamos rellenar
                if entry_name and hasattr(self, entry_name):
                    entry = getattr(self, entry_name)
                    
                    # Tratamiento para DatePicker (si el widget expone `set_date`)
                    if hasattr(entry, 'set_date'):
                        try:
                            if valor is None or str(valor).strip() == '':
                                # Intentar limpiar el control de fecha
                                try:
                                    entry.set_date('')
                                except Exception:
                                    try:
                                        entry.configure(state='normal')
                                        entry.delete(0, ctk.END)
                                    except Exception:
                                        pass
                            else:
                                # Normalizar la fecha y establecerla
                                try:
                                    fecha_iso = parse_date_to_iso(valor)
                                    from datetime import datetime
                                    dt = datetime.strptime(fecha_iso, '%Y-%m-%d')
                                    entry.set_date(dt)
                                except Exception:
                                    try:
                                        entry.configure(state='normal')
                                        entry.delete(0, ctk.END)
                                        entry.insert(0, str(valor))
                                    except Exception:
                                        pass
                        except Exception:
                            pass
                        # continuar con la siguiente columna
                        continue

                    # Tratamiento especial para OptionMenu (Autorizacion)
                    if columna == 'autorizacion' and isinstance(entry, ctk.CTkOptionMenu):
                        # Interpretar distintos posibles formatos que pueda tener el valor
                        is_si = False
                        try:
                            if isinstance(valor, (int, float)):
                                is_si = int(valor) == 1
                            elif isinstance(valor, bool):
                                is_si = bool(valor)
                            elif isinstance(valor, str):
                                vnorm = valor.strip().upper()
                                is_si = vnorm in ('1', 'SI', 'S', 'TRUE', 'T')
                        except Exception:
                            is_si = False
                        entry.set('SI' if is_si else 'NO')
                        
                    # Tratamiento para Entry
                    elif isinstance(entry, ctk.CTkEntry):
                        # Configurar Entry
                        estado_original = entry.cget("state")  
                        entry.configure(state="normal")
                        
                        # CRÍTICO: Borrar el placeholder ANTES de insertar
                        entry.delete(0, ctk.END)
                        entry.insert(0, str(valor) if valor is not None else "")

                        # Si el campo es 'numero_documento_cliente', permitir edición
                        # cuando el valor contiene 'Email' o 'Telefonica' (case-insensitive).
                        try:
                            if columna == 'numero_documento_cliente':
                                valor_str = str(valor) if valor is not None else ''
                                if re.search(r"\b(e-?mail|telefonica|telefonico)\b", valor_str, re.IGNORECASE):
                                    # Mantener en editable
                                    entry.configure(state='normal')
                                else:
                                    entry.configure(state=estado_original)
                            else:
                                entry.configure(state=estado_original)
                        except Exception:
                            entry.configure(state=estado_original)

                    # Tratamiento para DatePicker (CTkDatePicker o widgets similares)
                    elif hasattr(entry, 'set_date'):
                        try:
                            # Si el valor es nulo, limpiar
                            if valor is None:
                                entry.set_date(None)
                            else:
                                entry.set_date(valor)
                        except Exception:
                            # Fallback: escribir directamente en el sub-entry si existe
                            try:
                                if hasattr(entry, 'date_entry'):
                                    entry.date_entry.configure(state='normal')
                                    entry.date_entry.delete(0, tk.END)
                                    entry.date_entry.insert(0, str(valor) if valor is not None else '')
                            except Exception:
                                pass
                        
                    # Tratamiento para OptionMenu General
                    elif isinstance(entry, ctk.CTkOptionMenu):
                        valores_actuales = list(entry.cget("values"))
                        valor_str = str(valor) if valor is not None else ""
                        
                        # Si el valor de la BD no está en las opciones actuales, añadirlo
                        if valor_str and valor_str not in valores_actuales:
                            valores_actuales.insert(0, valor_str)  # Añadir al principio
                            entry.configure(values=valores_actuales)
                        
                        # Establecer el valor
                        if valor_str:
                            entry.set(valor_str)
                    # Tratamiento para Textbox grande (Observaciones Técnicas u otros)
                    elif isinstance(entry, ctk.CTkTextbox):
                        try:
                            entry.delete("1.0", "end")
                            entry.insert("1.0", str(valor) if valor is not None else "")
                        except Exception:
                            # algunos widgets CTkTextbox pueden tener métodos distintos; ignorar si falla
                            pass
            
            # --- NUEVO: Actualizar la etiqueta del total ---
            precio_total = datos_maestro.get('precio_total_expediente', 0.0) # Obtener el valor
            self.lbl_precio_total.configure(text=f"{precio_total:.2f} €")
            
            # --- SINCRONIZACIÓN: Si hay fecha_autorizacion, marcar Autorización como SI ---
            try:
                fecha_aut = datos_maestro.get('fecha_autorizacion')
                if fecha_aut and str(fecha_aut).strip():
                    # Hay una fecha de autorización, asegurar que Autorización esté en "SI"
                    if hasattr(self, 'entry_Autorizacion'):
                        try:
                            self.entry_Autorizacion.set('SI')
                        except Exception as e:
                            print(f"Error al sincronizar Autorización: {e}")
            except Exception as e:
                print(f"Error en sincronización de autorización: {e}")
            
            # 4. Rellenar Artículos (self.articulos_data)
            self.articulos_data = articulos_db
            self.actualizar_listado_articulos()

        except Exception as e:
            print(f"Error al cargar datos del RMA ID {rma_id}: {e}")
            conn.close()
            
    def guardar_cambio_historial(self, rma_id, campo, valor_antiguo, valor_nuevo):
        """Registra un cambio de un campo en la tabla de historial SOLO si hay cambio real."""
        # Normalizar valores para comparación (convertir None a cadena vacía, strip)
        def normalizar_valor(val):
            if val is None:
                return ""
            return str(val).strip()
        
        val_antiguo_norm = normalizar_valor(valor_antiguo)
        val_nuevo_norm = normalizar_valor(valor_nuevo)
        
        # Si los valores son iguales después de normalizar, no registrar el cambio
        if val_antiguo_norm == val_nuevo_norm:
            return
        
        conn, cursor = self.master.conectar_db()
        if not conn: return

        cursor = conn.cursor()
        
        descripcion = f"Campo '{campo}' modificado: '{valor_antiguo}' -> '{valor_nuevo}'"
        # Si el campo modificado es 'Procesado Por' (o su variante), añadimos la nota requerida
        try:
            campo_norm = str(campo).lower().strip().replace(' ', '_')
            if campo_norm == 'procesado_por' or campo_norm == 'procesadopor':
                descripcion = descripcion + " - MATERIAL REVISADO"
        except Exception:
            # En caso de cualquier problema al normalizar, no bloqueamos el guardado del historial
            pass
        
        try:
            cursor.execute("""
                INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                VALUES (?, ?, ?, ?)
            """, (rma_id, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), self.username, descripcion))
            conn.commit()
        except sqlite3.Error as e:
            print(f"Error al registrar historial: {e}")
        finally:
            conn.close()

    def actualizar_rma(self):
        """Compara los datos, actualiza rma_maestro y rma_detalles, y registra los cambios."""
        conn, cursor = self.master.conectar_db()
        if not conn: return

        cursor = conn.cursor()
        rma_id = self.rma_actual_id
        
        # 1. Obtener datos antiguos de la DB
        cursor.execute("SELECT * FROM rma_maestro WHERE id = ?", (rma_id,))
        columnas_maestro_db = [col[0] for col in cursor.description]
        datos_antiguos = dict(zip(columnas_maestro_db, cursor.fetchone()))
        
        # 2. Obtener datos nuevos del formulario
        datos_nuevos = self.obtener_datos_actuales_maestro()
        
        # Si la recolección devolvió None, significa que había fechas inválidas y ya se mostró un error
        if datos_nuevos is None:
            conn.close()
            return

        # ⚠️ ADVERTENCIA ADMIN: Cambios en campos críticos
        if self.username.lower() == "admin":
            campos_criticos_modificados = []
            
            # Verificar cambios en Cliente
            cliente_antiguo = datos_antiguos.get('cliente', '')
            cliente_nuevo = datos_nuevos.get('cliente', '')
            if str(cliente_antiguo).strip() != str(cliente_nuevo).strip():
                campos_criticos_modificados.append(f"• Cliente: '{cliente_antiguo}' → '{cliente_nuevo}'")
            
            # Verificar cambios en Número de Documento
            num_doc_antiguo = datos_antiguos.get('numero_documento_cliente', '')
            num_doc_nuevo = datos_nuevos.get('numero_documento_cliente', '')
            if str(num_doc_antiguo).strip() != str(num_doc_nuevo).strip():
                campos_criticos_modificados.append(f"• Número Documento: '{num_doc_antiguo}' → '{num_doc_nuevo}'")
            
            # Verificar cambios en Fecha Emisión
            fecha_emision_antigua = datos_antiguos.get('fecha_emision', '')
            fecha_emision_nueva = datos_nuevos.get('fecha_emision', '')
            if str(fecha_emision_antigua).strip() != str(fecha_emision_nueva).strip():
                campos_criticos_modificados.append(f"• Fecha Emisión: '{fecha_emision_antigua}' → '{fecha_emision_nueva}'")
            
            # Verificar cambios en Creado Por
            creado_por_antiguo = datos_antiguos.get('creado_por', '')
            creado_por_nuevo = datos_nuevos.get('creado_por', '')
            if str(creado_por_antiguo).strip() != str(creado_por_nuevo).strip():
                campos_criticos_modificados.append(f"• Creado Por: '{creado_por_antiguo}' → '{creado_por_nuevo}'")
            
            # Si hay cambios críticos, mostrar advertencia
            if campos_criticos_modificados:
                mensaje_cambios = "\n".join(campos_criticos_modificados)
                respuesta = messagebox.askyesno(
                    "⚠️ ADVERTENCIA - Modificación de Campos Críticos",
                    f"ADMIN: Está modificando campos críticos del expediente:\n\n"
                    f"{mensaje_cambios}\n\n"
                    f"Estos cambios quedarán registrados en el historial.\n\n"
                    f"¿Confirmar modificación?"
                )
                
                if not respuesta:
                    conn.close()
                    return

        # Aviso: si el número de documento nuevo contiene valores provisionales, pedir confirmación
        try:
            num_doc_nuevo = str(datos_nuevos.get('numero_documento_cliente', '') or '').strip()
            if num_doc_nuevo and re.search(r"\b(e-?mail|telefonica|telefonico)\b", num_doc_nuevo, re.IGNORECASE):
                resp = messagebox.askyesno("Valor provisional detectado",
                                           "El campo 'Núm. Doc. Cliente' contiene un valor provisional (p.ej. 'Email' o 'Telefonica').\n¿Deseas continuar con la actualización?")
                if not resp:
                    conn.close()
                    return
        except Exception:
            pass

        # 2.1. INTEGRACIÓN DE LA TRAZABILIDAD - Calcular el nuevo estado
        estado_nuevo = self.determinar_estado_rma(datos_nuevos)
        
        # 2.2. CONFIRMACIONES INTELIGENTES: Verificar tareas pendientes antes de completar
        estado_anterior = datos_antiguos.get('estado', '')
        if estado_nuevo == 'Completado' and estado_anterior != 'Completado':
            # El expediente se está intentando completar
            count_tareas, tareas_pendientes = self.verificar_tareas_pendientes_expediente(rma_id)
            
            if count_tareas > 0:
                # Crear mensaje detallado con las tareas pendientes
                mensaje_tareas = "\n• ".join(tareas_pendientes[:5])  # Mostrar máximo 5 tareas
                if len(tareas_pendientes) > 5:
                    mensaje_tareas += f"\n• ... y {len(tareas_pendientes) - 5} tareas más"
                
                respuesta = messagebox.askyesno(
                    "⚠️ Expediente con tareas pendientes",
                    f"¿Seguro que deseas completar este expediente?\n\n"
                    f"Tiene {count_tareas} tarea(s) pendiente(s):\n\n• {mensaje_tareas}\n\n"
                    f"Se recomienda completar todas las tareas antes de cerrar el expediente.\n\n"
                    f"¿Continuar de todas formas?"
                )
                
                if not respuesta:
                    conn.close()
                    messagebox.showinfo("Operación cancelada", "El expediente no ha sido completado.")
                    return
        
        # Añadir el estado al diccionario de datos nuevos
        datos_nuevos['estado'] = estado_nuevo 
        
        # 3. Comparar y construir la consulta de actualización (UPDATE)
        campos_a_actualizar = []
        valores_a_actualizar = []
        
        def valores_son_diferentes(val1, val2):
            """Compara dos valores normalizando None y cadenas vacías."""
            # Normalizar None y cadenas vacías
            v1 = str(val1).strip() if val1 is not None else ""
            v2 = str(val2).strip() if val2 is not None else ""
            return v1 != v2
        
        for columna_db, valor_nuevo in datos_nuevos.items():
            valor_antiguo = datos_antiguos.get(columna_db)
            
            # SQLite almacena el boolean como int (1/0)
            if columna_db == 'autorizacion':
                if valor_nuevo != valor_antiguo:
                    self.guardar_cambio_historial(rma_id, "Autorización", "NO" if valor_antiguo == 0 else "SI", "NO" if valor_nuevo == 0 else "SI")
                    campos_a_actualizar.append(f"{columna_db} = ?")
                    valores_a_actualizar.append(valor_nuevo)
            
            # Comparación de campos normales (no boolean) usando la función de normalización
            elif valores_son_diferentes(valor_antiguo, valor_nuevo):
                self.guardar_cambio_historial(rma_id, columna_db.title().replace('_', ' '), str(valor_antiguo), str(valor_nuevo))
                campos_a_actualizar.append(f"{columna_db} = ?")
                valores_a_actualizar.append(valor_nuevo)
        
        columna_estado_ya_anadida = 'estado' in [c.split('=')[0].strip() for c in campos_a_actualizar]

        if not columna_estado_ya_anadida:
            # Si no se incluyó en el bucle anterior (lo que indica que el estado no cambió),
            # lo incluimos ahora, usando el nuevo valor calculado.
            campos_a_actualizar.append("estado = ?")
            valores_a_actualizar.append(estado_nuevo)

        # 4. Actualizar rma_maestro si hay cambios
        updated_any = False
        if campos_a_actualizar:
            valores_a_actualizar.append(rma_id)
            set_clause = ", ".join(campos_a_actualizar)
            
            try:
                cursor.execute(f"UPDATE rma_maestro SET {set_clause} WHERE id = ?", tuple(valores_a_actualizar))
                updated_any = True
            except sqlite3.Error as e:
                print(f"Error al actualizar maestro: {e}")
                conn.rollback()
                conn.close()
                return

        # 5. Actualizar rma_detalles (Borrar antiguos e Insertar nuevos) - OPTIMIZADO
        try:
            # Primero obtener los artículos antiguos para comparar
            cursor.execute("SELECT * FROM rma_detalles WHERE rma_id = ? ORDER BY referencia_articulo", (rma_id,))
            articulos_antiguos = cursor.fetchall()
            
            # Comparar si los artículos realmente cambiaron
            articulos_cambiaron = False
            if len(articulos_antiguos) != len(self.articulos_data):
                articulos_cambiaron = True
            else:
                # Comparar cada artículo (simplificado: comparar campos clave)
                for i, art_antiguo in enumerate(articulos_antiguos):
                    if i < len(self.articulos_data):
                        art_nuevo = self.articulos_data[i]
                        # Comparar campos principales (ajustar según estructura de tu tabla)
                        # Asumiendo que la posición 1 es referencia, 2 es cantidad, etc.
                        if (str(art_antiguo[1] if len(art_antiguo) > 1 else '') != str(art_nuevo.get('referencia_articulo', '')) or
                            str(art_antiguo[2] if len(art_antiguo) > 2 else '') != str(art_nuevo.get('cantidad_entregada', '')) or
                            str(art_antiguo[3] if len(art_antiguo) > 3 else '') != str(art_nuevo.get('precio_unitario', ''))):
                            articulos_cambiaron = True
                            break
            
            # Solo actualizar y registrar en historial si realmente cambiaron
            if articulos_cambiaron:
                # Borrar todos los detalles existentes
                cursor.execute("DELETE FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                
                # Insertar los detalles de la lista temporal del formulario - BATCH
                if self.articulos_data:
                    primer_articulo = self.articulos_data[0].copy()
                    primer_articulo['rma_id'] = rma_id

                    columnas_detalle = ', '.join(primer_articulo.keys())
                    placeholders_detalle = ', '.join('?' * len(primer_articulo))

                    # Preparar lista de valores para executemany
                    valores_batch = []
                    for articulo in self.articulos_data:
                        articulo_copia = articulo.copy()
                        articulo_copia['rma_id'] = rma_id
                        valores_batch.append(tuple(articulo_copia.values()))

                    # Insertar todos los artículos en batch (mucho más rápido)
                    cursor.executemany(f"""
                        INSERT INTO rma_detalles ({columnas_detalle}) 
                        VALUES ({placeholders_detalle})
                    """, valores_batch)

                    # Solo registrar en historial si hubo cambios
                    self.guardar_cambio_historial(rma_id, "Detalle Artículos", f"{len(articulos_antiguos)} items", f"{len(self.articulos_data)} items")

                    # Recalcular el precio total a partir de los artículos insertados y actualizar rma_maestro
                    try:
                        precio_total_recalc = 0.0
                        for item in self.articulos_data:
                            cantidad = item.get('cantidad_entregada', 0) or 0
                            if isinstance(cantidad, str):
                                cantidad = float(cantidad.replace(',', '.')) if cantidad.strip() else 0.0
                            else:
                                cantidad = float(cantidad)

                            precio = item.get('precio_unitario', 0.0) or 0.0
                            if isinstance(precio, str):
                                precio = float(precio.replace(',', '.')) if precio.strip() else 0.0
                            else:
                                precio = float(precio)

                            precio_total_recalc += cantidad * precio

                        cursor.execute("UPDATE rma_maestro SET precio_total_expediente = ? WHERE id = ?", (precio_total_recalc, rma_id))
                    except Exception as e:
                        print(f"Error al recalcular/guardar precio_total_expediente en actualizar_rma: {e}")

                elif not self.articulos_data and len(articulos_antiguos) > 0: # Si borramos y no insertamos nada
                    # Borrar artículos
                    cursor.execute("DELETE FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                    self.guardar_cambio_historial(rma_id, "Detalle Artículos", f"{len(articulos_antiguos)} items", "0 items - Artículos eliminados")

                updated_any = True

        except sqlite3.Error as e:
            print(f"Error al actualizar detalles: {e}")
            conn.rollback()
            conn.close()
            return
            
        # 6. Commit final e invalidar caché
        conn.commit()
        conn.close()
        
        # Invalidar caché de estados (puede que se haya actualizado el estado)
        invalidate_cache('estados_rma')

        # Mostrar un único mensaje de éxito y mantener la ficha abierta
        if updated_any:
            try:
                messagebox.showinfo("Expediente actualizado", "El expediente se ha actualizado correctamente.")
            except Exception:
                pass
    
    def mostrar_historial(self, parent_frame):
        """Muestra la lista de registros de cambios para el RMA actual."""
        
        # Destruye el 'historial_scroll_frame' anterior y todos sus hijos.
        for widget in parent_frame.winfo_children():
            widget.destroy()
        
        # Marco scrollable para contener la lista de eventos
        historial_scroll_frame = ctk.CTkScrollableFrame(parent_frame, label_text="Detalle de Cambios en el Expediente")
        historial_scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Encabezados
        header_font = ctk.CTkFont(weight="bold")
        historial_scroll_frame.grid_columnconfigure(2, weight=1) # Descripción se expande
        ctk.CTkLabel(historial_scroll_frame, text="FECHA/HORA", font=header_font).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(historial_scroll_frame, text="USUARIO", font=header_font).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(historial_scroll_frame, text="DESCRIPCIÓN DEL CAMBIO", font=header_font).grid(row=0, column=2, padx=5, pady=5, sticky="w")
        
        conn, cursor = self.master.conectar_db()
        if not conn: return
        cursor = conn.cursor()

        try:
            cursor.execute("""
                SELECT fecha_cambio, usuario, descripcion_cambio 
                FROM rma_historial 
                WHERE rma_id = ? 
                ORDER BY id DESC
            """, (self.rma_actual_id,))
            
            registros = cursor.fetchall()
            conn.close()

            if not registros:
                ctk.CTkLabel(historial_scroll_frame, text="No hay registros de historial para este RMA.", text_color="gray").grid(row=1, column=0, columnspan=3, padx=10, pady=20)
                return
            
            # Mostrar los registros
            for i, reg in enumerate(registros):
                fecha, usuario, descripcion = reg
                row = i + 1
                
                ctk.CTkLabel(historial_scroll_frame, text=fecha).grid(row=row, column=0, padx=5, pady=2, sticky="w")
                ctk.CTkLabel(historial_scroll_frame, text=usuario).grid(row=row, column=1, padx=5, pady=2, sticky="w")
                
                # Usamos wrap para que el texto de la descripción no se salga
                ctk.CTkLabel(historial_scroll_frame, text=descripcion, wraplength=500, justify="left").grid(row=row, column=2, padx=5, pady=2, sticky="w")

        except Exception as e:
            print(f"Error al cargar historial: {e}")
            if conn: conn.close()
            ctk.CTkLabel(historial_scroll_frame, text="Error al cargar el historial.", text_color="red").grid(row=1, column=0, columnspan=3, padx=10, pady=20)


    def determinar_estado_rma(self, datos_maestro):
        """Calcula el estado del RMA basándose en las fechas clave."""
        
        # Obtener fechas del diccionario de datos (asegúrate que las claves son minúsculas como en la DB)
        fecha_recepcion = datos_maestro.get('fecha_recepcion')
        fecha_gestion = datos_maestro.get('fecha_gestion')
        fecha_autorizacion = datos_maestro.get('fecha_autorizacion')
        fecha_proceso = datos_maestro.get('fecha_proceso')
        fecha_emision = datos_maestro.get('fecha_emision')
        
        # 6. Estado 'Completado' (Último paso)
        if fecha_gestion:
            return "Completado"
            
        # 5. Estado 'En Trámite' (Cuando se ingresa fecha de proceso)
        elif fecha_proceso:
            return "En Trámite"
            
        # 4. Estado 'Recibido'
        elif fecha_recepcion:
            return "Recibido"
        
        # 3. Estado 'Autorizado' (Debe ir antes que fecha_emision)
        elif fecha_autorizacion:
            return "Autorizado"
            
        # 2. Estado 'Pendiente de Autorizacion' (Si solo existe la emisión)
        elif fecha_emision:
            return "Pendiente de Autorizacion"
            
        # 1. Estado por defecto
        else:
            return "Pendiente de Autorizacion"
    
    def aplicar_filtros_rma(self):
        """Lee los valores de los filtros y recarga la lista."""
        texto_busqueda = self.entry_busqueda.get()
        estado_filtro = self.filtro_estado.get()
        año_filtro = self.filtro_año.get()
        
        self.cargar_lista_rma(texto_busqueda, estado_filtro, año_filtro)
    
    def crear_tabla_adjuntos(self):
        """Crea la tabla rma_adjuntos si no existe."""
        conn, cursor = self.master.conectar_db()
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS rma_adjuntos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    rma_id INTEGER NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    ruta_relativa TEXT NOT NULL,
                    fecha_subida TEXT,
                    usuario_subida TEXT,
                    tipo_almacenamiento TEXT DEFAULT 'local',
                    FOREIGN KEY (rma_id) REFERENCES rma_maestro (id)
                )
            """)
            
            # Verificar si podemos usar la nueva columna
            self._verificar_columna_tipo_almacenamiento(cursor)
            
            conn.commit()
            if getattr(self, '_usar_tipo_almacenamiento', False):
                print("✓ Sistema de adjuntos configurado con esquema nuevo (Dropbox/Local tracking)")
            else:
                print("✓ Sistema de adjuntos configurado con esquema clásico (compatible)")
        except sqlite3.Error as e:
            print(f"Error al crear la tabla 'rma_adjuntos': {e}")
        finally:
            conn.close()

    def _verificar_columna_tipo_almacenamiento(self, cursor):
        """Verifica si la columna tipo_almacenamiento existe y funciona."""
        try:
            # Intentar hacer un SELECT con la columna
            cursor.execute("SELECT tipo_almacenamiento FROM rma_adjuntos LIMIT 1")
            self._usar_tipo_almacenamiento = True
            print("✓ Esquema nuevo detectado - usando columna tipo_almacenamiento")
        except Exception as e:
            error_str = str(e).lower()
            if ("no column named tipo_almacenamiento" in error_str or 
                "table rma_adjuntos has no column named tipo_almacenamiento" in error_str or
                "no such table" in error_str or
                "no such column: tipo_almacenamiento" in error_str):
                self._usar_tipo_almacenamiento = False
                print("✓ Esquema clásico detectado - usando compatibilidad con BD existente")
            else:
                # Si es otro error, usar esquema compatible por seguridad
                self._usar_tipo_almacenamiento = False
                print(f"✓ Usando esquema compatible por seguridad: {e}")

    def crear_tabla_tareas(self):
        """Crea la tabla 'tareas' si no existe. Asociada a RMA por código o libre."""
        conn, cursor = self.master.conectar_db()
        if not conn:
            return
        try:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS tareas (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    codigo_rma TEXT,
                    titulo TEXT NOT NULL,
                    descripcion TEXT,
                    fecha_vencimiento TEXT,
                    estado TEXT DEFAULT 'Pendiente',
                    creado_por TEXT,
                    creado_en TEXT,
                    notificado INTEGER DEFAULT 0
                )
            ''')
            conn.commit()
            print("Tabla 'tareas' verificada/creada.")
        except sqlite3.Error as e:
            print(f"Error al crear la tabla 'tareas': {e}")
        finally:
            conn.close()
    
    def crear_carpeta_adjuntos_rma(self, codigo_rma):
        """
        Crea la carpeta específica para el RMA en Dropbox o localmente.
        Retorna la ruta (para Dropbox será la ruta remota, para local será la ruta física).
        """
        if usar_dropbox():
            return self._crear_carpeta_dropbox(codigo_rma)
        else:
            return self._crear_carpeta_local(codigo_rma)
    
    def _crear_carpeta_dropbox(self, codigo_rma):
        """Crea una carpeta en Dropbox para el RMA."""
        dbx = get_dropbox_client()
        if not dbx:
            # Fallback a almacenamiento local si Dropbox falla
            print("Dropbox no disponible, usando almacenamiento local")
            return self._crear_carpeta_local(codigo_rma)
        
        # Ruta en Dropbox: /Adjuntos_RMA/RMA25001
        carpeta_rma = f"{DROPBOX_ROOT_FOLDER}/{codigo_rma}"
        carpeta_rma = normalizar_ruta_dropbox(carpeta_rma)
        
        try:
            # Verificar si la carpeta ya existe
            dbx.files_get_metadata(carpeta_rma)
            print(f"Carpeta Dropbox ya existe: {carpeta_rma}")
        except ApiError as e:
            # Verificar si el error es porque la carpeta no existe
            error_details = str(e)
            if "not_found" in error_details.lower() or "path_not_found" in error_details.lower():
                # La carpeta no existe, crearla
                try:
                    dbx.files_create_folder_v2(carpeta_rma)
                    print(f"Carpeta Dropbox creada: {carpeta_rma}")
                except ApiError as create_error:
                    print(f"Error creando carpeta en Dropbox: {create_error}")
                    # Fallback a almacenamiento local
                    return self._crear_carpeta_local(codigo_rma)
            else:
                print(f"Error verificando carpeta Dropbox: {e}")
                return self._crear_carpeta_local(codigo_rma)
        except Exception as e:
            print(f"Error inesperado verificando carpeta Dropbox: {e}")
            return self._crear_carpeta_local(codigo_rma)
        
        return carpeta_rma
    
    def _crear_carpeta_local(self, codigo_rma):
        """Crea una carpeta local para el RMA (implementación original)."""
        # 1. Asegurarse de que la carpeta raíz exista (Adjuntos_RMA)
        if not os.path.exists(ADJUNTOS_ROOT_DIR):
            os.makedirs(ADJUNTOS_ROOT_DIR)
            
        # 2. Crear la carpeta específica del RMA
        ruta_rma = os.path.join(ADJUNTOS_ROOT_DIR, codigo_rma)
        os.makedirs(ruta_rma, exist_ok=True)
        return ruta_rma
    
    def abrir_dialogo_adjunto(self, modo_abrir_carpeta=False):
        """Abre el diálogo de selección de archivo y lo sube al sistema."""
        # 1. Verificar si el RMA ya está guardado (si current_rma_id tiene valor)
        if not self.current_rma_id:
            messagebox.showwarning("Advertencia", "Debe guardar el RMA al menos una vez antes de adjuntar archivos.")
            return
        
        # Obtener el código RMA
        texto_completo = self.lbl_codigo_rma.cget("text") 
        codigo_rma = texto_completo.split(": ")[1] 
        
        # -----------------------------------------------------------------
        # LÓGICA DE ABRIR CARPETA (Modo Informe)
        # -----------------------------------------------------------------
        if modo_abrir_carpeta:
            if usar_dropbox():
                self._abrir_carpeta_dropbox(codigo_rma)
            else:
                self._abrir_carpeta_local(codigo_rma)
            return

        # -----------------------------------------------------------------
        # LÓGICA DE SUBIDA DE ARCHIVO
        # -----------------------------------------------------------------
        # 2. Abrir diálogo para seleccionar archivo(s) - MÚLTIPLE SELECCIÓN
        filepaths = filedialog.askopenfilenames(  # Cambio a askopenfilenames para múltiples
            title="Seleccionar Archivo(s) a Adjuntar - ¡Puedes seleccionar varias imágenes!",
            filetypes=(
                ("Todos los archivos", "*.*"), 
                ("Imágenes", "*.jpg;*.jpeg;*.png;*.bmp;*.gif;*.tiff;*.webp;*.heic"),
                ("Documentos PDF", "*.pdf")
            )
        )
        
        if not filepaths:
            return  # El usuario canceló

        # 3. Procesar cada archivo seleccionado
        total_archivos = len(filepaths)
        archivos_exitosos = 0
        
        # Crear una única ventana de progreso para múltiples archivos
        ventana_progreso_general = None
        if total_archivos > 1:
            ventana_progreso_general = ctk.CTkToplevel(self)
            ventana_progreso_general.title(f"📁 Procesando {total_archivos} archivos")
            ventana_progreso_general.geometry("450x130")
            ventana_progreso_general.transient(self)
            ventana_progreso_general.grab_set()
            
            # Centrar ventana
            ventana_progreso_general.update_idletasks()
            x = (ventana_progreso_general.winfo_screenwidth() // 2) - (450 // 2)
            y = (ventana_progreso_general.winfo_screenheight() // 2) - (130 // 2)
            ventana_progreso_general.geometry(f"450x130+{x}+{y}")
            
            label_archivo_actual = ctk.CTkLabel(ventana_progreso_general, text="", wraplength=420)
            label_archivo_actual.pack(pady=(10, 5))
            
            barra_general = ctk.CTkProgressBar(ventana_progreso_general, width=400)
            barra_general.pack(pady=5)
            barra_general.set(0)
        
        for i, filepath in enumerate(filepaths, 1):
            nombre_original = os.path.basename(filepath)
            
            # Actualizar progreso general si hay múltiples archivos
            if ventana_progreso_general:
                label_archivo_actual.configure(text=f"📁 Procesando {i}/{total_archivos}: {nombre_original}")
                barra_general.set((i-1) / total_archivos)
                ventana_progreso_general.update()
            else:
                # Mostrar progreso en consola para archivo único
                print(f"📁 Procesando archivo: {nombre_original}")
            
            # Subir archivo (Dropbox o local) con compresión automática para imágenes
            if usar_dropbox():
                exito, ruta_relativa = self._subir_archivo_dropbox(filepath, codigo_rma, nombre_original, ventana_progreso_general)
            else:
                exito, ruta_relativa = self._subir_archivo_local(filepath, codigo_rma, nombre_original)
            
            if not exito:
                continue  # Error ya mostrado, continuar con el siguiente archivo
            
            # 4. Insertar registro en la base de datos para este archivo
            self.crear_tabla_adjuntos()
            
            conn, cursor = self.master.conectar_db()
            try:
                if getattr(self, '_usar_tipo_almacenamiento', False):
                    # Usar esquema nuevo con tipo_almacenamiento
                    tipo_almacenamiento = 'dropbox' if usar_dropbox() else 'local'
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento) 
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        os.path.basename(ruta_relativa),  # Usar el nombre del archivo final (podría ser _optimizada.jpg)
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username,
                        tipo_almacenamiento
                    ))
                else:
                    # Usar esquema antiguo sin tipo_almacenamiento
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        os.path.basename(ruta_relativa),  # Usar el nombre del archivo final
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username
                    ))
                
                conn.commit()
                archivos_exitosos += 1
                
            except Exception as e:
                print(f"Error insertando adjunto {nombre_original} en BD: {e}")
                messagebox.showerror("Error de BD", f"No se pudo registrar el adjunto {nombre_original}: {e}")
            finally:
                conn.close()
        
        # Cerrar ventana de progreso general y mostrar resumen
        if ventana_progreso_general:
            barra_general.set(1.0)
            label_archivo_actual.configure(text=f"✅ Completado: {archivos_exitosos}/{total_archivos} archivos procesados")
            ventana_progreso_general.update()
            ventana_progreso_general.after(2000, lambda: ventana_progreso_general.destroy())
        
        # 5. Mostrar mensaje final y recargar adjuntos
        if archivos_exitosos == total_archivos:
            if total_archivos == 1:
                messagebox.showinfo("Éxito", f"Archivo procesado y subido correctamente.")
            else:
                messagebox.showinfo("Éxito", f"¡Todos los archivos procesados correctamente!\n{archivos_exitosos} archivos subidos.")
        elif archivos_exitosos > 0:
            messagebox.showwarning("Parcialmente completado", f"Se procesaron {archivos_exitosos} de {total_archivos} archivos.\nRevisa los errores en la consola.")
        else:
            messagebox.showerror("Error", "No se pudo procesar ningún archivo.")
        
        # Recargar la lista de adjuntos
        try:
            self.cargar_lista_adjuntos(self.current_rma_id)
        except Exception as e:
            print(f"Error recargando adjuntos: {e}")
        
        # 5. Mostrar resultado final
        if total_archivos == 1:
            if archivos_exitosos == 1:
                mensaje = f"✅ Archivo adjuntado correctamente"
                if es_imagen(filepaths[0]):
                    mensaje += " (imagen optimizada automáticamente)"
                elif es_video(filepaths[0]):
                    mensaje += " (video optimizado automáticamente)"
                messagebox.showinfo("Éxito", mensaje)
            # Si falla, el error ya se mostró arriba
        else:
            # Múltiples archivos
            if archivos_exitosos == total_archivos:
                mensaje = f"✅ Todos los archivos ({total_archivos}) adjuntados correctamente"
                imagenes_count = sum(1 for fp in filepaths if es_imagen(fp))
                videos_count = sum(1 for fp in filepaths if es_video(fp))
                if imagenes_count > 0:
                    mensaje += f"\n🖼️ {imagenes_count} imagen(es) optimizada(s) automáticamente"
                if videos_count > 0:
                    mensaje += f"\n🎬 {videos_count} video(s) optimizado(s) automáticamente"
                messagebox.showinfo("Éxito", mensaje)
            elif archivos_exitosos > 0:
                messagebox.showwarning("Parcialmente exitoso", 
                    f"Se adjuntaron {archivos_exitosos} de {total_archivos} archivos.\n"
                    f"Revisa los mensajes de error anteriores.")
            else:
                messagebox.showerror("Error", "No se pudo adjuntar ningún archivo.")
        
        # 6. Recargar lista de adjuntos si hubo éxitos
        if archivos_exitosos > 0:
            self.cargar_lista_adjuntos(self.current_rma_id)
    
    def _abrir_carpeta_dropbox(self, codigo_rma):
        """Maneja la apertura de carpeta en modo Dropbox."""
        # Para Dropbox, podemos abrir el URL web o crear una carpeta local temporal
        messagebox.showinfo("Dropbox", 
            f"Los adjuntos están almacenados en Dropbox.\n"
            f"Carpeta: {DROPBOX_ROOT_FOLDER}/{codigo_rma}\n\n"
            f"Para acceder, ve a tu Dropbox web o aplicación.")
    
    def _abrir_carpeta_local(self, codigo_rma):
        """Maneja la apertura de carpeta en modo local (implementación original)."""
        ruta_destino_base = self.crear_carpeta_adjuntos_rma(codigo_rma)
        
        try:
            if os.name == 'nt':
                os.startfile(ruta_destino_base)
            elif sys.platform == 'darwin':
                subprocess.Popen(['open', ruta_destino_base])
            else:
                subprocess.Popen(['xdg-open', ruta_destino_base])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta:\n{ruta_destino_base}\nError: {e}")
    
    def _subir_archivo_dropbox(self, filepath, codigo_rma, nombre_archivo, ventana_progreso_externa=None):
        """
        Sube un archivo a Dropbox con compresión inteligente para imágenes.
        Retorna: (éxito: bool, ruta_relativa: str)
        """
        dbx = get_dropbox_client()
        if not dbx:
            messagebox.showerror("Error", "No se puede conectar con Dropbox. Usando almacenamiento local.")
            return self._subir_archivo_local(filepath, codigo_rma, nombre_archivo)
        
        archivo_a_subir = filepath
        archivo_temporal = None
        nombre_archivo_final = nombre_archivo
        
        # ===== COMPRESIÓN DE IMÁGENES =====
        if es_imagen(filepath):
            try:
                ventana_progreso = None
                label_progreso = None
                barra_progreso = None
                
                # Solo crear ventana de progreso si no hay una externa (archivo único)
                if not ventana_progreso_externa:
                    ventana_progreso = ctk.CTkToplevel(self)
                    ventana_progreso.title("🖼️ Optimizando imagen")
                    ventana_progreso.geometry("400x120")
                    ventana_progreso.transient(self)
                    ventana_progreso.grab_set()
                    
                    # Centrar ventana
                    ventana_progreso.update_idletasks()
                    x = (ventana_progreso.winfo_screenwidth() // 2) - (400 // 2)
                    y = (ventana_progreso.winfo_screenheight() // 2) - (120 // 2)
                    ventana_progreso.geometry(f"400x120+{x}+{y}")
                    
                    label_progreso = ctk.CTkLabel(ventana_progreso, text="Preparando compresión...", wraplength=380)
                    label_progreso.pack(pady=(20, 10))
                    
                    barra_progreso = ctk.CTkProgressBar(ventana_progreso, width=350)
                    barra_progreso.pack(pady=10)
                    barra_progreso.set(0.1)
                
                # Función callback para actualizar progreso
                def actualizar_progreso(mensaje):
                    if ventana_progreso:
                        # Ventana individual
                        label_progreso.configure(text=mensaje)
                        ventana_progreso.update()
                        if barra_progreso.get() < 0.9:
                            barra_progreso.set(barra_progreso.get() + 0.15)
                    else:
                        # Solo log para ventana externa
                        print(f"  🎨 {mensaje}")
                
                if ventana_progreso:
                    ventana_progreso.update()
                
                # Comprimir imagen
                resultado = comprimir_imagen_inteligente(filepath, callback_progreso=actualizar_progreso)
                archivo_comprimido, tamaño_original, tamaño_final = resultado
                
                if archivo_comprimido and archivo_comprimido != filepath:
                    archivo_a_subir = archivo_comprimido
                    archivo_temporal = archivo_comprimido
                    
                    # Cambiar extensión a .jpg si se comprimió
                    nombre_base = os.path.splitext(nombre_archivo)[0]
                    nombre_archivo_final = f"{nombre_base}_optimizada.jpg"
                    
                    # Mostrar resultado final
                    if ventana_progreso:
                        barra_progreso.set(1.0)
                        if tamaño_original > tamaño_final:
                            actualizar_progreso(f"✅ ¡Imagen optimizada! {tamaño_original:.1f}MB → {tamaño_final:.1f}MB")
                        else:
                            actualizar_progreso(f"✅ Imagen procesada ({tamaño_original:.1f}MB)")
                
                # Cerrar ventana individual después de un tiempo
                if ventana_progreso:
                    ventana_progreso.after(1500, lambda: ventana_progreso.destroy())
                
            except Exception as e:
                # Si falla la compresión, usar archivo original
                print(f"Error en compresión de imagen: {e}")
                if ventana_progreso:
                    ventana_progreso.destroy()
        
        # ===== COMPRESIÓN DE VIDEOS =====
        if es_video(filepath):
            try:
                ventana_progreso = None
                label_progreso = None
                barra_progreso = None
                
                # Solo crear ventana de progreso si no hay una externa (archivo único)
                if not ventana_progreso_externa:
                    ventana_progreso = ctk.CTkToplevel(self)
                    ventana_progreso.title("🎬 Optimizando video")
                    ventana_progreso.geometry("450x120")
                    ventana_progreso.transient(self)
                    ventana_progreso.grab_set()
                    
                    # Centrar ventana
                    ventana_progreso.update_idletasks()
                    x = (ventana_progreso.winfo_screenwidth() // 2) - (450 // 2)
                    y = (ventana_progreso.winfo_screenheight() // 2) - (120 // 2)
                    ventana_progreso.geometry(f"450x120+{x}+{y}")
                    
                    label_progreso = ctk.CTkLabel(ventana_progreso, text="Preparando compresión...", wraplength=430)
                    label_progreso.pack(pady=(20, 10))
                    
                    barra_progreso = ctk.CTkProgressBar(ventana_progreso, width=400)
                    barra_progreso.pack(pady=10)
                    barra_progreso.set(0.1)
                
                # Función callback para actualizar progreso
                def actualizar_progreso_video(mensaje):
                    if ventana_progreso:
                        # Ventana individual
                        label_progreso.configure(text=mensaje)
                        ventana_progreso.update()
                        # Incrementar barra basándose en el mensaje
                        if "%" in mensaje:
                            try:
                                porcentaje = int(mensaje.split("%")[0].split()[-1])
                                barra_progreso.set(porcentaje / 100)
                            except:
                                pass
                        elif barra_progreso.get() < 0.9:
                            barra_progreso.set(barra_progreso.get() + 0.1)
                    else:
                        # Solo log para ventana externa
                        print(f"  🎬 {mensaje}")
                
                if ventana_progreso:
                    ventana_progreso.update()
                
                # Comprimir video
                resultado = comprimir_video_inteligente(filepath, callback_progreso=actualizar_progreso_video)
                
                # Verificar que el resultado sea válido
                if resultado and len(resultado) == 3:
                    archivo_comprimido, tamaño_original, tamaño_final = resultado
                    
                    if archivo_comprimido and archivo_comprimido != filepath:
                        archivo_a_subir = archivo_comprimido
                        archivo_temporal = archivo_comprimido
                        
                        # Cambiar extensión a .mp4 si se comprimió
                        nombre_base = os.path.splitext(nombre_archivo)[0]
                        nombre_archivo_final = f"{nombre_base}_optimizado.mp4"
                        
                        # Mostrar resultado final
                        if ventana_progreso:
                            barra_progreso.set(1.0)
                            if tamaño_original > tamaño_final:
                                actualizar_progreso_video(f"✅ ¡Video optimizado! {tamaño_original:.1f}MB → {tamaño_final:.1f}MB")
                            else:
                                actualizar_progreso_video(f"✅ Video procesado ({tamaño_original:.1f}MB)")
                
                # Cerrar ventana individual después de un tiempo
                if ventana_progreso:
                    ventana_progreso.after(2000, lambda: ventana_progreso.destroy())
                
            except Exception as e:
                # Si falla la compresión, usar archivo original
                print(f"Error en compresión de video: {e}")
                if 'ventana_progreso' in locals() and ventana_progreso:
                    try:
                        ventana_progreso.destroy()
                    except:
                        pass
        
        # ===== SUBIDA A DROPBOX =====
        # Crear la carpeta si no existe
        ruta_carpeta = self.crear_carpeta_adjuntos_rma(codigo_rma)
        
        # Ruta completa en Dropbox
        ruta_dropbox = f"{ruta_carpeta}/{nombre_archivo_final}"
        ruta_dropbox = normalizar_ruta_dropbox(ruta_dropbox)
        
        try:
            # Leer el archivo (original o comprimido) y subirlo
            with open(archivo_a_subir, 'rb') as f:
                dbx.files_upload(
                    f.read(), 
                    ruta_dropbox, 
                    mode=dropbox.files.WriteMode('overwrite')
                )
            
            # Limpiar archivo temporal si existe
            if archivo_temporal:
                try:
                    os.unlink(archivo_temporal)
                except:
                    pass
            
            # La ruta relativa para BD será: RMA25001/archivo.pdf
            ruta_relativa = f"{codigo_rma}/{nombre_archivo_final}"
            return True, ruta_relativa
            
        except Exception as e:
            # Limpiar archivo temporal en caso de error
            if archivo_temporal:
                try:
                    os.unlink(archivo_temporal)
                except:
                    pass
            messagebox.showerror("Error Dropbox", f"No se pudo subir el archivo a Dropbox: {e}")
            return False, ""
    
    def _subir_archivo_local(self, filepath, codigo_rma, nombre_archivo):
        """
        Sube un archivo al almacenamiento local con compresión inteligente para imágenes y videos.
        Retorna: (éxito: bool, ruta_relativa: str)
        """
        archivo_a_subir = filepath
        archivo_temporal = None
        nombre_archivo_final = nombre_archivo
        
        # ===== COMPRESIÓN DE IMÁGENES =====
        if es_imagen(filepath):
            try:
                ventana_progreso = ctk.CTkToplevel(self)
                ventana_progreso.title("🖼️ Optimizando imagen")
                ventana_progreso.geometry("400x120")
                ventana_progreso.transient(self)
                ventana_progreso.grab_set()
                
                # Centrar ventana
                ventana_progreso.update_idletasks()
                x = (ventana_progreso.winfo_screenwidth() // 2) - (400 // 2)
                y = (ventana_progreso.winfo_screenheight() // 2) - (120 // 2)
                ventana_progreso.geometry(f"400x120+{x}+{y}")
                
                label_progreso = ctk.CTkLabel(ventana_progreso, text="Preparando compresión...", wraplength=380)
                label_progreso.pack(pady=(20, 10))
                
                barra_progreso = ctk.CTkProgressBar(ventana_progreso, width=350)
                barra_progreso.pack(pady=10)
                barra_progreso.set(0.1)
                
                # Función callback para actualizar progreso
                def actualizar_progreso(mensaje):
                    label_progreso.configure(text=mensaje)
                    ventana_progreso.update()
                    if barra_progreso.get() < 0.9:
                        barra_progreso.set(barra_progreso.get() + 0.15)
                
                ventana_progreso.update()
                
                # Comprimir imagen
                resultado = comprimir_imagen_inteligente(filepath, callback_progreso=actualizar_progreso)
                archivo_comprimido, tamaño_original, tamaño_final = resultado
                
                if archivo_comprimido and archivo_comprimido != filepath:
                    archivo_a_subir = archivo_comprimido
                    archivo_temporal = archivo_comprimido
                    
                    # Cambiar extensión a .jpg si se comprimió
                    nombre_base = os.path.splitext(nombre_archivo)[0]
                    nombre_archivo_final = f"{nombre_base}_optimizada.jpg"
                    
                    # Mostrar resultado final
                    barra_progreso.set(1.0)
                    if tamaño_original > tamaño_final:
                        actualizar_progreso(f"✅ ¡Imagen optimizada! {tamaño_original:.1f}MB → {tamaño_final:.1f}MB")
                    else:
                        actualizar_progreso(f"✅ Imagen procesada ({tamaño_original:.1f}MB)")
                
                # Cerrar ventana después de un tiempo
                ventana_progreso.after(1500, lambda: ventana_progreso.destroy())
                
            except Exception as e:
                print(f"Error en compresión de imagen: {e}")
                if 'ventana_progreso' in locals():
                    ventana_progreso.destroy()
        
        # ===== COMPRESIÓN DE VIDEOS =====
        if es_video(filepath):
            try:
                ventana_progreso = ctk.CTkToplevel(self)
                ventana_progreso.title("🎬 Optimizando video")
                ventana_progreso.geometry("450x120")
                ventana_progreso.transient(self)
                ventana_progreso.grab_set()
                
                # Centrar ventana
                ventana_progreso.update_idletasks()
                x = (ventana_progreso.winfo_screenwidth() // 2) - (450 // 2)
                y = (ventana_progreso.winfo_screenheight() // 2) - (120 // 2)
                ventana_progreso.geometry(f"450x120+{x}+{y}")
                
                label_progreso = ctk.CTkLabel(ventana_progreso, text="Preparando compresión...", wraplength=430)
                label_progreso.pack(pady=(20, 10))
                
                barra_progreso = ctk.CTkProgressBar(ventana_progreso, width=400)
                barra_progreso.pack(pady=10)
                barra_progreso.set(0.1)
                
                # Función callback para actualizar progreso
                def actualizar_progreso_video(mensaje):
                    label_progreso.configure(text=mensaje)
                    ventana_progreso.update()
                    # Incrementar barra basándose en el mensaje
                    if "%" in mensaje:
                        try:
                            porcentaje = int(mensaje.split("%")[0].split()[-1])
                            barra_progreso.set(porcentaje / 100)
                        except:
                            pass
                    elif barra_progreso.get() < 0.9:
                        barra_progreso.set(barra_progreso.get() + 0.1)
                
                ventana_progreso.update()
                
                # Comprimir video
                resultado = comprimir_video_inteligente(filepath, callback_progreso=actualizar_progreso_video)
                
                # Verificar que el resultado sea válido
                if resultado and len(resultado) == 3:
                    archivo_comprimido, tamaño_original, tamaño_final = resultado
                    
                    if archivo_comprimido and archivo_comprimido != filepath:
                        archivo_a_subir = archivo_comprimido
                        archivo_temporal = archivo_comprimido
                        
                        # Cambiar extensión a .mp4 si se comprimió
                        nombre_base = os.path.splitext(nombre_archivo)[0]
                        nombre_archivo_final = f"{nombre_base}_optimizado.mp4"
                        
                        # Mostrar resultado final
                        barra_progreso.set(1.0)
                        if tamaño_original > tamaño_final:
                            actualizar_progreso_video(f"✅ ¡Video optimizado! {tamaño_original:.1f}MB → {tamaño_final:.1f}MB")
                        else:
                            actualizar_progreso_video(f"✅ Video procesado ({tamaño_original:.1f}MB)")
                
                # Cerrar ventana después de un tiempo
                ventana_progreso.after(2000, lambda: ventana_progreso.destroy())
                
            except Exception as e:
                print(f"Error en compresión de video: {e}")
                if 'ventana_progreso' in locals() and ventana_progreso:
                    try:
                        ventana_progreso.destroy()
                    except:
                        pass
        
        # ===== COPIA AL ALMACENAMIENTO LOCAL =====
        try:
            ruta_destino_dir = self.crear_carpeta_adjuntos_rma(codigo_rma)
            ruta_destino_completa = os.path.join(ruta_destino_dir, nombre_archivo_final)
            
            # Copiar archivo (original o comprimido)
            shutil.copy2(archivo_a_subir, ruta_destino_completa)
            
            # Limpiar archivo temporal si existe
            if archivo_temporal:
                try:
                    os.unlink(archivo_temporal)
                except:
                    pass
            
            # Ruta relativa para BD
            ruta_relativa = os.path.join(codigo_rma, nombre_archivo_final)
            return True, ruta_relativa
            
        except Exception as e:
            # Limpiar archivo temporal en caso de error
            if archivo_temporal:
                try:
                    os.unlink(archivo_temporal)
                except:
                    pass
            messagebox.showerror("Error de Copia", f"No se pudo copiar el archivo: {e}")
            return False, ""
    
    def _limpiar_archivo_subido(self, ruta_relativa):
        """Intenta eliminar un archivo subido si falla la inserción en BD."""
        if usar_dropbox():
            dbx = get_dropbox_client()
            if dbx:
                try:
                    ruta_dropbox = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{ruta_relativa}")
                    dbx.files_delete_v2(ruta_dropbox)
                except:
                    pass  # No importa si falla la limpieza
        else:
            try:
                ruta_completa = os.path.join(ADJUNTOS_ROOT_DIR, ruta_relativa)
                if os.path.exists(ruta_completa):
                    os.remove(ruta_completa)
            except:
                pass  # No importa si falla la limpieza
    def cargar_lista_adjuntos(self, rma_id):
        """Consulta y muestra el listado de adjuntos para un RMA específico."""
        
        # Verificar que el frame existe y la aplicación sigue activa
        try:
            if not hasattr(self, 'adjuntos_list_frame') or not self.adjuntos_list_frame.winfo_exists():
                return
        except Exception:
            return
        
        # Limpiar el frame antes de cargar la nueva lista
        try:
            for widget in self.adjuntos_list_frame.winfo_children():
                widget.destroy()
        except Exception as e:
            print(f"Error limpiando widgets: {e}")
            return

        try:
            conn, cursor = self.master.conectar_db()
            cursor.execute("SELECT id, nombre_archivo, ruta_relativa FROM rma_adjuntos WHERE rma_id = ?", (rma_id,))
            adjuntos = cursor.fetchall()
            conn.close()
        except Exception as e:
            print(f"Error cargando adjuntos: {e}")
            return

        if not adjuntos:
            ctk.CTkLabel(self.adjuntos_list_frame, text="No hay archivos adjuntos para este expediente.").pack(pady=10)
            return

        for i, adjunto in enumerate(adjuntos):
            adjunto_id, nombre, ruta = adjunto

            item_frame = ctk.CTkFrame(self.adjuntos_list_frame)
            item_frame.pack(fill='x', padx=5, pady=2)

            # Etiqueta del nombre del archivo
            ctk.CTkLabel(item_frame, text=nombre, width=250, anchor='w').pack(side='left', padx=5)

            # Botón Visualizar (solo lectura)
            ctk.CTkButton(
                item_frame, 
                text="👁️ Ver", 
                width=70, 
                command=lambda r=ruta: self.abrir_adjunto(r)
            ).pack(side='right', padx=2)

            # Botón Editar (descarga, edita y resube)
            if usar_dropbox():  # Solo mostrar editar en modo Dropbox
                ctk.CTkButton(
                    item_frame, 
                    text="� Editar", 
                    width=70,
                    fg_color="#2B7A0B",
                    hover_color="#1F5F08",
                    command=lambda r=ruta, aid=adjunto_id: self.editar_adjunto(r, aid)
                ).pack(side='right', padx=2)

            # Botón Eliminar
            # El comando usa lambda para pasar el ID del adjunto y la ruta
            ctk.CTkButton(
                item_frame, 
                text="🗑️ Eliminar", 
                width=80, 
                fg_color="red", 
                hover_color="darkred",
                command=lambda aid=adjunto_id, r=ruta: self.confirmar_eliminar_adjunto(aid, r)
            ).pack(side='right', padx=5)
            
        # Llamamos a esta función dentro de abrir_dialogo_adjunto() para que se recargue después de subir un archivo.
    def abrir_adjunto(self, ruta_relativa):
        """Abre el archivo adjunto desde Dropbox o almacenamiento local."""
        if usar_dropbox():
            self._abrir_adjunto_dropbox(ruta_relativa)
        else:
            self._abrir_adjunto_local(ruta_relativa)
    
    def _abrir_adjunto_dropbox(self, ruta_relativa):
        """Descarga temporalmente un archivo de Dropbox y lo abre."""
        dbx = get_dropbox_client()
        if not dbx:
            messagebox.showerror("Error", "No se puede conectar con Dropbox.")
            return
        
        # Construir ruta en Dropbox
        ruta_dropbox = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{ruta_relativa}")
        
        try:
            # Crear archivo temporal
            nombre_archivo = os.path.basename(ruta_relativa)
            with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{nombre_archivo}") as temp_file:
                temp_path = temp_file.name
                
                # Descargar archivo de Dropbox
                metadata, response = dbx.files_download(ruta_dropbox)
                temp_file.write(response.content)
            
            # Abrir archivo temporal
            self._abrir_archivo_sistema(temp_path)
            
            # Programar eliminación del archivo temporal después de un tiempo
            # (El usuario tendrá tiempo para abrirlo en su programa)
            def limpiar_temp():
                try:
                    os.unlink(temp_path)
                except:
                    pass
            
            # Limpiar después de 60 segundos (suficiente tiempo para que se abra)
            threading.Timer(60.0, limpiar_temp).start()
            
        except ApiError as e:
            error_details = str(e)
            if "not_found" in error_details.lower() or "path_not_found" in error_details.lower():
                messagebox.showerror("Error", f"Archivo no encontrado en Dropbox: {ruta_relativa}")
            else:
                messagebox.showerror("Error", f"Error descargando de Dropbox: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Error procesando archivo de Dropbox: {e}")
    
    def _abrir_adjunto_local(self, ruta_relativa):
        """Abre un archivo del almacenamiento local (implementación original)."""
        ruta_completa = os.path.join(ADJUNTOS_ROOT_DIR, ruta_relativa)
        
        if not os.path.exists(ruta_completa):
            messagebox.showerror("Error", f"Archivo no encontrado: {ruta_completa}")
            return
        
        self._abrir_archivo_sistema(ruta_completa)
    
    def _abrir_archivo_sistema(self, ruta_archivo):
        """Abre un archivo con el programa predeterminado del sistema."""
        try:
            if sys.platform == "win32":
                os.startfile(ruta_archivo)
            elif sys.platform == "darwin":
                subprocess.call(['open', ruta_archivo])
            else:
                subprocess.call(['xdg-open', ruta_archivo])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el archivo: {e}")
    
    def _abrir_archivo_sistema(self, ruta_archivo):
        """Abre un archivo con el programa predeterminado del sistema."""
        try:
            if sys.platform.startswith('win'):
                os.startfile(ruta_archivo)
            elif sys.platform.startswith('darwin'):  # macOS
                os.system(f'open "{ruta_archivo}"')
            else:  # Linux y otros
                os.system(f'xdg-open "{ruta_archivo}"')
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el archivo: {e}")
    
    def editar_adjunto(self, ruta_relativa, adjunto_id):
        """
        Descarga un archivo de Dropbox, permite editarlo y lo resube automáticamente.
        """
        if not usar_dropbox():
            messagebox.showinfo("Información", "La función de editar solo está disponible con archivos de Dropbox.")
            return
            
        dbx = get_dropbox_client()
        if not dbx:
            messagebox.showerror("Error", "No se puede conectar con Dropbox.")
            return
            
        # Construir ruta en Dropbox
        ruta_dropbox = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{ruta_relativa}")
        nombre_archivo = os.path.basename(ruta_relativa)
        
        try:
            # 1. Crear archivo temporal para edición
            temp_dir = tempfile.mkdtemp(prefix="dropbox_edit_")
            temp_path = os.path.join(temp_dir, nombre_archivo)
            
            # 2. Descargar archivo de Dropbox
            print(f"Descargando {nombre_archivo} para edición...")
            metadata, response = dbx.files_download(ruta_dropbox)
            with open(temp_path, 'wb') as temp_file:
                temp_file.write(response.content)
            
            # 3. Mostrar diálogo informativo
            respuesta = messagebox.askyesno(
                "Editar Archivo",
                f"Se va a abrir '{nombre_archivo}' para edición.\n\n"
                f"IMPORTANTE:\n"
                f"• El archivo se descargará temporalmente\n"
                f"• Podrás editarlo con el programa predeterminado\n"
                f"• Cuando GUARDES y CIERRES el programa, se resubirá automáticamente\n"
                f"• Los cambios se sincronizarán con Dropbox\n\n"
                f"¿Continuar?"
            )
            
            if not respuesta:
                # Limpiar archivo temporal si el usuario cancela
                try:
                    os.remove(temp_path)
                    os.rmdir(temp_dir)
                except:
                    pass
                return
            
            # 4. Obtener tiempo de modificación inicial
            tiempo_inicial = os.path.getmtime(temp_path)
            
            # 5. Abrir archivo para edición
            self._abrir_archivo_sistema(temp_path)
            
            # 6. Crear diálogo de seguimiento
            self._crear_dialogo_seguimiento_edicion(temp_path, ruta_dropbox, tiempo_inicial, temp_dir, nombre_archivo)
            
        except ApiError as e:
            error_details = str(e)
            if "not_found" in error_details.lower():
                messagebox.showerror("Error", f"Archivo no encontrado en Dropbox: {ruta_relativa}")
            else:
                messagebox.showerror("Error", f"Error descargando de Dropbox: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Error procesando archivo para edición: {e}")

    def _crear_dialogo_seguimiento_edicion(self, temp_path, ruta_dropbox, tiempo_inicial, temp_dir, nombre_archivo):
        """Crea un diálogo para hacer seguimiento del proceso de edición."""
        
        # Crear ventana de seguimiento
        dialogo = Toplevel(self)
        dialogo.title("Editando archivo...")
        dialogo.geometry("500x300")
        dialogo.resizable(False, False)
        dialogo.transient(self)
        dialogo.grab_set()
        
        # Centrar en pantalla
        dialogo.update_idletasks()
        x = (dialogo.winfo_screenwidth() // 2) - (500 // 2)
        y = (dialogo.winfo_screenheight() // 2) - (300 // 2)
        dialogo.geometry(f"500x300+{x}+{y}")
        
        # Frame principal
        main_frame = ctk.CTkFrame(dialogo)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        titulo = ctk.CTkLabel(main_frame, text=f"📝 Editando: {nombre_archivo}", 
                             font=ctk.CTkFont(size=16, weight="bold"))
        titulo.pack(pady=(10, 20))
        
        # Estado
        self.estado_label = ctk.CTkLabel(main_frame, 
                                        text="🟡 Archivo abierto para edición...\nGuarda los cambios y cierra el programa cuando termines.",
                                        font=ctk.CTkFont(size=12))
        self.estado_label.pack(pady=10)
        
        # Botones
        botones_frame = ctk.CTkFrame(main_frame)
        botones_frame.pack(pady=20, fill="x")
        
        # Botón para verificar cambios manualmente
        btn_verificar = ctk.CTkButton(botones_frame, text="🔄 Verificar cambios",
                                     command=lambda: self._verificar_cambios_manual(temp_path, tiempo_inicial))
        btn_verificar.pack(side="left", padx=10, pady=10)
        
        # Botón para subir cambios
        self.btn_subir = ctk.CTkButton(botones_frame, text="⬆️ Subir cambios", 
                                      state="disabled",
                                      command=lambda: self._subir_cambios_editados(temp_path, ruta_dropbox, temp_dir, dialogo))
        self.btn_subir.pack(side="left", padx=10, pady=10)
        
        # Botón cancelar
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                    fg_color="#D32F2F", hover_color="#B71C1C",
                                    command=lambda: self._cancelar_edicion(temp_path, temp_dir, dialogo))
        btn_cancelar.pack(side="right", padx=10, pady=10)

        # Variables de estado
        dialogo.tiempo_inicial = tiempo_inicial
        dialogo.temp_path = temp_path
        dialogo.cambios_detectados = False
        
        # Iniciar verificación automática cada 3 segundos
        self._verificar_cambios_automatico(dialogo, temp_path, tiempo_inicial)

    def _verificar_cambios_automatico(self, dialogo, temp_path, tiempo_inicial):
        """Verifica automáticamente si el archivo ha sido modificado."""
        try:
            if not os.path.exists(temp_path) or not dialogo.winfo_exists():
                return
                
            tiempo_actual = os.path.getmtime(temp_path)
            
            if tiempo_actual > tiempo_inicial and not dialogo.cambios_detectados:
                # ¡Cambios detectados!
                dialogo.cambios_detectados = True
                self.estado_label.configure(
                    text="✅ ¡Cambios detectados!\nPuedes subir los cambios a Dropbox ahora.",
                    text_color="green"
                )
                self.btn_subir.configure(state="normal", fg_color="#2E7D32", hover_color="#1B5E20")
                dialogo.tiempo_inicial = tiempo_actual  # Actualizar para futuras verificaciones
            
            # Programar próxima verificación
            dialogo.after(3000, lambda: self._verificar_cambios_automatico(dialogo, temp_path, tiempo_inicial))
            
        except Exception as e:
            print(f"Error verificando cambios: {e}")
            
    def _verificar_cambios_manual(self, temp_path, tiempo_inicial):
        """Verificación manual de cambios."""
        try:
            if not os.path.exists(temp_path):
                self.estado_label.configure(text="❌ Error: Archivo temporal no encontrado", text_color="red")
                return
                
            tiempo_actual = os.path.getmtime(temp_path)
            
            if tiempo_actual > tiempo_inicial:
                self.estado_label.configure(
                    text="✅ ¡Cambios detectados!\nPuedes subir los cambios a Dropbox.",
                    text_color="green"
                )
                self.btn_subir.configure(state="normal", fg_color="#2E7D32", hover_color="#1B5E20")
            else:
                self.estado_label.configure(
                    text="ℹ️ No se detectaron cambios aún.\nGuarda el archivo en tu programa de edición.",
                    text_color="blue"
                )
        except Exception as e:
            self.estado_label.configure(text=f"❌ Error verificando cambios: {e}", text_color="red")

    def _subir_cambios_editados(self, temp_path, ruta_dropbox, temp_dir, dialogo):
        """Sube los cambios editados de vuelta a Dropbox."""
        try:
            if not os.path.exists(temp_path):
                messagebox.showerror("Error", "Archivo temporal no encontrado.")
                return
                
            dbx = get_dropbox_client()
            if not dbx:
                messagebox.showerror("Error", "No se puede conectar con Dropbox.")
                return
            
            # Leer archivo modificado
            with open(temp_path, 'rb') as archivo:
                contenido = archivo.read()
            
            # Subir a Dropbox (sobrescribir)
            dbx.files_upload(contenido, ruta_dropbox, mode=dropbox.files.WriteMode('overwrite'))
            
            # Limpiar archivos temporales
            try:
                os.remove(temp_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            # Cerrar diálogo y mostrar éxito
            dialogo.destroy()
            messagebox.showinfo("Éxito", "¡Archivo editado y sincronizado con Dropbox correctamente!")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error subiendo cambios a Dropbox: {e}")

    def _cancelar_edicion(self, temp_path, temp_dir, dialogo):
        """Cancela la edición y limpia archivos temporales."""
        respuesta = messagebox.askyesno(
            "Cancelar edición", 
            "¿Estás seguro de que quieres cancelar?\nSe perderán todos los cambios no subidos."
        )
        
        if respuesta:
            # Limpiar archivos temporales
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                os.rmdir(temp_dir)
            except Exception as e:
                print(f"Error limpiando archivos temporales: {e}")
            
            # Cerrar diálogo
            dialogo.destroy()

    def confirmar_eliminar_adjunto(self, adjunto_id, ruta_relativa):
        """Pide confirmación antes de eliminar el registro y el archivo."""
        try:
            # Verificar que la aplicación sigue activa
            if not hasattr(self, 'master') or not self.master.winfo_exists():
                return
            
            if messagebox.askyesno("Confirmar Eliminación", 
                                 "¿Está seguro de que desea eliminar este adjunto? Esta acción es irreversible y también eliminará el archivo del disco."):
                self.eliminar_adjunto(adjunto_id, ruta_relativa)
        except Exception as e:
            print(f"Error en confirmación de eliminación: {e}")

    def eliminar_adjunto(self, adjunto_id, ruta_relativa):
        """Elimina el registro de la base de datos y el archivo físico."""
        try:
            conn, cursor = self.master.conectar_db()
        except Exception as e:
            messagebox.showerror("Error", f"Error conectando a la base de datos: {e}")
            return
        
        try:
            # 1. Eliminar archivo físico primero
            if usar_dropbox():
                exito_archivo = self._eliminar_archivo_dropbox(ruta_relativa)
            else:
                exito_archivo = self._eliminar_archivo_local(ruta_relativa)
            
            # 2. Eliminar registro de la BD (incluso si el archivo falló)
            cursor.execute("DELETE FROM rma_adjuntos WHERE id = ?", (adjunto_id,))
            conn.commit()
            
            if exito_archivo:
                messagebox.showinfo("Éxito", "Adjunto eliminado correctamente.")
            else:
                messagebox.showwarning("Parcial", "Registro eliminado de la base de datos, pero hubo problemas eliminando el archivo.")
            
            # Recargar el listado solo si la aplicación sigue activa
            try:
                if hasattr(self, 'adjuntos_list_frame') and self.adjuntos_list_frame.winfo_exists():
                    self.cargar_lista_adjuntos(self.current_rma_id)
            except Exception as e:
                print(f"No se pudo recargar lista de adjuntos: {e}")
            
        except Exception as e:
            # Manejar rollback de forma compatible con diferentes tipos de BD
            try:
                if hasattr(conn, 'rollback'):
                    conn.rollback()
            except Exception:
                pass  # Ignorar errores de rollback en Turso
            messagebox.showerror("Error", f"Error al eliminar el adjunto: {e}")
        finally:
            try:
                conn.close()
            except Exception:
                pass
    
    def _eliminar_archivo_dropbox(self, ruta_relativa):
        """
        Elimina un archivo de Dropbox.
        Retorna True si fue exitoso, False si hubo error.
        """
        dbx = get_dropbox_client()
        if not dbx:
            print("No se puede conectar con Dropbox para eliminar archivo")
            return False
        
        ruta_dropbox = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{ruta_relativa}")
        
        try:
            dbx.files_delete_v2(ruta_dropbox)
            return True
        except ApiError as e:
            error_details = str(e)
            if "not_found" in error_details.lower() or "path_not_found" in error_details.lower():
                print(f"Archivo no encontrado en Dropbox (ya eliminado?): {ruta_dropbox}")
                return True  # Considerarlo exitoso si ya no existe
            else:
                print(f"Error eliminando archivo de Dropbox: {e}")
                return False
        except Exception as e:
            print(f"Error eliminando archivo de Dropbox: {e}")
            return False
    
    def _eliminar_archivo_local(self, ruta_relativa):
        """
        Elimina un archivo del almacenamiento local.
        Retorna True si fue exitoso, False si hubo error.
        """
        ruta_completa = os.path.join(ADJUNTOS_ROOT_DIR, ruta_relativa)
        
        try:
            if os.path.exists(ruta_completa):
                os.remove(ruta_completa)
                return True
            else:
                print(f"Archivo local no encontrado (ya eliminado?): {ruta_completa}")
                return True  # Considerarlo exitoso si ya no existe
        except Exception as e:
            print(f"Error eliminando archivo local: {e}")
            return False
    
    # Dentro de la clase VentanaPrincipal

    # Dentro de la clase VentanaPrincipal

    def generar_informe_dinamico(self):
        """
        Genera un informe dinámico usando python-docx, lo guarda en Dropbox 
        y lo registra en la base de datos.
        """
        # 1. Validaciones y Obtención de Datos
        if not self.current_rma_id:
            messagebox.showerror("Error", "Debe cargar un RMA guardado para generar el informe.")
            return

        # Asumimos que self.datos_rma_maestro contiene los datos del RMA cargado
        # Esta variable debe llenarse cuando llamas a self.cargar_datos_rma(rma_id)
        datos = self.datos_rma_maestro 
        
        # 🚨 Verificación: Asegúrate de que los datos clave existen
        codigo_rma = datos.get('codigo_rma')
        nombre_cliente = datos.get('cliente')
        
        if not codigo_rma or not nombre_cliente:
             messagebox.showerror("Error", "Los datos del RMA no están cargados. Intente recargar el expediente.")
             return
        
        # 1.5. Obtener datos de artículos desde rma_detalles
        conn, cursor = self.master.conectar_db()
        if not conn:
            messagebox.showerror("Error", "No se pudo conectar a la base de datos.")
            return
        
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT referencia_articulo, cantidad_entregada 
                FROM rma_detalles 
                WHERE rma_id = ?
                ORDER BY id
            """, (self.current_rma_id,))
            articulos_data = cursor.fetchall()
            conn.close()
        except Exception as e:
            conn.close()
            messagebox.showerror("Error", f"Error al obtener datos de artículos: {e}")
            return

        # 2. Rutas - Plantilla sigue siendo local
        plantilla_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas", "Plantilla_RMA.docx")
        
        # Nombre del archivo final: Ej. RMA2024-001_Informe_20240920.docx
        fecha_str = datetime.datetime.now().strftime("%Y%m%d")
        nombre_archivo_final = f"{codigo_rma}_Informe_{fecha_str}.docx"

        try:
            # 3. Cargar la plantilla y preparar datos para mapeo
            document = docx.Document(plantilla_path)
            
            # 3.1. Preparar lista de artículos para los marcadores en formato estructurado
            if articulos_data:
                # Crear formato de lista estructurada para artículos
                lista_articulos = []
                for i, (ref_articulo, cantidad) in enumerate(articulos_data, 1):
                    ref_str = str(ref_articulo) if ref_articulo else 'N/A'
                    cant_str = str(cantidad) if cantidad else '0'
                    lista_articulos.append(f"{i}. {ref_str} - Cantidad: {cant_str}")
                
                # Unir con saltos de línea para formato de lista
                articulos_formateados = '\n'.join(lista_articulos)
                
                # También crear versiones simples por compatibilidad
                referencias_articulos = [str(ref) if ref else 'N/A' for ref, _ in articulos_data]
                cantidades_articulos = [str(cant) if cant else '0' for _, cant in articulos_data]
            else:
                # Valores por defecto si no hay artículos
                articulos_formateados = "No se han registrado artículos para este RMA."
                referencias_articulos = ['N/A']
                cantidades_articulos = ['N/A']
            
            # 3.2. Mapeo expandido: [Marcador en Word]: [Valor a insertar]
            mapeo = {
                # Campos existentes
                '[[CODIGO_RMA]]': codigo_rma,
                '[[CLIENTE]]': nombre_cliente,
                '[[FECHA_EMISION]]': datos.get('fecha_emision', 'N/A'),
                '[[FECHA_RECEPCION]]': datos.get('fecha_recepcion', 'N/A'),
                '[[ESTADO_ACTUAL]]': datos.get('estado', 'N/A'),
                '[[USUARIO_CREADOR]]': datos.get('creado_por', self.username),
                
                # Campos nuevos de rma_maestro
                '[[NUMERO_DOC]]': datos.get('numero_documento_cliente', 'N/A'),
                '[[MOTIVO]]': datos.get('motivo', 'N/A'),
                '[[NUMERO_ALBARAN]]': datos.get('numero_albaran', 'N/A'),
                
                # Campos de artículos mejorados
                '[[LISTA_ARTICULOS]]': articulos_formateados,  # Lista estructurada completa
                '[[REF_ARTICULO]]': ', '.join(referencias_articulos[:3]),  # Mantener compatibilidad
                '[[CANTIDAD]]': ', '.join(cantidades_articulos[:3]),  # Mantener compatibilidad
                
                # Campos adicionales para flexibilidad
                '[[TOTAL_ARTICULOS]]': str(len(articulos_data)) if articulos_data else '0'
            }
            
            # 4. Función auxiliar para reemplazar texto preservando formato
            def reemplazar_texto_preservando_formato(document, mapeo):
                """Reemplaza texto manteniendo el formato original de la plantilla (fuentes, tamaños, etc.)"""
                for paragraph in document.paragraphs:
                    for run in paragraph.runs:
                        for clave, valor in mapeo.items():
                            if clave in run.text:
                                # Reemplazar manteniendo el formato del run original
                                run.text = run.text.replace(clave, str(valor))
                
                # También procesar tablas si las hay
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    for clave, valor in mapeo.items():
                                        if clave in run.text:
                                            run.text = run.text.replace(clave, str(valor))
            
            # 4.1. Aplicar reemplazos preservando formato
            reemplazar_texto_preservando_formato(document, mapeo)
            
            # 5. Guardar temporalmente para subirlo a Dropbox
            temp_dir = tempfile.mkdtemp(prefix="informe_rma_")
            temp_file_path = os.path.join(temp_dir, nombre_archivo_final)
            document.save(temp_file_path)
            
            # 6. Decidir dónde guardar (Dropbox o local)
            if usar_dropbox():
                # Subir a Dropbox
                exito, ruta_relativa = self._subir_archivo_dropbox(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'dropbox'
                ubicacion_desc = "Dropbox"
            else:
                # Guardar localmente (fallback)
                exito, ruta_relativa = self._subir_archivo_local(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'local'
                ubicacion_desc = "local"
            
            # 7. Limpiar archivo temporal
            try:
                os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            if not exito:
                messagebox.showerror("Error", f"No se pudo guardar el informe en {ubicacion_desc}.")
                return
            
            # 8. Registrar en la Base de Datos
            conn, cursor = self.master.conectar_db()
            try:
                # Verificar esquema de BD antes de insertar
                self._verificar_columna_tipo_almacenamiento(cursor)
                
                # Preparar inserción con o sin tipo_almacenamiento según el esquema
                if getattr(self, '_usar_tipo_almacenamiento', False):
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento) 
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username,
                        tipo_almacenamiento
                    ))
                else:
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username
                    ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().isoformat(),
                    f"Generado documento de Informe: {nombre_archivo_final} ({'☁️ Dropbox' if usar_dropbox() else '💾 Local'})", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                
                try:
                    # Actualizar historial si está visible
                    if hasattr(self, 'historial_tab'):
                        self.mostrar_historial(self.historial_tab)
                except AttributeError:
                    # Si la pestaña historial_tab no está definida, ignoramos.
                    pass
                
                # Mensaje personalizado según donde se guardó
                if usar_dropbox():
                    messagebox.showinfo("Éxito", f"✅ Informe '{nombre_archivo_final}' generado y subido a Dropbox correctamente.\n\n📁 Ubicación: {ruta_relativa}")
                else:
                    messagebox.showinfo("Éxito", f"✅ Informe '{nombre_archivo_final}' generado y guardado localmente.")
                
            except Exception as db_e:
                conn.rollback()
                messagebox.showerror("Error DB", f"Informe generado, pero error al registrar en DB.\nError: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el informe dinámico. Asegúrese de que la plantilla existe y python-docx está instalado.\nError: {e}")
    
    # Dentro de la clase VentanaPrincipal

    def generar_reposicion_devolucion(self):
        """
        Genera el documento de Reposición/Devolución usando la plantilla
        "Reposicion_RMA.docx", lo guarda en Dropbox y lo registra como adjunto.
        """
        # 1. Validaciones y Obtención de Datos
        if not self.current_rma_id:
            messagebox.showerror("Error", "Debe cargar un RMA guardado para generar el documento de Reposición/Devolución.")
            return

        # Asumimos que self.datos_rma_maestro contiene los datos del RMA cargado
        datos = self.datos_rma_maestro 
        
        # Obtener datos clave para el archivo y la ruta
        codigo_rma = datos.get('codigo_rma')
        nombre_cliente = datos.get('cliente')
        
        if not codigo_rma or not nombre_cliente:
             messagebox.showerror("Error", "No se pudieron cargar los datos clave del RMA. Intente recargar el expediente.")
             return

        # 2. Rutas y Nombres de Archivo - Plantilla sigue siendo local
        nombre_plantilla = "Reposicion_RMA.docx" 
        plantilla_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas", nombre_plantilla)
        
        # Nombre del archivo final: Ej. RMA2024-001_Reposicion_20251016.docx
        fecha_str = datetime.datetime.now().strftime("%Y%m%d")
        nombre_archivo_final = f"{codigo_rma}_Reposicion_{fecha_str}.docx"

        # 3. Verificar la Plantilla
        if not os.path.exists(plantilla_path):
            messagebox.showerror("Error", f"No se encontró la plantilla requerida en:\n{plantilla_path}")
            return
            
        try:
            # 4. Cargar la plantilla y definir mapeo de marcadores
            document = docx.Document(plantilla_path)
            
            # Mapeo: Reutilizamos el mapeo existente
            mapeo = {
                '[[CODIGO_RMA]]': codigo_rma,
                '[[CLIENTE]]': nombre_cliente,
                '[[FECHA_EMISION]]': datos.get('fecha_emision', 'N/A'),
                '[[ESTADO_ACTUAL]]': datos.get('estado', 'N/A'),
                '[[USUARIO_CREADOR]]': datos.get('creado_por', self.username)
            }
            
            # 5. Reemplazar marcadores en párrafos
            for p in document.paragraphs:
                for clave, valor in mapeo.items():
                    valor_a_insertar = str(valor) if valor is not None else "" 
                    if clave in p.text:
                        p.text = p.text.replace(clave, valor_a_insertar)
            
            # 6. Guardar temporalmente para subirlo a Dropbox
            temp_dir = tempfile.mkdtemp(prefix="reposicion_rma_")
            temp_file_path = os.path.join(temp_dir, nombre_archivo_final)
            document.save(temp_file_path)
            
            # 7. Decidir dónde guardar (Dropbox o local)
            if usar_dropbox():
                # Subir a Dropbox
                exito, ruta_relativa = self._subir_archivo_dropbox(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'dropbox'
                ubicacion_desc = "Dropbox"
            else:
                # Guardar localmente (fallback)
                exito, ruta_relativa = self._subir_archivo_local(temp_file_path, codigo_rma, nombre_archivo_final)
                tipo_almacenamiento = 'local'
                ubicacion_desc = "local"
            
            # 8. Limpiar archivo temporal
            try:
                os.remove(temp_file_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            if not exito:
                messagebox.showerror("Error", f"No se pudo guardar el documento de reposición en {ubicacion_desc}.")
                return
            
            # 9. Registrar en la Base de Datos
            conn, cursor = self.master.conectar_db()
            try:
                # Verificar esquema de BD antes de insertar
                self._verificar_columna_tipo_almacenamiento(cursor)
                
                # Preparar inserción con o sin tipo_almacenamiento según el esquema
                if getattr(self, '_usar_tipo_almacenamiento', False):
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida, tipo_almacenamiento) 
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username,
                        tipo_almacenamiento
                    ))
                else:
                    cursor.execute("""
                        INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                        VALUES (?, ?, ?, ?, ?)
                    """, (
                        self.current_rma_id, 
                        nombre_archivo_final, 
                        ruta_relativa, 
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                        self.username
                    ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().isoformat(),
                    f"Generado documento de Reposición/Devolución: {nombre_archivo_final} ({'☁️ Dropbox' if usar_dropbox() else '💾 Local'})", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                
                try:
                    # Actualizar historial si está visible
                    if hasattr(self, 'historial_tab'):
                        self.mostrar_historial(self.historial_tab)
                except AttributeError:
                    # Si la pestaña historial_tab no está definida, ignoramos.
                    pass
                
                # Mensaje personalizado según donde se guardó
                if usar_dropbox():
                    messagebox.showinfo("Éxito", f"✅ Documento de Reposición/Devolución '{nombre_archivo_final}' generado y subido a Dropbox correctamente.\n\n📁 Ubicación: {ruta_relativa}")
                else:
                    messagebox.showinfo("Éxito", f"✅ Documento de Reposición/Devolución '{nombre_archivo_final}' generado y guardado localmente.")
                
            except Exception as db_e:
                conn.rollback()
                messagebox.showerror("Error DB", f"Documento generado, pero error al registrar en DB/Historial. Error: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el documento. Asegúrese de que la plantilla existe y es un archivo .docx válido.\nError: {e}")
            conn, cursor = self.master.conectar_db()
            try:
                # Registro en rma_adjuntos
                cursor.execute("""
                    INSERT INTO rma_adjuntos (rma_id, nombre_archivo, ruta_relativa, fecha_subida, usuario_subida) 
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    nombre_archivo_final, 
                    ruta_relativa, 
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                    self.username
                ))
                
                # Registro en el historial
                cursor.execute("""
                    INSERT INTO rma_historial (rma_id, fecha_cambio, descripcion_cambio, usuario)
                    VALUES (?, ?, ?, ?)
                """, (
                    self.current_rma_id, 
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                    f"Generado documento de Reposición/Devolución: {nombre_archivo_final}", 
                    self.username
                ))
                
                conn.commit()
                self.cargar_lista_adjuntos(self.current_rma_id) # Refresca la lista de adjuntos
                try:
                    # Llamamos al método que recarga el contenido de la pestaña
                    self.mostrar_historial(self.historial_tab) 
                except AttributeError:
                    # Si la pestaña historial_tab no está definida (ej. en modo "nuevo"), ignoramos.
                    pass
                
                messagebox.showinfo("Éxito", f"Documento de Reposición/Devolución '{nombre_archivo_final}' generado y adjuntado correctamente.")
                
            except Exception as db_e:
                conn.rollback()
                messagebox.showerror("Error DB", f"Documento generado, pero error al registrar en DB/Historial. Error: {db_e}")
            finally:
                conn.close()

        except Exception as e:
            messagebox.showerror("Error de Generación", f"No se pudo generar el documento. Asegúrese de que la plantilla existe y es un archivo .docx válido.\nError: {e}")
    
    
    
    def abrir_plantilla_informe_manual_sinusoactualmente(self):
        """
        Abre la plantilla de informe de Word y la carpeta de destino de los adjuntos.
        """
        # Verificación inicial: ¿Estamos editando un RMA?
        if not self.rma_actual_id:
            messagebox.showerror("Error", "Primero debe guardar el expediente o cargar uno existente para generar un informe.")
            return

        # 1. Definir rutas
        # Obtiene la ruta absoluta de la carpeta 'plantillas'
        plantilla_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantillas")
        plantilla_path = os.path.join(plantilla_dir, "Plantilla_RMA.docx") # Asegúrate de que el nombre coincide
        
        # 2. Verificar que la plantilla existe
        if not os.path.exists(plantilla_path):
            messagebox.showerror("Error", f"No se encontró la plantilla de informe en:\n{plantilla_path}")
            return
            
        # 3. Abrir la plantilla con el programa asociado (Word)
        try:
            # os.startfile es la forma más limpia en Windows para abrir archivos
            if os.name == 'nt': 
                 os.startfile(plantilla_path)
            else:
                 # Común para macOS o Linux (puede requerir ajustes según la distribución)
                 subprocess.Popen(['xdg-open', plantilla_path]) 

            messagebox.showinfo("Instrucción", 
                "Se ha abierto la plantilla de Word.\n"
                "A continuación se abrirá la carpeta de destino. Por favor, "
                "**GUARDE el documento de Word FINAL ahí**.")
            
            # 4. Abrir la carpeta de destino para que el usuario guarde el resultado
            self.abrir_dialogo_adjunto(modo_abrir_carpeta=True)

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la plantilla.\nError: {e}")
    
    def mostrar_menu_admin(self):
        """Muestra un menú desplegable con funciones de administración."""
        if str(self.rol).strip().lower() not in ("admin", "administrador"):
            messagebox.showerror("Acceso Denegado", "Solo los administradores tienen acceso a estas funciones.")
            return
            
        # Crear ventana de menú
        menu_window = ctk.CTkToplevel(self)
        menu_window.title("Menú de Administración")
        menu_window.geometry("300x430")
        menu_window.resizable(False, False)
        
        # Centrar la ventana
        menu_window.transient(self)
        menu_window.grab_set()
        
        # Título
        titulo = ctk.CTkLabel(menu_window, text="🔧 Funciones de Administración", 
                             font=ctk.CTkFont(size=16, weight="bold"))
        titulo.pack(pady=20)
        
        # Frame para botones
        buttons_frame = ctk.CTkFrame(menu_window)
        buttons_frame.pack(pady=10, padx=20, fill="both", expand=True)
        
        # Botón Eliminar RMA
        btn_eliminar = ctk.CTkButton(buttons_frame,
                                    text="🗑️ Eliminar RMA",
                                    width=240,
                                    height=40,
                                    fg_color="#dc3545",
                                    hover_color="#c82333",
                                    command=lambda: [menu_window.destroy(), self.mostrar_eliminar_rma()])
        btn_eliminar.pack(pady=10)
        
        # Botón Generar Número Manual
        btn_numero_manual = ctk.CTkButton(buttons_frame,
                                         text="🔢 Generar Número Manual",
                                         width=240,
                                         height=40,
                                         fg_color="#007bff",
                                         hover_color="#0056b3",
                                         command=lambda: [menu_window.destroy(), self.mostrar_generar_numero_manual()])
        btn_numero_manual.pack(pady=10)
        
        # Botón Administrar Avisos
        btn_avisos = ctk.CTkButton(buttons_frame,
                                  text="📢 Administrar Avisos",
                                  width=240,
                                  height=40,
                                  fg_color="#16a34a",
                                  hover_color="#15803d",
                                  command=lambda: [menu_window.destroy(), self.mostrar_admin_avisos()])
        btn_avisos.pack(pady=10)
        
        # Botón Gestionar Estados de Artículos
        btn_estados = ctk.CTkButton(buttons_frame,
                                    text="📋 Gestionar Estados Artículos",
                                    width=240,
                                    height=40,
                                    fg_color="#8b5cf6",
                                    hover_color="#7c3aed",
                                    command=lambda: [menu_window.destroy(), self.mostrar_gestor_estados()])
        btn_estados.pack(pady=10)
        
        # Botón Gestionar Personas
        btn_personas = ctk.CTkButton(buttons_frame,
                                     text="👥 Gestionar Personas",
                                     width=240,
                                     height=40,
                                     fg_color="#06b6d4",
                                     hover_color="#0891b2",
                                     command=lambda: [menu_window.destroy(), self.mostrar_gestor_personas()])
        btn_personas.pack(pady=10)
        
        # Botón Gestionar Resultado Expediente
        btn_resultado = ctk.CTkButton(buttons_frame,
                                      text="📊 Gestionar Resultado Expediente",
                                      width=240,
                                      height=40,
                                      fg_color="#10b981",
                                      hover_color="#059669",
                                      command=lambda: [menu_window.destroy(), self.mostrar_gestor_resultado_expediente()])
        btn_resultado.pack(pady=10)
        
        # Botón Cerrar
        btn_cerrar = ctk.CTkButton(buttons_frame,
                                  text="❌ Cerrar",
                                  width=240,
                                  height=32,
                                  fg_color="gray",
                                  command=menu_window.destroy)
        btn_cerrar.pack(pady=(20, 10))
    
    def mostrar_admin_avisos(self):
        """Muestra el panel de administración de avisos (solo para admin)."""
        if str(self.rol).strip().lower() not in ("admin", "administrador"):
            messagebox.showerror("Acceso Denegado", "Solo los administradores pueden gestionar avisos.")
            return
        
        try:
            self.avisos_manager.mostrar_panel_admin(self)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el panel de avisos:\n{e}")
    
    def mostrar_gestor_estados(self):
        """Muestra la ventana de gestión de estados de artículos"""
        if str(self.rol).strip().lower() not in ("admin", "administrador"):
            messagebox.showerror("Acceso Denegado", "Solo los administradores pueden gestionar estados.")
            return
        
        try:
            # Crear ventana
            ventana = ctk.CTkToplevel(self)
            ventana.title("Gestionar Estados de Artículos")
            ventana.geometry("700x600")
            ventana.transient(self)
            ventana.grab_set()
            
            # Frame principal
            main_frame = ctk.CTkFrame(ventana)
            main_frame.pack(fill="both", expand=True, padx=15, pady=15)
            
            # Título
            titulo = ctk.CTkLabel(main_frame, text="📋 Gestión de Estados de Artículos",
                                 font=ctk.CTkFont(size=18, weight="bold"))
            titulo.pack(pady=(0, 15))
            
            # Frame de controles
            controles_frame = ctk.CTkFrame(main_frame)
            controles_frame.pack(fill="x", pady=(0, 10))
            
            # Entrada para nuevo estado
            entry_nuevo = ctk.CTkEntry(controles_frame, placeholder_text="Nuevo estado...", width=400)
            entry_nuevo.pack(side="left", padx=(10, 5), pady=10)
            
            # Botón añadir
            btn_anadir = ctk.CTkButton(controles_frame, text="➕ Añadir", width=100)
            btn_anadir.pack(side="left", padx=5, pady=10)
            
            # Lista de estados
            lista_frame = ctk.CTkFrame(main_frame)
            lista_frame.pack(fill="both", expand=True, pady=(0, 10))
            
            # Header
            header_frame = ctk.CTkFrame(lista_frame, fg_color=("gray80", "gray25"), height=35)
            header_frame.pack(fill="x", padx=2, pady=(2, 0))
            header_frame.pack_propagate(False)
            
            ctk.CTkLabel(header_frame, text="Estado", font=ctk.CTkFont(weight="bold"), width=450, anchor="w").pack(side="left", padx=10, pady=5)
            ctk.CTkLabel(header_frame, text="Acciones", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=5, pady=5)
            
            # Scrollable frame
            scroll_frame = ctk.CTkScrollableFrame(lista_frame)
            scroll_frame.pack(fill="both", expand=True, padx=2, pady=2)
            
            # Frame de botones inferiores
            botones_frame = ctk.CTkFrame(main_frame)
            botones_frame.pack(fill="x")
            
            btn_recargar = ctk.CTkButton(botones_frame, text="↻ Recargar", width=120)
            btn_recargar.pack(side="left", padx=5, pady=5)
            
            btn_cerrar = ctk.CTkButton(botones_frame, text="Cerrar", width=100, 
                                       command=ventana.destroy)
            btn_cerrar.pack(side="right", padx=5, pady=5)
            
            # Manager
            estados_mgr = EstadosArticuloManager()
            
            # Funciones
            def cargar_estados():
                """Carga y muestra los estados"""
                # Limpiar lista
                for widget in scroll_frame.winfo_children():
                    widget.destroy()
                
                estados = estados_mgr.cargar_estados()
                
                for i, estado in enumerate(estados):
                    fila = ctk.CTkFrame(scroll_frame, fg_color=("gray90", "gray20"))
                    fila.pack(fill="x", padx=2, pady=2)
                    
                    # Estado (vacío se muestra como texto especial)
                    texto_estado = "(Vacío)" if estado == "" else estado
                    lbl_estado = ctk.CTkLabel(fila, text=texto_estado, width=450, anchor="w")
                    lbl_estado.pack(side="left", padx=10, pady=8)
                    
                    # Botones
                    acciones = ctk.CTkFrame(fila, fg_color="transparent")
                    acciones.pack(side="left", padx=5)
                    
                    # Editar
                    btn_editar = ctk.CTkButton(acciones, text="✏", width=30, height=25,
                                               command=lambda e=estado: editar_estado(e))
                    btn_editar.pack(side="left", padx=2)
                    Tooltip(btn_editar, "Editar este estado")
                    
                    # Mover arriba
                    if i > 0:
                        btn_arriba = ctk.CTkButton(acciones, text="▲", width=30, height=25,
                                                   command=lambda e=estado: mover_estado(e, 'arriba'))
                        btn_arriba.pack(side="left", padx=2)
                        Tooltip(btn_arriba, "Mover hacia arriba")
                    
                    # Mover abajo
                    if i < len(estados) - 1:
                        btn_abajo = ctk.CTkButton(acciones, text="▼", width=30, height=25,
                                                  command=lambda e=estado: mover_estado(e, 'abajo'))
                        btn_abajo.pack(side="left", padx=2)
                        Tooltip(btn_abajo, "Mover hacia abajo")
                    
                    # Eliminar
                    btn_eliminar = ctk.CTkButton(acciones, text="🗑", width=30, height=25,
                                                 fg_color="red", hover_color="darkred",
                                                 command=lambda e=estado: eliminar_estado(e))
                    btn_eliminar.pack(side="left", padx=2)
                    Tooltip(btn_eliminar, "Eliminar este estado")
            
            def anadir_estado():
                """Añade un nuevo estado"""
                nuevo = entry_nuevo.get().strip()
                if not nuevo:
                    messagebox.showwarning("Aviso", "Escribe un estado para añadir")
                    return
                
                success, msg = estados_mgr.añadir_estado(nuevo)
                if success:
                    entry_nuevo.delete(0, 'end')
                    cargar_estados()
                    recargar_estados_app()
                    messagebox.showinfo("Éxito", msg)
                else:
                    messagebox.showerror("Error", msg)
            
            def editar_estado(estado_antiguo):
                """Edita un estado existente"""
                # Ventana de edición
                ventana_edit = ctk.CTkToplevel(ventana)
                ventana_edit.title("Editar Estado")
                ventana_edit.geometry("500x150")
                ventana_edit.transient(ventana)
                ventana_edit.grab_set()
                
                ctk.CTkLabel(ventana_edit, text="Nuevo valor:", font=ctk.CTkFont(size=12)).pack(pady=(20, 5))
                
                entry_edit = ctk.CTkEntry(ventana_edit, width=400)
                entry_edit.insert(0, estado_antiguo)
                entry_edit.pack(pady=5)
                
                def guardar_edicion():
                    nuevo_valor = entry_edit.get().strip()
                    if not nuevo_valor:
                        messagebox.showwarning("Aviso", "El estado no puede estar vacío")
                        return
                    
                    success, msg = estados_mgr.editar_estado(estado_antiguo, nuevo_valor)
                    if success:
                        ventana_edit.destroy()
                        cargar_estados()
                        recargar_estados_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
                
                ctk.CTkButton(ventana_edit, text="Guardar", command=guardar_edicion).pack(pady=10)
            
            def eliminar_estado(estado):
                """Elimina un estado"""
                texto_estado = "(Vacío)" if estado == "" else estado
                if not messagebox.askyesno("Confirmar", f"¿Eliminar el estado '{texto_estado}'?"):
                    return
                
                success, msg = estados_mgr.eliminar_estado(estado)
                if success:
                    cargar_estados()
                    recargar_estados_app()
                    messagebox.showinfo("Éxito", msg)
                else:
                    messagebox.showerror("Error", msg)
            
            def mover_estado(estado, direccion):
                """Mueve un estado arriba o abajo"""
                success, msg = estados_mgr.mover_estado(estado, direccion)
                if success:
                    cargar_estados()
                    recargar_estados_app()
                else:
                    messagebox.showinfo("Info", msg)
            
            def recargar_estados_app():
                """Recarga los estados en la aplicación principal"""
                self.OPCIONES["Estado_Producto"] = estados_mgr.cargar_estados()
            
            # Conectar botones
            btn_anadir.configure(command=anadir_estado)
            btn_recargar.configure(command=cargar_estados)
            entry_nuevo.bind("<Return>", lambda e: anadir_estado())
            
            # Cargar estados inicialmente
            cargar_estados()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el gestor de estados:\n{e}")
            import traceback
            traceback.print_exc()
    
    def mostrar_gestor_personas(self):
        """
        Abre una ventana para gestionar la lista de personas
        """
        from lib.safe_toplevel import SafeCTkToplevel
        from lib.personas_manager import PersonasManager
        
        try:
            # Crear ventana de gestión
            gestor_window = SafeCTkToplevel(self)
            gestor_window.title("Gestión de Personas")
            gestor_window.geometry("600x500")
            gestor_window.transient(self)
            gestor_window.grab_set()
            
            # Frame principal con padding
            main_frame = ctk.CTkFrame(gestor_window)
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Título
            title_label = ctk.CTkLabel(main_frame,
                                       text="👥 Gestión de Personas",
                                       font=("Arial", 18, "bold"))
            title_label.pack(pady=(0, 20))
            
            # Frame para agregar nueva persona
            add_frame = ctk.CTkFrame(main_frame)
            add_frame.pack(fill="x", pady=(0, 20))
            
            entry_label = ctk.CTkLabel(add_frame,
                                       text="Nueva Persona:",
                                       font=("Arial", 12))
            entry_label.pack(side="left", padx=(10, 5))
            
            entry_persona = ctk.CTkEntry(add_frame,
                                         width=300,
                                         placeholder_text="Ingrese el nombre de la persona")
            entry_persona.pack(side="left", padx=5)
            
            btn_agregar = ctk.CTkButton(add_frame,
                                        text="➕ Agregar",
                                        width=100)
            btn_agregar.pack(side="left", padx=5)
            
            # Frame scrollable para la lista de personas
            list_frame = ctk.CTkScrollableFrame(main_frame,
                                                label_text="Personas Registradas")
            list_frame.pack(fill="both", expand=True)
            
            # Instanciar el gestor
            personas_mgr = PersonasManager()
            
            def cargar_personas():
                """Carga y muestra la lista de personas"""
                # Limpiar lista actual
                for widget in list_frame.winfo_children():
                    widget.destroy()
                
                # Obtener personas
                personas = personas_mgr.cargar_personas()
                
                if not personas:
                    no_data_label = ctk.CTkLabel(list_frame,
                                                 text="No hay personas registradas",
                                                 text_color="gray")
                    no_data_label.pack(pady=20)
                    return
                
                # Mostrar cada persona
                for i, persona in enumerate(personas):
                    persona_frame = ctk.CTkFrame(list_frame)
                    persona_frame.pack(fill="x", pady=5, padx=5)
                    
                    # Nombre de la persona
                    nombre_label = ctk.CTkLabel(persona_frame,
                                               text=persona,
                                               font=("Arial", 12),
                                               anchor="w")
                    nombre_label.pack(side="left", padx=10, fill="x", expand=True)
                    
                    # Botones de acción
                    btn_frame = ctk.CTkFrame(persona_frame, fg_color="transparent")
                    btn_frame.pack(side="right", padx=5)
                    
                    # Botón Editar
                    btn_editar = ctk.CTkButton(btn_frame,
                                              text="✏️",
                                              width=40,
                                              command=lambda p=persona: editar_persona(p))
                    btn_editar.pack(side="left", padx=2)
                    Tooltip(btn_editar, "Editar persona")
                    
                    # Botón Subir
                    if i > 0:  # No mostrar para el primero
                        btn_subir = ctk.CTkButton(btn_frame,
                                                text="⬆️",
                                                width=40,
                                                command=lambda p=persona: mover_persona(p, "arriba"))
                        btn_subir.pack(side="left", padx=2)
                        Tooltip(btn_subir, "Mover arriba")
                    
                    # Botón Bajar
                    if i < len(personas) - 1:  # No mostrar para el último
                        btn_bajar = ctk.CTkButton(btn_frame,
                                                text="⬇️",
                                                width=40,
                                                command=lambda p=persona: mover_persona(p, "abajo"))
                        btn_bajar.pack(side="left", padx=2)
                        Tooltip(btn_bajar, "Mover abajo")
                    
                    # Botón Eliminar
                    btn_eliminar = ctk.CTkButton(btn_frame,
                                               text="🗑️",
                                               width=40,
                                               fg_color="red",
                                               hover_color="darkred",
                                               command=lambda p=persona: eliminar_persona(p))
                    btn_eliminar.pack(side="left", padx=2)
                    Tooltip(btn_eliminar, "Eliminar persona")
            
            def anadir_persona():
                """Agrega una nueva persona a la lista"""
                nueva_persona = entry_persona.get().strip()
                if nueva_persona:
                    success, msg = personas_mgr.añadir_persona(nueva_persona)
                    if success:
                        entry_persona.delete(0, "end")
                        cargar_personas()
                        recargar_personas_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
                else:
                    messagebox.showwarning("Advertencia", "Por favor ingrese un nombre")
            
            def editar_persona(persona_actual):
                """Abre un diálogo para editar una persona"""
                from tkinter import simpledialog
                nueva_persona = simpledialog.askstring("Editar Persona",
                                                       f"Editar nombre de la persona:\n\n'{persona_actual}'",
                                                       initialvalue=persona_actual,
                                                       parent=gestor_window)
                if nueva_persona and nueva_persona.strip():
                    success, msg = personas_mgr.editar_persona(persona_actual, nueva_persona.strip())
                    if success:
                        cargar_personas()
                        recargar_personas_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
            
            def eliminar_persona(persona):
                """Elimina una persona de la lista"""
                if messagebox.askyesno("Confirmar eliminación",
                                      f"¿Está seguro de eliminar la persona?\n\n'{persona}'",
                                      parent=gestor_window):
                    
                    success, msg = personas_mgr.eliminar_persona(persona)
                    if success:
                        cargar_personas()
                        recargar_personas_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
            
            def mover_persona(persona, direccion):
                """Mueve una persona arriba o abajo"""
                success, msg = personas_mgr.mover_persona(persona, direccion)
                if success:
                    cargar_personas()
                    recargar_personas_app()
                else:
                    messagebox.showinfo("Info", msg)
            
            def recargar_personas_app():
                """Recarga las personas en la aplicación principal"""
                personas_actualizadas = personas_mgr.cargar_personas()
                self.OPCIONES["Autorizado_Por"] = personas_actualizadas
                self.OPCIONES["Gestionado_Por"] = personas_actualizadas
            
            # Conectar botones
            btn_agregar.configure(command=anadir_persona)
            entry_persona.bind("<Return>", lambda e: anadir_persona())
            
            # Cargar personas inicialmente
            cargar_personas()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el gestor de personas:\n{e}")
            import traceback
            traceback.print_exc()
    
    def mostrar_gestor_resultado_expediente(self):
        """
        Abre una ventana para gestionar la lista de resultados de expedientes
        """
        from lib.safe_toplevel import SafeCTkToplevel
        from lib.resultado_expediente_manager import ResultadoExpedienteManager
        
        try:
            # Crear ventana de gestión
            gestor_window = SafeCTkToplevel(self)
            gestor_window.title("Gestión de Resultado Expediente")
            gestor_window.geometry("600x500")
            gestor_window.transient(self)
            gestor_window.grab_set()
            
            # Frame principal con padding
            main_frame = ctk.CTkFrame(gestor_window)
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Título
            title_label = ctk.CTkLabel(main_frame,
                                       text="📊 Gestión de Resultado Expediente",
                                       font=("Arial", 18, "bold"))
            title_label.pack(pady=(0, 20))
            
            # Frame para agregar nuevo resultado
            add_frame = ctk.CTkFrame(main_frame)
            add_frame.pack(fill="x", pady=(0, 20))
            
            entry_label = ctk.CTkLabel(add_frame,
                                       text="Nuevo Resultado:",
                                       font=("Arial", 12))
            entry_label.pack(side="left", padx=(10, 5))
            
            entry_resultado = ctk.CTkEntry(add_frame,
                                           width=300,
                                           placeholder_text="Ingrese el resultado del expediente")
            entry_resultado.pack(side="left", padx=5)
            
            btn_agregar = ctk.CTkButton(add_frame,
                                        text="➕ Agregar",
                                        width=100)
            btn_agregar.pack(side="left", padx=5)
            
            # Frame scrollable para la lista de resultados
            list_frame = ctk.CTkScrollableFrame(main_frame,
                                                label_text="Resultados Registrados")
            list_frame.pack(fill="both", expand=True)
            
            # Instanciar el gestor
            resultados_mgr = ResultadoExpedienteManager()
            
            def cargar_resultados():
                """Carga y muestra la lista de resultados"""
                # Limpiar lista actual
                for widget in list_frame.winfo_children():
                    widget.destroy()
                
                # Obtener resultados
                resultados = resultados_mgr.cargar_resultados()
                
                if not resultados:
                    no_data_label = ctk.CTkLabel(list_frame,
                                                 text="No hay resultados registrados",
                                                 text_color="gray")
                    no_data_label.pack(pady=20)
                    return
                
                # Mostrar cada resultado
                for i, resultado in enumerate(resultados):
                    resultado_frame = ctk.CTkFrame(list_frame)
                    resultado_frame.pack(fill="x", pady=5, padx=5)
                    
                    # Texto del resultado (mostrar etiqueta si está vacío)
                    texto_mostrar = resultado if resultado else "(vacío)"
                    color_texto = "gray" if not resultado else None
                    
                    nombre_label = ctk.CTkLabel(resultado_frame,
                                               text=texto_mostrar,
                                               font=("Arial", 12),
                                               anchor="w",
                                               text_color=color_texto)
                    nombre_label.pack(side="left", padx=10, fill="x", expand=True)
                    
                    # Botones de acción
                    btn_frame = ctk.CTkFrame(resultado_frame, fg_color="transparent")
                    btn_frame.pack(side="right", padx=5)
                    
                    # Botón Editar
                    btn_editar = ctk.CTkButton(btn_frame,
                                              text="✏️",
                                              width=40,
                                              command=lambda r=resultado: editar_resultado(r))
                    btn_editar.pack(side="left", padx=2)
                    Tooltip(btn_editar, "Editar resultado")
                    
                    # Botón Subir
                    if i > 0:  # No mostrar para el primero
                        btn_subir = ctk.CTkButton(btn_frame,
                                                text="⬆️",
                                                width=40,
                                                command=lambda r=resultado: mover_resultado(r, "arriba"))
                        btn_subir.pack(side="left", padx=2)
                        Tooltip(btn_subir, "Mover arriba")
                    
                    # Botón Bajar
                    if i < len(resultados) - 1:  # No mostrar para el último
                        btn_bajar = ctk.CTkButton(btn_frame,
                                                text="⬇️",
                                                width=40,
                                                command=lambda r=resultado: mover_resultado(r, "abajo"))
                        btn_bajar.pack(side="left", padx=2)
                        Tooltip(btn_bajar, "Mover abajo")
                    
                    # Botón Eliminar
                    btn_eliminar = ctk.CTkButton(btn_frame,
                                               text="🗑️",
                                               width=40,
                                               fg_color="red",
                                               hover_color="darkred",
                                               command=lambda r=resultado: eliminar_resultado(r))
                    btn_eliminar.pack(side="left", padx=2)
                    Tooltip(btn_eliminar, "Eliminar resultado")
            
            def anadir_resultado():
                """Agrega un nuevo resultado a la lista"""
                nuevo_resultado = entry_resultado.get().strip()
                if nuevo_resultado:
                    success, msg = resultados_mgr.añadir_resultado(nuevo_resultado)
                    if success:
                        entry_resultado.delete(0, "end")
                        cargar_resultados()
                        recargar_resultados_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
                else:
                    messagebox.showwarning("Advertencia", "Por favor ingrese un resultado")
            
            def editar_resultado(resultado_actual):
                """Abre un diálogo para editar un resultado"""
                from tkinter import simpledialog
                
                texto_mostrar = resultado_actual if resultado_actual else "(vacío)"
                nuevo_resultado = simpledialog.askstring("Editar Resultado",
                                                         f"Editar resultado del expediente:\n\n'{texto_mostrar}'",
                                                         initialvalue=resultado_actual,
                                                         parent=gestor_window)
                if nuevo_resultado is not None and nuevo_resultado.strip():
                    success, msg = resultados_mgr.editar_resultado(resultado_actual, nuevo_resultado.strip())
                    if success:
                        cargar_resultados()
                        recargar_resultados_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
            
            def eliminar_resultado(resultado):
                """Elimina un resultado de la lista"""
                texto_mostrar = resultado if resultado else "(vacío)"
                if messagebox.askyesno("Confirmar eliminación",
                                      f"¿Está seguro de eliminar el resultado?\n\n'{texto_mostrar}'",
                                      parent=gestor_window):
                    
                    success, msg = resultados_mgr.eliminar_resultado(resultado)
                    if success:
                        cargar_resultados()
                        recargar_resultados_app()
                        messagebox.showinfo("Éxito", msg)
                    else:
                        messagebox.showerror("Error", msg)
            
            def mover_resultado(resultado, direccion):
                """Mueve un resultado arriba o abajo"""
                success, msg = resultados_mgr.mover_resultado(resultado, direccion)
                if success:
                    cargar_resultados()
                    recargar_resultados_app()
                else:
                    messagebox.showinfo("Info", msg)
            
            def recargar_resultados_app():
                """Recarga los resultados en la aplicación principal"""
                self.OPCIONES["Resultado_Expediente"] = resultados_mgr.cargar_resultados()
            
            # Conectar botones
            btn_agregar.configure(command=anadir_resultado)
            entry_resultado.bind("<Return>", lambda e: anadir_resultado())
            
            # Cargar resultados inicialmente
            cargar_resultados()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el gestor de resultado expediente:\n{e}")
            import traceback
            traceback.print_exc()
    
    def mostrar_gestion_usuarios(self):
        """Muestra la ventana de gestión de usuarios con opciones para añadir/editar usuarios."""
        # Solo admin tiene permisos para gestionar usuarios
        if str(self.rol).strip().lower() not in ("admin", "administrador"):
            messagebox.showerror("Error", "Solo el administrador puede gestionar usuarios.")
            return

        # Crear ventana modal
        ventana = ctk.CTkToplevel(self)
        ventana.title("Gestión de Usuarios")
        ventana.geometry("600x500")
        
        # Configurar para permitir minimización
        ventana.resizable(True, True)
        ventana.attributes('-topmost', False)
        ventana.minsize(500, 400)
        # No usar grab_set para permitir minimización completa
        try:
            ventana.focus_set()  # Solo dar foco sin modalidad
        except:
            pass
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        ventana.attributes('-topmost', True)   # Temporalmente al frente
        ventana.lift()
        ventana.focus_force()
        ventana.after(500, lambda: ventana.attributes('-topmost', False))  # Quitar topmost después de 500ms
        
        # Agregar icono personalizado
        try:
            ventana.iconbitmap("Icono_Ilutrek.ico")
        except Exception:
            pass

        # Frame principal
        frame = ctk.CTkFrame(ventana)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Título
        ctk.CTkLabel(frame, text="Gestión de Usuarios", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(0, 20))

        # Frame para el formulario de nuevo usuario
        form_frame = ctk.CTkFrame(frame)
        form_frame.pack(fill="x", padx=10, pady=10)

        # Campos del formulario
        ctk.CTkLabel(form_frame, text="Nombre de Usuario:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        entry_usuario = ctk.CTkEntry(form_frame)
        entry_usuario.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        ctk.CTkLabel(form_frame, text="Contraseña:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        entry_password = ctk.CTkEntry(form_frame, show="*")
        entry_password.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        ctk.CTkLabel(form_frame, text="Rol:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        # Lista de roles disponibles
        roles_disponibles = [
            "usuario",        # Usuario básico
            "admin",         # Administrador total
            "Dpto. Tecnico", # Departamento técnico
            "Administracion",# Administración
            "Contabilidad",  # Contabilidad
            "Almacen"        # Almacén
        ]
        combo_rol = ctk.CTkOptionMenu(form_frame, values=roles_disponibles)
        combo_rol.grid(row=2, column=1, padx=5, pady=5, sticky="ew")
        combo_rol.set("usuario")  # Valor por defecto

        form_frame.grid_columnconfigure(1, weight=1)

        def agregar_usuario():
            username = entry_usuario.get().strip()
            password = entry_password.get().strip()
            rol = combo_rol.get()

            if not username or not password:
                messagebox.showerror("Error", "Por favor, complete todos los campos.")
                return

            try:
                # Generar hash de la contraseña y guardarlo como str UTF-8
                password_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode('utf-8')

                conn = connect_db()
                cursor = conn.cursor()

                # Verificar si el usuario ya existe
                cursor.execute("SELECT nombre_usuario FROM usuarios WHERE nombre_usuario = ?", (username,))
                if cursor.fetchone():
                    messagebox.showerror("Error", "El nombre de usuario ya existe.")
                    return

                # Insertar nuevo usuario
                cursor.execute(
                    "INSERT INTO usuarios (nombre_usuario, password_hash, rol) VALUES (?, ?, ?)",
                    (username, password_hash, rol)
                )
                conn.commit()
                messagebox.showinfo("Éxito", "Usuario creado correctamente.")
                
                # Limpiar campos
                entry_usuario.delete(0, 'end')
                entry_password.delete(0, 'end')
                combo_rol.set("usuario")
                
                # Actualizar lista de usuarios
                actualizar_lista_usuarios()

            except sqlite3.Error as e:
                messagebox.showerror("Error", f"Error al crear usuario: {e}")
            finally:
                conn.close()

        # Botón para agregar usuario
        ctk.CTkButton(form_frame, text="Agregar Usuario", command=agregar_usuario).grid(row=3, column=0, columnspan=2, pady=20)

        # Frame para la lista de usuarios
        list_frame = ctk.CTkFrame(frame)
        list_frame.pack(fill="both", expand=True, padx=10, pady=(20,10))

        # Título de la lista
        ctk.CTkLabel(list_frame, text="Usuarios Existentes", font=ctk.CTkFont(weight="bold")).pack(pady=10)

        # Crear un frame scrollable para la lista de usuarios
        scroll_frame = ctk.CTkScrollableFrame(list_frame)
        scroll_frame.pack(fill="both", expand=True, padx=5, pady=5)

        def actualizar_lista_usuarios():
            # Limpiar lista actual
            for widget in scroll_frame.winfo_children():
                widget.destroy()

            try:
                conn = connect_db()
                cursor = conn.cursor()
                # Asegurarnos de que la columna email exista (añadir si no)
                try:
                    cursor.execute("PRAGMA table_info('usuarios')")
                    cols = [r[1] for r in cursor.fetchall()]
                    if 'email' not in cols:
                        cursor.execute("ALTER TABLE usuarios ADD COLUMN email TEXT")
                        conn.commit()
                except Exception:
                    # Ignorar si ya existe o si el backend no permite ALTER (p.ej. versiones restringidas)
                    pass
                # Seleccionamos también el email si existe
                try:
                    cursor.execute("SELECT nombre_usuario, rol, COALESCE(email, '') FROM usuarios")
                except Exception:
                    # Fallback si la columna no existe por alguna razón
                    cursor.execute("SELECT nombre_usuario, rol FROM usuarios")
                usuarios = cursor.fetchall()

                for i, row in enumerate(usuarios):
                    # row can be (usuario, rol) or (usuario, rol, email)
                    usuario = row[0] if len(row) > 0 else ''
                    rol = row[1] if len(row) > 1 else ''
                    display_email = row[2] if len(row) > 2 else ''

                    row_frame = ctk.CTkFrame(scroll_frame)
                    row_frame.pack(fill="x", padx=5, pady=2)

                    label_text = f"{usuario} ({rol})"
                    if display_email:
                        label_text += f" — {display_email}"
                    lbl = ctk.CTkLabel(row_frame, text=label_text)
                    lbl.pack(side="left", padx=5)

                    # Hacer clic en el label abre editor de usuario, excepto para admin
                    if usuario != "admin":
                        def make_editor(u=usuario):
                            return lambda e=None: editar_usuario(u)
                        try:
                            lbl.bind("<Button-1>", make_editor(usuario))
                            lbl.configure(cursor="hand2")
                        except Exception:
                            pass
                    else:
                        # Mostrar cursor normal y no permitir edición
                        try:
                            lbl.configure(cursor="")
                        except Exception:
                            pass
                    
                    if usuario != "admin":  # No permitir eliminar al usuario admin
                        def make_delete(u=usuario):
                            return lambda: eliminar_usuario(u)
                        
                        ctk.CTkButton(row_frame, text="❌", width=30, command=make_delete(usuario)).pack(side="right", padx=5)

            except sqlite3.Error as e:
                messagebox.showerror("Error", f"Error al cargar usuarios: {e}")
            finally:
                conn.close()

        def editar_usuario(username):
            """Abrir dialogo para editar rol, password y email del usuario seleccionado."""
            # Proteger al usuario admin: no se puede editar
            if username == "admin":
                messagebox.showerror("Acceso denegado", "El usuario 'admin' no se puede editar ni eliminar.")
                return
            try:
                conn = connect_db()
                cur = conn.cursor()
                # Intentar obtener email si existe
                try:
                    cur.execute("SELECT nombre_usuario, rol, COALESCE(email, '') FROM usuarios WHERE nombre_usuario = ?", (username,))
                    row = cur.fetchone()
                    if row and len(row) >= 3:
                        _, rol_actual, email_actual = row
                    else:
                        cur.execute("SELECT nombre_usuario, rol FROM usuarios WHERE nombre_usuario = ?", (username,))
                        row2 = cur.fetchone()
                        rol_actual = row2[1] if row2 else ''
                        email_actual = ''
                except Exception:
                    # fallback
                    cur.execute("SELECT nombre_usuario, rol FROM usuarios WHERE nombre_usuario = ?", (username,))
                    row2 = cur.fetchone()
                    rol_actual = row2[1] if row2 else ''
                    email_actual = ''
                conn.close()
            except Exception:
                rol_actual = ''
                email_actual = ''

            # Editor modal
            ed = ctk.CTkToplevel(ventana)
            ed.title(f"Editar usuario: {username}")
            ed.geometry("420x260")
            ed.grab_set()

            f = ctk.CTkFrame(ed)
            f.pack(fill="both", expand=True, padx=12, pady=12)

            ctk.CTkLabel(f, text=f"Usuario: {username}", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0,8))

            ctk.CTkLabel(f, text="Rol:").grid(row=1, column=0, sticky="e", padx=5, pady=6)
            rol_menu = ctk.CTkOptionMenu(f, values=roles_disponibles)
            rol_menu.grid(row=1, column=1, sticky="ew", padx=5, pady=6)
            try:
                rol_menu.set(rol_actual or "usuario")
            except Exception:
                rol_menu.set("usuario")

            ctk.CTkLabel(f, text="Nueva contraseña:").grid(row=2, column=0, sticky="e", padx=5, pady=6)
            new_pass = ctk.CTkEntry(f, show="*")
            new_pass.grid(row=2, column=1, sticky="ew", padx=5, pady=6)

            ctk.CTkLabel(f, text="Email:").grid(row=3, column=0, sticky="e", padx=5, pady=6)
            email_entry = ctk.CTkEntry(f)
            email_entry.grid(row=3, column=1, sticky="ew", padx=5, pady=6)
            try:
                email_entry.insert(0, email_actual)
            except Exception:
                pass

            f.grid_columnconfigure(1, weight=1)

            def guardar_cambios():
                nuevo_rol = rol_menu.get()
                nueva_pass = new_pass.get().strip()
                nuevo_email = email_entry.get().strip()
                try:
                    conn2 = connect_db()
                    cur2 = conn2.cursor()
                    # Si se proporcionó nueva contraseña, hashearla
                    if nueva_pass:
                        # Guardar el hash como string UTF-8 para mantener compatibilidad
                        hashed = bcrypt.hashpw(nueva_pass.encode(), bcrypt.gensalt()).decode('utf-8')
                        try:
                            cur2.execute("UPDATE usuarios SET password_hash = ? WHERE nombre_usuario = ?", (hashed, username))
                        except Exception:
                            pass
                    # Actualizar rol y email
                    try:
                        cur2.execute("UPDATE usuarios SET rol = ?, email = ? WHERE nombre_usuario = ?", (nuevo_rol, nuevo_email, username))
                    except Exception:
                        # Si UPDATE falla porque la columna email no existe, intentar crearla y reintentar
                        try:
                            cur2.execute("PRAGMA table_info('usuarios')")
                            cols = [r[1] for r in cur2.fetchall()]
                            if 'email' not in cols:
                                cur2.execute("ALTER TABLE usuarios ADD COLUMN email TEXT")
                            cur2.execute("UPDATE usuarios SET rol = ?, email = ? WHERE nombre_usuario = ?", (nuevo_rol, nuevo_email, username))
                        except Exception:
                            pass
                    conn2.commit()
                    conn2.close()
                    messagebox.showinfo("Éxito", "Usuario actualizado correctamente.")
                    ed.destroy()
                    actualizar_lista_usuarios()
                except sqlite3.Error as e:
                    messagebox.showerror("Error", f"No se pudo actualizar usuario: {e}")

            btn_frame = ctk.CTkFrame(f)
            btn_frame.grid(row=6, column=0, columnspan=2, pady=(12,0))
            ctk.CTkButton(btn_frame, text="Guardar", command=guardar_cambios).pack(side="left", padx=6)
            ctk.CTkButton(btn_frame, text="Cancelar", command=ed.destroy).pack(side="left", padx=6)

        def eliminar_usuario(username):
            if messagebox.askyesno("Confirmar", f"¿Está seguro de eliminar al usuario {username}?"):
                try:
                    conn = connect_db()
                    cursor = conn.cursor()
                    cursor.execute("DELETE FROM usuarios WHERE nombre_usuario = ?", (username,))
                    conn.commit()
                    messagebox.showinfo("Éxito", "Usuario eliminado correctamente.")
                    actualizar_lista_usuarios()
                except sqlite3.Error as e:
                    messagebox.showerror("Error", f"Error al eliminar usuario: {e}")
                finally:
                    conn.close()

        # Cargar lista inicial de usuarios
        actualizar_lista_usuarios()
        # Cargar lista inicial de usuarios
        actualizar_lista_usuarios()

    def mostrar_gestion_tareas(self):
        """Ventana para crear y listar tareas relacionadas con expedientes."""
        # Actualizar badge antes de mostrar la ventana
        self.actualizar_badge_tareas()
        
        # Mostrar solo el listado de tareas y filtros (la creación se hace desde la ficha del expediente)
        ventana = ctk.CTkToplevel(self)
        ventana.title("Listado de Tareas")
        ventana.geometry("700x550")
        
        # Configurar para permitir minimización
        ventana.resizable(True, True)
        ventana.attributes('-topmost', False)
        ventana.minsize(600, 400)
        # No usar transient ni grab_set para permitir minimización completa
        try:
            ventana.focus_set()  # Dar foco sin bloquear
        except:
            pass
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        ventana.attributes('-topmost', True)   # Temporalmente al frente
        ventana.lift()
        try:
            ventana.focus_force()
        except:
            pass
        
        def quitar_topmost_ventana():
            try:
                if ventana.winfo_exists():
                    ventana.attributes('-topmost', False)
            except:
                pass
        
        ventana.after(500, quitar_topmost_ventana)
        ventana.focus_force()
        ventana.after(500, lambda: ventana.attributes('-topmost', False))  # Quitar topmost después de 500ms

        frame = ctk.CTkFrame(ventana)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ctk.CTkLabel(frame, text="Listado de Tareas", font=ctk.CTkFont(size=18, weight="bold")).pack(pady=(0,10))

        # Filtros y lista
        controls = ctk.CTkFrame(frame)
        controls.pack(fill="x", padx=5, pady=(0,10))

        ctk.CTkLabel(controls, text="Filtrar por estado:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        filtro_estado = ctk.CTkOptionMenu(controls, values=["Todos", "Pendiente", "En Progreso", "Completado"]) 
        filtro_estado.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        filtro_estado.set("Todos")

        # Filtro por creador de la tarea
        try:
            conn_tmp = connect_db()
            cur_tmp = conn_tmp.cursor()
            cur_tmp.execute("SELECT DISTINCT creado_por FROM tareas")
            creadores_rows = cur_tmp.fetchall()
            try:
                conn_tmp.close()
            except Exception:
                pass
            creadores = [r[0] for r in creadores_rows if r and r[0]]
        except Exception:
            creadores = []

        valores_creador = ["Todos"] + creadores
        ctk.CTkLabel(controls, text="Filtrar por Usuario:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
        filtro_creador = ctk.CTkOptionMenu(controls, values=valores_creador)
        filtro_creador.grid(row=0, column=3, padx=5, pady=5, sticky="w")
        filtro_creador.set("Todos")

        list_frame = ctk.CTkFrame(frame)
        list_frame.pack(fill="both", expand=True, padx=5, pady=5)

        scroll = ctk.CTkScrollableFrame(list_frame)
        scroll.pack(fill="both", expand=True, padx=5, pady=5)

        def abrir_expediente_por_codigo(codigo):
            try:
                conn = connect_db()
                cur = conn.cursor()
                cur.execute("SELECT id FROM rma_maestro WHERE codigo_rma = ?", (codigo,))
                row = cur.fetchone()
                conn.close()
                if row:
                    rma_id = row[0]
                    # Abrir en ventana independiente
                    self._abrir_editor_rma(rma_id=rma_id)
                    # No cerrar la ventana de tareas para poder seguir trabajando
                else:
                    messagebox.showwarning("No encontrado", f"No se encontró expediente con código {codigo}")
            except sqlite3.Error as e:
                messagebox.showerror("Error BD", f"Error buscando expediente: {e}")

        def actualizar_lista_tareas():
            for w in scroll.winfo_children():
                w.destroy()
            try:
                conn = connect_db()
                cur = conn.cursor()
                estado = filtro_estado.get()
                creador = filtro_creador.get()

                base_sql = "SELECT id, codigo_rma, titulo, descripcion, fecha_vencimiento, estado, creado_por FROM tareas"
                where_clauses = []
                params = []
                if estado and estado != "Todos":
                    where_clauses.append("estado = ?")
                    params.append(estado)
                if creador and creador != "Todos":
                    where_clauses.append("creado_por = ?")
                    params.append(creador)

                order_clause = " ORDER BY fecha_vencimiento IS NULL, fecha_vencimiento ASC"

                if where_clauses:
                    sql = base_sql + " WHERE " + " AND ".join(where_clauses) + order_clause
                    cur.execute(sql, tuple(params))
                else:
                    cur.execute(base_sql + order_clause)
                filas = cur.fetchall()
                conn.close()

                # Encabezados de tabla con estructura en columnas
                header = ctk.CTkFrame(scroll)
                header.pack(fill="x", padx=5, pady=(0, 5))
                header.grid_columnconfigure(0, weight=2, minsize=200)  # Título
                header.grid_columnconfigure(1, weight=1, minsize=100)  # Código
                header.grid_columnconfigure(2, weight=1, minsize=120)  # Estado
                header.grid_columnconfigure(3, weight=1, minsize=120)  # Vencimiento
                header.grid_columnconfigure(4, weight=0, minsize=80)   # Acciones
                
                header_font = ctk.CTkFont(weight="bold")
                ctk.CTkLabel(header, text="TÍTULO", font=header_font, anchor="w").grid(row=0, column=0, padx=5, sticky="w")
                ctk.CTkLabel(header, text="CÓDIGO", font=header_font, anchor="w").grid(row=0, column=1, padx=5, sticky="w")
                ctk.CTkLabel(header, text="ESTADO", font=header_font, anchor="w").grid(row=0, column=2, padx=5, sticky="w")
                ctk.CTkLabel(header, text="VENCIMIENTO", font=header_font, anchor="w").grid(row=0, column=3, padx=5, sticky="w")
                ctk.CTkLabel(header, text="ACCIONES", font=header_font, anchor="center").grid(row=0, column=4, padx=5)

                # Filas de datos
                for tid, codigo, titulo, desc, fecha_v, estado_tarea, creador in filas:
                    row = ctk.CTkFrame(scroll, fg_color="transparent")
                    row.pack(fill="x", padx=5, pady=2)
                    
                    # Configurar columnas igual que header
                    row.grid_columnconfigure(0, weight=2, minsize=200)
                    row.grid_columnconfigure(1, weight=1, minsize=100)
                    row.grid_columnconfigure(2, weight=1, minsize=120)
                    row.grid_columnconfigure(3, weight=1, minsize=120)
                    row.grid_columnconfigure(4, weight=0, minsize=80)
                    
                    # Título - clickeable
                    titulo_lbl = ctk.CTkLabel(row, text=titulo, anchor="w", cursor="hand2")
                    titulo_lbl.grid(row=0, column=0, padx=5, sticky="w")
                    
                    # Código RMA - clickeable
                    codigo_lbl = ctk.CTkLabel(row, text=codigo, anchor="w", cursor="hand2")
                    codigo_lbl.grid(row=0, column=1, padx=5, sticky="w")
                    
                    # Estado
                    estado_lbl = ctk.CTkLabel(row, text=estado_tarea, anchor="w")
                    estado_lbl.grid(row=0, column=2, padx=5, sticky="w")
                    
                    # Fecha vencimiento
                    fecha_lbl = ctk.CTkLabel(row, text=fecha_v if fecha_v else "Sin fecha", anchor="w")
                    fecha_lbl.grid(row=0, column=3, padx=5, sticky="w")
                    
                    # Efectos hover para toda la fila
                    def on_enter(e, r=row):
                        try:
                            r.configure(fg_color="transparent")
                        except Exception:
                            pass
                    def on_leave(e, r=row):
                        try:
                            r.configure(fg_color="transparent")
                        except Exception:
                            pass
                    
                    row.bind("<Enter>", on_enter)
                    row.bind("<Leave>", on_leave)
                    titulo_lbl.bind("<Enter>", on_enter)
                    titulo_lbl.bind("<Leave>", on_leave)
                    codigo_lbl.bind("<Enter>", on_enter)
                    codigo_lbl.bind("<Leave>", on_leave)
                    
                    # Click para abrir expediente
                    titulo_lbl.bind("<Button-1>", lambda e, c=codigo: abrir_expediente_por_codigo(c))
                    codigo_lbl.bind("<Button-1>", lambda e, c=codigo: abrir_expediente_por_codigo(c))
                    
                    # Botones de acción
                    acciones_frame = ctk.CTkFrame(row, fg_color="transparent")
                    acciones_frame.grid(row=0, column=4, padx=5)
                    
                    def make_done(tid=tid):
                        return lambda: marcar_completada(tid)
                    def make_delete(tid=tid):
                        return lambda: eliminar_tarea(tid)
                    
                    if estado_tarea != "Completado":
                        ctk.CTkButton(acciones_frame, text="✅", width=30, command=make_done(tid)).pack(side="left", padx=2)
                    ctk.CTkButton(acciones_frame, text="❌", width=30, command=make_delete(tid)).pack(side="left", padx=2)

            except sqlite3.Error as e:
                messagebox.showerror("Error", f"Error cargando tareas: {e}")

        def marcar_completada(task_id):
            try:
                conn = connect_db()
                cur = conn.cursor()
                cur.execute("UPDATE tareas SET estado = 'Completado', notificado = 1 WHERE id = ?", (task_id,))
                conn.commit()
                conn.close()
                actualizar_lista_tareas()
                # Actualizar badge de tareas
                self.actualizar_badge_tareas()
            except sqlite3.Error as e:
                messagebox.showerror("Error", f"No se pudo actualizar la tarea: {e}")

        def eliminar_tarea(task_id):
            if not messagebox.askyesno("Confirmar", "¿Eliminar esta tarea?"):
                return
            try:
                conn = connect_db()
                cur = conn.cursor()
                cur.execute("DELETE FROM tareas WHERE id = ?", (task_id,))
                conn.commit()
                conn.close()
                actualizar_lista_tareas()
            except sqlite3.Error as e:
                messagebox.showerror("Error", f"No se pudo eliminar la tarea: {e}")

        filtro_estado.configure(command=lambda v=None: actualizar_lista_tareas())
        actualizar_lista_tareas()

    def mostrar_gestion_rmp(self):
        """Ventana para listar proveedores (rmp_proveedores) y permitir acciones: búsqueda, orden, editar estado y ver expedientes asociados."""
        # Evitar abrir múltiples instancias
        if hasattr(self, 'gestion_rmp_window') and getattr(self, 'gestion_rmp_window').winfo_exists():
            getattr(self, 'gestion_rmp_window').focus()
            return

        self.gestion_rmp_window = ctk.CTkToplevel(self)
        win = self.gestion_rmp_window
        win.title("Gestión RMP - Proveedores")
        win.geometry("1000x650")
        
        # Configurar para permitir minimización
        win.resizable(True, True)
        win.attributes('-topmost', False)
        win.minsize(800, 500)
        # No usar transient ni grab_set para permitir minimización completa
        win.focus_set()  # Dar foco sin bloquear
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        win.attributes('-topmost', True)   # Temporalmente al frente
        win.lift()
        win.focus_force()
        win.after(500, lambda: win.attributes('-topmost', False))  # Quitar topmost después de 500ms
        
        # Agregar icono personalizado
        try:
            win.iconbitmap("Icono_Ilutrek.ico")
        except Exception:
            pass

        main = ctk.CTkFrame(win)
        main.pack(fill="both", expand=True, padx=12, pady=12)

        header = ctk.CTkFrame(main)
        header.pack(fill="x", pady=(0,8))

        ctk.CTkLabel(header, text="Listado de RMA - Proveedores", font=ctk.CTkFont(size=18, weight="bold")).grid(row=0, column=0, sticky="w")

        # Búsqueda y filtro por estado
        ctk.CTkLabel(header, text="Buscar:").grid(row=1, column=0, sticky="w", pady=(8,0))
        entry_buscar = ctk.CTkEntry(header, placeholder_text="Escriba parte del numero RMA proveedor...")
        entry_buscar.grid(row=1, column=1, sticky="ew", padx=(8,0), pady=(8,0))
        # Filtro por estado
        ctk.CTkLabel(header, text="Filtrar estado:").grid(row=1, column=2, sticky="w", padx=(12,6), pady=(8,0))
        # Incluir 'Exportado' como una opción de filtro de estado
        estado_filter = ctk.CTkOptionMenu(header, values=["Todos", "En Progreso", "Enviado", "Completado", "Exportado"])
        estado_filter.set("Todos")
        estado_filter.grid(row=1, column=3, sticky="w", padx=(0,6), pady=(8,0))
        header.grid_columnconfigure(1, weight=1)

        btn_buscar = ctk.CTkButton(header, text="🔎", width=40)
        btn_buscar.grid(row=1, column=4, padx=(8,0), pady=(8,0))

        list_frame = ctk.CTkFrame(main)
        list_frame.pack(fill="both", expand=True)

        scroll = ctk.CTkScrollableFrame(list_frame)
        scroll.pack(fill="both", expand=True, padx=5, pady=5)

        # Estado de ordenación
        sort_state = {'col': 'nombre', 'dir': 'asc'}

        # Controles de paginación para proveedores
        prov_page = {'page': 1, 'page_size': 20, 'total': 0}
        prov_pframe = ctk.CTkFrame(main)
        prov_pframe.pack(fill="x", padx=12, pady=(6,8))
        prov_prev = ctk.CTkButton(prov_pframe, text="◀ Anterior", width=100)
        prov_prev.pack(side="left", padx=(0,8))
        prov_page_lbl = ctk.CTkLabel(prov_pframe, text=f"Página {prov_page['page']}")
        prov_page_lbl.pack(side="left")
        prov_next = ctk.CTkButton(prov_pframe, text="Siguiente ▶", width=100)
        prov_next.pack(side="left", padx=8)
        prov_size_opt = ctk.CTkOptionMenu(prov_pframe, values=["10","20","50","100"], width=80)
        prov_size_opt.set(str(prov_page['page_size']))
        prov_size_opt.pack(side="right")
        ctk.CTkLabel(prov_pframe, text="Registros por página:").pack(side="right", padx=(0,8))

        # Handlers mínimos: actualizan prov_page y recargan la lista
        def _prov_prev():
            if prov_page['page'] > 1:
                prov_page['page'] -= 1
                try:
                    cargar_proveedores()
                except Exception:
                    pass

        def _prov_next():
            # no conocemos aún prov_page['total'], simplemente incrementamos y dejamos que la consulta limite
            prov_page['page'] += 1
            try:
                cargar_proveedores()
            except Exception:
                pass

        def _prov_size_changed(v):
            try:
                prov_page['page_size'] = int(v)
            except Exception:
                prov_page['page_size'] = 20
            prov_page['page'] = 1
            try:
                cargar_proveedores()
            except Exception:
                pass

        prov_prev.configure(command=_prov_prev)
        prov_next.configure(command=_prov_next)
        prov_size_opt.configure(command=_prov_size_changed)

        # Función para mostrar historial completo de un proveedor y añadir comentarios
        def mostrar_historial_proveedor(proveedor_nombre):
            try:
                vent_hist = ctk.CTkToplevel(self)
                vent_hist.title(f"Historial - {proveedor_nombre}")
                vent_hist.geometry("700x500")
                
                # Configurar para permitir minimización
                vent_hist.resizable(True, True)
                vent_hist.attributes('-topmost', False)
                vent_hist.minsize(600, 400)
                # No usar transient para permitir minimización completa
                vent_hist.focus_set()  # Dar foco sin bloquear
                
                # Forzar aparición al frente (incluso si la principal está maximizada)
                vent_hist.attributes('-topmost', True)   # Temporalmente al frente
                vent_hist.lift()
                try:
                    vent_hist.focus_force()
                except:
                    pass
                
                def quitar_topmost_vent_hist():
                    try:
                        if vent_hist.winfo_exists():
                            vent_hist.attributes('-topmost', False)
                    except:
                        pass
                
                vent_hist.after(500, quitar_topmost_vent_hist)

                cont = ctk.CTkFrame(vent_hist)
                cont.pack(fill="both", expand=True, padx=10, pady=10)

                sf = ctk.CTkScrollableFrame(cont)
                sf.pack(fill="both", expand=True)

                # Cargar historial
                try:
                    connh = connect_db()
                    curh = connh.cursor()
                    curh.execute("SELECT fecha, usuario, estado, comentario FROM rma_proveedor_hist WHERE lower(proveedor)=? OR proveedor=? ORDER BY fecha DESC", (proveedor_nombre.lower(), proveedor_nombre))
                    hist_rows = curh.fetchall()
                    connh.close()
                except Exception:
                    hist_rows = []

                
                if not hist_rows:
                    try:
                        import sqlite3 as _sqlite
                        local_db = os.path.join(os.path.dirname(__file__), DB_NAME)
                        if os.path.exists(local_db):
                            conn_local = _sqlite.connect(local_db)
                            cur_local = conn_local.cursor()
                            cur_local.execute("SELECT fecha, usuario, estado, comentario FROM rma_proveedor_hist WHERE lower(proveedor)=? OR proveedor=? ORDER BY fecha DESC", (proveedor_nombre.lower(), proveedor_nombre))
                            hist_rows = cur_local.fetchall()
                            conn_local.close()
                    except Exception:
                        hist_rows = []

                # Mostrar cada entrada
                for idx, (fecha, usuario, estado_h, comentario) in enumerate(hist_rows):
                    # Usar fondo por defecto del tema en lugar de cebra
                    rowf = ctk.CTkFrame(sf, fg_color="transparent")
                    rowf.pack(fill="x", padx=5, pady=3)
                    txt = f"{fecha} - {usuario} - {estado_h or ''}"
                    ctk.CTkLabel(rowf, text=txt, font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=6, pady=(4,0))
                    if comentario:
                        ctk.CTkLabel(rowf, text=comentario, wraplength=650).pack(anchor="w", padx=6, pady=(0,6))

                # Area para añadir comentario
                ctk.CTkLabel(cont, text="Añadir comentario:").pack(anchor="w", pady=(8,2))
                comment_box = ctk.CTkTextbox(cont, height=80)
                comment_box.pack(fill="x", pady=(0,8))

                def _add_comment():
                    text = comment_box.get("0.0", "end").strip()
                    if not text:
                        messagebox.showwarning("Vacío", "Escribe un comentario antes de añadirlo.")
                        return
                    try:
                        connc = connect_db()
                        curc = connc.cursor()
                        # Asegurar la tabla existe (por si acaso)
                        curc.execute("CREATE TABLE IF NOT EXISTS rma_proveedor_hist (id INTEGER PRIMARY KEY, proveedor TEXT, estado TEXT, comentario TEXT, usuario TEXT, fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
                        curc.execute("INSERT INTO rma_proveedor_hist (proveedor, estado, comentario, usuario) VALUES (?, ?, ?, ?)", (proveedor_nombre, '', text, getattr(self, 'username', 'unknown')))
                        connc.commit()
                        connc.close()
                        messagebox.showinfo("Añadido", "Comentario añadido al historial.")
                        # refrescar la ventana
                        vent_hist.destroy()
                        mostrar_historial_proveedor(proveedor_nombre)
                        # refrescar la lista principal
                        try:
                            cargar_proveedores()
                        except Exception:
                            pass
                    except Exception as e:
                        messagebox.showerror("Error BD", f"No se pudo añadir el comentario: {e}")

                ctk.CTkButton(cont, text="Añadir comentario", command=_add_comment).pack(anchor="e")
            except Exception as e:
                messagebox.showerror("Error", f"Error mostrando historial: {e}")

        def cargar_proveedores():
            """Carga la lista de proveedores EXTRAÍDOS de rma_maestro.rma_proveedor (DISTINCT)."""
            for w in scroll.winfo_children():
                w.destroy()
            filtro = entry_buscar.get().strip()
            try:
                conn = connect_db()
                cur = conn.cursor()

                # Asegurar que exista la tabla de proveedores para persistir estados
                try:
                    cur.execute("CREATE TABLE IF NOT EXISTS rma_proveedor (id INTEGER PRIMARY KEY, proveedor TEXT UNIQUE, estado TEXT, factura_abono TEXT)")
                    # Tabla de historial de proveedores: proveedor, estado, comentario, usuario, fecha
                    cur.execute(
                        "CREATE TABLE IF NOT EXISTS rma_proveedor_hist (id INTEGER PRIMARY KEY, proveedor TEXT, estado TEXT, comentario TEXT, usuario TEXT, fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"
                    )
                    conn.commit()
                    # Nota: La columna factura_abono ya está incluida en la creación de la tabla arriba
                except Exception:
                    # Si no podemos crearla en Turso u otro backend, seguimos sin persistencia
                    pass

                # Construir consulta: obtenemos proveedores distintos de rma_maestro
                # y left-join con rma_proveedor para traer el estado y factura_abono si existe.
                params = []
                search_clause = ""
                if filtro:
                    search_clause = " AND lower(rma_proveedor) LIKE ?"
                    params.append(f"%{filtro.lower()}%")

                direction = 'ASC' if sort_state.get('dir', 'asc') == 'asc' else 'DESC'

                sql = (
                    "SELECT p.proveedor, COALESCE(r.estado, '') as estado, COALESCE(r.factura_abono, '') as factura_abono "
                    "FROM (SELECT DISTINCT rma_proveedor AS proveedor FROM rma_maestro WHERE rma_proveedor IS NOT NULL AND rma_proveedor != ''"
                    + search_clause + ") p "
                    "LEFT JOIN rma_proveedor r ON lower(p.proveedor) = lower(r.proveedor) "
                )

                # Aplicar filtro por estado si no es 'Todos'
                estado_sel = None
                try:
                    estado_sel = estado_filter.get()
                except Exception:
                    estado_sel = "Todos"

                if estado_sel and estado_sel != "Todos":
                    sql += " WHERE r.estado = ?"
                    params.append(estado_sel)

                # Primero calcular total (COUNT) usando la misma subconsulta
                try:
                    count_sql = "SELECT COUNT(*) FROM (" + sql + ") as _sub"
                    cur.execute(count_sql, tuple(params))
                    raw_total = cur.fetchone()[0]
                    try:
                        total = int(raw_total) if raw_total is not None else 0
                    except Exception:
                        # Algunos adaptadores devuelven resultados como strings
                        try:
                            total = int(str(raw_total))
                        except Exception:
                            total = 0
                except Exception:
                    total = 0
                prov_page['total'] = int(total)

                # Aplicar orden y paginación (LIMIT/OFFSET)
                sql += f" ORDER BY p.proveedor {direction}"

                # Asegurar página válida
                try:
                    page_size = int(prov_page.get('page_size', 20))
                except Exception:
                    page_size = 20
                if page_size <= 0:
                    page_size = 20

                total_pages = max(1, (total + page_size - 1) // page_size)
                if prov_page['page'] > total_pages:
                    prov_page['page'] = total_pages
                if prov_page['page'] < 1:
                    prov_page['page'] = 1

                offset = (prov_page['page'] - 1) * page_size
                sql += " LIMIT ? OFFSET ?"
                params_with_limit = list(params) + [page_size, offset]

                cur.execute(sql, tuple(params_with_limit))
                rows = cur.fetchall()

                conn.close()

                # Encabezado simple
                header_row = ctk.CTkFrame(scroll)
                header_row.pack(fill="x", padx=5, pady=(0,5))
                # Ajuste de anchos: reducir PROVEEDOR a la mitad y dar espacio al HISTORIAL
                # PROVEEDOR ~180 (antes 360), ACCIONES ~260, HISTORIAL ~440 (antes 260)
                header_row.grid_columnconfigure(0, weight=1, minsize=180)
                header_row.grid_columnconfigure(1, weight=0, minsize=260)
                header_row.grid_columnconfigure(2, weight=0, minsize=440)
                # Columna extra para botón/ver historial
                header_row.grid_columnconfigure(3, weight=0, minsize=120)

                hf = ctk.CTkFont(weight="bold")
                lbl_nom = ctk.CTkLabel(header_row, text="RMP", font=hf, anchor="w", cursor="hand2")
                lbl_nom.grid(row=0, column=0, padx=5, sticky="w")
                ctk.CTkLabel(header_row, text="ACCIONES", font=hf, anchor="center").grid(row=0, column=1, padx=5)
                ctk.CTkLabel(header_row, text="HISTORIAL", font=hf, anchor="center").grid(row=0, column=2, padx=5)
                ctk.CTkLabel(header_row, text="VER", font=hf, anchor="center").grid(row=0, column=3, padx=5)

                # Alternar orden al pinchar encabezado
                def toggle_sort():
                    sort_state['dir'] = 'desc' if sort_state.get('dir', 'asc') == 'asc' else 'asc'
                    cargar_proveedores()

                lbl_nom.bind("<Button-1>", lambda e: toggle_sort())

                # Filas con estado editable (si es posible)
                colors = ("#FFFFFF", "#F3F4F6")
                # Actualizar etiqueta y estados de botones
                try:
                    total_pages = max(1, (prov_page['total'] + page_size - 1) // page_size)
                    prov_page_lbl.configure(text=f"Página {prov_page['page']} de {total_pages} ({prov_page['total']})")
                    prov_prev.configure(state="normal" if prov_page['page'] > 1 else "disabled")
                    prov_next.configure(state="normal" if prov_page['page'] < total_pages else "disabled")
                except Exception:
                    prov_page_lbl.configure(text=f"Página {prov_page.get('page',1)}")

                for idx, (prov, estado_actual, factura_actual) in enumerate(rows):
                    nombre = prov
                    # Usar fondo por defecto del tema en lugar de cebra
                    bg = "transparent"
                    row = ctk.CTkFrame(scroll, fg_color="transparent")
                    row.pack(fill="x", padx=5, pady=2)
                    # Mantener los minsize solicitados para columnas (consistentes con header)
                    row.grid_columnconfigure(0, weight=1, minsize=180)
                    row.grid_columnconfigure(1, weight=0, minsize=260)
                    row.grid_columnconfigure(2, weight=0, minsize=440)
                    row.grid_columnconfigure(3, weight=0, minsize=120)

                    lbl_nombre = ctk.CTkLabel(row, text=nombre or "-", anchor="w", cursor="hand2")
                    lbl_nombre.grid(row=0, column=0, padx=5, sticky="w")
                    opciones_estado = ["", "En Progreso", "Enviado", "Completado", "Exportado"]
                    try:
                        opt = ctk.CTkOptionMenu(row, values=opciones_estado)
                        # Si no hay estado, dejamos en vacío (primer opción)
                        opt.set(estado_actual if estado_actual in opciones_estado else "")
                    except Exception:
                        opt = ctk.CTkEntry(row)
                        opt.insert(0, estado_actual)
                    opt.grid(row=0, column=1, padx=5, sticky="w")

                    # Última entrada de historial (mostrar último comentario/estado corto)
                    try:
                        conn_h = connect_db()
                        cur_h = conn_h.cursor()
                        cur_h.execute(
                            "SELECT estado, comentario, usuario, fecha FROM rma_proveedor_hist WHERE lower(proveedor)=? OR proveedor=? ORDER BY fecha DESC LIMIT 1",
                            (nombre.lower(), nombre)
                        )
                        last_hist = cur_h.fetchone()
                        conn_h.close()
                    except Exception:
                        last_hist = None

                    
                    if not last_hist:
                        try:
                            import sqlite3 as _sqlite
                            local_db = os.path.join(os.path.dirname(__file__), DB_NAME)
                            if os.path.exists(local_db):
                                conn_local = _sqlite.connect(local_db)
                                cur_local = conn_local.cursor()
                                cur_local.execute(
                                    "SELECT estado, comentario, usuario, fecha FROM rma_proveedor_hist WHERE lower(proveedor)=? OR proveedor=? ORDER BY fecha DESC LIMIT 1",
                                    (nombre.lower(), nombre)
                                )
                                last_hist = cur_local.fetchone()
                                conn_local.close()
                        except Exception:
                            pass

                    hist_text = ""
                    if last_hist:
                        lh_estado, lh_coment, lh_user, lh_fecha = last_hist
                        if lh_coment:
                            hist_text = f"{lh_fecha} - {lh_user}: {lh_coment}"
                        else:
                            hist_text = f"{lh_fecha} - {lh_user}: {lh_estado}"

                    # Mostrar columna de historial con acceso al historial completo
                    # Mostrar el texto completo pero permitiendo wrap para no descuadrar columnas
                    hist_lbl_text = hist_text or ""
                    hist_lbl = ctk.CTkLabel(row, text=hist_lbl_text, anchor="w", cursor="hand2", wraplength=420)
                    hist_lbl.grid(row=0, column=2, padx=5, sticky="w")
                    hist_lbl.bind("<Button-1>", lambda e, n=nombre: mostrar_historial_proveedor(n))

                    # Botón para abrir la ventana completa de historial
                    try:
                        btn_hist = ctk.CTkButton(row, text="Ver historial", width=110, command=lambda n=nombre: mostrar_historial_proveedor(n))
                        btn_hist.grid(row=0, column=3, padx=5)
                    except Exception:
                        # Si falla la creación del botón, ignoramos para no romper el listado
                        pass

                    # Hover: usar fondo por defecto (transparent) para no añadir colores propios
                    def on_enter(e, r=row):
                        try:
                            r.configure(fg_color="transparent")
                        except Exception:
                            pass
                    def on_leave(e, r=row, original_bg=bg):
                        try:
                            r.configure(fg_color="transparent")
                        except Exception:
                            pass

                    row.bind("<Enter>", on_enter)
                    row.bind("<Leave>", on_leave)
                    lbl_nombre.bind("<Enter>", on_enter)
                    lbl_nombre.bind("<Leave>", on_leave)

                    # Click en nombre abre detalle del proveedor
                    lbl_nombre.bind("<Button-1>", lambda e, nombre=nombre, estado=estado_actual, factura=factura_actual: mostrar_expedientes_proveedor(nombre, estado, factura))

                    # Persistir cambio de estado (upsert) y anotar en historial
                    def make_state_updater(proveedor_nombre, prev_estado):
                        def updater(selected_value):
                            val = selected_value
                            # Si no hay cambio, no hacemos nada
                            try:
                                if (prev_estado is not None) and (str(val) == str(prev_estado)):
                                    return
                            except Exception:
                                pass
                            try:
                                conn3 = connect_db()
                                cur3 = conn3.cursor()
                                # Usar upsert: insertar o actualizar el estado para el proveedor
                                try:
                                    cur3.execute(
                                        "INSERT INTO rma_proveedor (proveedor, estado) VALUES (?, ?) ON CONFLICT(proveedor) DO UPDATE SET estado=excluded.estado",
                                        (proveedor_nombre, val)
                                    )
                                except sqlite3.Error:
                                    # Fallback: intentar UPDATE, si no existe, INSERT
                                    cur3.execute("UPDATE rma_proveedor SET estado = ? WHERE proveedor = ?", (val, proveedor_nombre))
                                    if getattr(cur3, 'rowcount', 0) == 0:
                                        cur3.execute("INSERT INTO rma_proveedor (proveedor, estado) VALUES (?, ?)", (proveedor_nombre, val))

                                # Añadir entrada en historial de proveedor
                                try:
                                    # Asegurar tabla existe
                                    cur3.execute("CREATE TABLE IF NOT EXISTS rma_proveedor_hist (id INTEGER PRIMARY KEY, proveedor TEXT, estado TEXT, comentario TEXT, usuario TEXT, fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
                                except Exception:
                                    pass
                                try:
                                    usuario = getattr(self, 'username', 'unknown')
                                    # Guardar comentario explicativo del cambio de estado
                                    if prev_estado is None or prev_estado == "":
                                        comentario_text = f"Estado establecido a: {val}"
                                    else:
                                        comentario_text = f"Cambio de estado de '{prev_estado}' a '{val}'"
                                    cur3.execute("INSERT INTO rma_proveedor_hist (proveedor, estado, comentario, usuario) VALUES (?, ?, ?, ?)", (proveedor_nombre, val, comentario_text, usuario))
                                except Exception:
                                    # Si falla en este backend, intentamos igual con conexión nueva/SQLite local (no preferido)
                                    pass

                                try:
                                    conn3.commit()
                                except Exception:
                                    pass
                                try:
                                    conn3.close()
                                except Exception:
                                    pass

                                # Refrescar la lista para mostrar el último historial
                                try:
                                    cargar_proveedores()
                                except Exception:
                                    pass
                            except Exception as e:
                                messagebox.showerror("Error BD", f"No se pudo guardar el estado: {e}")
                        return updater

                    try:
                        # CTkOptionMenu passes the selected value to the command
                        opt.configure(command=make_state_updater(nombre, estado_actual))
                    except Exception:
                        if isinstance(opt, ctk.CTkEntry):
                            opt.bind("<FocusOut>", lambda e, nombre=nombre, w=opt: make_state_updater(nombre, estado_actual)(w.get()))

            except sqlite3.Error as e:
                messagebox.showerror("Error BD", f"No se pudo cargar lista de proveedores: {e}")

        def mostrar_expedientes_proveedor(proveedor_nombre, estado_actual='', factura_actual=''):
            # Ventana detallada del proveedor: info, expedientes, historial y comentarios
            vent = ctk.CTkToplevel(self)
            vent.title(f"Detalle RMP - {proveedor_nombre}")
            vent.geometry("1200x900")
            
            # Configurar para permitir minimización
            vent.resizable(True, True)
            vent.attributes('-topmost', False)
            vent.minsize(900, 600)
            # No usar transient para permitir minimización completa
            vent.focus_set()  # Dar foco sin bloquear
            
            # Forzar aparición al frente (incluso si la principal está maximizada)
            vent.attributes('-topmost', True)   # Temporalmente al frente
            vent.lift()
            try:
                vent.focus_force()
            except:
                pass
            
            def quitar_topmost_vent():
                try:
                    if vent.winfo_exists():
                        vent.attributes('-topmost', False)
                except:
                    pass
            
            vent.after(500, quitar_topmost_vent)

            cont = ctk.CTkFrame(vent)
            cont.pack(fill="both", expand=True, padx=12, pady=12)

            # ===== SECCIÓN 1: ENCABEZADO CON INFORMACIÓN DEL PROVEEDOR =====
            header_frame = ctk.CTkFrame(cont, fg_color="#4A90E2", corner_radius=8)
            header_frame.pack(fill="x", pady=(0,10))
            header_frame.grid_columnconfigure(0, weight=1)
            header_frame.grid_columnconfigure(1, weight=1)
            header_frame.grid_columnconfigure(2, weight=1)

            ctk.CTkLabel(header_frame, text="RMP", font=ctk.CTkFont(size=11, weight="bold"), text_color="white").grid(row=0, column=0, padx=15, pady=(8,2), sticky="w")
            ctk.CTkLabel(header_frame, text=proveedor_nombre, font=ctk.CTkFont(size=14), text_color="white").grid(row=1, column=0, padx=15, pady=(0,8), sticky="w")

            ctk.CTkLabel(header_frame, text="ESTADO", font=ctk.CTkFont(size=11, weight="bold"), text_color="white").grid(row=0, column=1, padx=15, pady=(8,2), sticky="w")
            
            # Variable para rastrear el estado actual
            estado_var = {'actual': estado_actual or ''}
            
            # CTkOptionMenu para editar el estado
            opciones_estado = ["", "En Progreso", "Enviado", "Completado", "Exportado"]
            estado_menu = ctk.CTkOptionMenu(
                header_frame, 
                values=opciones_estado,
                fg_color="white",
                button_color="#4A90E2",
                button_hover_color="#357ABD",
                text_color="#212529",
                dropdown_fg_color="white",
                dropdown_text_color="#212529"
            )
            estado_menu.set(estado_actual if estado_actual in opciones_estado else "")
            estado_menu.grid(row=1, column=1, padx=15, pady=(0,8), sticky="ew")

            ctk.CTkLabel(header_frame, text="FACTURA ABONO", font=ctk.CTkFont(size=11, weight="bold"), text_color="white").grid(row=0, column=2, padx=15, pady=(8,2), sticky="w")
            factura_entry = ctk.CTkEntry(header_frame, placeholder_text="Ej: FA2025001")
            factura_entry.insert(0, factura_actual or "")
            factura_entry.grid(row=1, column=2, padx=15, pady=(0,8), sticky="ew")

            # Función para actualizar estado
            def actualizar_estado(nuevo_estado):
                estado_anterior = estado_var['actual']
                # Si no hay cambio, no hacemos nada
                if estado_anterior == nuevo_estado:
                    return
                
                try:
                    conn_e = connect_db()
                    cur_e = conn_e.cursor()
                    # Actualizar estado en la tabla
                    try:
                        cur_e.execute(
                            "INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?) ON CONFLICT(proveedor) DO UPDATE SET estado=excluded.estado",
                            (proveedor_nombre, nuevo_estado, factura_actual or '')
                        )
                    except Exception:
                        # Fallback para backends sin ON CONFLICT
                        cur_e.execute("UPDATE rma_proveedor SET estado = ? WHERE proveedor = ?", (nuevo_estado, proveedor_nombre))
                        if getattr(cur_e, 'rowcount', 0) == 0:
                            cur_e.execute("INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?)", (proveedor_nombre, nuevo_estado, factura_actual or ''))
                    
                    # Añadir al historial
                    usuario = getattr(self, 'username', 'unknown')
                    if estado_anterior == '' or estado_anterior is None:
                        comentario_text = f"Estado establecido a: {nuevo_estado}"
                    else:
                        comentario_text = f"Cambio de estado de '{estado_anterior}' a '{nuevo_estado}'"
                    
                    cur_e.execute(
                        "INSERT INTO rma_proveedor_hist (proveedor, estado, comentario, usuario) VALUES (?, ?, ?, ?)",
                        (proveedor_nombre, nuevo_estado, comentario_text, usuario)
                    )
                    
                    conn_e.commit()
                    conn_e.close()
                    
                    # Actualizar variable local
                    estado_var['actual'] = nuevo_estado
                    
                    # Refrescar historial en esta ventana
                    cargar_historial()
                    
                    # Refrescar lista principal
                    try:
                        cargar_proveedores()
                    except Exception:
                        pass
                    
                    messagebox.showinfo("Actualizado", "Estado actualizado correctamente.")
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo actualizar el estado: {e}")
            
            # Configurar el comando del menú de estado
            estado_menu.configure(command=actualizar_estado)

            # Botón para guardar cambios de factura_abono
            def guardar_factura():
                nueva_factura = factura_entry.get().strip()
                try:
                    conn_f = connect_db()
                    cur_f = conn_f.cursor()
                    # Actualizar factura_abono (usar el estado actual de estado_var)
                    try:
                        cur_f.execute(
                            "INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?) ON CONFLICT(proveedor) DO UPDATE SET factura_abono=excluded.factura_abono",
                            (proveedor_nombre, estado_var['actual'], nueva_factura)
                        )
                    except Exception:
                        # Fallback para backends sin ON CONFLICT
                        cur_f.execute("UPDATE rma_proveedor SET factura_abono = ? WHERE proveedor = ?", (nueva_factura, proveedor_nombre))
                        if getattr(cur_f, 'rowcount', 0) == 0:
                            cur_f.execute("INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?)", (proveedor_nombre, estado_var['actual'], nueva_factura))
                    conn_f.commit()
                    conn_f.close()
                    messagebox.showinfo("Guardado", "Factura de abono actualizada correctamente.")
                    # Refrescar lista principal
                    try:
                        cargar_proveedores()
                    except Exception:
                        pass
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo guardar la factura: {e}")

            ctk.CTkButton(header_frame, text="💾 Guardar Factura", command=guardar_factura, width=130).grid(row=1, column=3, padx=(5,15), pady=(0,8))

            # ===== SECCIÓN 2: LISTADO DE EXPEDIENTES =====
            exp_frame = ctk.CTkFrame(cont)
            exp_frame.pack(fill="both", expand=True, pady=(0,10))

            ctk.CTkLabel(exp_frame, text="Expedientes Asociados", font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=5, pady=(5,5))

            sf_exp = ctk.CTkScrollableFrame(exp_frame, height=250)
            sf_exp.pack(fill="both", expand=True, padx=5, pady=(0,5))

            # Cargar expedientes del proveedor
            try:
                conn = connect_db()
                cur = conn.cursor()
                cur.execute(
                    "SELECT id, codigo_rma, cliente, numero_documento_cliente, modelo, ref_proveedor, fecha_emision, estado "
                    "FROM rma_maestro WHERE lower(Rma_Proveedor)=? OR Rma_Proveedor=? ORDER BY fecha_emision DESC",
                    (proveedor_nombre.lower(), proveedor_nombre)
                )
                filas = cur.fetchall()
                conn.close()
            except Exception as e:
                messagebox.showerror("Error BD", f"No se pudieron cargar expedientes: {e}")
                filas = []

            # Función para exportar a Excel
            def export_to_excel():
                try:
                    if not filas:
                        messagebox.showinfo('Exportar', 'No hay expedientes para exportar.')
                        return
                    
                    data = []
                    for r in filas:
                        (_id, codigo_rma, cliente, num_doc, modelo, ref_prov, fecha_emision, estado) = r
                        data.append({
                            'Nº Expediente': codigo_rma,
                            'Proveedor': proveedor_nombre,
                            'Cliente': cliente or '',
                            'Numero Documento Cliente': num_doc or '',
                            'Descripcion Articulo': modelo or '',
                            'Referencia': ref_prov or ''
                        })

                    df = pd.DataFrame(data)
                    base_dir = os.path.join(os.path.dirname(__file__), 'Adjuntos_RMA')
                    rmp_dir = os.path.join(base_dir, 'RMP')
                    os.makedirs(rmp_dir, exist_ok=True)

                    safe_name = ''.join(c for c in proveedor_nombre if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    safe_name = safe_name.replace(' ', '_')
                    file_path = os.path.join(rmp_dir, f"{safe_name}.xlsx")

                    if os.path.exists(file_path):
                        if not messagebox.askyesno('Exportar', f'El archivo {os.path.basename(file_path)} ya existe. ¿Desea sobreescribirlo?'):
                            return

                    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                        df.to_excel(writer, index=False, sheet_name='Expedientes')
                        workbook = writer.book
                        worksheet = writer.sheets['Expedientes']
                        from openpyxl.utils import get_column_letter
                        for i, col in enumerate(df.columns):
                            if df.empty:
                                max_len = len(str(col))
                            else:
                                col_max = df[col].astype(str).map(len).max()
                                max_len = max(int(col_max) if col_max is not None else 0, len(str(col)))
                            adjusted_width = min(max_len + 2, 60)
                            worksheet.column_dimensions[get_column_letter(i+1)].width = adjusted_width

                    messagebox.showinfo('Exportar', f'Exportado correctamente: {file_path}')

                    # Subir archivo a Dropbox
                    if usar_dropbox():
                        try:
                            # Crear ruta en Dropbox: /RMP/{nombre_proveedor}.xlsx
                            dropbox_path = f"/RMP/{safe_name}.xlsx"
                            
                            dbx_client = get_dropbox_client()
                            with open(file_path, 'rb') as f:
                                # Subir archivo a Dropbox (sobreescribir si existe)
                                dbx_client.files_upload(
                                    f.read(),
                                    dropbox_path,
                                    mode=dropbox.files.WriteMode('overwrite')
                                )
                            
                            print(f"✅ Excel RMP subido a Dropbox: {dropbox_path}")
                            # Opcional: mostrar confirmación al usuario
                            # messagebox.showinfo('Dropbox', f'Archivo también guardado en Dropbox: {dropbox_path}')
                            
                        except Exception as e:
                            print(f"⚠️ Error subiendo Excel RMP a Dropbox: {e}")
                            # No mostrar error al usuario para no interrumpir el flujo
                    else:
                        print("ℹ️ Dropbox no configurado, Excel solo guardado localmente")

                    # Añadir a historial
                    try:
                        connh = connect_db()
                        curh = connh.cursor()
                        rma_codes = [str(r[1]) for r in filas if len(r) > 1 and r[1] is not None]
                        count = len(filas)
                        codes_str = ', '.join(rma_codes)
                        if len(codes_str) > 500:
                            codes_str = codes_str[:500] + '...'
                        comentario = f'Exportado {count} expedientes a Excel: {os.path.basename(file_path)}'
                        if codes_str:
                            comentario += f' (RMAs: {codes_str})'
                        usuario = getattr(self, 'username', 'unknown')
                        curh.execute(
                            "INSERT INTO rma_proveedor_hist (proveedor, estado, comentario, usuario) VALUES (?, ?, ?, ?)",
                            (proveedor_nombre, 'Exportado', comentario, usuario)
                        )
                        try:
                            curh.execute(
                                "INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?) ON CONFLICT(proveedor) DO UPDATE SET estado=excluded.estado",
                                (proveedor_nombre, 'Exportado', factura_actual or '')
                            )
                        except Exception:
                            curh.execute("UPDATE rma_proveedor SET estado = ? WHERE proveedor = ?", ('Exportado', proveedor_nombre))
                            if getattr(curh, 'rowcount', 0) == 0:
                                curh.execute("INSERT INTO rma_proveedor (proveedor, estado, factura_abono) VALUES (?, ?, ?)", (proveedor_nombre, 'Exportado', factura_actual or ''))
                        connh.commit()
                        connh.close()
                        # Refrescar historial en esta ventana
                        cargar_historial()
                        cargar_proveedores()
                    except Exception as e:
                        print(f'Warning: error en historial: {e}')
                except Exception as e:
                    messagebox.showerror('Exportar', f'Error exportando a Excel: {e}')

            # Botón exportar
            btn_frame = ctk.CTkFrame(exp_frame)
            btn_frame.pack(fill="x", padx=5, pady=(0,5))
            ctk.CTkButton(btn_frame, text="📊 Exportar a Excel", command=export_to_excel, width=150).pack(side="right")

            # Encabezado de expedientes
            head = ctk.CTkFrame(sf_exp)
            head.pack(fill="x", padx=5, pady=(0,5))
            head.grid_columnconfigure(0, weight=1, minsize=150)
            head.grid_columnconfigure(1, weight=2, minsize=250)
            head.grid_columnconfigure(2, weight=1, minsize=120)
            head.grid_columnconfigure(3, weight=1, minsize=120)
            head.grid_columnconfigure(4, weight=0, minsize=80)

            hf = ctk.CTkFont(weight="bold")
            ctk.CTkLabel(head, text="CÓDIGO", font=hf).grid(row=0, column=0, padx=5, sticky="w")
            ctk.CTkLabel(head, text="CLIENTE", font=hf).grid(row=0, column=1, padx=5, sticky="w")
            ctk.CTkLabel(head, text="FECHA", font=hf).grid(row=0, column=2, padx=5, sticky="w")
            ctk.CTkLabel(head, text="ESTADO", font=hf).grid(row=0, column=3, padx=5, sticky="w")
            ctk.CTkLabel(head, text="", font=hf).grid(row=0, column=4, padx=5)

            # Filas de expedientes
            colors = ("#FFFFFF", "#F7F7F7")
            for idx, r in enumerate(filas):
                try:
                    rma_id, codigo, cliente, num_doc, modelo, ref_prov, fecha, estado = r
                except Exception:
                    vals = list(r)
                    rma_id = vals[0] if len(vals) > 0 else None
                    codigo = vals[1] if len(vals) > 1 else ''
                    cliente = vals[2] if len(vals) > 2 else ''
                    fecha = vals[6] if len(vals) > 6 else ''
                    estado = vals[7] if len(vals) > 7 else ''

                # Usar fondo por defecto del tema en lugar de cebra
                bg = "transparent"
                row = ctk.CTkFrame(sf_exp, fg_color="transparent")
                row.pack(fill="x", padx=5, pady=2)
                row.grid_columnconfigure(0, weight=1, minsize=150)
                row.grid_columnconfigure(1, weight=2, minsize=250)
                row.grid_columnconfigure(2, weight=1, minsize=120)
                row.grid_columnconfigure(3, weight=1, minsize=120)
                row.grid_columnconfigure(4, weight=0, minsize=80)

                lbl_codigo = ctk.CTkLabel(row, text=codigo, anchor="w", cursor="hand2")
                lbl_codigo.grid(row=0, column=0, padx=5, sticky="w")
                lbl_cliente = ctk.CTkLabel(row, text=cliente if cliente else "-", anchor="w")
                lbl_cliente.grid(row=0, column=1, padx=5, sticky="w")
                lbl_fecha = ctk.CTkLabel(row, text=fecha if fecha else "-", anchor="w")
                lbl_fecha.grid(row=0, column=2, padx=5, sticky="w")
                lbl_estado = ctk.CTkLabel(row, text=estado if estado else "-", anchor="w")
                lbl_estado.grid(row=0, column=3, padx=5, sticky="w")

                ctk.CTkButton(row, text="Editar", width=70, command=lambda rid=rma_id: (self.mostrar_nuevo_rma(rma_id=rid), vent.destroy())).grid(row=0, column=4, padx=5)

                # Hover
                def on_enter(e, r=row):
                    try:
                        r.configure(fg_color="transparent")
                    except Exception:
                        pass
                def on_leave(e, r=row, original_bg=bg):
                    try:
                        r.configure(fg_color="transparent")
                    except Exception:
                        pass

                row.bind("<Enter>", on_enter)
                row.bind("<Leave>", on_leave)

                # Doble clic abre editor
                row.bind("<Double-Button-1>", lambda e, rid=rma_id: (self._abrir_editor_rma(rma_id=rid), vent.destroy()))
                lbl_codigo.bind("<Double-Button-1>", lambda e, rid=rma_id: (self._abrir_editor_rma(rma_id=rid), vent.destroy()))

            # ===== SECCIÓN 3: HISTORIAL DEL PROVEEDOR =====
            hist_frame = ctk.CTkFrame(cont)
            hist_frame.pack(fill="both", expand=True, pady=(0,10))

            ctk.CTkLabel(hist_frame, text="Historial del Proveedor", font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=5, pady=(5,5))

            sf_hist = ctk.CTkScrollableFrame(hist_frame, height=180)
            sf_hist.pack(fill="both", expand=True, padx=5, pady=(0,5))

            def cargar_historial():
                for w in sf_hist.winfo_children():
                    w.destroy()
                try:
                    conn_h = connect_db()
                    cur_h = conn_h.cursor()
                    cur_h.execute(
                        "SELECT fecha, usuario, estado, comentario FROM rma_proveedor_hist WHERE lower(proveedor)=? OR proveedor=? ORDER BY fecha DESC",
                        (proveedor_nombre.lower(), proveedor_nombre)
                    )
                    hist_rows = cur_h.fetchall()
                    conn_h.close()

                    if not hist_rows:
                        ctk.CTkLabel(sf_hist, text="No hay historial registrado.", text_color="gray").pack(anchor="w", padx=5, pady=10)
                    else:
                        for idx, (fecha, usuario, estado_h, comentario) in enumerate(hist_rows):
                            # Usar fondo por defecto del tema en lugar de filas cebra
                            rowf = ctk.CTkFrame(sf_hist, fg_color="transparent", corner_radius=6)
                            rowf.pack(fill="x", padx=3, pady=3)
                            txt = f"📅 {fecha} | 👤 {usuario}"
                            if estado_h:
                                txt += f" | 🏷️ {estado_h}"
                            ctk.CTkLabel(rowf, text=txt, font=ctk.CTkFont(weight="bold", size=11)).pack(anchor="w", padx=8, pady=(6,2))
                            if comentario:
                                ctk.CTkLabel(rowf, text=comentario, wraplength=1100, anchor="w").pack(anchor="w", padx=8, pady=(0,6))
                except Exception as e:
                    ctk.CTkLabel(sf_hist, text=f"Error cargando historial: {e}", text_color="red").pack(anchor="w", padx=5, pady=10)

            cargar_historial()

            # ===== SECCIÓN 4: AÑADIR COMENTARIO =====
            comment_frame = ctk.CTkFrame(cont)
            comment_frame.pack(fill="x")

            ctk.CTkLabel(comment_frame, text="Añadir comentario al historial:", font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=5, pady=(5,3))
            comment_box = ctk.CTkTextbox(comment_frame, height=60)
            comment_box.pack(fill="x", padx=5, pady=(0,5))

            def add_comment():
                text = comment_box.get("0.0", "end").strip()
                if not text:
                    messagebox.showwarning("Vacío", "Escribe un comentario antes de añadirlo.")
                    return
                try:
                    connc = connect_db()
                    curc = connc.cursor()
                    curc.execute("CREATE TABLE IF NOT EXISTS rma_proveedor_hist (id INTEGER PRIMARY KEY, proveedor TEXT, estado TEXT, comentario TEXT, usuario TEXT, fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
                    curc.execute("INSERT INTO rma_proveedor_hist (proveedor, estado, comentario, usuario) VALUES (?, ?, ?, ?)", 
                                (proveedor_nombre, '', text, getattr(self, 'username', 'unknown')))
                    connc.commit()
                    connc.close()
                    messagebox.showinfo("Añadido", "Comentario añadido al historial.")
                    comment_box.delete("0.0", "end")
                    cargar_historial()
                    cargar_proveedores()
                except Exception as e:
                    messagebox.showerror("Error BD", f"No se pudo añadir el comentario: {e}")

            ctk.CTkButton(comment_frame, text="💬 Añadir Comentario", command=add_comment, width=180).pack(anchor="e", padx=5, pady=(0,5))

        # Vincular búsqueda
        btn_buscar.configure(command=cargar_proveedores)
        try:
            estado_filter.configure(command=lambda v=None: cargar_proveedores())
        except Exception:
            pass
        entry_buscar.bind("<Return>", lambda e: cargar_proveedores())

        # Cargar inicialmente
        cargar_proveedores()

    def mostrar_articulos_window(self):
        """Muestra una ventana con el listado de artículos y la cantidad de expedientes relacionados."""
        # Evitar múltiples instancias
        if hasattr(self, 'articulos_window') and getattr(self, 'articulos_window').winfo_exists():
            getattr(self, 'articulos_window').focus()
            return

        self.articulos_window = ctk.CTkToplevel(self)
        win = self.articulos_window
        win.title("Artículos")
        win.geometry("800x600")
        
        # Configurar para permitir minimización
        win.resizable(True, True)
        win.attributes('-topmost', False)
        win.minsize(600, 400)
        # No usar transient ni grab_set para permitir minimización completa
        win.focus_set()  # Dar foco sin bloquear
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        win.attributes('-topmost', True)   # Temporalmente al frente
        win.lift()
        win.focus_force()
        win.after(500, lambda: win.attributes('-topmost', False))  # Quitar topmost después de 500ms

        main = ctk.CTkFrame(win)
        main.pack(fill="both", expand=True, padx=12, pady=12)

        header = ctk.CTkFrame(main)
        header.pack(fill="x", pady=(0,8))
        ctk.CTkLabel(header, text="Listado de Artículos", font=ctk.CTkFont(size=18, weight="bold")).grid(row=0, column=0, sticky="w")

        # Frame para la lista con scroll (con búsqueda, paginación y carga en background)
        list_frame = ctk.CTkFrame(main)
        list_frame.pack(fill="both", expand=True)

        canvas = ctk.CTkCanvas(list_frame, borderwidth=0, highlightthickness=0)
        try:
            from tkinter import Canvas as _Canvas
            canvas = _Canvas(list_frame, borderwidth=0, highlightthickness=0)
        except Exception:
            pass

        sb = ctk.CTkScrollbar(list_frame, orientation="vertical", command=lambda *args: canvas.yview(*args))
        canvas.configure(yscrollcommand=lambda *args: sb.set(*args))
        sb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        sf = ctk.CTkFrame(canvas)
        try:
            window_id = canvas.create_window((0,0), window=sf, anchor="nw")
        except Exception:
            window_id = canvas.create_window((0,0), window=sf, anchor="nw")

        def on_sf_configure(event):
            try:
                canvas.configure(scrollregion=canvas.bbox("all"))
            except Exception:
                pass

        def on_canvas_config(event):
            try:
                canvas.itemconfig(window_id, width=event.width)
            except Exception:
                pass

        sf.bind("<Configure>", on_sf_configure)
        canvas.bind("<Configure>", on_canvas_config)
        
        # Habilitar scroll con rueda del mouse
        def on_mousewheel(event):
            try:
                canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except Exception:
                pass
        
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        # Limpiar binding cuando se cierre la ventana
        win.bind("<Destroy>", lambda e: canvas.unbind_all("<MouseWheel>"))

        # Controles de búsqueda y paginación en el header
        search_var = tk.StringVar()
        page_state = {"page": 1, "total_pages": 1}

        ctk.CTkLabel(header, text="Buscar:").grid(row=1, column=0, sticky="w", pady=(6,0))
        search_entry = ctk.CTkEntry(header, textvariable=search_var, placeholder_text="Referencia...", width=300)
        search_entry.grid(row=1, column=1, sticky="w", padx=(8,0), pady=(6,0))

        page_size_opt = ctk.CTkOptionMenu(header, values=["10","25","50","100"])
        page_size_opt.set("50")
        page_size_opt.grid(row=1, column=2, padx=(12,0), pady=(6,0))

        btn_prev = ctk.CTkButton(header, text="◀", width=40)
        btn_prev.grid(row=1, column=3, padx=(12,2), pady=(6,0))
        page_label = ctk.CTkLabel(header, text="Página 1/1")
        page_label.grid(row=1, column=4, padx=(2,6), pady=(6,0))
        btn_next = ctk.CTkButton(header, text="▶", width=40)
        btn_next.grid(row=1, column=5, padx=(2,0), pady=(6,0))

        # Progress indicator removed for cleaner UI (was showing next to pagination controls)
        # If needed in the future, re-add a subtle indicator and start/stop it when loading asynchronously.

        # Container for rows so we can clear it on reload
        rows_container = ctk.CTkFrame(sf)
        rows_container.pack(fill="both", expand=True)

        def render_rows(filas, page, total_count, page_size):
            # Clear previous rows
            for w in rows_container.winfo_children():
                w.destroy()

            hf = ctk.CTkFont(weight="bold")
            header_frame = ctk.CTkFrame(rows_container)
            header_frame.pack(fill="x", padx=5, pady=(0,4))
            header_frame.grid_columnconfigure(0, weight=3, minsize=300)
            header_frame.grid_columnconfigure(1, weight=1, minsize=80)
            ctk.CTkLabel(header_frame, text="REFERENCIA ARTÍCULO", font=hf).grid(row=0, column=0, padx=5, sticky="w")
            ctk.CTkLabel(header_frame, text="EXPEDIENTES", font=hf).grid(row=0, column=1, padx=5, sticky="w")

            colors = ("#FFFFFF", "#F3F4F6")
            for idx, row in enumerate(filas):
                try:
                    referencia, cnt = row[0], row[1]
                except Exception:
                    vals = list(row)
                    referencia = vals[0] if len(vals) > 0 else ''
                    cnt = vals[1] if len(vals) > 1 else 0

                # Usar fondo por defecto del tema en lugar de cebra
                bg = "transparent"
                rf = ctk.CTkFrame(rows_container, fg_color="transparent")
                rf.pack(fill="x", padx=5, pady=2)
                rf.grid_columnconfigure(0, weight=3, minsize=300)
                rf.grid_columnconfigure(1, weight=1, minsize=80)

                lbl_ref = ctk.CTkLabel(rf, text=referencia or '-', anchor="w", cursor="hand2")
                lbl_ref.grid(row=0, column=0, padx=5, sticky="w")
                lbl_cnt = ctk.CTkLabel(rf, text=str(cnt), anchor="w")
                lbl_cnt.grid(row=0, column=1, padx=5, sticky="w")

                ctk.CTkButton(rf, text="Ver Expedientes", width=140, command=lambda r=referencia: self.mostrar_expedientes_por_articulo(r)).grid(row=0, column=2, padx=6)

                def on_enter(e, w=rf):
                    try:
                        w.configure(fg_color="transparent")
                    except Exception:
                        pass
                def on_leave(e, w=rf, original=bg):
                    try:
                        w.configure(fg_color="transparent")
                    except Exception:
                        pass

                rf.bind("<Enter>", on_enter)
                rf.bind("<Leave>", on_leave)
                # Doble clic ahora muestra los estados del artículo con sus totales
                lbl_ref.bind("<Double-Button-1>", lambda e, r=referencia: self.mostrar_estados_por_articulo(r))

            # Ensure numeric types for pagination calculation
            try:
                total_count = int(total_count or 0)
            except Exception:
                try:
                    total_count = int(float(total_count))
                except Exception:
                    total_count = 0
            try:
                page_size = int(page_size)
            except Exception:
                page_size = 50

            # Update page info label
            total_pages = max(1, (total_count + page_size - 1) // page_size)
            page_state["page"] = page
            page_state["total_pages"] = total_pages
            page_label.configure(text=f"Página {page}/{total_pages}")

        def load_articles_thread(page=1):
            try:
                # Progressbar removed from UI; directly schedule data rendering when ready.
                search = search_var.get().strip()
                page_size = int(page_size_opt.get())
                offset = (page - 1) * page_size

                conn = connect_db()
                cur = conn.cursor()

                # Count total
                try:
                    if search:
                        cur.execute("SELECT COUNT(DISTINCT referencia_articulo) FROM rma_detalles WHERE referencia_articulo LIKE ?", (f"%{search}%",))
                    else:
                        cur.execute("SELECT COUNT(DISTINCT referencia_articulo) FROM rma_detalles")
                    total = cur.fetchone()[0]
                    total = int(total) if total is not None else 0
                except Exception:
                    total = 0

                # Main query with limit/offset
                params = []
                where = ""
                if search:
                    where = "WHERE referencia_articulo LIKE ?"
                    params.append(f"%{search}%")

                sql = f"SELECT referencia_articulo, COUNT(DISTINCT rma_maestro.id) as expedientes_count FROM rma_detalles INNER JOIN rma_maestro ON rma_detalles.rma_id = rma_maestro.id {where} GROUP BY referencia_articulo ORDER BY expedientes_count DESC, referencia_articulo ASC LIMIT ? OFFSET ?"
                params.extend([page_size, offset])
                cur.execute(sql, tuple(params))
                filas = cur.fetchall()
                conn.close()
                # Schedule UI update (render rows directly)
                self.after(0, lambda: render_rows(filas, page, total, page_size))
            except Exception as e:
                # No progressbar to stop; show error to user
                self.after(0, lambda: messagebox.showerror("Error BD", f"No se pudieron cargar los artículos: {e}"))

        def start_load(page=1):
            threading.Thread(target=load_articles_thread, args=(page,), daemon=True).start()

        def goto_prev():
            p = max(1, page_state.get("page", 1) - 1)
            start_load(p)

        def goto_next():
            p = page_state.get("page", 1) + 1
            if page_state.get("total_pages", 1) and p > page_state.get("total_pages"):
                return
            start_load(p)

        btn_prev.configure(command=goto_prev)
        btn_next.configure(command=goto_next)
        search_entry.bind("<Return>", lambda e: start_load(1))

        # Cargar inicialmente
        start_load(1)

    def mostrar_estados_por_articulo(self, referencia):
        """Muestra una ventana con los estados que ha tenido un artículo, con suma de cantidades y total en euros por estado."""
        if not referencia:
            messagebox.showinfo("Info", "Referencia vacía.")
            return

        from lib.articulo_utils import VentanaEstadosArticulo
        VentanaEstadosArticulo(self, referencia, connect_db)

    def mostrar_expedientes_por_articulo(self, referencia):
        """Muestra una ventana con los expedientes asociados a una referencia de artículo."""
        if not referencia:
            messagebox.showinfo("Info", "Referencia vacía.")
            return

        # Evitar múltiples instancias por la misma referencia
        name = f"exp_{referencia}"
        # No strict unique naming; we just open a new window
        vent = ctk.CTkToplevel(self)
        vent.title(f"Expedientes - {referencia}")
        vent.geometry("900x600")
        
        # Configurar para permitir minimización
        vent.resizable(True, True)
        vent.attributes('-topmost', False)
        vent.minsize(700, 450)
        # No usar transient para permitir minimización completa
        try:
            vent.focus_set()  # Dar foco sin bloquear
        except:
            pass
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        vent.attributes('-topmost', True)   # Temporalmente al frente
        vent.lift()
        try:
            vent.focus_force()
        except:
            pass
        
        # Función segura para quitar topmost
        def quitar_topmost():
            try:
                if vent.winfo_exists():
                    vent.attributes('-topmost', False)
            except:
                pass
        
        vent.after(500, quitar_topmost)

        main = ctk.CTkFrame(vent)
        main.pack(fill="both", expand=True, padx=12, pady=12)

        header = ctk.CTkFrame(main)
        header.pack(fill="x", pady=(0,8))
        ctk.CTkLabel(header, text=f"Expedientes asociados a: {referencia}", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, sticky="w")

        list_frame = ctk.CTkFrame(main)
        list_frame.pack(fill="both", expand=True)

        canvas = ctk.CTkCanvas(list_frame, borderwidth=0, highlightthickness=0)
        try:
            from tkinter import Canvas as _Canvas
            canvas = _Canvas(list_frame, borderwidth=0, highlightthickness=0)
        except Exception:
            pass
        sb = ctk.CTkScrollbar(list_frame, orientation="vertical", command=lambda *args: canvas.yview(*args))
        canvas.configure(yscrollcommand=lambda *args: sb.set(*args))
        sb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        sf = ctk.CTkFrame(canvas)
        try:
            window_id = canvas.create_window((0,0), window=sf, anchor="nw")
        except Exception:
            window_id = canvas.create_window((0,0), window=sf, anchor="nw")

        def on_sf_config(event):
            try:
                canvas.configure(scrollregion=canvas.bbox("all"))
            except Exception:
                pass

        def on_canvas_cfg(event):
            try:
                canvas.itemconfig(window_id, width=event.width)
            except Exception:
                pass

        sf.bind("<Configure>", on_sf_config)
        canvas.bind("<Configure>", on_canvas_cfg)

        # Añadimos búsqueda, paginación y carga en background para la lista de expedientes
        search_var = tk.StringVar()
        page_state = {"page": 1, "total_pages": 1}

        ctk.CTkLabel(header, text="Buscar (RMA/Cliente/Doc):").grid(row=1, column=0, sticky="w", pady=(6,0))
        search_entry = ctk.CTkEntry(header, textvariable=search_var, placeholder_text="Texto a buscar...", width=350)
        search_entry.grid(row=1, column=1, sticky="w", padx=(8,0), pady=(6,0))

        page_size_opt = ctk.CTkOptionMenu(header, values=["10","25","50","100"])
        page_size_opt.set("50")
        page_size_opt.grid(row=1, column=2, padx=(12,0), pady=(6,0))

        btn_prev = ctk.CTkButton(header, text="◀", width=40)
        btn_prev.grid(row=1, column=3, padx=(12,2), pady=(6,0))
        page_label = ctk.CTkLabel(header, text="Página 1/1")
        page_label.grid(row=1, column=4, padx=(2,6), pady=(6,0))
        btn_next = ctk.CTkButton(header, text="▶", width=40)
        btn_next.grid(row=1, column=5, padx=(2,0), pady=(6,0))

        rows_container = ctk.CTkFrame(sf)
        rows_container.pack(fill="both", expand=True)

        def render_rows_expedientes(filas, page, total_count, page_size):
            for w in rows_container.winfo_children():
                w.destroy()

            head = ctk.CTkFrame(rows_container)
            head.pack(fill="x", padx=5, pady=(0,4))
            hf = ctk.CTkFont(weight="bold")
            head.grid_columnconfigure(0, weight=1, minsize=160)
            head.grid_columnconfigure(1, weight=2, minsize=300)
            head.grid_columnconfigure(2, weight=1, minsize=140)
            head.grid_columnconfigure(3, weight=1, minsize=140)
            ctk.CTkLabel(head, text="RMA", font=hf).grid(row=0, column=0, padx=5, sticky="w")
            ctk.CTkLabel(head, text="CLIENTE", font=hf).grid(row=0, column=1, padx=5, sticky="w")
            ctk.CTkLabel(head, text="DOCUMENTO", font=hf).grid(row=0, column=2, padx=5, sticky="w")
            ctk.CTkLabel(head, text="ESTADO ARTÍCULO", font=hf).grid(row=0, column=3, padx=5, sticky="w")

            colors = ("#FFFFFF", "#F3F4F6")
            for idx, r in enumerate(filas):
                try:
                    rma_id, codigo, cliente, num_doc, estado = r
                except Exception:
                    vals = list(r)
                    rma_id = vals[0] if len(vals) > 0 else None
                    codigo = vals[1] if len(vals) > 1 else ''
                    cliente = vals[2] if len(vals) > 2 else ''
                    num_doc = vals[3] if len(vals) > 3 else ''
                    estado = vals[4] if len(vals) > 4 else ''

                # Usar fondo por defecto del tema en lugar de cebra
                bg = "transparent"
                rowf = ctk.CTkFrame(rows_container, fg_color="transparent")
                rowf.pack(fill="x", padx=5, pady=2)
                rowf.grid_columnconfigure(0, weight=1, minsize=160)
                rowf.grid_columnconfigure(1, weight=2, minsize=300)
                rowf.grid_columnconfigure(2, weight=1, minsize=140)
                rowf.grid_columnconfigure(3, weight=1, minsize=140)

                lbl_codigo = ctk.CTkLabel(rowf, text=codigo or '-', anchor="w", cursor="hand2")
                lbl_codigo.grid(row=0, column=0, padx=5, sticky="w")
                lbl_cliente = ctk.CTkLabel(rowf, text=cliente or '-', anchor="w")
                lbl_cliente.grid(row=0, column=1, padx=5, sticky="w")
                lbl_doc = ctk.CTkLabel(rowf, text=num_doc or '-', anchor="w")
                lbl_doc.grid(row=0, column=2, padx=5, sticky="w")
                lbl_estado = ctk.CTkLabel(rowf, text=estado or '-', anchor="w")
                lbl_estado.grid(row=0, column=3, padx=5, sticky="w")

                acciones = ctk.CTkFrame(rowf, fg_color="transparent")
                acciones.grid(row=0, column=4, padx=5)
                ctk.CTkButton(acciones, text="Abrir", width=90, command=lambda rid=rma_id: (self._abrir_editor_rma(rma_id=rid), vent.destroy())).pack(side="left", padx=4)

                def on_ent(e, r=rowf):
                    try:
                        r.configure(fg_color="transparent")
                    except Exception:
                        pass
                def on_lve(e, r=rowf, original=bg):
                    try:
                        r.configure(fg_color=original)
                    except Exception:
                        pass

                rowf.bind("<Enter>", on_ent)
                rowf.bind("<Leave>", on_lve)
                rowf.bind("<Double-Button-1>", lambda e, rid=rma_id: (self._abrir_editor_rma(rma_id=rid), vent.destroy()))
                lbl_codigo.bind("<Double-Button-1>", lambda e, rid=rma_id: (self._abrir_editor_rma(rma_id=rid), vent.destroy()))

            # Ensure numeric types for pagination calculation
            try:
                total_count = int(total_count or 0)
            except Exception:
                try:
                    total_count = int(float(total_count))
                except Exception:
                    total_count = 0
            try:
                page_size = int(page_size)
            except Exception:
                page_size = 50

            total_pages = max(1, (total_count + page_size - 1) // page_size)
            page_state["page"] = page
            page_state["total_pages"] = total_pages
            page_label.configure(text=f"Página {page}/{total_pages}")

        def load_expedientes_thread(page=1):
            try:
                search = search_var.get().strip()
                page_size = int(page_size_opt.get())
                offset = (page - 1) * page_size

                conn = connect_db()
                cur = conn.cursor()

                # Count total
                try:
                    if search:
                        cnt_sql = "SELECT COUNT(DISTINCT T2.id) FROM rma_detalles T1 JOIN rma_maestro T2 ON T1.rma_id = T2.id WHERE T1.referencia_articulo = ? AND (T2.codigo_rma LIKE ? OR T2.cliente LIKE ? OR T2.numero_documento_cliente LIKE ?)"
                        cur.execute(cnt_sql, (referencia, f"%{search}%", f"%{search}%", f"%{search}%"))
                    else:
                        cnt_sql = "SELECT COUNT(DISTINCT T2.id) FROM rma_detalles T1 JOIN rma_maestro T2 ON T1.rma_id = T2.id WHERE T1.referencia_articulo = ?"
                        cur.execute(cnt_sql, (referencia,))
                    total = cur.fetchone()[0]
                    total = int(total) if total is not None else 0
                except Exception:
                    total = 0

                # Main query
                params = [referencia]
                where = ""
                if search:
                    where = "AND (T2.codigo_rma LIKE ? OR T2.cliente LIKE ? OR T2.numero_documento_cliente LIKE ?)"
                    params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])

                sql = f"SELECT DISTINCT T2.id, T2.codigo_rma, T2.cliente, T2.numero_documento_cliente, T1.estado_producto FROM rma_detalles T1 JOIN rma_maestro T2 ON T1.rma_id = T2.id WHERE T1.referencia_articulo = ? {where} ORDER BY T2.fecha_emision DESC LIMIT ? OFFSET ?"
                params.extend([page_size, offset])
                cur.execute(sql, tuple(params))
                filas = cur.fetchall()
                conn.close()
                self.after(0, lambda: render_rows_expedientes(filas, page, total, page_size))
            except Exception as e:
                self.after(0, lambda: messagebox.showerror("Error BD", f"No se pudieron cargar expedientes: {e}"))

        def start_load_expedientes(page=1):
            threading.Thread(target=load_expedientes_thread, args=(page,), daemon=True).start()

        def prev_page():
            p = max(1, page_state.get("page", 1) - 1)
            start_load_expedientes(p)

        def next_page():
            p = page_state.get("page", 1) + 1
            if page_state.get("total_pages", 1) and p > page_state.get("total_pages"):
                return
            start_load_expedientes(p)

        btn_prev.configure(command=prev_page)
        btn_next.configure(command=next_page)
        search_entry.bind("<Return>", lambda e: start_load_expedientes(1))

        # Cargar inicialmente
        start_load_expedientes(1)

    def comprobar_tareas_vencidas(self):
        """Comprueba tareas vencidas para el usuario actual y muestra notificaciones (sistema si es posible)."""
        try:
            hoy = datetime.date.today().strftime("%Y-%m-%d")
            conn = connect_db()
            cur = conn.cursor()
            cur.execute("""
                SELECT id, codigo_rma, titulo, fecha_vencimiento 
                FROM tareas 
                WHERE creado_por = ? 
                AND notificado = 0 
                AND fecha_vencimiento IS NOT NULL 
                AND fecha_vencimiento <= ?
                AND estado NOT IN ('Completado', 'Completada', 'Finalizada')
            """, (self.username, hoy))
            filas = cur.fetchall()
            if filas:
                mensajes = []
                ids = []
                for tid, codigo, titulo, fecha in filas:
                    mensajes.append(f"{fecha} - {titulo} [{codigo}]")
                    ids.append(tid)
                texto = "Tienes tareas vencidas o para hoy:\n" + "\n".join(mensajes)
                # Intentar notificación nativa en Windows
                try:
                    if hasattr(self, 'toaster') and self.toaster:
                        # win10toast limita longitud, mostramos un resumen
                        resumen = "; ".join(mensajes[:5])
                        self.toaster.show_toast("Tareas vencidas/hoy", resumen, duration=10, threaded=True)
                    else:
                        raise Exception("No toaster")
                except Exception:
                    # Fallback a messagebox si no hay toaster
                    messagebox.showinfo("Tareas vencidas / hoy", texto)

                # Marcar como notificado para no repetir
                for tid in ids:
                    cur.execute("UPDATE tareas SET notificado = 1 WHERE id = ?", (tid,))
                conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error comprobando tareas vencidas: {e}")

    def programar_chequeo_tareas(self, intervalo_ms=1_800_000):
        """Programa la comprobación periódica de tareas vencidas (por defecto cada 30 minutos)."""
        try:
            # Llamamos a la comprobación
            self.comprobar_tareas_vencidas()
            # Reprogramar
            self.after(intervalo_ms, lambda: self.programar_chequeo_tareas(intervalo_ms))
        except Exception as e:
            print(f"Error programando chequeo tareas: {e}")
    def mostrar_ventana_estadisticas(self):
        """Crea y muestra la ventana Toplevel y la estructura de navegación interna para estadísticas."""
        
        # 1. Prevenir que se abra más de una vez (si ya existe, simplemente enfocarla)
        if hasattr(self, 'stats_window') and self.stats_window.winfo_exists():
            self.stats_window.focus()
            return

        # Crear una nueva ventana modal pero independiente
        self.stats_window = ctk.CTkToplevel(self)
        self.stats_window.title("Módulo de Estadísticas")
        self.stats_window.geometry("1600x900")
        self.stats_window.minsize(800, 600)  # Tamaño mínimo para asegurar legibilidad
        
        # Configurar para permitir minimización
        self.stats_window.resizable(True, True)
        self.stats_window.attributes('-topmost', False)
        
        # Centrar la ventana en la pantalla
        screen_width = self.stats_window.winfo_screenwidth()
        screen_height = self.stats_window.winfo_screenheight()
        x = (screen_width - 1600) // 2
        y = (screen_height - 900) // 2
        self.stats_window.geometry(f"1600x900+{x}+{y}")

        # Marco principal que contendrá la navegación y el contenido
        stats_frame = ctk.CTkFrame(self.stats_window)
        stats_frame.pack(fill="both", expand=True)

        # Configurar para que los marcos se expandan correctamente
        stats_frame.grid_columnconfigure(1, weight=1)
        stats_frame.grid_rowconfigure(0, weight=1)

        # 2. Marco de Navegación Lateral Interna (Columna 0)
        nav_frame = ctk.CTkFrame(stats_frame, width=220, corner_radius=0)
        nav_frame.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(nav_frame, text="INFORMES", font=ctk.CTkFont(weight="bold", size=16)).pack(pady=20, padx=10)

        # 3. Marco de Contenido Principal para la Estadística (Columna 1)
        self.main_stats_frame = ctk.CTkFrame(stats_frame)
        self.main_stats_frame.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        self.main_stats_frame.grid_columnconfigure(0, weight=1) # Permite que el contenido se expanda
        
        # 4. Definición de las estadísticas y sus métodos (Aún no creados)
        # Se elimina la estadística 'Abonos por Cliente y Periodo' (no funcional).
        self.botones_stats = {
            "Rentabilidad por Cliente": self.mostrar_expedientes_completados,
            "Referencia (Incidencia)": self.mostrar_articulos_incidencia,
            "⏱️ Tiempos de Tramitación": self.mostrar_estadisticas_tiempos,
            "📦 Artículos - Estadísticas": self.mostrar_estadisticas_articulos_menu,
            "📋 Resolución de Expedientes": self.mostrar_estadisticas_resolucion_menu,
            "⚠️ Recepciones Anticipadas": self.mostrar_recepciones_anticipadas
        }
        
        # 5. Creación de los botones del menú interno
        for text, command in self.botones_stats.items():
            ctk.CTkButton(
                nav_frame, 
                text=text, 
                anchor="w",
                command=lambda cmd=command: self.cargar_estadistica(cmd)
            ).pack(fill="x", padx=10, pady=5)

        # 6. Cargar la primera estadística por defecto
        #self.cargar_estadistica(self.mostrar_expedientes_cliente)
        self.cargar_estadistica(self.mostrar_expedientes_completados)

    def cargar_estadistica(self, func_callback):
        """
        Método central que limpia el marco principal y llama a la función 
        de la estadística seleccionada (callback).
        """
        # Limpiar el marco principal
        for widget in self.main_stats_frame.winfo_children():
            widget.destroy()
        
        # Llamar a la función que dibuja la estadística
        func_callback()
    
    def validar_fecha_entrada(self, fecha_ddmmyyyy):
        """Función auxiliar para validar el formato de fecha DD/MM/AAAA."""
        if not fecha_ddmmyyyy:
            return None 
        try:
            datetime.strptime(fecha_ddmmyyyy, '%d/%m/%Y')
            return fecha_ddmmyyyy 
        except ValueError:
            messagebox.showerror("Error de Formato de Fecha", 
                                 f"El formato de la fecha '{fecha_ddmmyyyy}' es incorrecto. Debe ser DD/MM/AAAA (ej: 01/01/2024).")
            return False 
        except Exception:
            return None
    
    def mostrar_expedientes_completados(self):
        """Wrapper que delega la creación de la estadística de rentabilidad
        al módulo externo `lib.client_rentability`.
        """
        try:
            from lib.client_rentability import mostrar_rentabilidad_clientes
            mostrar_rentabilidad_clientes(self)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo cargar la estadística de rentabilidad: {e}")

    def mostrar_articulos_incidencia(self):
        """
        Estadística con filtros avanzados: Muestra el top de artículos según 
        múltiples 'estado_producto', rango de 'fecha_gestion' y búsqueda por 'referencia_articulo'.
        """
        
        # Lista de estados fijos proporcionada por el usuario
        self.ESTADOS_DISPONIBLES = [
            "EN PERFECTO ESTADO ; ABONAR", "FUNCIONA PERFECTAMENTE ; ABONAR", 
            "SOBRANTE DE OBRA ; ABONAR", "NO FUNCIONA, ABONAR", 
            "FUNCIONA PERFECTAMENTE ; NO ABONAR", "NO FUNCIONA ; NO ABONAR", 
            "REPOSICION FALLO PRODUCTO", "REPOSICION ; ABONAR", 
            "MERCANCIA ENVIADA POR ERROR", "MALA MANIPULACION ; NO ABONAR",
            "EN PERFECTO ESTADO ; ABONAR 10% DEPRECIACION", "FALLO SOLDADURA ; ABONAR", 
            "FALLO SOLDADURA ; NO ABONAR", "FALLO MODULO ; ABONAR", 
            "MAL MANIPULACION ; ABONAR", "DANA", "CAMBIO DE PRODUCTO"
        ]
        
        # 1. Preparación de la interfaz
        self.limpiar_marco_stats()
        if not self.main_stats_frame: return
        
        ctk.CTkLabel(self.main_stats_frame, 
                     text="ARTÍCULOS CON MÁS INCIDENCIA", 
                     font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=20)
        
        # --- 2. Marco de Controles Principales (Filtros y Búsqueda) ---
        controles_frame = ctk.CTkFrame(self.main_stats_frame)
        controles_frame.pack(padx=20, pady=(0, 10), fill="x")
        
        # Configuración de columnas para la primera fila (Fechas)
        controles_frame.grid_columnconfigure((0, 2, 4), weight=0) # Etiquetas y Botón (fijo)
        controles_frame.grid_columnconfigure((1, 3), weight=1)    # Entries de Fecha (expandible)

        # FILTRO 1: Fecha Inicial (DD/MM/AAAA) - Fila 0
        ctk.CTkLabel(controles_frame, text="Fecha Inicial (DD/MM/AAAA):").grid(row=0, column=0, padx=(0, 5), pady=5, sticky="w")
        self.fecha_inicial_entry = ctk.CTkEntry(controles_frame, placeholder_text="Ej: 01/01/2024")
        self.fecha_inicial_entry.grid(row=0, column=1, padx=(0, 10), pady=5, sticky="ew")

        # FILTRO 2: Fecha Final (DD/MM/AAAA) - Fila 0
        ctk.CTkLabel(controles_frame, text="Fecha Final (DD/MM/AAAA):").grid(row=0, column=2, padx=(10, 5), pady=5, sticky="w")
        self.fecha_final_entry = ctk.CTkEntry(controles_frame, placeholder_text="Ej: 31/12/2024")
        self.fecha_final_entry.grid(row=0, column=3, padx=(0, 10), pady=5, sticky="ew")
        
        # Botón para aplicar filtros - Fila 0
        ctk.CTkButton(
            controles_frame, 
            text="Aplicar Filtros", 
            command=self._cargar_datos_articulos_incidencia
        ).grid(row=0, column=4, padx=(10, 0), pady=5)
        
        # --- NUEVOS WIDGETS DE BÚSQUEDA POR REFERENCIA (Fila 1) ---
        controles_frame.grid_columnconfigure(5, weight=0) # Columna del botón Limpiar

        ctk.CTkLabel(controles_frame, text="Buscar Referencia:").grid(row=1, column=0, padx=(0, 5), pady=5, sticky="w")
        self.referencia_entry = ctk.CTkEntry(controles_frame, placeholder_text="Escriba parte de la referencia...")
        self.referencia_entry.grid(row=1, column=1, columnspan=3, padx=(0, 10), pady=5, sticky="ew") # Ocupa 3 columnas
        
        ctk.CTkButton(
            controles_frame, 
            text="Limpiar Búsqueda", 
            command=self._limpiar_filtro_referencia_y_recargar,
            fg_color="gray"
        ).grid(row=1, column=4, padx=(10, 0), pady=5)


        # --- 3. Marco Contenedor de Listado de Estados y Resultados ---
        content_frame = ctk.CTkFrame(self.main_stats_frame)
        content_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        # Configuración clave de expansión: la tabla (col 1) se expande
        content_frame.grid_columnconfigure(0, weight=0, minsize=300) 
        content_frame.grid_columnconfigure(1, weight=1)             
        content_frame.grid_rowconfigure(0, weight=1)

        # --- 4. Listado de Estados con Checkboxes (Panel Izquierdo) ---
        estados_panel = ctk.CTkFrame(content_frame)
        estados_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        estados_panel.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(estados_panel, 
                     text="ESTADOS DE INCIDENCIA", 
                     font=ctk.CTkFont(size=14, weight="bold")
        ).pack(pady=(10, 5))
        
        scroll_frame = ctk.CTkScrollableFrame(estados_panel)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        scroll_frame.grid_columnconfigure(0, weight=1)

        # Variables para almacenar el estado de cada checkbox
        self.estado_vars = {} 
        
        # Creación dinámica de las Checkboxes
        for idx, estado in enumerate(self.ESTADOS_DISPONIBLES):
            var = ctk.StringVar(value="0") 
            chk = ctk.CTkCheckBox(
                scroll_frame, 
                text=estado, 
                variable=var,
                onvalue="1", offvalue="0"
            )
            chk.grid(row=idx, column=0, padx=5, pady=2, sticky="w")
            self.estado_vars[estado] = var

        # Checkbox para seleccionar/deseleccionar todos
        all_var = ctk.StringVar(value="0")
        def toggle_all_states():
            new_value = all_var.get()
            for var in self.estado_vars.values():
                var.set(new_value)

        ctk.CTkCheckBox(
            estados_panel, 
            text="SELECCIONAR TODOS", 
            variable=all_var,
            command=toggle_all_states,
            onvalue="1", offvalue="0",
            font=ctk.CTkFont(weight="bold")
        ).pack(pady=(5, 10))


        # --- 5. Marco donde se dibujará la tabla de resultados (Panel Derecho) ---
        self.tabla_resultados_frame = ctk.CTkFrame(content_frame)
        self.tabla_resultados_frame.grid(row=0, column=1, sticky="nsew") 

        # 6. Cargar los datos iniciales
        self._cargar_datos_articulos_incidencia()
    
    def _cargar_datos_articulos_incidencia(self, event=None):
        """
        Consulta la base de datos, aplica filtros de estado (múltiples), fecha,
        y búsqueda LIKE por referencia de artículo.
        """
        from datetime import datetime
        
        # 1. Obtener los valores de los filtros
        
        # 1.1. Obtener los estados seleccionados (Lista de strings)
        estados_seleccionados = [
            estado for estado, var in self.estado_vars.items() if var.get() == "1"
        ]
        
        # 1.2. Obtener filtro de Referencia
        referencia_filtro = self.referencia_entry.get().strip()

        # 1.3. Obtener y validar las fechas (DD/MM/AAAA)
        fecha_inicial_str = self.fecha_inicial_entry.get().strip()
        fecha_final_str = self.fecha_final_entry.get().strip()
        
        # Función que valida DD/MM/AAAA y devuelve el string si es válido
        def validar_fecha_entrada(fecha_ddmmyyyy):
            if not fecha_ddmmyyyy:
                return None 
            try:
                datetime.strptime(fecha_ddmmyyyy, '%d/%m/%Y')
                return fecha_ddmmyyyy 
            except ValueError:
                messagebox.showerror("Error de Formato de Fecha", 
                                     f"El formato de la fecha '{fecha_ddmmyyyy}' es incorrecto. Debe ser DD/MM/AAAA (ej: 01/01/2024).")
                return False 
            except Exception:
                return None
        
        fecha_inicial_db = validar_fecha_entrada(fecha_inicial_str)
        fecha_final_db = validar_fecha_entrada(fecha_final_str)
        
        if fecha_inicial_db is False or fecha_final_db is False:
            return

        # 2. Conexión y Limpieza del Marco
        conn, cursor = self.master.conectar_db()
        if not conn: 
            messagebox.showerror("Error", "No se pudo conectar a la base de datos.")
            return

        for widget in self.tabla_resultados_frame.winfo_children():
            widget.destroy()

        # 3. Construcción dinámica de la consulta SQL y los parámetros
        
        sql_query_base = """
            SELECT 
                T2.codigo_rma, 
                T1.referencia_articulo,
                T1.estado_producto,
                T2.fecha_gestion,
                COUNT(T1.id) AS total_incidencias
            FROM 
                rma_detalles AS T1
            INNER JOIN 
                rma_maestro AS T2 ON T1.rma_id = T2.id
            WHERE 1=1 
        """
        
        condiciones = []
        parametros = []
        
        # 3.1. Filtro por Estado de Producto 
        if estados_seleccionados:
            placeholders = ', '.join('?' for _ in estados_seleccionados)
            condiciones.append(f"T1.estado_producto IN ({placeholders})")
            parametros.extend(estados_seleccionados)
            
        # 3.2. Filtro por Referencia de Artículo
        if referencia_filtro:
            condiciones.append("T1.referencia_articulo LIKE ?")
            parametros.append(f"%{referencia_filtro}%")
            
        # 3.3. Filtro por Rango de Fechas (Convierte la fecha de la DB a YYYY-MM-DD para comparación)
        if fecha_inicial_db:
            # Reorganiza: T2.fecha_gestion (DD/MM/AAAA) -> YYYY-MM-DD
            condiciones.append("SUBSTR(T2.fecha_gestion, 7, 4) || '-' || SUBSTR(T2.fecha_gestion, 4, 2) || '-' || SUBSTR(T2.fecha_gestion, 1, 2) >= ?")
            # El parámetro se convierte a YYYY-MM-DD para la comparación
            parametros.append(datetime.strptime(fecha_inicial_db, '%d/%m/%Y').strftime('%Y-%m-%d'))

        if fecha_final_db:
            condiciones.append("SUBSTR(T2.fecha_gestion, 7, 4) || '-' || SUBSTR(T2.fecha_gestion, 4, 2) || '-' || SUBSTR(T2.fecha_gestion, 1, 2) <= ?")
            parametros.append(datetime.strptime(fecha_final_db, '%d/%m/%Y').strftime('%Y-%m-%d'))

        # 4. Ensamblaje y Ejecución
        
        if condiciones:
            sql_query_base += " AND " + " AND ".join(condiciones)

        sql_query_final = sql_query_base + """
            GROUP BY 
                T2.codigo_rma, T1.referencia_articulo, T1.estado_producto, T2.fecha_gestion
            ORDER BY 
                total_incidencias DESC
            LIMIT 50; 
        """

        try:
            cursor.execute(sql_query_final, parametros)
            datos_raw = cursor.fetchall()

            # Los datos están en formato DD/MM/AAAA, no se necesita conversión de salida
            datos_formateados = list(datos_raw) 

            # 5. Dibujar la tabla
            self.mostrar_tabla_estadistica(
                datos_formateados, 
                columnas=["CÓDIGO RMA", "REFERENCIA ARTÍCULO", "ESTADO PRODUCTO", "FECHA GESTIÓN", "TOTAL INCIDENCIAS"],
                export_filename="Articulos_Incidencia_Filtrado",
                frame=self.tabla_resultados_frame, 
                formato_moneda=False 
            )

            if not datos_formateados:
                 ctk.CTkLabel(self.tabla_resultados_frame, 
                              text="No se encontraron datos con los filtros aplicados.",
                              font=ctk.CTkFont(size=14)
                 ).pack(pady=40)

        except sqlite3.Error as e:
            messagebox.showerror("Error de Base de Datos", f"Error SQL: {e}")
        except Exception as e:
            messagebox.showerror("Error", f"Error inesperado: {e}")
        finally:
            if conn:
                conn.close()
    
    def _limpiar_filtro_referencia_y_recargar(self):
        """Limpia el campo de búsqueda de referencia y recarga los datos."""
        # Se asegura de que el atributo existe antes de intentar acceder a él
        if hasattr(self, 'referencia_entry'):
            self.referencia_entry.delete(0, 'end')
        # Recarga los datos para aplicar el filtro limpio (que es no filtrar)
        self._cargar_datos_articulos_incidencia()
    
    def mostrar_estadisticas_tiempos(self):
        """Muestra estadísticas de tiempos de tramitación por cliente."""
        self.limpiar_marco_stats()
        
        ctk.CTkLabel(self.main_stats_frame, 
                     text="⏱️ TIEMPOS DE TRAMITACIÓN POR CLIENTE", 
                     font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=20)
        
        # Obtener conexión a la base de datos
        conn, cursor = self.master.conectar_db()
        if not conn:
            ctk.CTkLabel(self.main_stats_frame, text="Error al conectar con la base de datos.", 
                        text_color="red").pack(pady=20)
            return
        
        try:
            # Obtener todos los clientes únicos con expedientes cerrados
            cursor.execute("""
                SELECT DISTINCT cliente
                FROM rma_maestro
                WHERE fecha_gestion IS NOT NULL AND fecha_gestion != ''
                AND cliente IS NOT NULL AND cliente != ''
                ORDER BY cliente ASC
            """)
            
            clientes = [fila[0] for fila in cursor.fetchall()]
            
            if not clientes:
                ctk.CTkLabel(self.main_stats_frame, 
                           text="No se encontraron expedientes cerrados para calcular estadísticas.", 
                           text_color="gray").pack(pady=20)
                conn.close()
                return
            
            # Marco scrollable para la tabla
            scroll_frame = ctk.CTkScrollableFrame(self.main_stats_frame)
            scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
            
            # Encabezados de la tabla
            header_font = ctk.CTkFont(weight="bold", size=12)
            headers = ["CLIENTE", "EXPEDIENTES", "DÍAS PROMEDIO TOTAL", "E→A", "A→R", "R→P", "P→C"]
            
            for col, header in enumerate(headers):
                lbl = ctk.CTkLabel(scroll_frame, text=header, font=header_font)
                lbl.grid(row=0, column=col, padx=10, pady=10, sticky="w")
            
            # Calcular estadísticas para cada cliente
            for i, cliente in enumerate(clientes):
                promedio_info = obtener_promedio_cliente(cliente, conn)
                
                row = i + 1
                
                # Columna: Cliente
                ctk.CTkLabel(scroll_frame, text=cliente).grid(row=row, column=0, padx=10, pady=5, sticky="w")
                
                # Columna: Número de expedientes
                ctk.CTkLabel(scroll_frame, text=str(promedio_info['total_expedientes'])).grid(
                    row=row, column=1, padx=10, pady=5, sticky="w")
                
                # Columna: Días promedio total
                if promedio_info['promedio_total'] is not None:
                    dias_total = int(promedio_info['promedio_total'])
                    color_total = obtener_color_tiempo(dias_total)
                    ctk.CTkLabel(scroll_frame, text=f"{dias_total} días", text_color=color_total).grid(
                        row=row, column=2, padx=10, pady=5, sticky="w")
                else:
                    ctk.CTkLabel(scroll_frame, text="-", text_color="gray").grid(
                        row=row, column=2, padx=10, pady=5, sticky="w")
                
                # Columnas: Tiempos entre fases (E→A, A→R, R→P, P→C)
                fases = ['promedio_e_a', 'promedio_a_r', 'promedio_r_p', 'promedio_p_c']
                for col_idx, fase in enumerate(fases, start=3):
                    if promedio_info[fase] is not None:
                        dias = int(promedio_info[fase])
                        color = obtener_color_tiempo(dias)
                        ctk.CTkLabel(scroll_frame, text=f"{dias}d", text_color=color).grid(
                            row=row, column=col_idx, padx=10, pady=5, sticky="w")
                    else:
                        ctk.CTkLabel(scroll_frame, text="-", text_color="gray").grid(
                            row=row, column=col_idx, padx=10, pady=5, sticky="w")
            
            # Leyenda de colores
            leyenda_frame = ctk.CTkFrame(self.main_stats_frame)
            leyenda_frame.pack(pady=10)
            
            ctk.CTkLabel(leyenda_frame, text="Leyenda: ", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=5)
            ctk.CTkLabel(leyenda_frame, text="🟢 <10 días", text_color="#22c55e").pack(side="left", padx=5)
            ctk.CTkLabel(leyenda_frame, text="🟡 10-20 días", text_color="#eab308").pack(side="left", padx=5)
            ctk.CTkLabel(leyenda_frame, text="🟠 20-30 días", text_color="#f97316").pack(side="left", padx=5)
            ctk.CTkLabel(leyenda_frame, text="🔴 >30 días", text_color="#ef4444").pack(side="left", padx=5)
            
            conn.close()
            
        except Exception as e:
            print(f"Error al generar estadísticas de tiempos: {e}")
            ctk.CTkLabel(self.main_stats_frame, 
                        text=f"Error al cargar estadísticas: {e}", 
                        text_color="red").pack(pady=20)
            if conn:
                conn.close()
    
    def mostrar_estadisticas_articulos_menu(self):
        """Wrapper que delega la creación de la estadística de artículos
        al módulo externo `lib.articulos_estadisticas`.
        """
        try:
            from lib.articulos_estadisticas import mostrar_estadisticas_articulos
            mostrar_estadisticas_articulos(self)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo cargar la estadística de artículos: {e}")
    
    def mostrar_estadisticas_resolucion_menu(self):
        """Wrapper que delega la creación de la estadística de resolución de expedientes
        al módulo externo `lib.resolucion_estadisticas`.
        """
        try:
            from lib.resolucion_estadisticas import mostrar_estadisticas_resolucion
            mostrar_estadisticas_resolucion(self)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo cargar la estadística de resolución: {e}")
    
    def mostrar_recepciones_anticipadas(self):
        """Muestra estadísticas de recepciones anticipadas (recepción antes de autorización)."""
        from lib.recepciones_anticipadas import obtener_recepciones_anticipadas, obtener_expedientes_anticipados_por_cliente, buscar_clientes_anticipados, ordenar_resultados
        
        self.limpiar_marco_stats()
        
        # Título
        ctk.CTkLabel(self.main_stats_frame, 
                     text="⚠️ RECEPCIONES ANTICIPADAS", 
                     font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=20)
        
        ctk.CTkLabel(self.main_stats_frame,
                     text="Clientes que reciben productos antes de autorizarlos",
                     font=ctk.CTkFont(size=12),
                     text_color="gray"
        ).pack(pady=(0, 10))
        
        # Frame de controles (búsqueda y ordenamiento)
        controles_frame = ctk.CTkFrame(self.main_stats_frame)
        controles_frame.pack(fill="x", padx=20, pady=10)
        
        # Búsqueda
        ctk.CTkLabel(controles_frame, text="Buscar cliente:").pack(side="left", padx=5)
        entry_buscar = ctk.CTkEntry(controles_frame, placeholder_text="Nombre del cliente...", width=250)
        entry_buscar.pack(side="left", padx=5)
        
        # Ordenamiento
        ctk.CTkLabel(controles_frame, text="Ordenar por:").pack(side="left", padx=(20, 5))
        combo_orden = ctk.CTkOptionMenu(controles_frame, 
                                        values=["Cantidad", "Cliente", "Media días"],
                                        width=150)
        combo_orden.set("Cantidad")
        combo_orden.pack(side="left", padx=5)
        
        # Frame scrollable para resultados
        scroll_frame = ctk.CTkScrollableFrame(self.main_stats_frame, height=400)
        scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        def actualizar_listado():
            # Limpiar resultados anteriores
            for widget in scroll_frame.winfo_children():
                widget.destroy()
            
            # Obtener datos
            conn, cursor = self.master.conectar_db()
            if not conn:
                ctk.CTkLabel(scroll_frame, text="Error al conectar con la base de datos.",
                            text_color="red").pack(pady=20)
                return
            
            try:
                termino = entry_buscar.get().strip()
                if termino:
                    resultados = buscar_clientes_anticipados(conn, termino)
                else:
                    resultados = obtener_recepciones_anticipadas(conn)
                
                conn.close()
                
                if not resultados:
                    ctk.CTkLabel(scroll_frame, 
                               text="No se encontraron recepciones anticipadas.",
                               text_color="gray").pack(pady=20)
                    return
                
                # Ordenar según criterio
                criterio_map = {
                    "Cantidad": "cantidad",
                    "Cliente": "cliente",
                    "Media días": "media_dias"
                }
                criterio = criterio_map.get(combo_orden.get(), "cantidad")
                resultados = ordenar_resultados(resultados, criterio)
                
                # Encabezados
                header_frame = ctk.CTkFrame(scroll_frame)
                header_frame.pack(fill="x", pady=(0, 10))
                header_frame.grid_columnconfigure(0, weight=3)
                header_frame.grid_columnconfigure(1, weight=1)
                header_frame.grid_columnconfigure(2, weight=1)
                
                header_font = ctk.CTkFont(weight="bold", size=12)
                ctk.CTkLabel(header_frame, text="CLIENTE", font=header_font, anchor="w").grid(row=0, column=0, padx=10, pady=5, sticky="w")
                ctk.CTkLabel(header_frame, text="EXPEDIENTES", font=header_font, anchor="center").grid(row=0, column=1, padx=10, pady=5)
                ctk.CTkLabel(header_frame, text="MEDIA DÍAS", font=header_font, anchor="center").grid(row=0, column=2, padx=10, pady=5)
                
                # Filas de datos
                for cliente, cantidad, media_dias in resultados:
                    row_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
                    row_frame.pack(fill="x", pady=2)
                    row_frame.grid_columnconfigure(0, weight=3)
                    row_frame.grid_columnconfigure(1, weight=1)
                    row_frame.grid_columnconfigure(2, weight=1)
                    
                    # Cliente (clickeable)
                    cliente_btn = ctk.CTkButton(row_frame, 
                                               text=cliente,
                                               anchor="w",
                                               command=lambda c=cliente: self.mostrar_detalle_recepciones_cliente(c))
                    cliente_btn.grid(row=0, column=0, padx=10, pady=2, sticky="ew")
                    
                    # Cantidad
                    ctk.CTkLabel(row_frame, text=str(cantidad), anchor="center").grid(row=0, column=1, padx=10, pady=2)
                    
                    # Media días
                    media_texto = f"{media_dias:.1f}" if media_dias is not None else "-"
                    color = "red" if media_dias and media_dias > 7 else "orange" if media_dias and media_dias > 3 else "green"
                    ctk.CTkLabel(row_frame, text=media_texto, text_color=color, anchor="center").grid(row=0, column=2, padx=10, pady=2)
                
            except Exception as e:
                print(f"Error al cargar recepciones anticipadas: {e}")
                ctk.CTkLabel(scroll_frame,
                            text=f"Error: {e}",
                            text_color="red").pack(pady=20)
        
        # Botones de acción
        btn_actualizar = ctk.CTkButton(controles_frame, text="🔍 Buscar", command=actualizar_listado, width=100)
        btn_actualizar.pack(side="left", padx=10)
        
        # Botones de exportación e impresión
        btn_frame_export = ctk.CTkFrame(self.main_stats_frame)
        btn_frame_export.pack(pady=10)
        
        def exportar_listado():
            # Obtener datos actuales
            conn, cursor = self.master.conectar_db()
            if not conn:
                return
            termino = entry_buscar.get().strip()
            if termino:
                resultados = buscar_clientes_anticipados(conn, termino)
            else:
                resultados = obtener_recepciones_anticipadas(conn)
            conn.close()
            
            if resultados:
                criterio_map = {"Cantidad": "cantidad", "Cliente": "cliente", "Media días": "media_dias"}
                criterio = criterio_map.get(combo_orden.get(), "cantidad")
                resultados = ordenar_resultados(resultados, criterio)
                self._exportar_listado_recepciones_excel(resultados)
        
        def imprimir_listado():
            # Obtener datos actuales
            conn, cursor = self.master.conectar_db()
            if not conn:
                return
            termino = entry_buscar.get().strip()
            if termino:
                resultados = buscar_clientes_anticipados(conn, termino)
            else:
                resultados = obtener_recepciones_anticipadas(conn)
            conn.close()
            
            if resultados:
                criterio_map = {"Cantidad": "cantidad", "Cliente": "cliente", "Media días": "media_dias"}
                criterio = criterio_map.get(combo_orden.get(), "cantidad")
                resultados = ordenar_resultados(resultados, criterio)
                self._imprimir_listado_recepciones_pdf(resultados)
        
        ctk.CTkButton(btn_frame_export, text="📄 Imprimir PDF", 
                     command=imprimir_listado, width=130).pack(side="left", padx=5)
        
        ctk.CTkButton(btn_frame_export, text="📊 Exportar Excel", 
                     command=exportar_listado, width=130).pack(side="left", padx=5)
        
        # Bind enter en búsqueda
        entry_buscar.bind("<Return>", lambda e: actualizar_listado())
        combo_orden.configure(command=lambda v: actualizar_listado())
        
        # Cargar datos iniciales
        actualizar_listado()
    
    def mostrar_detalle_recepciones_cliente(self, cliente):
        """Muestra ventana con detalle de expedientes con recepción anticipada de un cliente."""
        from lib.recepciones_anticipadas import obtener_expedientes_anticipados_por_cliente
        
        # Crear ventana
        ventana = ctk.CTkToplevel(self)
        ventana.title(f"Recepciones Anticipadas - {cliente}")
        ventana.geometry("900x600")
        ventana.transient(self)
        
        try:
            ventana.iconbitmap("Icono_Ilutrek.ico")
        except Exception:
            pass
        
        # Título
        ctk.CTkLabel(ventana,
                     text=f"📋 Expedientes con recepción anticipada",
                     font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)
        
        ctk.CTkLabel(ventana,
                     text=f"Cliente: {cliente}",
                     font=ctk.CTkFont(size=14),
                     text_color="gray"
        ).pack(pady=(0, 10))
        
        # Frame scrollable
        scroll_frame = ctk.CTkScrollableFrame(ventana)
        scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Obtener datos
        conn, cursor = self.master.conectar_db()
        if not conn:
            ctk.CTkLabel(scroll_frame, text="Error al conectar con la base de datos.",
                        text_color="red").pack(pady=20)
            return
        
        try:
            expedientes = obtener_expedientes_anticipados_por_cliente(conn, cliente)
            conn.close()
            
            if not expedientes:
                ctk.CTkLabel(scroll_frame,
                           text="No se encontraron expedientes.",
                           text_color="gray").pack(pady=20)
                return
            
            # Encabezados
            header_frame = ctk.CTkFrame(scroll_frame)
            header_frame.pack(fill="x", pady=(0, 10))
            header_frame.grid_columnconfigure(0, weight=2)
            header_frame.grid_columnconfigure(1, weight=2)
            header_frame.grid_columnconfigure(2, weight=2)
            header_frame.grid_columnconfigure(3, weight=1)
            header_frame.grid_columnconfigure(4, weight=1)
            
            header_font = ctk.CTkFont(weight="bold", size=11)
            ctk.CTkLabel(header_frame, text="CÓDIGO RMA", font=header_font, anchor="w").grid(row=0, column=0, padx=10, pady=5, sticky="w")
            ctk.CTkLabel(header_frame, text="FECHA RECEPCIÓN", font=header_font, anchor="center").grid(row=0, column=1, padx=10, pady=5)
            ctk.CTkLabel(header_frame, text="FECHA AUTORIZACIÓN", font=header_font, anchor="center").grid(row=0, column=2, padx=10, pady=5)
            ctk.CTkLabel(header_frame, text="DÍAS", font=header_font, anchor="center").grid(row=0, column=3, padx=10, pady=5)
            ctk.CTkLabel(header_frame, text="ESTADO", font=header_font, anchor="center").grid(row=0, column=4, padx=10, pady=5)
            
            # Filas de expedientes
            for exp_id, codigo_rma, fecha_recep, fecha_auto, dias, estado in expedientes:
                row_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
                row_frame.pack(fill="x", pady=2)
                row_frame.grid_columnconfigure(0, weight=2)
                row_frame.grid_columnconfigure(1, weight=2)
                row_frame.grid_columnconfigure(2, weight=2)
                row_frame.grid_columnconfigure(3, weight=1)
                row_frame.grid_columnconfigure(4, weight=1)
                
                # Código (clickeable)
                codigo_btn = ctk.CTkButton(row_frame,
                                          text=codigo_rma,
                                          anchor="w",
                                          command=lambda eid=exp_id: [self._abrir_editor_rma(rma_id=eid), ventana.destroy()])
                codigo_btn.grid(row=0, column=0, padx=10, pady=2, sticky="ew")
                
                # Fechas
                ctk.CTkLabel(row_frame, text=fecha_recep or "-", anchor="center").grid(row=0, column=1, padx=10, pady=2)
                ctk.CTkLabel(row_frame, text=fecha_auto or "-", anchor="center").grid(row=0, column=2, padx=10, pady=2)
                
                # Días de adelanto
                dias_texto = f"{int(dias)}" if dias else "-"
                color = "red" if dias and dias > 7 else "orange" if dias and dias > 3 else "green"
                ctk.CTkLabel(row_frame, text=dias_texto, text_color=color, anchor="center", font=ctk.CTkFont(weight="bold")).grid(row=0, column=3, padx=10, pady=2)
                
                # Estado
                ctk.CTkLabel(row_frame, text=estado or "-", anchor="center", text_color="gray").grid(row=0, column=4, padx=10, pady=2)
        
        except Exception as e:
            print(f"Error al cargar detalle de recepciones: {e}")
            ctk.CTkLabel(scroll_frame,
                        text=f"Error: {e}",
                        text_color="red").pack(pady=20)
        
        # Botones de acción
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(pady=10)
        
        ctk.CTkButton(botones_frame, text="📄 Imprimir PDF", 
                     command=lambda: self._imprimir_recepciones_anticipadas(cliente, expedientes),
                     width=120).pack(side="left", padx=5)
        
        ctk.CTkButton(botones_frame, text="📊 Exportar Excel", 
                     command=lambda: self._exportar_recepciones_excel(cliente, expedientes),
                     width=120).pack(side="left", padx=5)
        
        ctk.CTkButton(botones_frame, text="Cerrar", command=ventana.destroy, width=100).pack(side="left", padx=5)
            

    def limpiar_marco_stats(self):
        """Método auxiliar necesario para limpiar el marco principal antes de cargar una estadística."""
        for widget in self.main_stats_frame.winfo_children():
            widget.destroy()
    
    def _exportar_recepciones_excel(self, cliente, expedientes):
        """Exporta el listado de recepciones anticipadas a Excel con formato."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from tkinter import filedialog
        except ImportError:
            messagebox.showerror("Error", "OpenPyXL no está instalado. Instale openpyxl:\npip install openpyxl")
            return
        
        if not expedientes:
            messagebox.showwarning("Exportar", "No hay datos para exportar.")
            return
        
        # Solicitar ubicación para guardar
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            initialfile=f"Recepciones_Anticipadas_{cliente.replace(' ', '_')}.xlsx"
        )
        
        if not filename:
            return
        
        try:
            # Crear libro de trabajo
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Recepciones Anticipadas"
            
            # Estilos
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Título
            ws.merge_cells('A1:F1')
            title_cell = ws['A1']
            title_cell.value = f"RECEPCIONES ANTICIPADAS - {cliente}"
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            title_cell.fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Fecha de generación
            ws.merge_cells('A2:F2')
            fecha_cell = ws['A2']
            fecha_cell.value = f"Generado el: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
            fecha_cell.alignment = Alignment(horizontal="center")
            fecha_cell.font = Font(italic=True, size=10)
            
            # Encabezados
            headers = ["Código RMA", "Fecha Recepción", "Fecha Autorización", "Días Adelanto", "Estado", "Observaciones"]
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=4, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # Datos
            for row_idx, (exp_id, codigo_rma, fecha_recep, fecha_auto, dias, estado) in enumerate(expedientes, start=5):
                ws.cell(row=row_idx, column=1, value=codigo_rma).border = border
                ws.cell(row=row_idx, column=2, value=fecha_recep or "-").border = border
                ws.cell(row=row_idx, column=3, value=fecha_auto or "-").border = border
                
                # Días con formato condicional
                dias_cell = ws.cell(row=row_idx, column=4, value=int(dias) if dias else 0)
                dias_cell.border = border
                dias_cell.alignment = Alignment(horizontal="center")
                
                if dias and dias > 7:
                    dias_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    dias_cell.font = Font(color="9C0006", bold=True)
                elif dias and dias > 3:
                    dias_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
                    dias_cell.font = Font(color="9C6500", bold=True)
                else:
                    dias_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                    dias_cell.font = Font(color="006100", bold=True)
                
                ws.cell(row=row_idx, column=5, value=estado or "-").border = border
                ws.cell(row=row_idx, column=6, value="").border = border
            
            # Ajustar anchos de columna
            ws.column_dimensions['A'].width = 18
            ws.column_dimensions['B'].width = 18
            ws.column_dimensions['C'].width = 20
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 18
            ws.column_dimensions['F'].width = 30
            
            # Estadísticas al final
            last_row = len(expedientes) + 6
            ws.merge_cells(f'A{last_row}:B{last_row}')
            stats_cell = ws[f'A{last_row}']
            stats_cell.value = "Total expedientes:"
            stats_cell.font = Font(bold=True)
            stats_cell.alignment = Alignment(horizontal="right")
            
            ws[f'C{last_row}'] = len(expedientes)
            ws[f'C{last_row}'].font = Font(bold=True)
            ws[f'C{last_row}'].fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Media de días
            if expedientes:
                media_dias = sum(dias for _, _, _, _, dias, _ in expedientes if dias) / len([d for _, _, _, _, d, _ in expedientes if d])
                ws.merge_cells(f'A{last_row+1}:B{last_row+1}')
                media_cell = ws[f'A{last_row+1}']
                media_cell.value = "Media días adelanto:"
                media_cell.font = Font(bold=True)
                media_cell.alignment = Alignment(horizontal="right")
                
                ws[f'C{last_row+1}'] = f"{media_dias:.1f}"
                ws[f'C{last_row+1}'].font = Font(bold=True)
                ws[f'C{last_row+1}'].fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Guardar
            wb.save(filename)
            messagebox.showinfo("Éxito", f"Archivo exportado correctamente:\n{filename}")
            
            # Abrir archivo
            import os
            os.startfile(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar a Excel:\n{e}")
    
    def _exportar_listado_recepciones_excel(self, resultados):
        """Exporta el listado completo de clientes con recepciones anticipadas a Excel."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from tkinter import filedialog
        except ImportError:
            messagebox.showerror("Error", "OpenPyXL no está instalado. Instale openpyxl:\npip install openpyxl")
            return
        
        if not resultados:
            messagebox.showwarning("Exportar", "No hay datos para exportar.")
            return
        
        # Solicitar ubicación para guardar
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            initialfile=f"Recepciones_Anticipadas_Listado.xlsx"
        )
        
        if not filename:
            return
        
        try:
            # Crear libro de trabajo
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Recepciones Anticipadas"
            
            # Estilos
            header_font = Font(bold=True, color="FFFFFF", size=12)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Título
            ws.merge_cells('A1:C1')
            title_cell = ws['A1']
            title_cell.value = "RECEPCIONES ANTICIPADAS - LISTADO DE CLIENTES"
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            title_cell.fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Fecha de generación
            ws.merge_cells('A2:C2')
            fecha_cell = ws['A2']
            fecha_cell.value = f"Generado el: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
            fecha_cell.alignment = Alignment(horizontal="center")
            fecha_cell.font = Font(italic=True, size=10)
            
            # Encabezados
            headers = ["Cliente", "Nº Expedientes", "Media Días Adelanto"]
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=4, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # Datos
            for row_idx, (cliente, cantidad, media_dias) in enumerate(resultados, start=5):
                ws.cell(row=row_idx, column=1, value=cliente).border = border
                ws.cell(row=row_idx, column=2, value=cantidad).border = border
                ws.cell(row=row_idx, column=2).alignment = Alignment(horizontal="center")
                
                # Media días con formato condicional
                media_cell = ws.cell(row=row_idx, column=3, value=f"{media_dias:.1f}" if media_dias else "0")
                media_cell.border = border
                media_cell.alignment = Alignment(horizontal="center")
                
                if media_dias and media_dias > 7:
                    media_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    media_cell.font = Font(color="9C0006", bold=True)
                elif media_dias and media_dias > 3:
                    media_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
                    media_cell.font = Font(color="9C6500", bold=True)
                else:
                    media_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                    media_cell.font = Font(color="006100", bold=True)
            
            # Ajustar anchos de columna
            ws.column_dimensions['A'].width = 40
            ws.column_dimensions['B'].width = 18
            ws.column_dimensions['C'].width = 22
            
            # Estadísticas al final
            last_row = len(resultados) + 6
            ws[f'A{last_row}'] = "Total clientes:"
            ws[f'A{last_row}'].font = Font(bold=True)
            ws[f'A{last_row}'].alignment = Alignment(horizontal="right")
            
            ws[f'B{last_row}'] = len(resultados)
            ws[f'B{last_row}'].font = Font(bold=True)
            ws[f'B{last_row}'].fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Total expedientes
            total_expedientes = sum(cantidad for _, cantidad, _ in resultados)
            ws[f'A{last_row+1}'] = "Total expedientes:"
            ws[f'A{last_row+1}'].font = Font(bold=True)
            ws[f'A{last_row+1}'].alignment = Alignment(horizontal="right")
            
            ws[f'B{last_row+1}'] = total_expedientes
            ws[f'B{last_row+1}'].font = Font(bold=True)
            ws[f'B{last_row+1}'].fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            
            # Guardar
            wb.save(filename)
            messagebox.showinfo("Éxito", f"Archivo exportado correctamente:\n{filename}")
            
            # Abrir archivo
            import os
            os.startfile(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al exportar a Excel:\n{e}")
    
    def _imprimir_listado_recepciones_pdf(self, resultados):
        """Genera un PDF con el listado completo de clientes con recepciones anticipadas."""
        from tkinter import filedialog
        import os
        
        if not resultados:
            messagebox.showwarning("Imprimir", "No hay datos para imprimir.")
            return
        
        # Solicitar ubicación para guardar PDF
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"Recepciones_Anticipadas_Listado.pdf"
        )
        
        if not filename:
            return
        
        try:
            # Importar reportlab
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_CENTER
            except ImportError:
                messagebox.showerror("Error", "ReportLab no está instalado. Instale reportlab:\npip install reportlab")
                return
            
            # Crear PDF
            doc = SimpleDocTemplate(filename, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#366092'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            elements.append(Paragraph("RECEPCIONES ANTICIPADAS - LISTADO DE CLIENTES", title_style))
            
            # Fecha
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey,
                spaceAfter=20,
                alignment=TA_CENTER
            )
            elements.append(Paragraph(f"Generado el: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle_style))
            
            # Tabla de datos
            data = [['Cliente', 'Nº Expedientes', 'Media Días']]
            
            for cliente, cantidad, media_dias in resultados:
                data.append([
                    cliente,
                    str(cantidad),
                    f"{media_dias:.1f}" if media_dias else "0"
                ])
            
            # Crear tabla
            table = Table(data, colWidths=[10*cm, 3*cm, 3*cm])
            
            # Estilo de tabla
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('ALIGN', (0, 1), (0, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')])
            ])
            
            # Colorear columna de media días según valor
            for i, (_, _, media_dias) in enumerate(resultados, start=1):
                if media_dias and media_dias > 7:
                    table_style.add('BACKGROUND', (2, i), (2, i), colors.HexColor('#FFC7CE'))
                    table_style.add('TEXTCOLOR', (2, i), (2, i), colors.HexColor('#9C0006'))
                elif media_dias and media_dias > 3:
                    table_style.add('BACKGROUND', (2, i), (2, i), colors.HexColor('#FFEB9C'))
                    table_style.add('TEXTCOLOR', (2, i), (2, i), colors.HexColor('#9C6500'))
                else:
                    table_style.add('BACKGROUND', (2, i), (2, i), colors.HexColor('#C6EFCE'))
                    table_style.add('TEXTCOLOR', (2, i), (2, i), colors.HexColor('#006100'))
            
            table.setStyle(table_style)
            elements.append(table)
            
            # Estadísticas
            elements.append(Spacer(1, 0.5*cm))
            total_expedientes = sum(cantidad for _, cantidad, _ in resultados)
            stats_text = f"<b>Total clientes:</b> {len(resultados)}  |  <b>Total expedientes:</b> {total_expedientes}"
            elements.append(Paragraph(stats_text, styles['Normal']))
            
            # Generar PDF
            doc.build(elements)
            messagebox.showinfo("Éxito", f"PDF generado correctamente:\n{filename}")
            
            # Abrir PDF
            os.startfile(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar PDF:\n{e}")
            print(f"Error detallado: {e}")
    
    def _imprimir_recepciones_anticipadas(self, cliente, expedientes):
        """Genera un PDF con el listado de recepciones anticipadas para imprimir."""
        from tkinter import filedialog
        import os
        
        if not expedientes:
            messagebox.showwarning("Imprimir", "No hay datos para imprimir.")
            return
        
        # Solicitar ubicación para guardar PDF
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=f"Recepciones_Anticipadas_{cliente.replace(' ', '_')}.pdf"
        )
        
        if not filename:
            return
        
        try:
            # Importar reportlab
            try:
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.lib import colors
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_CENTER, TA_LEFT
            except ImportError:
                messagebox.showerror("Error", "ReportLab no está instalado. Instale reportlab:\npip install reportlab")
                return
            
            # Crear PDF
            doc = SimpleDocTemplate(filename, pagesize=landscape(A4))
            elements = []
            styles = getSampleStyleSheet()
            
            # Título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#366092'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            elements.append(Paragraph(f"RECEPCIONES ANTICIPADAS - {cliente}", title_style))
            
            # Fecha
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey,
                spaceAfter=20,
                alignment=TA_CENTER
            )
            elements.append(Paragraph(f"Generado el: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle_style))
            
            # Tabla de datos
            data = [['Código RMA', 'Fecha Recepción', 'Fecha Autorización', 'Días', 'Estado']]
            
            for exp_id, codigo_rma, fecha_recep, fecha_auto, dias, estado in expedientes:
                data.append([
                    codigo_rma,
                    fecha_recep or "-",
                    fecha_auto or "-",
                    str(int(dias)) if dias else "0",
                    estado or "-"
                ])
            
            # Crear tabla
            table = Table(data, colWidths=[4*cm, 4*cm, 4*cm, 2*cm, 4*cm])
            
            # Estilo de tabla
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')])
            ])
            
            # Colorear columna de días según valor
            for i, row in enumerate(expedientes, start=1):
                dias = row[4]
                if dias and dias > 7:
                    table_style.add('BACKGROUND', (3, i), (3, i), colors.HexColor('#FFC7CE'))
                    table_style.add('TEXTCOLOR', (3, i), (3, i), colors.HexColor('#9C0006'))
                elif dias and dias > 3:
                    table_style.add('BACKGROUND', (3, i), (3, i), colors.HexColor('#FFEB9C'))
                    table_style.add('TEXTCOLOR', (3, i), (3, i), colors.HexColor('#9C6500'))
                else:
                    table_style.add('BACKGROUND', (3, i), (3, i), colors.HexColor('#C6EFCE'))
                    table_style.add('TEXTCOLOR', (3, i), (3, i), colors.HexColor('#006100'))
            
            table.setStyle(table_style)
            elements.append(table)
            
            # Estadísticas
            elements.append(Spacer(1, 0.5*cm))
            stats_text = f"<b>Total expedientes:</b> {len(expedientes)}"
            if expedientes:
                media_dias = sum(dias for _, _, _, _, dias, _ in expedientes if dias) / len([d for _, _, _, _, d, _ in expedientes if d])
                stats_text += f"  |  <b>Media días adelanto:</b> {media_dias:.1f}"
            
            elements.append(Paragraph(stats_text, styles['Normal']))
            
            # Generar PDF
            doc.build(elements)
            messagebox.showinfo("Éxito", f"PDF generado correctamente:\n{filename}")
            
            # Abrir PDF
            os.startfile(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar PDF:\n{e}")
            print(f"Error detallado: {e}")
    
    def exportar_a_excel(self, datos, columnas, filename):
        """Exporta los datos de la estadística actual a un archivo Excel."""
        if not datos:
            messagebox.showwarning("Exportar", "No hay datos para exportar.")
            return

        # 1. Crear el DataFrame de Pandas
        df = pd.DataFrame(datos, columns=columnas)
        
        # 2. Abrir diálogo para guardar archivo
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx", 
            filetypes=[("Archivos Excel", "*.xlsx")],
            initialfile=filename
        )
        
        if path:
            try:
                # 3. Guardar en Excel (openpyxl se usa internamente por pandas)
                df.to_excel(path, index=False) # index=False evita añadir la columna de índice de pandas
                messagebox.showinfo("Exportación Exitosa", f"Datos exportados a:\n{path}")
            except Exception as e:
                messagebox.showerror("Error de Exportación", f"No se pudo guardar el archivo Excel.\nError: {e}")

    def mostrar_tabla_estadistica(self, datos, columnas, export_filename, frame, formato_moneda=False):
        """
        Dibuja un listado genérico de resultados usando un CTkScrollableFrame 
        y añade el botón de exportación.
        """
        
        if not datos:
            ctk.CTkLabel(frame, text="No se encontraron datos para esta estadística.", text_color="gray").pack(pady=20)
            return

        # Función para abrir el expediente al hacer clic en el código RMA
        def abrir_expediente(codigo_rma):
            """Abre el expediente en una nueva ventana independiente para consulta/edición.

            La ventana muestra un resumen y permite abrir el expediente en el panel principal si
            el usuario prefiere editar en la interfaz habitual.
            """
            conn, cursor = self.master.conectar_db()
            if not conn:
                return

            try:
                # Obtener datos maestro
                cursor.execute("SELECT id, codigo_rma, cliente, fecha_gestion, motivo FROM rma_maestro WHERE codigo_rma = ?", (codigo_rma,))
                maestro = cursor.fetchone()
                if not maestro:
                    messagebox.showerror("Error", f"No se encontró el expediente con código {codigo_rma}")
                    return

                rma_id = maestro[0]
                codigo = maestro[1]
                cliente = maestro[2]
                fecha = maestro[3]
                motivo = maestro[4]

                # Crear ventana independiente
                vent = ctk.CTkToplevel(self)
                vent.title(f"Expediente {codigo}")
                vent.geometry("900x700")
                vent.minsize(700, 500)

                # Contenedor principal
                cont = ctk.CTkFrame(vent)
                cont.pack(fill="both", expand=True, padx=12, pady=12)

                # Cabecera - campos editables
                header_frame = ctk.CTkFrame(cont)
                header_frame.pack(fill="x", pady=(0,6))

                ctk.CTkLabel(header_frame, text=f"Nº EXPEDIENTE: {codigo}", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0,6))

                ctk.CTkLabel(header_frame, text="Cliente:").grid(row=1, column=0, sticky="w")
                entry_cliente = ctk.CTkEntry(header_frame)
                entry_cliente.grid(row=1, column=1, sticky="ew", padx=(8,0))
                entry_cliente.insert(0, str(cliente) if cliente is not None else "")

                ctk.CTkLabel(header_frame, text="Nº Documento:").grid(row=2, column=0, sticky="w")
                entry_num_doc = ctk.CTkEntry(header_frame)
                entry_num_doc.grid(row=2, column=1, sticky="ew", padx=(8,0))
                # intentar rellenar si existe
                try:
                    entry_num_doc.insert(0, str(maestro[3]) if maestro[3] is not None else "")
                except Exception:
                    pass

                ctk.CTkLabel(header_frame, text="Fecha Gestión:").grid(row=3, column=0, sticky="w")
                entry_fecha = ctk.CTkEntry(header_frame)
                entry_fecha.grid(row=3, column=1, sticky="ew", padx=(8,0))
                entry_fecha.insert(0, str(fecha) if fecha is not None else "")

                ctk.CTkLabel(header_frame, text="Email contacto:").grid(row=4, column=0, sticky="w")
                entry_email = ctk.CTkEntry(header_frame)
                entry_email.grid(row=4, column=1, sticky="ew", padx=(8,0))
                # maestro[6] será email si está presente
                try:
                    entry_email.insert(0, str(maestro[6]) if len(maestro) > 6 and maestro[6] is not None else "")
                except Exception:
                    pass

                header_frame.grid_columnconfigure(1, weight=1)

                ctk.CTkLabel(cont, text="Motivo:").pack(anchor="w")
                txt_motivo = ctk.CTkTextbox(cont, height=80)
                txt_motivo.pack(fill="x", pady=(0,8))
                try:
                    txt_motivo.insert("1.0", str(motivo) if motivo is not None else "")
                except Exception:
                    pass

                # Mostrar artículos relacionados (lista simple)
                ctk.CTkLabel(cont, text="Artículos asociados:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(6,4))
                art_frame = ctk.CTkScrollableFrame(cont)
                art_frame.pack(fill="both", expand=True, pady=(0,8))

                # Obtener detalles incluyendo su id para permitir edición
                cursor.execute("SELECT id, referencia_articulo, cantidad_segun_documento, cantidad_entregada, estado_producto FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                detalles = cursor.fetchall()
                if detalles:
                    for d in detalles:
                        det_id = d[0]
                        ref = d[1]
                        cant_doc = d[2]
                        cant_ent = d[3]
                        estado = d[4]

                        row_fr = ctk.CTkFrame(art_frame)
                        row_fr.pack(fill="x", padx=4, pady=2)

                        lbl = ctk.CTkLabel(row_fr, text=f"{ref} — Doc: {cant_doc} — Ent: {cant_ent} — Estado: {estado}")
                        lbl.pack(side="left", anchor="w")

                        def make_editar(did, lbl_widget):
                            def editar():
                                ed_win = ctk.CTkToplevel(vent)
                                ed_win.title(f"Editar artículo {ref}")
                                ed_win.geometry("420x200")

                                ctk.CTkLabel(ed_win, text=f"Referencia: {ref}").pack(anchor="w", padx=8, pady=(8,4))
                                ctk.CTkLabel(ed_win, text="Cantidad Entregada:").pack(anchor="w", padx=8)
                                ent_cant = ctk.CTkEntry(ed_win)
                                ent_cant.pack(fill="x", padx=8, pady=(0,6))
                                ent_cant.insert(0, str(cant_ent) if cant_ent is not None else "")

                                ctk.CTkLabel(ed_win, text="Estado Producto:").pack(anchor="w", padx=8)
                                ent_estado = ctk.CTkEntry(ed_win)
                                ent_estado.pack(fill="x", padx=8, pady=(0,6))
                                ent_estado.insert(0, str(estado) if estado is not None else "")

                                def guardar_detalle():
                                    new_cant = ent_cant.get().strip()
                                    new_estado = ent_estado.get().strip()
                                    try:
                                        conn2 = connect_db()
                                        cur2 = conn2.cursor()
                                        cur2.execute("UPDATE rma_detalles SET cantidad_entregada = ?, estado_producto = ? WHERE id = ?", (new_cant, new_estado, did))
                                        conn2.commit()
                                        conn2.close()
                                        # Actualizar etiqueta
                                        lbl_widget.configure(text=f"{ref} — Doc: {cant_doc} — Ent: {new_cant} — Estado: {new_estado}")
                                        ed_win.destroy()
                                        # Refrescar los datos de la tabla
                                        refrescar_estadisticas()
                                    except sqlite3.Error as e:
                                        messagebox.showerror("Error BD", f"No se pudo actualizar el detalle: {e}")

                                ctk.CTkButton(ed_win, text="Guardar", command=guardar_detalle).pack(side="left", padx=12, pady=12)
                                ctk.CTkButton(ed_win, text="Cerrar", command=ed_win.destroy).pack(side="right", padx=12, pady=12)

                            return editar

                        btn_ed = ctk.CTkButton(row_fr, text="Editar", command=make_editar(det_id, lbl))
                        btn_ed.pack(side="right")
                else:
                    ctk.CTkLabel(art_frame, text="No hay artículos registrados.").pack(pady=8)

                # Botones en el footer de la ventana
                footer = ctk.CTkFrame(cont)
                footer.pack(fill="x", pady=(8,0))

                # Botones en el footer de la ventana
                # Función para refrescar la tabla de estadísticas
                def refrescar_estadisticas():
                    """Refresca los datos de la tabla de estadísticas."""
                    # Limpiar y recargar los datos
                    try:
                        # Re-ejecutar la carga de datos actual
                        self._cargar_datos_articulos_incidencia()
                    except Exception as e:
                        print(f"Error al refrescar estadísticas: {e}")

                def guardar_maestro():
                    new_cliente = entry_cliente.get().strip()
                    new_num_doc = entry_num_doc.get().strip()
                    new_fecha = entry_fecha.get().strip()
                    new_email = entry_email.get().strip().lower()  # Convertir a minúsculas
                    new_motivo = txt_motivo.get("1.0", "end-1c").strip()

                    try:
                        conn2 = connect_db()
                        cur2 = conn2.cursor()
                        cur2.execute(
                            "UPDATE rma_maestro SET cliente = ?, numero_documento_cliente = ?, fecha_gestion = ?, motivo = ?, email_de_contacto = ? WHERE id = ?",
                            (new_cliente, new_num_doc, new_fecha, new_motivo, new_email, rma_id)
                        )
                        conn2.commit()
                        conn2.close()
                        messagebox.showinfo("Guardado", "Expediente actualizado correctamente.")
                        # Refrescar datos de la tabla
                        refrescar_estadisticas()
                    except sqlite3.Error as e:
                        messagebox.showerror("Error BD", f"No se pudo actualizar el expediente: {e}")

                def abrir_en_panel():
                    try:
                        self.mostrar_nuevo_rma(rma_id)
                        vent.destroy()  # Cerrar la ventana actual al abrir en panel principal
                    except Exception as e:
                        messagebox.showerror("Error", f"No se pudo abrir en el panel principal: {e}")

                ctk.CTkButton(footer, text="Guardar cambios", command=guardar_maestro).pack(side="left")
                ctk.CTkButton(footer, text="✏️ Abrir en panel principal", command=abrir_en_panel).pack(side="left", padx=8)
                ctk.CTkButton(footer, text="Cerrar", command=vent.destroy).pack(side="right")

            except sqlite3.Error as e:
                messagebox.showerror("Error de Base de Datos", f"Error al consultar el expediente: {e}")
            finally:
                conn.close()

        # Mensaje informativo para el usuario
        ctk.CTkLabel(frame, text="Pulse el Nº EXPEDIENTE (primera columna) para abrirlo en una nueva ventana.", text_color="gray").pack(pady=(6,4))

        # 1. Botón de Exportar
        btn_export = ctk.CTkButton(
            frame, 
            text="💾 Exportar a Excel", 
            command=lambda: self.exportar_a_excel(datos, columnas, export_filename) # Llama al método de exportación
        )
        btn_export.pack(pady=(5, 8))

        # 2. Contenedor de la Tabla
        tabla_scroll_frame = ctk.CTkScrollableFrame(frame, label_text="Resultados")
        tabla_scroll_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        
        # La columna 0 (típicamente el nombre del cliente/artículo) se expande
        for col_idx in range(len(columnas)):
            # Damos un peso 3 a la primera columna (Referencia Artículo)
            # y peso 1 al resto para asegurar que la Referencia se estire más.
            if col_idx == 0:
                tabla_scroll_frame.grid_columnconfigure(col_idx, weight=3) 
            else:
                tabla_scroll_frame.grid_columnconfigure(col_idx, weight=1) 
        
        # 3. Encabezados
        header_font = ctk.CTkFont(weight="bold")
        for col_idx, col_name in enumerate(columnas):
            # Alineación del encabezado
            sticky = "w" if col_idx == 0 else "e"
            ctk.CTkLabel(tabla_scroll_frame, text=col_name, font=header_font).grid(row=0, column=col_idx, padx=10, pady=5, sticky=sticky)
            
        # 4. Datos
        for row_idx, row_data in enumerate(datos):
            for col_idx, cell_value in enumerate(row_data):
                
                text_to_display = str(cell_value if cell_value is not None else "N/A")
                
                # 🚨 Aplicar formato de moneda si se especifica y estamos en la columna de valor (la última)
                if formato_moneda and col_idx == len(row_data) - 1 and cell_value is not None:
                    try:
                        # Intenta usar locale para el formato de moneda (€)
                        text_to_display = locale.currency(float(cell_value), grouping=True, symbol=True)
                    except (ValueError, TypeError, locale.Error):
                        # Fallback si no es un número o si falla el locale
                        text_to_display = f"{float(cell_value):,.2f} €"

                # Alineación del contenido
                sticky = "w" if col_idx == 0 else "e"
                
                # Crear el label como widget para poder bindear eventos
                label_widget = ctk.CTkLabel(
                    tabla_scroll_frame,
                    text=text_to_display,
                    justify="left"
                )
                label_widget.grid(row=row_idx + 1, column=col_idx, padx=10, pady=2, sticky=sticky)

                # Si es la primera columna (código RMA), hacerlo clicable
                if col_idx == 0:
                    # Cambiar cursor a mano y subrayar/colorear al hover
                    try:
                        original_color = label_widget.cget("text_color")
                    except Exception:
                        original_color = None

                    # Set cursor to hand2 if supported
                    try:
                        label_widget.configure(cursor="hand2")
                    except Exception:
                        pass

                    def on_enter(e):
                        try:
                            label_widget.configure(text_color="#1a73e8")
                        except Exception:
                            pass

                    def on_leave(e):
                        try:
                            if original_color is not None:
                                label_widget.configure(text_color=original_color)
                        except Exception:
                            pass

                    # Bind click to abrir_expediente using the displayed text (codigo)
                    label_widget.bind("<Button-1>", lambda e, code=text_to_display: abrir_expediente(code))
                    label_widget.bind("<Enter>", on_enter)
                    label_widget.bind("<Leave>", on_leave)
    
    def abrir_cliente_correo_con_mailto(self, email_acontacto, asunto, cuerpo):
        """Prepara el enlace mailto y lo abre con el cliente de correo por defecto."""
        
        # Necesitamos la librería para codificar el Asunto y el Cuerpo para la URL
        import urllib.parse
        
        # Codificar el asunto y el cuerpo (necesario para URLs mailto con espacios o caracteres especiales)
        asunto_codificado = urllib.parse.quote(asunto)
        cuerpo_codificado = urllib.parse.quote(cuerpo)
        
        # Construir el enlace mailto
        mailto_link = f"mailto:{email_acontacto}?subject={asunto_codificado}&body={cuerpo_codificado}"
        
        try:
            # Abrir el enlace usando el cliente de correo/navegador por defecto
            import webbrowser
            webbrowser.open(mailto_link)
            return True # Éxito
        except Exception as e:
            messagebox.showerror("Error de Email", "No se pudo abrir el cliente de correo por defecto. Por favor, asegúrate de que tienes un cliente de correo configurado.")
            return False # Fallo
    
    def enviar_email_contacto(self):
        """Abre el cliente de correo para enviar un email al contacto registrado en el RMA actual."""
        
        # 1. Verificar que haya un RMA actual abierto
        if not self.rma_actual_id:
            messagebox.showwarning("Advertencia", "Debe seleccionar o tener abierto un expediente RMA para enviar un email.")
            return

        # 2. Mostrar diálogo para seleccionar adjuntos (si hay disponibles)
        self.mostrar_dialogo_adjuntos_email()

    def mostrar_dialogo_adjuntos_email(self):
        """Muestra un diálogo para seleccionar adjuntos antes de enviar el email."""
        
        conn = connect_db()
        cursor = conn.cursor()
        
        try:
            # Consultar datos del RMA
            cursor.execute("SELECT email_de_contacto, cliente, numero_documento_cliente, codigo_rma FROM rma_maestro WHERE id = ?", (self.rma_actual_id,))
            resultado = cursor.fetchone()
            
            if not resultado:
                messagebox.showerror("Error", f"No se encontró el expediente RMA: {self.rma_actual_id}")
                return
            
            email_acontacto, nombre_cliente, numero_documento_cliente, numero_rma = resultado
            
            # Verificar si el email_contacto está vacío o es nulo
            if not email_acontacto:
                messagebox.showwarning("Advertencia", f"El campo 'email_de_contacto' del expediente {self.rma_actual_id} está vacío.")
                return

            # Obtener lista de adjuntos disponibles
            cursor.execute("SELECT id, ruta_relativa, nombre_archivo FROM rma_adjuntos WHERE rma_id = ?", (self.rma_actual_id,))
            adjuntos_disponibles = cursor.fetchall()
            
            # Crear ventana de diálogo
            dialogo = Toplevel(self)
            dialogo.title("📧 Preparar Email con Adjuntos")
            dialogo.geometry("700x600")
            dialogo.resizable(True, True)
            dialogo.transient(self)
            dialogo.grab_set()
            
            # Centrar en pantalla
            dialogo.update_idletasks()
            x = (dialogo.winfo_screenwidth() // 2) - (700 // 2)
            y = (dialogo.winfo_screenheight() // 2) - (600 // 2)
            dialogo.geometry(f"700x600+{x}+{y}")
            
            # Frame principal
            main_frame = ctk.CTkFrame(dialogo)
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Título
            titulo = ctk.CTkLabel(main_frame, text=f"📧 Email para: {nombre_cliente}", 
                                 font=ctk.CTkFont(size=18, weight="bold"))
            titulo.pack(pady=(10, 5))
            
            # Información del RMA
            info_frame = ctk.CTkFrame(main_frame)
            info_frame.pack(fill="x", pady=(0, 20))
            
            ctk.CTkLabel(info_frame, text=f"RMA: {numero_rma}", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(pady=5)
            ctk.CTkLabel(info_frame, text=f"Email: {email_acontacto}", 
                        font=ctk.CTkFont(size=12)).pack(pady=2)
            ctk.CTkLabel(info_frame, text=f"Doc. Cliente: {numero_documento_cliente}", 
                        font=ctk.CTkFont(size=12)).pack(pady=2)
            
            # Sección de adjuntos
            adjuntos_frame = ctk.CTkFrame(main_frame)
            adjuntos_frame.pack(fill="both", expand=True, pady=(0, 20))
            
            ctk.CTkLabel(adjuntos_frame, text="📎 Seleccionar adjuntos para el email:", 
                        font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 10))
            
            # Variables para almacenar selecciones
            self.adjuntos_seleccionados = []
            self.checkboxes_adjuntos = []
            
            if adjuntos_disponibles:
                # Frame scrollable para adjuntos
                scroll_frame = ctk.CTkScrollableFrame(adjuntos_frame, height=200)
                scroll_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
                
                for adjunto_id, ruta_relativa, nombre_archivo in adjuntos_disponibles:
                    # Frame para cada adjunto
                    adj_frame = ctk.CTkFrame(scroll_frame)
                    adj_frame.pack(fill="x", pady=2, padx=5)
                    
                    # Variable de checkbox
                    var_checkbox = ctk.BooleanVar()
                    
                    # Checkbox
                    checkbox = ctk.CTkCheckBox(
                        adj_frame, 
                        text="",
                        variable=var_checkbox,
                        width=20
                    )
                    checkbox.pack(side="left", padx=(10, 5), pady=10)
                    
                    # Información del archivo
                    info_label = ctk.CTkLabel(
                        adj_frame, 
                        text=f"{nombre_archivo or os.path.basename(ruta_relativa)}",
                        font=ctk.CTkFont(size=11)
                    )
                    info_label.pack(side="left", padx=(0, 10), pady=10, fill="x", expand=True)
                    
                    # Indicador de ubicación (Dropbox/Local)
                    ubicacion = "☁️ Dropbox" if usar_dropbox() else "💾 Local"
                    ubicacion_label = ctk.CTkLabel(
                        adj_frame,
                        text=ubicacion,
                        font=ctk.CTkFont(size=10),
                        text_color="gray"
                    )
                    ubicacion_label.pack(side="right", padx=(5, 10), pady=10)
                    
                    # Guardar referencia
                    self.checkboxes_adjuntos.append((var_checkbox, adjunto_id, ruta_relativa, nombre_archivo))
                
                # Botones de selección rápida
                botones_sel_frame = ctk.CTkFrame(adjuntos_frame)
                botones_sel_frame.pack(fill="x", padx=10, pady=(0, 10))
                
                ctk.CTkButton(botones_sel_frame, text="✅ Seleccionar todos",
                             command=self._seleccionar_todos_adjuntos,
                             width=120).pack(side="left", padx=5, pady=5)
                
                ctk.CTkButton(botones_sel_frame, text="❌ Deseleccionar todos",
                             command=self._deseleccionar_todos_adjuntos,
                             width=120).pack(side="left", padx=5, pady=5)
            else:
                # No hay adjuntos
                no_adj_label = ctk.CTkLabel(adjuntos_frame, 
                                           text="📋 No hay adjuntos disponibles en este RMA",
                                           font=ctk.CTkFont(size=12),
                                           text_color="gray")
                no_adj_label.pack(pady=20)
            
            # Frame de botones principales
            botones_frame = ctk.CTkFrame(main_frame)
            botones_frame.pack(fill="x", pady=(0, 10))
            
            # Botón para continuar con email
            btn_email = ctk.CTkButton(
                botones_frame, 
                text="📧 Abrir cliente de correo",
                command=lambda: self._procesar_email_con_adjuntos(dialogo, email_acontacto, nombre_cliente, numero_documento_cliente, numero_rma),
                font=ctk.CTkFont(size=12, weight="bold"),
                height=40
            )
            btn_email.pack(side="left", padx=(10, 5), pady=10, fill="x", expand=True)
            
            # Botón cancelar
            btn_cancelar = ctk.CTkButton(
                botones_frame, 
                text="❌ Cancelar",
                command=dialogo.destroy,
                fg_color="#D32F2F", 
                hover_color="#B71C1C",
                width=100
            )
            btn_cancelar.pack(side="right", padx=(5, 10), pady=10)
            
            # Nota informativa
            nota_frame = ctk.CTkFrame(main_frame)
            nota_frame.pack(fill="x")
            
            nota_text = (
                "💡 Nota: Los archivos seleccionados se descargarán temporalmente para adjuntarlos al email.\n"
                "El cliente de correo se abrirá con el email preparado y los archivos listos para adjuntar."
            )
            
            ctk.CTkLabel(nota_frame, text=nota_text, 
                        font=ctk.CTkFont(size=10), 
                        text_color="gray",
                        wraplength=650).pack(pady=10, padx=10)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error preparando email: {e}")
        finally:
            conn.close()
    
    def _seleccionar_todos_adjuntos(self):
        """Selecciona todos los checkboxes de adjuntos."""
        for var_checkbox, _, _, _ in self.checkboxes_adjuntos:
            var_checkbox.set(True)
    
    def _deseleccionar_todos_adjuntos(self):
        """Deselecciona todos los checkboxes de adjuntos."""
        for var_checkbox, _, _, _ in self.checkboxes_adjuntos:
            var_checkbox.set(False)
    
    def _procesar_email_con_adjuntos(self, dialogo, email_acontacto, nombre_cliente, numero_documento_cliente, numero_rma):
        """Procesa los adjuntos seleccionados y abre el cliente de correo."""
        
        # Recopilar adjuntos seleccionados
        adjuntos_para_email = []
        for var_checkbox, adjunto_id, ruta_relativa, nombre_archivo in self.checkboxes_adjuntos:
            if var_checkbox.get():
                adjuntos_para_email.append((adjunto_id, ruta_relativa, nombre_archivo))
        
        # Cerrar diálogo
        dialogo.destroy()
        
        # Continuar con el proceso de email
        self._enviar_email_con_adjuntos_seleccionados(
            email_acontacto, nombre_cliente, numero_documento_cliente, numero_rma, adjuntos_para_email
        )
    
    def _enviar_email_con_adjuntos_seleccionados(self, email_acontacto, nombre_cliente, numero_documento_cliente, numero_rma, adjuntos_seleccionados):
        """Envía email con los adjuntos seleccionados."""
        
        try:
            # Lista para almacenar rutas temporales de archivos descargados
            archivos_temporales = []
            
            if adjuntos_seleccionados and usar_dropbox():
                # Descargar archivos seleccionados de Dropbox
                dbx = get_dropbox_client()
                if not dbx:
                    messagebox.showerror("Error", "No se puede conectar con Dropbox para descargar adjuntos.")
                    return
                
                # Crear directorio temporal
                temp_dir = tempfile.mkdtemp(prefix="email_attachments_")
                
                for adjunto_id, ruta_relativa, nombre_archivo in adjuntos_seleccionados:
                    try:
                        # Construir ruta en Dropbox
                        ruta_dropbox = normalizar_ruta_dropbox(f"{DROPBOX_ROOT_FOLDER}/{ruta_relativa}")
                        
                        # Nombre del archivo temporal
                        nombre_archivo_final = nombre_archivo or os.path.basename(ruta_relativa)
                        temp_file_path = os.path.join(temp_dir, nombre_archivo_final)
                        
                        # Descargar archivo
                        print(f"Descargando {nombre_archivo_final} para adjuntar al email...")
                        metadata, response = dbx.files_download(ruta_dropbox)
                        
                        # Guardar en archivo temporal
                        with open(temp_file_path, 'wb') as temp_file:
                            temp_file.write(response.content)
                        
                        archivos_temporales.append(temp_file_path)
                        print(f"✓ Descargado: {nombre_archivo_final}")
                        
                    except Exception as e:
                        print(f"Error descargando {nombre_archivo_final}: {e}")
                        messagebox.showwarning("Advertencia", f"No se pudo descargar el archivo '{nombre_archivo_final}': {e}")
            
            # Definir Asunto y Cuerpo del email
            asunto_base = f"ILUTREK - Nº {numero_rma} ; DOCUMENTO CLIENTE: {numero_documento_cliente}"
            
            # Texto base del email
            cuerpo_base = (
                f"Buenos dias,\n\n"
                f"Se adjunta resolución sobre el expediente abierto a su número de devolución:\n"
                f"{numero_documento_cliente}.\n"
                f"Para saber el estado de este informe, puede responder a este mismo correo.\n\n"
                f"Transcurridos 15 días del envío de este correo, se dará por cerrado el expediente, no aceptando ningún tipo de no conformidad a esta resolución.\n\n\n"
                f"Dpto. Tecnico Ilutrek."
            )
            
            # Añadir información sobre adjuntos al cuerpo del email
            if adjuntos_seleccionados:
                cuerpo_base += f"\n\n--- Adjuntos incluidos ---\n"
                for _, _, nombre_archivo in adjuntos_seleccionados:
                    archivo_nombre = nombre_archivo or "Archivo adjunto"
                    cuerpo_base += f"• {archivo_nombre}\n"
            
            # Mostrar diálogo informativo sobre los archivos descargados
            if archivos_temporales:
                archivos_list = "\n".join([f"• {os.path.basename(f)}" for f in archivos_temporales])
                info_message = (
                    f"✅ Archivos descargados y listos para adjuntar:\n\n{archivos_list}\n\n"
                    f"📁 Ubicación temporal: {os.path.dirname(archivos_temporales[0])}\n\n"
                    f"Se abrirá tu cliente de correo. Adjunta manualmente estos archivos al email."
                )
                
                respuesta = messagebox.askokcancel(
                    "Archivos preparados", 
                    info_message + "\n\n¿Abrir cliente de correo ahora?",
                    icon="info"
                )
                
                if not respuesta:
                    # Usuario canceló, limpiar archivos temporales
                    self._limpiar_archivos_temporales(archivos_temporales)
                    return
            
            # Abrir cliente de correo con mailto
            email_abierto_ok = self.abrir_cliente_correo_con_mailto(email_acontacto, asunto_base, cuerpo_base)
            
            if email_abierto_ok:
                # Mostrar información sobre archivos temporales si los hay
                if archivos_temporales:
                    # Crear ventana de información persistente
                    self._crear_ventana_archivos_temporales(archivos_temporales, temp_dir)
                
                # Registrar la acción en el historial
                self._registrar_accion_email_historial(numero_rma, email_acontacto, adjuntos_seleccionados)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error enviando email con adjuntos: {e}")
    
    def _crear_ventana_archivos_temporales(self, archivos_temporales, temp_dir):
        """Crea una ventana informativa sobre los archivos temporales descargados."""
        
        ventana_info = Toplevel(self)
        ventana_info.title("📎 Archivos temporales para email")
        ventana_info.geometry("500x400")
        ventana_info.resizable(False, False)
        
        # Mantener ventana siempre visible
        ventana_info.transient(self)
        ventana_info.lift()
        ventana_info.focus_force()
        
        # Centrar
        ventana_info.update_idletasks()
        x = (ventana_info.winfo_screenwidth() // 2) - (500 // 2)
        y = (ventana_info.winfo_screenheight() // 2) - (400 // 2)
        ventana_info.geometry(f"500x400+{x}+{y}")
        
        # Frame principal
        main_frame = ctk.CTkFrame(ventana_info)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Título
        ctk.CTkLabel(main_frame, text="📎 Archivos listos para adjuntar", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(10, 20))
        
        # Lista de archivos
        ctk.CTkLabel(main_frame, text="Archivos descargados:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10)
        
        scroll_frame = ctk.CTkScrollableFrame(main_frame, height=150)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        for archivo in archivos_temporales:
            archivo_frame = ctk.CTkFrame(scroll_frame)
            archivo_frame.pack(fill="x", pady=2)
            
            ctk.CTkLabel(archivo_frame, text=f"📄 {os.path.basename(archivo)}", 
                        font=ctk.CTkFont(size=11)).pack(side="left", padx=10, pady=5)
            
            # Botón para abrir carpeta
            ctk.CTkButton(archivo_frame, text="📁", width=30,
                         command=lambda f=archivo: self._abrir_carpeta_archivo(f)).pack(side="right", padx=5, pady=5)
        
        # Información
        info_text = (
            f"📁 Ubicación: {temp_dir}\n\n"
            f"💡 Instrucciones:\n"
            f"1. Ve a tu cliente de correo que se acaba de abrir\n"
            f"2. Adjunta manualmente los archivos mostrados arriba\n"
            f"3. Envía el email cuando esté listo\n"
            f"4. Los archivos temporales se eliminarán automáticamente"
        )
        
        ctk.CTkLabel(main_frame, text=info_text, 
                    font=ctk.CTkFont(size=10), 
                    text_color="gray",
                    justify="left").pack(pady=10, padx=10, anchor="w")
        
        # Botones
        botones_frame = ctk.CTkFrame(main_frame)
        botones_frame.pack(fill="x", pady=10)
        
        ctk.CTkButton(botones_frame, text="📁 Abrir carpeta",
                     command=lambda: self._abrir_carpeta_archivo(archivos_temporales[0])).pack(side="left", padx=10)
        
        ctk.CTkButton(botones_frame, text="🗑️ Limpiar archivos",
                     command=lambda: self._confirmar_limpiar_archivos(archivos_temporales, ventana_info),
                     fg_color="#D32F2F", hover_color="#B71C1C").pack(side="right", padx=10)
    
    def _abrir_carpeta_archivo(self, ruta_archivo):
        """Abre la carpeta que contiene el archivo."""
        try:
            folder_path = os.path.dirname(ruta_archivo)
            if sys.platform.startswith('win'):
                os.startfile(folder_path)
            elif sys.platform.startswith('darwin'):  # macOS
                os.system(f'open "{folder_path}"')
            else:  # Linux
                os.system(f'xdg-open "{folder_path}"')
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta: {e}")
    
    def _confirmar_limpiar_archivos(self, archivos_temporales, ventana):
        """Confirma y limpia los archivos temporales."""
        respuesta = messagebox.askyesno(
            "Limpiar archivos", 
            "¿Estás seguro de que quieres eliminar los archivos temporales?\n\n"
            "Solo hazlo después de haber adjuntado los archivos a tu email."
        )
        
        if respuesta:
            self._limpiar_archivos_temporales(archivos_temporales)
            ventana.destroy()
            messagebox.showinfo("Limpieza completada", "Archivos temporales eliminados correctamente.")
    
    def _limpiar_archivos_temporales(self, archivos_temporales):
        """Limpia los archivos temporales descargados."""
        for archivo in archivos_temporales:
            try:
                if os.path.exists(archivo):
                    os.remove(archivo)
            except Exception as e:
                print(f"Error eliminando archivo temporal {archivo}: {e}")
        
        # Intentar eliminar el directorio temporal si está vacío
        try:
            if archivos_temporales:
                temp_dir = os.path.dirname(archivos_temporales[0])
                if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                    os.rmdir(temp_dir)
        except Exception as e:
            print(f"Error eliminando directorio temporal: {e}")
    
    def _registrar_accion_email_historial(self, numero_rma, email_destinatario, adjuntos_seleccionados):
        """Registra la acción de envío de email en el historial con nombres de adjuntos."""
        try:
            conn = connect_db()
            cursor = conn.cursor()
            
            descripcion = f"Email enviado a {email_destinatario}"
            
            if adjuntos_seleccionados:
                nombres_adjuntos = []
                for _, _, nombre_archivo in adjuntos_seleccionados:
                    nombre = nombre_archivo or "Archivo sin nombre"
                    nombres_adjuntos.append(nombre)
                
                if len(nombres_adjuntos) == 1:
                    descripcion += f" con adjunto: {nombres_adjuntos[0]}"
                else:
                    # Para múltiples adjuntos, mostrar lista
                    lista_adjuntos = ", ".join(nombres_adjuntos)
                    descripcion += f" con adjuntos: {lista_adjuntos}"
            
            cursor.execute("""
                INSERT INTO rma_historial (rma_id, fecha_cambio, usuario, descripcion_cambio)
                VALUES (?, ?, ?, ?)
            """, (self.rma_actual_id, datetime.datetime.now().isoformat(), self.username, descripcion))
            
            conn.commit()
            conn.close()
            
            # Actualizar historial en la interfaz si está visible
            if hasattr(self, 'historial_tab'):
                self.mostrar_historial(self.historial_tab)
                
        except Exception as e:
            print(f"Error registrando email en historial: {e}")
    
    # Nota: la implementación anterior de 'mostrar_formulario_github' era un duplicado
    # y ha sido eliminada para evitar confusión. La implementación activa está más
    # abajo en el archivo y contiene el botón de envío correctamente visible.
        
    def mostrar_formulario_github(self):
        """Muestra un formulario para crear un issue en GitHub."""
        # Crear una nueva ventana modal con tamaño fijo (no maximizable)
        ventana_issue = ctk.CTkToplevel(self)
        ventana_issue.title("Crear Issue en GitHub")
        ventana_issue.geometry("500x600")  # Altura aumentada para asegurar que todo sea visible
        ventana_issue.resizable(False, False)  # Ventana de tamaño fijo
        ventana_issue.grab_set()  # Hace la ventana modal

        # Marco con scroll para el contenido
        scroll_frame = ctk.CTkScrollableFrame(ventana_issue)
        scroll_frame.pack(fill="both", expand=True, padx=12, pady=12)

        # Marco para el contenido dentro del scroll
        content_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
        content_frame.pack(fill="both", expand=True)

        # Campo de nombre (autocompletado con el usuario actual)
        ctk.CTkLabel(content_frame, text="Nombre:", font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=5, pady=(4,2))
        nombre_entry = ctk.CTkEntry(content_frame)
        nombre_entry.insert(0, self.username)
        nombre_entry.pack(fill="x", padx=5, pady=(0,12))

        # Campo de fecha (autocompletado con la fecha actual)
        ctk.CTkLabel(content_frame, text="Fecha:").pack(anchor="w", padx=5, pady=(4,2))
        fecha_entry = ctk.CTkEntry(content_frame)
        fecha_entry.insert(0, datetime.datetime.now().strftime("%Y-%m-%d"))
        fecha_entry.pack(fill="x", padx=5, pady=(0,8))

        # Tipo de issue (Bug/Sugerencia)
        ctk.CTkLabel(content_frame, text="Tipo:").pack(anchor="w", padx=5, pady=(4,2))
        tipo_var = tk.StringVar(value="Sugerencia")
        tipo_frame = ctk.CTkFrame(content_frame)
        tipo_frame.pack(fill="x", padx=5, pady=(0,8))
        ctk.CTkRadioButton(tipo_frame, text="Bug", variable=tipo_var, value="Bug").pack(side="left", padx=12)
        ctk.CTkRadioButton(tipo_frame, text="Sugerencia", variable=tipo_var, value="Sugerencia").pack(side="left", padx=12)

        # Campo de descripción
        ctk.CTkLabel(content_frame, text="Descripción:").pack(anchor="w", padx=5, pady=(4,2))
        desc_text = ctk.CTkTextbox(content_frame, height=160)
        desc_text.pack(fill="both", expand=True, padx=5, pady=(0,8))

        # Footer con el botón para asegurarnos que está siempre visible
        footer = ctk.CTkFrame(ventana_issue, fg_color="transparent", height=50)
        footer.pack(side="bottom", fill="x", padx=12, pady=8)
        footer.pack_propagate(False)  # Mantiene la altura fija del footer

        # Función para enviar el issue
        def enviar_issue():
            nombre = nombre_entry.get().strip()
            fecha = fecha_entry.get().strip()
            tipo = tipo_var.get()
            descripcion = desc_text.get("1.0", "end-1c").strip()

            if not nombre or not fecha or not descripcion:
                messagebox.showerror("Error", "Por favor, complete todos los campos.")
                return

            titulo = f"[{tipo}] Reporte de {nombre}"
            cuerpo = (
                f"Reporte creado por: {nombre}\n"
                f"Fecha: {fecha}\n"
                f"Tipo: {tipo}\n"
                f"Rol del usuario: {self.rol}\n"
                f"Versión de la App: {APP_VERSION}\n\n"
                f"Descripción:\n{descripcion}\n"
            )

            token = os.getenv('GITHUB_TOKEN')
            if not token:
                messagebox.showerror("Error", "No se encontró el token de GitHub. Contacte al administrador del sistema.")
                return

            try:
                url = "https://api.github.com/repos/ilutreksl/Gestor_Expedientes/issues"
                headers = {
                    "Authorization": f"token {token}",
                    "Accept": "application/vnd.github.v3+json"
                }
                data = {"title": titulo, "body": cuerpo, "labels": [tipo]}

                response = requests.post(url, headers=headers, json=data)
                response.raise_for_status()

                issue_number = response.json().get('number')
                message = f"Issue creado correctamente en GitHub." if not issue_number else f"El issue #{issue_number} ha sido creado correctamente en GitHub."
                messagebox.showinfo("Éxito", message)
                ventana_issue.destroy()

            except requests.exceptions.RequestException as e:
                messagebox.showerror("Error", f"Error al crear el issue en GitHub: {e}")
            except Exception as e:
                messagebox.showerror("Error", f"Error al enviar el reporte: {e}")

        # Botón de enviar en el footer (posicionamiento absoluto)
        btn_enviar = ctk.CTkButton(
            footer, 
            text="✉️ Enviar Issue", 
            command=enviar_issue,
            width=120,  # Ancho fijo para el botón
            height=32   # Altura fija para el botón
        )
        btn_enviar.place(relx=0.95, rely=0.5, anchor="e")  # Posicionamiento relativo a la derecha

    def abrir_formulario_email(self):
        """Abre el cliente de correo por defecto con un enlace mailto preconfigurado."""
        
        # 🚨 ¡IMPORTANTE! Reemplaza 'tu.email@empresa.com' por la dirección correcta
        email_destino = "carlos@ilutrek.es"
        
        # Generamos un asunto útil que incluye el nombre de la aplicación y el usuario
        asunto = f"Bug/Sugerencia de {self.username} ({self.rol}) en Gestión RMA"
        
        # Creamos un cuerpo predefinido para guiar al compañero
        cuerpo = f"""
Hola,
        
Por favor, describe el problema o sugerencia a continuación:
        
----------------------------------------------------------------------
        
Tipo: [BUG / SUGERENCIA]
Página/Función: [Ej: Listado RMA / Formulario Edición]
Descripción: 
Pasos para reproducir (solo bugs):
        
----------------------------------------------------------------------
        
Información del Sistema (NO MODIFICAR):
Usuario: {self.username}
Versión de la App: {APP_VERSION}
"""
        
        # URL-encode el asunto y el cuerpo (necesario para URLs mailto)
        import urllib.parse
        asunto_codificado = urllib.parse.quote(asunto)
        cuerpo_codificado = urllib.parse.quote(cuerpo)
        
        # Construir el enlace mailto
        mailto_link = f"mailto:{email_destino}?subject={asunto_codificado}&body={cuerpo_codificado}"
        
        try:
            # Abrir el enlace usando el cliente de correo/navegador por defecto
            webbrowser.open(mailto_link)
        except Exception as e:
            # Si falla (ej. en un entorno sin navegador o cliente de correo configurado)
            print(f"Error al abrir el cliente de correo: {e}")
            messagebox.showerror("Error de Email", "No se pudo abrir el cliente de correo por defecto. Por favor, envía un email manualmente a " + email_destino)
    
    def mostrar_gestor_backups(self):
        """Muestra la ventana de gestión de backups en Backblaze B2"""
        try:
            from tkinter import filedialog
            import threading
            
            # Crear ventana
            ventana = ctk.CTkToplevel(self)
            ventana.title("Gestor de Backups - Backblaze B2")
            ventana.geometry("1000x600")
            ventana.transient(self)
            
            # Inicializar manager
            manager = BackupManagerB2()
            
            # Variable para almacenar archivos
            archivos_actuales = []
            
            # Frame principal
            main_frame = ctk.CTkFrame(ventana)
            main_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            # ===== PARTE SUPERIOR: CONTROLES =====
            control_frame = ctk.CTkFrame(main_frame)
            control_frame.pack(fill="x", padx=5, pady=(5, 10))
            
            # Fila 1: Búsqueda y filtros
            filtros_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
            filtros_frame.pack(fill="x", padx=5, pady=5)
            
            # Búsqueda
            ctk.CTkLabel(filtros_frame, text="Buscar:").pack(side="left", padx=(0, 5))
            entry_busqueda = ctk.CTkEntry(filtros_frame, placeholder_text="Nombre de archivo...", width=250)
            entry_busqueda.pack(side="left", padx=5)
            
            # Filtro por tipo
            ctk.CTkLabel(filtros_frame, text="Tipo:").pack(side="left", padx=(20, 5))
            filtro_tipo = ctk.CTkOptionMenu(filtros_frame, values=["Todos", ".db", ".sql"], width=100)
            filtro_tipo.set("Todos")
            filtro_tipo.pack(side="left", padx=5)
            
            # Filtro por ubicación
            ctk.CTkLabel(filtros_frame, text="Ubicación:").pack(side="left", padx=(20, 5))
            filtro_ubicacion = ctk.CTkOptionMenu(filtros_frame, values=["Todos", "Raíz", "Archivo/"], width=120)
            filtro_ubicacion.set("Todos")
            filtro_ubicacion.pack(side="left", padx=5)
            
            # Fila 2: Botones de acción
            botones_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
            botones_frame.pack(fill="x", padx=5, pady=5)
            
            btn_crear_backup = ctk.CTkButton(botones_frame, text="🔄 Crear Backup Ahora", width=150)
            btn_crear_backup.pack(side="left", padx=5)
            Tooltip(btn_crear_backup, "Ejecuta el script de backup para crear una nueva copia de seguridad de la base de datos")
            
            btn_actualizar = ctk.CTkButton(botones_frame, text="↻ Actualizar Lista", width=120)
            btn_actualizar.pack(side="left", padx=5)
            Tooltip(btn_actualizar, "Recarga la lista de archivos desde Backblaze B2")
            
            # Label de estado
            lbl_estado = ctk.CTkLabel(botones_frame, text="Listo", text_color="gray")
            lbl_estado.pack(side="left", padx=20)
            
            # ===== PARTE CENTRAL: LISTADO DE ARCHIVOS =====
            lista_frame = ctk.CTkFrame(main_frame)
            lista_frame.pack(fill="both", expand=True, padx=5, pady=5)
            
            # Header del listado
            header_frame = ctk.CTkFrame(lista_frame, fg_color=("gray80", "gray25"), height=35)
            header_frame.pack(fill="x", padx=2, pady=(2, 0))
            header_frame.pack_propagate(False)
            
            ctk.CTkLabel(header_frame, text="Nombre", font=ctk.CTkFont(weight="bold"), width=350).pack(side="left", padx=10, pady=5)
            ctk.CTkLabel(header_frame, text="Tamaño", font=ctk.CTkFont(weight="bold"), width=100).pack(side="left", padx=5, pady=5)
            ctk.CTkLabel(header_frame, text="Fecha", font=ctk.CTkFont(weight="bold"), width=150).pack(side="left", padx=5, pady=5)
            ctk.CTkLabel(header_frame, text="Ubicación", font=ctk.CTkFont(weight="bold"), width=100).pack(side="left", padx=5, pady=5)
            ctk.CTkLabel(header_frame, text="Acciones", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=5, pady=5)
            
            # Scrollable frame para archivos
            scroll_frame = ctk.CTkScrollableFrame(lista_frame)
            scroll_frame.pack(fill="both", expand=True, padx=2, pady=2)
            
            # ===== PARTE INFERIOR: INFORMACIÓN =====
            info_frame = ctk.CTkFrame(main_frame, height=40)
            info_frame.pack(fill="x", padx=5, pady=(5, 5))
            info_frame.pack_propagate(False)
            
            lbl_info = ctk.CTkLabel(info_frame, text="Total: 0 archivos | Espacio: 0 B")
            lbl_info.pack(side="left", padx=10, pady=10)
            
            # ===== FUNCIONES =====
            
            def actualizar_estado(mensaje, color="gray"):
                lbl_estado.configure(text=mensaje, text_color=color)
                ventana.update()
            
            def cargar_archivos():
                """Carga la lista de archivos desde B2"""
                actualizar_estado("Autenticando...", "blue")
                
                # Autenticar
                success, msg = manager.autenticar()
                if not success:
                    messagebox.showerror("Error", f"Error de autenticación: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
                    return
                
                # Obtener bucket
                success, msg = manager.obtener_bucket_id()
                if not success:
                    messagebox.showerror("Error", f"Error al obtener bucket: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
                    return
                
                actualizar_estado("Cargando archivos...", "blue")
                
                # Listar archivos
                archivos, msg = manager.listar_archivos()
                if archivos is None:
                    messagebox.showerror("Error", f"Error al listar archivos: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
                    return
                
                nonlocal archivos_actuales
                archivos_actuales = archivos
                
                actualizar_estado(f"Cargados {len(archivos)} archivos", "green")
                mostrar_archivos()
            
            def mostrar_archivos():
                """Muestra los archivos filtrados en la lista"""
                # Limpiar lista actual
                for widget in scroll_frame.winfo_children():
                    widget.destroy()
                
                # Aplicar filtros
                busqueda = entry_busqueda.get().lower()
                tipo = filtro_tipo.get()
                ubicacion = filtro_ubicacion.get()
                
                archivos_filtrados = []
                for archivo in archivos_actuales:
                    nombre = archivo['fileName']
                    
                    # Filtro de búsqueda
                    if busqueda and busqueda not in nombre.lower():
                        continue
                    
                    # Filtro de tipo
                    if tipo != "Todos":
                        if not nombre.endswith(tipo):
                            continue
                    
                    # Filtro de ubicación
                    if ubicacion == "Raíz" and nombre.startswith("Archivo/"):
                        continue
                    elif ubicacion == "Archivo/" and not nombre.startswith("Archivo/"):
                        continue
                    
                    archivos_filtrados.append(archivo)
                
                # Mostrar archivos
                if not archivos_filtrados:
                    lbl_vacio = ctk.CTkLabel(scroll_frame, text="No hay archivos que coincidan con los filtros", 
                                             text_color="gray")
                    lbl_vacio.pack(pady=20)
                else:
                    for archivo in archivos_filtrados:
                        crear_fila_archivo(archivo)
                
                # Actualizar información
                total_tamano = sum(a['contentLength'] for a in archivos_filtrados)
                lbl_info.configure(text=f"Total: {len(archivos_filtrados)} archivos | Espacio: {manager.formatear_tamaño(total_tamano)}")
            
            def crear_fila_archivo(archivo):
                """Crea una fila para mostrar un archivo"""
                fila = ctk.CTkFrame(scroll_frame, fg_color=("gray90", "gray20"))
                fila.pack(fill="x", padx=2, pady=2)
                
                nombre = archivo['fileName'].replace("Archivo/", "") if archivo['fileName'].startswith("Archivo/") else archivo['fileName']
                tamaño = manager.formatear_tamaño(archivo['contentLength'])
                fecha = manager.formatear_fecha(archivo['uploadTimestamp'])
                ubicacion = "Archivo/" if archivo['fileName'].startswith("Archivo/") else "Raíz"
                
                ctk.CTkLabel(fila, text=nombre, width=350, anchor="w").pack(side="left", padx=10, pady=5)
                ctk.CTkLabel(fila, text=tamaño, width=100, anchor="w").pack(side="left", padx=5, pady=5)
                ctk.CTkLabel(fila, text=fecha, width=150, anchor="w").pack(side="left", padx=5, pady=5)
                ctk.CTkLabel(fila, text=ubicacion, width=100, anchor="w").pack(side="left", padx=5, pady=5)
                
                # Botones de acción
                acciones_frame = ctk.CTkFrame(fila, fg_color="transparent")
                acciones_frame.pack(side="left", padx=5, pady=2)
                
                btn_descargar = ctk.CTkButton(acciones_frame, text="⬇", width=30, height=25,
                                              command=lambda: descargar_archivo(archivo))
                btn_descargar.pack(side="left", padx=2)
                Tooltip(btn_descargar, "Descargar este archivo a tu ordenador")
                
                if not archivo['fileName'].startswith("Archivo/"):
                    btn_mover = ctk.CTkButton(acciones_frame, text="📁", width=30, height=25,
                                             command=lambda: mover_a_archivo(archivo))
                    btn_mover.pack(side="left", padx=2)
                    Tooltip(btn_mover, "Mover este archivo a la carpeta Archivo/ (backups antiguos)")
                
                btn_eliminar = ctk.CTkButton(acciones_frame, text="🗑", width=30, height=25,
                                            fg_color="red", hover_color="darkred",
                                            command=lambda: eliminar_archivo(archivo))
                btn_eliminar.pack(side="left", padx=2)
                Tooltip(btn_eliminar, "Eliminar permanentemente este archivo de Backblaze B2")
            
            def descargar_archivo(archivo):
                """Descarga un archivo de B2"""
                destino = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(archivo['fileName'])[1],
                    initialfile=archivo['fileName'].replace("Archivo/", ""),
                    title="Guardar archivo"
                )
                
                if not destino:
                    return
                
                actualizar_estado("Descargando...", "blue")
                success, msg = manager.descargar_archivo(archivo['fileId'], archivo['fileName'], destino)
                
                if success:
                    messagebox.showinfo("Éxito", f"Archivo descargado en:\n{destino}")
                    actualizar_estado("Descarga completada", "green")
                else:
                    messagebox.showerror("Error", f"Error al descargar: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
            
            def mover_a_archivo(archivo):
                """Mueve un archivo a la carpeta Archivo/"""
                if not messagebox.askyesno("Confirmar", f"¿Mover '{archivo['fileName']}' a Archivo/?"):
                    return
                
                actualizar_estado("Moviendo...", "blue")
                success, msg = manager.mover_a_archivo(archivo['fileId'], archivo['fileName'])
                
                if success:
                    messagebox.showinfo("Éxito", "Archivo movido a Archivo/")
                    actualizar_estado("Archivo movido", "green")
                    cargar_archivos()  # Recargar lista
                else:
                    messagebox.showerror("Error", f"Error al mover: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
            
            def eliminar_archivo(archivo):
                """Elimina un archivo de B2"""
                if not messagebox.askyesno("Confirmar Eliminación", 
                                          f"¿Estás seguro de eliminar '{archivo['fileName']}'?\n\nEsta acción no se puede deshacer."):
                    return
                
                actualizar_estado("Eliminando...", "blue")
                success, msg = manager.eliminar_archivo(archivo['fileId'], archivo['fileName'])
                
                if success:
                    messagebox.showinfo("Éxito", "Archivo eliminado")
                    actualizar_estado("Archivo eliminado", "green")
                    cargar_archivos()  # Recargar lista
                else:
                    messagebox.showerror("Error", f"Error al eliminar: {msg}")
                    actualizar_estado(f"Error: {msg}", "red")
            
            def crear_backup():
                """Ejecuta el script de backup"""
                if not messagebox.askyesno("Crear Backup", "¿Crear una nueva copia de seguridad?\n\nEsto puede tardar varios minutos."):
                    return
                
                actualizar_estado("Creando backup...", "blue")
                btn_crear_backup.configure(state="disabled", text="Creando...")
                
                def ejecutar():
                    success, msg = manager.ejecutar_backup()
                    
                    # Actualizar UI en el hilo principal
                    ventana.after(0, lambda: backup_completado(success, msg))
                
                threading.Thread(target=ejecutar, daemon=True).start()
            
            def backup_completado(success, msg):
                """Callback cuando el backup termina"""
                btn_crear_backup.configure(state="normal", text="🔄 Crear Backup Ahora")
                
                if success:
                    messagebox.showinfo("Éxito", "Backup creado correctamente")
                    actualizar_estado("Backup creado", "green")
                    cargar_archivos()  # Recargar lista
                else:
                    messagebox.showerror("Error", f"Error al crear backup:\n{msg}")
                    actualizar_estado(f"Error: {msg}", "red")
            
            # Conectar botones
            btn_actualizar.configure(command=lambda: threading.Thread(target=cargar_archivos, daemon=True).start())
            btn_crear_backup.configure(command=crear_backup)
            
            # Conectar filtros para actualizar al cambiar
            entry_busqueda.bind("<KeyRelease>", lambda e: mostrar_archivos())
            filtro_tipo.configure(command=lambda _: mostrar_archivos())
            filtro_ubicacion.configure(command=lambda _: mostrar_archivos())
            
            # Cargar archivos inicialmente en un hilo
            threading.Thread(target=cargar_archivos, daemon=True).start()
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al abrir gestor de backups: {e}")
            import traceback
            traceback.print_exc()
    
    # ==================================================================
    # 🧑‍💼 GESTIÓN INDEPENDIENTE DE CLIENTES
    # ==================================================================
    
    def mostrar_clientes(self):
        """Muestra la gestión independiente de clientes."""
        
        # Mantener el tamaño actual de la ventana
        current_geometry = self.geometry()
        
        try:
            self.limpiar_contenido()
            
            # Asegurar que la ventana mantenga su tamaño
            self.geometry(current_geometry)
            
            # Header
            header_frame = ctk.CTkFrame(self.content_frame)
            header_frame.pack(fill="x", padx=10, pady=10)
            
            ctk.CTkLabel(header_frame, 
                        text="🧑‍💼 GESTIÓN DE CLIENTES", 
                        font=ctk.CTkFont(size=20, weight="bold")).pack(pady=10)
            
            # Botones de acción principal
            botones_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            botones_frame.pack(fill="x", padx=10, pady=5)
            
            btn_nuevo = ctk.CTkButton(botones_frame, text="➕ Nuevo Cliente", 
                                     command=self.nuevo_cliente, 
                                     width=140, height=35)
            btn_nuevo.pack(side="left", padx=(0,10))
            Tooltip(btn_nuevo, "Crear un nuevo cliente en el sistema")
            
            btn_migrar = ctk.CTkButton(botones_frame, text="🔄 Migrar desde RMAs", 
                                      command=self.migrar_clientes_desde_rmas, 
                                      width=160, height=35)
            btn_migrar.pack(side="left", padx=(0,10))
            Tooltip(btn_migrar, "Importar clientes automáticamente desde los RMAs existentes")
            
            # Búsqueda y filtros
            busqueda_frame = ctk.CTkFrame(self.content_frame)
            busqueda_frame.pack(fill="x", padx=10, pady=5)
            
            ctk.CTkLabel(busqueda_frame, text="Buscar cliente:", 
                        font=ctk.CTkFont(size=12)).pack(side="left", padx=10, pady=10)
            
            self.entry_buscar_cliente = ctk.CTkEntry(busqueda_frame, 
                                                   placeholder_text="Nombre del cliente...",
                                                   width=200)
            self.entry_buscar_cliente.pack(side="left", padx=5, pady=10)
            self.entry_buscar_cliente.bind("<KeyRelease>", self.filtrar_clientes)
            Tooltip(self.entry_buscar_cliente, "Escriba para buscar clientes por nombre en tiempo real")
            
            # Filtro por tipo
            ctk.CTkLabel(busqueda_frame, text="Tipo:", 
                        font=ctk.CTkFont(size=12)).pack(side="left", padx=(20,5), pady=10)
            
            self.filtro_tipo_cliente = ctk.CTkOptionMenu(busqueda_frame, 
                                                       values=["Todos", "Regular", "Premium", "VIP"],
                                                       command=self.filtrar_clientes,
                                                       width=100)
            self.filtro_tipo_cliente.set("Todos")
            self.filtro_tipo_cliente.pack(side="left", padx=5, pady=10)
            Tooltip(self.filtro_tipo_cliente, "Filtrar clientes por tipo: Regular, Premium o VIP")
            
            # Lista de clientes con más altura para aprovechar mejor el espacio
            self.clientes_frame = ctk.CTkScrollableFrame(self.content_frame, height=500)
            self.clientes_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            # Cargar clientes
            self.cargar_lista_clientes()
            
        except Exception as e:
            print(f"❌ Error en mostrar_clientes: {e}")
            import traceback
            print(traceback.format_exc())
    
    def cargar_lista_clientes(self):
        """Carga y muestra la lista de clientes."""
        # Limpiar lista actual
        for widget in self.clientes_frame.winfo_children():
            widget.destroy()
        
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                print("❌ No se pudo conectar a la base de datos")
                return
            
            # Buscar clientes con estadísticas (consulta simplificada)
            filtro_busqueda = self.entry_buscar_cliente.get().strip() if hasattr(self, 'entry_buscar_cliente') else ""
            filtro_tipo = self.filtro_tipo_cliente.get() if hasattr(self, 'filtro_tipo_cliente') else "Todos"
            
            # Consulta con estadísticas reales
            query = """
                SELECT c.cliente_id, c.nombre, c.tipo_cliente, c.activo, c.fecha_registro,
                       COALESCE(COUNT(DISTINCT r.id), 0) as total_rmas,
                       CASE 
                           WHEN COUNT(DISTINCT r.id) > 0 THEN 
                               ROUND(CAST(SUM(CASE WHEN r.resultado_expediente = 'ABONAR' THEN 1 ELSE 0 END) AS FLOAT) * 100.0 / COUNT(DISTINCT r.id), 1)
                           ELSE 0 
                       END as tasa_exito,
                       MAX(r.fecha_emision) as ultimo_rma,
                       COALESCE(COUNT(DISTINCT con.contacto_id), 0) as total_contactos
                FROM clientes c
                LEFT JOIN rma_maestro r ON c.nombre = r.cliente
                LEFT JOIN contactos_cliente con ON c.cliente_id = con.cliente_id
                WHERE 1=1
            """
            params = []
            
            if filtro_busqueda:
                query += " AND c.nombre LIKE ?"
                params.append(f"%{filtro_busqueda}%")
            
            if filtro_tipo != "Todos":
                query += " AND c.tipo_cliente = ?"
                params.append(filtro_tipo)
            
            query += " GROUP BY c.cliente_id, c.nombre, c.tipo_cliente, c.activo, c.fecha_registro ORDER BY c.nombre"
            
            cursor.execute(query, params)
            clientes = cursor.fetchall()
            
            if not clientes:
                ctk.CTkLabel(self.clientes_frame, 
                           text="📋 No se encontraron clientes. Haz clic en 'Nuevo Cliente' para agregar uno.",
                           font=ctk.CTkFont(size=13)).pack(pady=30)
                conn.close()
                return
            
            # Mostrar cada cliente
            for i, cliente in enumerate(clientes):
                self.crear_item_cliente(cliente)
            
            conn.close()
            
        except Exception as e:
            error_msg = f"❌ Error cargando clientes: {str(e)}"
            print(error_msg)
            ctk.CTkLabel(self.clientes_frame, 
                       text=error_msg,
                       font=ctk.CTkFont(size=12), 
                       text_color="red").pack(pady=20)
            import traceback
            print(traceback.format_exc())
    
    def crear_item_cliente(self, cliente):
        """Crea un elemento visual para un cliente en la lista con diseño compacto."""
        cliente_id, nombre, tipo, activo, fecha_reg, total_rmas, tasa_exito, ultimo_rma, total_contactos = cliente
        
        # Frame principal del cliente con altura reducida y borde sutil
        cliente_frame = ctk.CTkFrame(self.clientes_frame, height=45, border_width=1)
        cliente_frame.pack(fill="x", padx=5, pady=1)
        cliente_frame.pack_propagate(False)  # Mantener altura fija
        
        # Configurar grid para el layout horizontal
        cliente_frame.grid_columnconfigure(0, weight=1)  # Info del cliente (expandible)
        
        # Información del cliente
        info_frame = ctk.CTkFrame(cliente_frame, fg_color="transparent")
        info_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=5)
        info_frame.grid_columnconfigure(0, weight=1)
        
        # Línea superior: Nombre, estado y tipo en una sola línea
        header_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew")
        
        # Nombre del cliente (con doble clic para abrir)
        nombre_label = ctk.CTkLabel(header_frame, text=f"🧑‍💼 {nombre}", 
                                  font=ctk.CTkFont(size=13, weight="bold"),
                                  cursor="hand2")
        nombre_label.pack(side="left")
        nombre_label.bind("<Double-Button-1>", lambda e: self.abrir_ficha_cliente(cliente_id))
        Tooltip(nombre_label, "Doble clic para abrir la ficha del cliente")
        
        # Estado a la derecha
        estado_color = "green" if activo else "red"
        estado_texto = "🟢" if activo else "🔴"
        estado_label = ctk.CTkLabel(header_frame, text=estado_texto, 
                                  font=ctk.CTkFont(size=10))
        estado_label.pack(side="right", padx=(5,0))
        
        # Línea inferior: Estadísticas compactas
        stats_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        stats_frame.grid(row=1, column=0, sticky="ew")
        
        # Convertir valores a números para el formateo
        try:
            total_rmas_num = int(total_rmas) if total_rmas else 0
            tasa_exito_num = float(tasa_exito) if tasa_exito else 0.0
            total_contactos_num = int(total_contactos) if total_contactos else 0
        except (ValueError, TypeError):
            total_rmas_num = 0
            tasa_exito_num = 0.0
            total_contactos_num = 0
        
        stats_text = f"📊 {total_rmas_num} RMAs • ✅ {tasa_exito_num:.1f}% • 👥 {total_contactos_num} contactos"
        if ultimo_rma:
            stats_text += f" • 📅 {ultimo_rma}"
        
        stats_label = ctk.CTkLabel(stats_frame, text=stats_text, 
                                 font=ctk.CTkFont(size=9), text_color="gray")
        stats_label.pack(side="left")
    
    def filtrar_clientes(self, event=None):
        """Filtra la lista de clientes según los criterios de búsqueda."""
        self.cargar_lista_clientes()
    
    def nuevo_cliente(self):
        """Muestra el formulario para crear un nuevo cliente."""
        ventana = ctk.CTkToplevel(self)
        ventana.title("Nuevo Cliente")
        ventana.geometry("500x600")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="➕ Nuevo Cliente", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=400)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Nombre (obligatorio)
        ctk.CTkLabel(form_frame, text="Nombre del Cliente *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_nombre = ctk.CTkEntry(form_frame, placeholder_text="Nombre completo del cliente")
        entry_nombre.pack(fill="x", pady=(0,10))
        
        # Tipo de cliente
        ctk.CTkLabel(form_frame, text="Tipo de Cliente", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        option_tipo = ctk.CTkOptionMenu(form_frame, values=["Regular", "Premium", "VIP"])
        option_tipo.set("Regular")
        option_tipo.pack(fill="x", pady=(0,10))
        
        # Dirección
        ctk.CTkLabel(form_frame, text="Dirección", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_direccion = ctk.CTkEntry(form_frame, placeholder_text="Dirección completa")
        entry_direccion.pack(fill="x", pady=(0,10))
        
        # Teléfono principal
        ctk.CTkLabel(form_frame, text="Teléfono Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_telefono = ctk.CTkEntry(form_frame, placeholder_text="Teléfono de contacto principal")
        entry_telefono.pack(fill="x", pady=(0,10))
        
        # Email principal
        ctk.CTkLabel(form_frame, text="Email Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_email = ctk.CTkEntry(form_frame, placeholder_text="Email de contacto principal")
        entry_email.pack(fill="x", pady=(0,10))
        
        # Notas generales
        ctk.CTkLabel(form_frame, text="Notas Generales", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        text_notas = ctk.CTkTextbox(form_frame, height=80)
        text_notas.pack(fill="x", pady=(0,10))
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        def guardar_cliente():
            nombre = entry_nombre.get().strip()
            if not nombre:
                messagebox.showerror("Error", "El nombre del cliente es obligatorio")
                return
            
            try:
                conn, cursor = self.master.conectar_db()
                if not conn: return
                
                cursor.execute("""
                    INSERT INTO clientes (nombre, tipo_cliente, direccion, telefono_principal, 
                                        email_principal, notas_generales)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (nombre, option_tipo.get(), entry_direccion.get().strip(),
                     entry_telefono.get().strip(), entry_email.get().strip(),
                     text_notas.get("1.0", "end-1c").strip()))
                
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Éxito", f"Cliente '{nombre}' creado correctamente")
                ventana.destroy()
                self.cargar_lista_clientes()
                
            except Exception as e:
                if "UNIQUE constraint failed" in str(e):
                    messagebox.showerror("Error", f"Ya existe un cliente con el nombre '{nombre}'")
                else:
                    messagebox.showerror("Error", f"Error al crear cliente: {str(e)}")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Guardar Cliente", 
                                  command=guardar_cliente,
                                  width=120)
        btn_guardar.pack(side="right")
    
    def migrar_clientes_desde_rmas(self):
        """Migra clientes existentes desde la tabla rma_maestro."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return
            
            # Contar clientes únicos en RMAs antes de migrar
            cursor.execute("""
                SELECT COUNT(DISTINCT cliente) 
                FROM rma_maestro 
                WHERE cliente IS NOT NULL AND cliente != ''
            """)
            total_clientes_en_rmas = cursor.fetchone()[0]
            
            # Contar clientes ya existentes en la tabla
            cursor.execute("SELECT COUNT(*) FROM clientes")
            clientes_antes = cursor.fetchone()[0]
            
            # Ejecutar migración
            cursor.execute("""
                INSERT OR IGNORE INTO clientes (nombre, fecha_registro, email_principal)
                SELECT DISTINCT 
                    Cliente,
                    MIN(fecha_emision) as fecha_registro,
                    Email_de_Contacto
                FROM rma_maestro 
                WHERE Cliente IS NOT NULL AND Cliente != ''
                GROUP BY Cliente
            """)
            
            # Contar después de migrar
            cursor.execute("SELECT COUNT(*) FROM clientes")
            clientes_despues = cursor.fetchone()[0]
            
            clientes_nuevos = int(clientes_despues) - int(clientes_antes)
            clientes_ya_existian = int(total_clientes_en_rmas) - clientes_nuevos
            
            # Migrar contactos
            cursor.execute("""
                INSERT OR IGNORE INTO contactos_cliente (cliente_id, nombre, email, es_principal)
                SELECT DISTINCT
                    c.cliente_id,
                    COALESCE(rm.Persona_de_Contacto, rm.Cliente) as nombre,
                    rm.Email_de_Contacto,
                    1 as es_principal
                FROM clientes c
                JOIN rma_maestro rm ON c.nombre = rm.Cliente
                WHERE rm.Persona_de_Contacto IS NOT NULL AND rm.Persona_de_Contacto != ''
                GROUP BY c.cliente_id, rm.Persona_de_Contacto
            """)
            
            contactos_migrados = cursor.rowcount
            
            conn.commit()
            conn.close()
            
            # Mensaje mejorado
            if clientes_nuevos > 0:
                mensaje = f"✅ Migración completada:\n\n"
                mensaje += f"• {clientes_nuevos} cliente(s) NUEVO(S) migrado(s)\n"
                if clientes_ya_existian > 0:
                    mensaje += f"• {clientes_ya_existian} cliente(s) ya existían (no duplicados)\n"
                mensaje += f"• {contactos_migrados} contacto(s) migrado(s)\n\n"
                mensaje += f"📊 Total de clientes en sistema: {clientes_despues}"
            else:
                mensaje = f"ℹ️ Migración completada:\n\n"
                mensaje += f"• Todos los clientes ({total_clientes_en_rmas}) ya estaban migrados\n"
                mensaje += f"• No se encontraron duplicados\n"
                mensaje += f"• {contactos_migrados} contacto(s) nuevos añadidos\n\n"
                mensaje += f"📊 Total de clientes en sistema: {clientes_despues}"
            
            messagebox.showinfo("Migración Completada", mensaje)
            
            self.cargar_lista_clientes()
            
        except Exception as e:
            messagebox.showerror("Error", f"Error en la migración: {str(e)}")
    
    def abrir_ficha_cliente(self, cliente_id):
        """Abre la ficha completa del cliente con todas sus pestañas."""
        # Obtener información del cliente
        cliente = self.obtener_cliente(cliente_id)
        if not cliente:
            messagebox.showerror("Error", "No se pudo cargar la información del cliente")
            return
        
        # Crear ventana principal
        ventana = ctk.CTkToplevel(self)
        ventana.title(f"Ficha Cliente: {cliente[1]}")  # cliente[1] es el nombre
        ventana.geometry("900x900")
        ventana.resizable(True, True)
        
        # Configurar para permitir minimización
        ventana.attributes('-topmost', False)
        ventana.minsize(700, 650)
        # No usar transient para permitir minimización completa
        ventana.focus_set()  # Dar foco sin bloquear
        
        # Forzar aparición al frente (incluso si la principal está maximizada)
        ventana.attributes('-topmost', True)   # Temporalmente al frente
        ventana.lift()
        ventana.focus_force()
        ventana.after(500, lambda: ventana.attributes('-topmost', False))  # Quitar topmost después de 500ms
        
        # Header con información básica
        header_frame = ctk.CTkFrame(ventana)
        header_frame.pack(fill="x", padx=10, pady=(10, 5))
        
        info_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=10, pady=5)
        
        # Nombre del cliente
        nombre_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        nombre_frame.pack(fill="x")
        
        ctk.CTkLabel(nombre_frame, text=f"🧑‍💼 {cliente[1]}", 
                    font=ctk.CTkFont(size=20, weight="bold")).pack(side="left")
        
        # Widget de estadísticas básicas
        from lib.cliente_estadisticas import calcular_estadisticas_basicas_cliente, crear_widget_estadisticas_basicas
        
        try:
            conn, cursor = self.master.conectar_db()
            if conn:
                stats = calcular_estadisticas_basicas_cliente(cliente[1], conn)
                crear_widget_estadisticas_basicas(info_frame, stats)
                conn.close()
        except Exception as e:
            print(f"Error cargando estadísticas: {e}")
        
        # Crear pestañas (sin expand para dejar espacio a los botones)
        tabview = ctk.CTkTabview(ventana, width=880, height=350)
        tabview.pack(fill="x", expand=False, padx=10, pady=5)
        
        # Pestaña 1: Información General (en modo edición)
        tab_info = tabview.add("📋 Información")
        widgets_info = self.crear_tab_informacion_cliente_editable(tab_info, cliente)
        
        # Pestaña 2: Contactos
        tab_contactos = tabview.add("👥 Contactos")
        self.crear_tab_contactos_cliente(tab_contactos, cliente_id)
        
        # Pestaña 3: Historial RMAs
        tab_rmas = tabview.add("📦 Historial RMAs")
        self.crear_tab_historial_rmas(tab_rmas, cliente_id)
        
        # Pestaña 4: Notas
        tab_notas = tabview.add("📝 Notas")
        self.crear_tab_notas_cliente(tab_notas, cliente_id)
        
        # Pestaña 5: Estadísticas
        tab_stats = tabview.add("📊 Estadísticas")
        self.crear_tab_estadisticas_cliente_completa(tab_stats, cliente_id, cliente[1])
        
        # Botones de acción ANTES de definir funciones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=10, pady=(5, 10))
        
        # Función para guardar cambios
        def guardar_cambios_cliente():
            datos = {
                'nombre': widgets_info['entry_nombre'].get().strip(),
                'direccion': widgets_info['entry_direccion'].get().strip(),
                'telefono_principal': widgets_info['entry_telefono'].get().strip(),
                'email_principal': widgets_info['entry_email'].get().strip(),
                'notas_generales': widgets_info['text_notas'].get("1.0", "end-1c").strip()
            }
            
            if not datos['nombre']:
                messagebox.showerror("Error", "El nombre del cliente es obligatorio")
                return
            
            if self.actualizar_cliente(cliente_id, datos):
                messagebox.showinfo("Éxito", f"Cliente '{datos['nombre']}' actualizado correctamente")
                # Recargar estadísticas
                try:
                    conn, cursor = self.master.conectar_db()
                    if conn:
                        stats = calcular_estadisticas_basicas_cliente(datos['nombre'], conn)
                        # Actualizar widget de estadísticas
                        for widget in info_frame.winfo_children():
                            if isinstance(widget, ctk.CTkFrame) and widget != nombre_frame:
                                widget.destroy()
                        crear_widget_estadisticas_basicas(info_frame, stats)
                        conn.close()
                except Exception as e:
                    print(f"Error recargando estadísticas: {e}")
            else:
                messagebox.showerror("Error", "Error al actualizar el cliente")
        
        # Configurar botones (frame ya creado arriba)
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Guardar Cambios", 
                                 command=guardar_cambios_cliente,
                                 width=140)
        btn_guardar.pack(side="left", padx=(0,10))
        
        btn_cerrar = ctk.CTkButton(botones_frame, text="❌ Cerrar", 
                                 command=ventana.destroy,
                                 width=100)
        btn_cerrar.pack(side="right")
    
    def crear_tab_informacion_cliente(self, tab_frame, cliente):
        """Crea la pestaña de información general del cliente."""
        # Frame scrollable para el formulario
        scroll_frame = ctk.CTkScrollableFrame(tab_frame, height=500)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Información básica
        ctk.CTkLabel(scroll_frame, text="📋 Información Básica", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(0,10))
        
        info_frame = ctk.CTkFrame(scroll_frame)
        info_frame.pack(fill="x", pady=(0,20))
        info_frame.grid_columnconfigure(1, weight=1)
        
        # Datos del cliente
        datos = [
            ("ID Cliente:", str(cliente[0])),
            ("Nombre:", cliente[1]),
            ("Tipo:", cliente[2]),
            ("Dirección:", cliente[3] or "No especificada"),
            ("Teléfono Principal:", cliente[4] or "No especificado"),
            ("Email Principal:", cliente[5] or "No especificado"),
            ("Fecha Registro:", cliente[8].split()[0] if cliente[8] else "No disponible"),
            ("Última Actualización:", cliente[9].split()[0] if cliente[9] else "No disponible")
        ]
        
        for i, (etiqueta, valor) in enumerate(datos):
            ctk.CTkLabel(info_frame, text=etiqueta, 
                        font=ctk.CTkFont(size=12, weight="bold")).grid(
                        row=i, column=0, sticky="w", padx=10, pady=5)
            ctk.CTkLabel(info_frame, text=valor, 
                        font=ctk.CTkFont(size=12)).grid(
                        row=i, column=1, sticky="w", padx=10, pady=5)
        
        # Notas generales
        if cliente[6]:  # cliente[6] son las notas generales
            ctk.CTkLabel(scroll_frame, text="📝 Notas Generales", 
                        font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(20,10))
            
            notas_frame = ctk.CTkFrame(scroll_frame)
            notas_frame.pack(fill="x", pady=(0,10))
            
            text_notas = ctk.CTkTextbox(notas_frame, height=100)
            text_notas.pack(fill="x", padx=10, pady=10)
            text_notas.insert("1.0", cliente[6])
            text_notas.configure(state="disabled")
    
    def crear_tab_informacion_cliente_editable(self, tab_frame, cliente):
        """Crea la pestaña de información general del cliente en modo edición."""
        # Frame scrollable para el formulario
        scroll_frame = ctk.CTkScrollableFrame(tab_frame, height=500)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Información básica
        ctk.CTkLabel(scroll_frame, text="📋 Información Básica (Editable)", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(0,10))
        
        info_frame = ctk.CTkFrame(scroll_frame)
        info_frame.pack(fill="x", pady=(0,20))
        info_frame.grid_columnconfigure(1, weight=1)
        
        # Widgets editables
        widgets = {}
        
        # ID Cliente (solo lectura)
        ctk.CTkLabel(info_frame, text="ID Cliente:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=0, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkLabel(info_frame, text=str(cliente[0]), 
                    font=ctk.CTkFont(size=12)).grid(
                    row=0, column=1, sticky="w", padx=10, pady=5)
        
        # Nombre (editable)
        ctk.CTkLabel(info_frame, text="Nombre: *", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=1, column=0, sticky="w", padx=10, pady=5)
        widgets['entry_nombre'] = ctk.CTkEntry(info_frame)
        widgets['entry_nombre'].grid(row=1, column=1, sticky="ew", padx=10, pady=5)
        widgets['entry_nombre'].insert(0, cliente[1])
        
        # Dirección (editable)
        ctk.CTkLabel(info_frame, text="Dirección:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=2, column=0, sticky="w", padx=10, pady=5)
        widgets['entry_direccion'] = ctk.CTkEntry(info_frame)
        widgets['entry_direccion'].grid(row=2, column=1, sticky="ew", padx=10, pady=5)
        widgets['entry_direccion'].insert(0, cliente[3] or "")
        
        # Teléfono (editable)
        ctk.CTkLabel(info_frame, text="Teléfono Principal:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=3, column=0, sticky="w", padx=10, pady=5)
        widgets['entry_telefono'] = ctk.CTkEntry(info_frame)
        widgets['entry_telefono'].grid(row=3, column=1, sticky="ew", padx=10, pady=5)
        widgets['entry_telefono'].insert(0, cliente[4] or "")
        
        # Email (editable)
        ctk.CTkLabel(info_frame, text="Email Principal:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=4, column=0, sticky="w", padx=10, pady=5)
        widgets['entry_email'] = ctk.CTkEntry(info_frame)
        widgets['entry_email'].grid(row=4, column=1, sticky="ew", padx=10, pady=5)
        widgets['entry_email'].insert(0, cliente[5] or "")
        
        # Fechas (solo lectura)
        ctk.CTkLabel(info_frame, text="Fecha Registro:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=5, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkLabel(info_frame, text=cliente[8].split()[0] if cliente[8] else "No disponible", 
                    font=ctk.CTkFont(size=12)).grid(
                    row=5, column=1, sticky="w", padx=10, pady=5)
        
        ctk.CTkLabel(info_frame, text="Última Actualización:", 
                    font=ctk.CTkFont(size=12, weight="bold")).grid(
                    row=6, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkLabel(info_frame, text=cliente[9].split()[0] if cliente[9] else "No disponible", 
                    font=ctk.CTkFont(size=12)).grid(
                    row=6, column=1, sticky="w", padx=10, pady=5)
        
        # Notas generales (editable)
        ctk.CTkLabel(scroll_frame, text="📝 Notas Generales", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(20,10))
        
        notas_frame = ctk.CTkFrame(scroll_frame)
        notas_frame.pack(fill="x", pady=(0,10))
        
        widgets['text_notas'] = ctk.CTkTextbox(notas_frame, height=100)
        widgets['text_notas'].pack(fill="x", padx=10, pady=10)
        if cliente[6]:
            widgets['text_notas'].insert("1.0", cliente[6])
        
        return widgets
    
    def crear_tab_contactos_cliente(self, tab_frame, cliente_id):
        """Crea la pestaña de contactos del cliente."""
        # Header con botón para agregar
        header_frame = ctk.CTkFrame(tab_frame)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(header_frame, text="👥 Contactos del Cliente", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(side="left", padx=10, pady=10)
        
        btn_nuevo_contacto = ctk.CTkButton(header_frame, text="➕ Nuevo Contacto", 
                                         command=lambda: self.nuevo_contacto_cliente(cliente_id),
                                         width=120)
        btn_nuevo_contacto.pack(side="right", padx=10, pady=10)
        
        # Lista de contactos
        contactos_frame = ctk.CTkScrollableFrame(tab_frame, height=450)
        contactos_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Cargar y mostrar contactos
        contactos = self.obtener_contactos_cliente(cliente_id)
        
        if not contactos:
            ctk.CTkLabel(contactos_frame, 
                       text="📭 No hay contactos registrados para este cliente.\nHaz clic en 'Nuevo Contacto' para agregar uno.",
                       font=ctk.CTkFont(size=13)).pack(pady=50)
        else:
            for contacto in contactos:
                self.crear_item_contacto(contactos_frame, contacto, cliente_id)
    
    def crear_tab_historial_rmas(self, tab_frame, cliente_id):
        """Crea la pestaña de historial de RMAs."""
        from lib.cliente_utils import obtener_años_rmas_cliente, obtener_historial_rmas_cliente
        
        # Header
        header_frame = ctk.CTkFrame(tab_frame)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(header_frame, text="📦 Historial de RMAs", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=10, pady=(10,5))
        
        # Frame de filtros
        filtros_frame = ctk.CTkFrame(header_frame)
        filtros_frame.pack(fill="x", padx=10, pady=5)
        
        # Búsqueda por número de RMA
        ctk.CTkLabel(filtros_frame, text="Buscar RMA:", font=ctk.CTkFont(size=11)).pack(side="left", padx=5)
        entry_buscar = ctk.CTkEntry(filtros_frame, placeholder_text="Número de RMA...", width=150)
        entry_buscar.pack(side="left", padx=5)
        
        # Filtro por año
        ctk.CTkLabel(filtros_frame, text="Año:", font=ctk.CTkFont(size=11)).pack(side="left", padx=(20,5))
        
        # Obtener años disponibles
        años_disponibles = obtener_años_rmas_cliente(cliente_id, self.master.conectar_db)
        option_año = ctk.CTkOptionMenu(filtros_frame, values=["Todos"] + años_disponibles, width=100)
        option_año.set("Todos")
        option_año.pack(side="left", padx=5)
        
        # Lista de RMAs
        rmas_frame = ctk.CTkScrollableFrame(tab_frame, height=450)
        rmas_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Función para filtrar y mostrar RMAs
        def actualizar_lista():
            for widget in rmas_frame.winfo_children():
                widget.destroy()
            
            busqueda = entry_buscar.get().strip().upper()
            año_filtro = option_año.get()
            
            rmas = obtener_historial_rmas_cliente(cliente_id, self.master.conectar_db, año_filtro if año_filtro != "Todos" else None, busqueda)
            
            if not rmas:
                ctk.CTkLabel(rmas_frame, 
                           text="📭 No se encontraron RMAs con los filtros especificados.",
                           font=ctk.CTkFont(size=13)).pack(pady=50)
            else:
                # Mostrar RMAs
                for rma in rmas:
                    numero_rma, fecha_emision, estado, motivo = rma
                    datos_rma = {
                        'info': rma,
                        'productos': []
                    }
                    self.crear_item_rma_historial(rmas_frame, numero_rma, datos_rma, cliente_id)
        
        # Botón filtrar
        btn_filtrar = ctk.CTkButton(filtros_frame, text="🔍 Filtrar", command=actualizar_lista, width=80)
        btn_filtrar.pack(side="left", padx=10)
        
        # Cargar lista inicial
        actualizar_lista()
    
    def crear_tab_notas_cliente(self, tab_frame, cliente_id):
        """Crea la pestaña de notas del cliente."""
        # Header con botón para agregar
        header_frame = ctk.CTkFrame(tab_frame)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(header_frame, text="📝 Notas del Cliente", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(side="left", padx=10, pady=10)
        
        btn_nueva_nota = ctk.CTkButton(header_frame, text="➕ Nueva Nota", 
                                     command=lambda: self.nueva_nota_cliente(cliente_id),
                                     width=120)
        btn_nueva_nota.pack(side="right", padx=10, pady=10)
        
        # Lista de notas
        notas_frame = ctk.CTkScrollableFrame(tab_frame, height=450)
        notas_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Cargar y mostrar notas
        notas = self.obtener_notas_cliente(cliente_id)
        
        if not notas:
            ctk.CTkLabel(notas_frame, 
                       text="📭 No hay notas registradas para este cliente.\nHaz clic en 'Nueva Nota' para agregar una.",
                       font=ctk.CTkFont(size=13)).pack(pady=50)
        else:
            for nota in notas:
                self.crear_item_nota(notas_frame, nota)
    
    def crear_tab_estadisticas_cliente_completa(self, tab_frame, cliente_id, nombre_cliente):
        """Crea la pestaña de estadísticas del cliente con filtros completos y exportación."""
        from lib.cliente_estadisticas import obtener_estadisticas_detalladas_cliente, exportar_estadisticas_cliente_excel
        from tkinter import filedialog
        from datetime import datetime
        
        # Header con controles
        header_frame = ctk.CTkFrame(tab_frame)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        # Título
        ctk.CTkLabel(header_frame, text="📊 Estadísticas Detalladas del Cliente", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=10, pady=(10,5))
        
        # Frame de filtros
        filtros_frame = ctk.CTkFrame(header_frame)
        filtros_frame.pack(fill="x", padx=10, pady=5)
        filtros_frame.grid_columnconfigure(1, weight=1)
        filtros_frame.grid_columnconfigure(3, weight=1)
        
        # Filtro de fecha desde
        ctk.CTkLabel(filtros_frame, text="Desde:", font=ctk.CTkFont(size=11)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        entry_fecha_desde = ctk.CTkEntry(filtros_frame, placeholder_text="YYYY-MM-DD", width=120)
        entry_fecha_desde.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Filtro de fecha hasta
        ctk.CTkLabel(filtros_frame, text="Hasta:", font=ctk.CTkFont(size=11)).grid(row=0, column=2, padx=5, pady=5, sticky="e")
        entry_fecha_hasta = ctk.CTkEntry(filtros_frame, placeholder_text="YYYY-MM-DD", width=120)
        entry_fecha_hasta.grid(row=0, column=3, padx=5, pady=5, sticky="w")
        
        # Filtro de resultado de expediente
        ctk.CTkLabel(filtros_frame, text="Resultado:", font=ctk.CTkFont(size=11)).grid(row=0, column=4, padx=5, pady=5, sticky="e")
        filtro_estado = ctk.CTkOptionMenu(filtros_frame, 
                                         values=["Todos", "NO ABONAR", "ABONAR", "ABONAR OK", "ABONAR FALLO", "REPOSICION"],
                                         width=140)
        filtro_estado.set("Todos")
        filtro_estado.grid(row=0, column=5, padx=5, pady=5, sticky="w")
        
        # Contenedor de resultados
        resultados_frame = ctk.CTkScrollableFrame(tab_frame, height=400)
        resultados_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Función para cargar datos
        def cargar_datos_estadisticas():
            # Limpiar resultados anteriores
            for widget in resultados_frame.winfo_children():
                widget.destroy()
            
            try:
                conn, cursor = self.master.conectar_db()
                if not conn:
                    return
                
                # Obtener filtros
                fecha_desde = entry_fecha_desde.get().strip() or None
                fecha_hasta = entry_fecha_hasta.get().strip() or None
                estado = filtro_estado.get()
                estado_filtro = None if estado == "Todos" else estado
                
                # Obtener datos
                datos = obtener_estadisticas_detalladas_cliente(
                    nombre_cliente, conn, fecha_desde, fecha_hasta, estado_filtro
                )
                
                if not datos:
                    ctk.CTkLabel(resultados_frame, 
                               text="📊 No hay datos para los filtros seleccionados.",
                               font=ctk.CTkFont(size=14)).pack(pady=50)
                    conn.close()
                    return
                
                # Calcular totales
                total_exp = len(datos)
                total_art = sum(row[4] for row in datos)
                total_coste = sum(row[5] for row in datos)
                
                # Mostrar resumen
                resumen_frame = ctk.CTkFrame(resultados_frame)
                resumen_frame.pack(fill="x", pady=(0,10))
                
                ctk.CTkLabel(resumen_frame, 
                           text=f"📦 Total: {total_exp} expedientes | 📦 {total_art} artículos | 💰 {total_coste:.2f} €",
                           font=ctk.CTkFont(size=13, weight="bold")).pack(pady=10)
                
                # Tabla de resultados
                tabla_frame = ctk.CTkFrame(resultados_frame)
                tabla_frame.pack(fill="both", expand=True)
                
                # Encabezados
                headers = ["Código RMA", "Fecha", "Estado", "Resultado", "Nº Art.", "Coste (€)"]
                header_row = ctk.CTkFrame(tabla_frame, fg_color=("#3B8ED0", "#1F6AA5"))
                header_row.pack(fill="x", padx=2, pady=2)
                
                for i, header in enumerate(headers):
                    ctk.CTkLabel(header_row, text=header, 
                               font=ctk.CTkFont(size=11, weight="bold"),
                               text_color="white").grid(row=0, column=i, padx=10, pady=5, sticky="w")
                
                # Datos
                for idx, row in enumerate(datos):
                    codigo, fecha, estado_exp, resultado, num_art, coste = row
                    
                    color_bg = ("#E8E8E8", "#2B2B2B") if idx % 2 == 0 else "transparent"
                    data_row = ctk.CTkFrame(tabla_frame, fg_color=color_bg)
                    data_row.pack(fill="x", padx=2, pady=1)
                    
                    ctk.CTkLabel(data_row, text=codigo, font=ctk.CTkFont(size=10)).grid(row=0, column=0, padx=10, pady=3, sticky="w")
                    ctk.CTkLabel(data_row, text=fecha or "", font=ctk.CTkFont(size=10)).grid(row=0, column=1, padx=10, pady=3, sticky="w")
                    ctk.CTkLabel(data_row, text=estado_exp or "", font=ctk.CTkFont(size=10)).grid(row=0, column=2, padx=10, pady=3, sticky="w")
                    ctk.CTkLabel(data_row, text=resultado or "", font=ctk.CTkFont(size=10)).grid(row=0, column=3, padx=10, pady=3, sticky="w")
                    ctk.CTkLabel(data_row, text=str(num_art), font=ctk.CTkFont(size=10)).grid(row=0, column=4, padx=10, pady=3, sticky="w")
                    ctk.CTkLabel(data_row, text=f"{coste:.2f} €", font=ctk.CTkFont(size=10)).grid(row=0, column=5, padx=10, pady=3, sticky="w")
                
                conn.close()
                
            except Exception as e:
                print(f"Error cargando estadísticas: {e}")
                messagebox.showerror("Error", f"Error al cargar estadísticas: {str(e)}")
        
        # Función para exportar a Excel
        def exportar_a_excel():
            try:
                conn, cursor = self.master.conectar_db()
                if not conn:
                    return
                
                # Obtener filtros
                fecha_desde = entry_fecha_desde.get().strip() or None
                fecha_hasta = entry_fecha_hasta.get().strip() or None
                estado = filtro_estado.get()
                estado_filtro = None if estado == "Todos" else estado
                
                # Obtener datos
                datos = obtener_estadisticas_detalladas_cliente(
                    nombre_cliente, conn, fecha_desde, fecha_hasta, estado_filtro
                )
                
                conn.close()
                
                if not datos:
                    messagebox.showinfo("Info", "No hay datos para exportar")
                    return
                
                # Solicitar ubicación de archivo
                fecha_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                nombre_archivo = f"Estadisticas_{nombre_cliente.replace(' ', '_')}_{fecha_str}.xlsx"
                
                ruta = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    initialfile=nombre_archivo
                )
                
                if ruta:
                    if exportar_estadisticas_cliente_excel(nombre_cliente, datos, ruta):
                        messagebox.showinfo("Éxito", f"Estadísticas exportadas correctamente a:\n{ruta}")
                    else:
                        messagebox.showerror("Error", "Error al exportar estadísticas")
                        
            except Exception as e:
                print(f"Error exportando: {e}")
                messagebox.showerror("Error", f"Error al exportar: {str(e)}")
        
        # Botones de acción
        botones_frame = ctk.CTkFrame(header_frame)
        botones_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkButton(botones_frame, text="🔍 Buscar", 
                     command=cargar_datos_estadisticas,
                     width=120).pack(side="left", padx=5)
        
        ctk.CTkButton(botones_frame, text="📊 Exportar Excel", 
                     command=exportar_a_excel,
                     width=140, 
                     fg_color="green", 
                     hover_color="dark green").pack(side="left", padx=5)
        
        # Cargar datos iniciales
        cargar_datos_estadisticas()
    
    def cargar_estadisticas_cliente(self, cliente_id):
        """Carga las estadísticas del cliente con filtros aplicados."""
        # Limpiar contenido anterior
        for widget in self.stats_container.winfo_children():
            widget.destroy()
        
        try:
            conn, cursor = self.master.conectar_db()
            if not conn:
                return
            
            # Obtener nombre del cliente
            cursor.execute("SELECT nombre FROM clientes WHERE cliente_id = ?", (cliente_id,))
            cliente_result = cursor.fetchone()
            if not cliente_result:
                return
            
            cliente_nombre = cliente_result[0]
            
            # Construir filtros de fecha
            filtro_año = self.filtro_año_stats.get()
            filtro_mes = self.filtro_mes_stats.get()
            
            where_fecha = ""
            params = [cliente_nombre]
            
            if filtro_año != "Todos":
                where_fecha += " AND strftime('%Y', fecha_emision) = ?"
                params.append(filtro_año)
            
            if filtro_mes != "Todos":
                mes_num = ["", "01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"][
                    ["Todos", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
                     "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"].index(filtro_mes)
                ]
                where_fecha += " AND strftime('%m', fecha_emision) = ?"
                params.append(mes_num)
            
            # Consulta principal: desglose por resultado_expediente
            query_resultados = f"""
                SELECT 
                    resultado_expediente,
                    COUNT(*) as cantidad,
                    SUM(COALESCE(precio_total_expediente, 0)) as total_importe,
                    AVG(COALESCE(precio_total_expediente, 0)) as importe_promedio
                FROM rma_maestro 
                WHERE cliente = ? {where_fecha}
                GROUP BY resultado_expediente
                ORDER BY cantidad DESC
            """
            
            cursor.execute(query_resultados, params)
            resultados = cursor.fetchall()
            
            if not resultados:
                ctk.CTkLabel(self.stats_container, 
                           text="📊 No hay datos para los filtros seleccionados.",
                           font=ctk.CTkFont(size=14)).pack(pady=50)
            else:
                self.mostrar_desglose_resultados(resultados)
                
                # Consulta adicional: estadísticas generales
                query_generales = f"""
                    SELECT 
                        COUNT(*) as total_expedientes,
                        SUM(COALESCE(precio_total_expediente, 0)) as importe_total,
                        AVG(COALESCE(precio_total_expediente, 0)) as importe_promedio,
                        MIN(fecha_emision) as primer_expediente,
                        MAX(fecha_emision) as ultimo_expediente
                    FROM rma_maestro 
                    WHERE cliente = ? {where_fecha}
                """
                
                cursor.execute(query_generales, params)
                generales = cursor.fetchone()
                
                if generales:
                    self.mostrar_estadisticas_generales(generales)
            
            conn.close()
            
        except Exception as e:
            print(f"Error cargando estadísticas: {e}")
            import traceback
            print(traceback.format_exc())
    
    def mostrar_desglose_resultados(self, resultados):
        """Muestra el desglose por tipo de resultado de expediente."""
        # Frame para desglose
        desglose_frame = ctk.CTkFrame(self.stats_container)
        desglose_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(desglose_frame, text="📋 Desglose por Resultado de Expediente", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        
        # Tabla
        tabla_frame = ctk.CTkFrame(desglose_frame)
        tabla_frame.pack(fill="x", padx=10, pady=10)
        
        # Headers
        headers = ["Resultado", "Cantidad", "Importe Total", "Importe Promedio"]
        for col, header in enumerate(headers):
            ctk.CTkLabel(tabla_frame, text=header, 
                        font=ctk.CTkFont(weight="bold")).grid(row=0, column=col, padx=10, pady=5, sticky="ew")
        
        # Datos
        total_cantidad = 0
        total_importe = 0
        
        for row, (resultado, cantidad, importe_total, importe_promedio) in enumerate(resultados, 1):
            resultado_texto = resultado if resultado else "Sin Resultado"
            
            # Convertir a tipos correctos
            cantidad = int(cantidad) if cantidad is not None else 0
            importe_total = float(importe_total) if importe_total is not None else 0.0
            importe_promedio = float(importe_promedio) if importe_promedio is not None else 0.0
            
            ctk.CTkLabel(tabla_frame, text=resultado_texto).grid(row=row, column=0, padx=10, pady=2, sticky="ew")
            ctk.CTkLabel(tabla_frame, text=str(cantidad)).grid(row=row, column=1, padx=10, pady=2, sticky="ew")
            ctk.CTkLabel(tabla_frame, text=f"€{importe_total:.2f}").grid(row=row, column=2, padx=10, pady=2, sticky="ew")
            ctk.CTkLabel(tabla_frame, text=f"€{importe_promedio:.2f}").grid(row=row, column=3, padx=10, pady=2, sticky="ew")
            
            total_cantidad += cantidad
            total_importe += importe_total
        
        # Totales
        separator_frame = ctk.CTkFrame(tabla_frame, height=2)
        separator_frame.grid(row=len(resultados)+1, column=0, columnspan=4, sticky="ew", pady=5)
        
        ctk.CTkLabel(tabla_frame, text="TOTAL", 
                    font=ctk.CTkFont(weight="bold")).grid(row=len(resultados)+2, column=0, padx=10, pady=5, sticky="ew")
        ctk.CTkLabel(tabla_frame, text=str(total_cantidad), 
                    font=ctk.CTkFont(weight="bold")).grid(row=len(resultados)+2, column=1, padx=10, pady=5, sticky="ew")
        ctk.CTkLabel(tabla_frame, text=f"€{total_importe:.2f}", 
                    font=ctk.CTkFont(weight="bold")).grid(row=len(resultados)+2, column=2, padx=10, pady=5, sticky="ew")
        
        # Configurar columnas
        for col in range(4):
            tabla_frame.grid_columnconfigure(col, weight=1)
    
    def mostrar_estadisticas_generales(self, generales):
        """Muestra estadísticas generales del cliente."""
        total_exp, importe_total, importe_prom, primer_exp, ultimo_exp = generales
        
        # Convertir a tipos correctos
        total_exp = int(total_exp) if total_exp is not None else 0
        importe_total = float(importe_total) if importe_total is not None else 0.0
        importe_prom = float(importe_prom) if importe_prom is not None else 0.0
        
        # Frame para estadísticas generales
        generales_frame = ctk.CTkFrame(self.stats_container)
        generales_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(generales_frame, text="📈 Estadísticas Generales", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        
        # Grid para estadísticas
        stats_grid = ctk.CTkFrame(generales_frame)
        stats_grid.pack(fill="x", padx=10, pady=10)
        stats_grid.grid_columnconfigure((0, 1, 2), weight=1)
        
        # Tarjetas de estadísticas
        self.crear_tarjeta_stat(stats_grid, "📊 Total Expedientes", str(total_exp), 0, 0)
        self.crear_tarjeta_stat(stats_grid, "💰 Importe Total", f"€{importe_total:.2f}", 0, 1)
        self.crear_tarjeta_stat(stats_grid, "📊 Importe Promedio", f"€{importe_prom:.2f}", 0, 2)
        
        if primer_exp:
            fecha_primer = str(primer_exp)[:10] if primer_exp else "N/A"
            self.crear_tarjeta_stat(stats_grid, "📅 Primer Expediente", fecha_primer, 1, 0)
        if ultimo_exp:
            fecha_ultimo = str(ultimo_exp)[:10] if ultimo_exp else "N/A"
            self.crear_tarjeta_stat(stats_grid, "📅 Último Expediente", fecha_ultimo, 1, 1)
    
    def crear_tarjeta_stat(self, parent, titulo, valor, row, col):
        """Crea una tarjeta de estadística."""
        tarjeta = ctk.CTkFrame(parent)
        tarjeta.grid(row=row, column=col, padx=5, pady=5, sticky="ew")
        
        ctk.CTkLabel(tarjeta, text=titulo, font=ctk.CTkFont(size=11)).pack(pady=(5,0))
        ctk.CTkLabel(tarjeta, text=valor, font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(0,5))
    
    def exportar_estadisticas_excel(self, cliente_id):
        """Exporta las estadísticas del cliente a Excel."""
        try:
            # Verificar si pandas está disponible
            if not HAS_PANDAS:
                messagebox.showerror("Error", "Pandas no está instalado. Instale pandas para exportar a Excel:\npip install pandas openpyxl")
                return
            
            # Verificar si openpyxl está disponible
            try:
                import openpyxl
            except ImportError:
                messagebox.showerror("Error", "OpenPyXL no está instalado. Instale openpyxl para exportar a Excel:\npip install openpyxl")
                return
            
            conn, cursor = self.master.conectar_db()
            if not conn:
                return
            
            # Obtener nombre del cliente
            cursor.execute("SELECT nombre FROM clientes WHERE cliente_id = ?", (cliente_id,))
            cliente_result = cursor.fetchone()
            if not cliente_result:
                return
            
            cliente_nombre = cliente_result[0]
            
            # Construir filtros de fecha
            filtro_año = self.filtro_año_stats.get()
            filtro_mes = self.filtro_mes_stats.get()
            
            where_fecha = ""
            params = [cliente_nombre]
            
            if filtro_año != "Todos":
                where_fecha += " AND strftime('%Y', fecha_emision) = ?"
                params.append(filtro_año)
            
            if filtro_mes != "Todos":
                mes_num = ["", "01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"][
                    ["Todos", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
                     "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"].index(filtro_mes)
                ]
                where_fecha += " AND strftime('%m', fecha_emision) = ?"
                params.append(mes_num)
            
            # Consulta para datos detallados
            query_detalle = f"""
                SELECT 
                    codigo_rma,
                    fecha_emision,
                    resultado_expediente,
                    estado,
                    precio_total_expediente,
                    obs_tecnica
                FROM rma_maestro 
                WHERE cliente = ? {where_fecha}
                ORDER BY fecha_emision DESC
            """
            
            cursor.execute(query_detalle, params)
            datos_detalle = cursor.fetchall()
            
            # Crear DataFrame
            import pandas as pd
            df = pd.DataFrame(datos_detalle, columns=[
                'Código RMA', 'Fecha Emisión', 'Resultado Expediente', 
                'Estado', 'Precio Total', 'Observaciones Técnicas'
            ])
            
            # Crear nombre de archivo
            filtro_texto = ""
            if filtro_año != "Todos":
                filtro_texto += f"_{filtro_año}"
            if filtro_mes != "Todos":
                filtro_texto += f"_{filtro_mes}"
            
            nombre_archivo = f"Estadisticas_{cliente_nombre.replace(' ', '_')}{filtro_texto}.xlsx"
            
            # Guardar archivo con formato mejorado
            with pd.ExcelWriter(nombre_archivo, engine='openpyxl') as writer:
                # Escribir hoja de expedientes
                df.to_excel(writer, sheet_name='Expedientes', index=False)
                
                # Crear hoja de resumen
                query_resumen = f"""
                    SELECT 
                        resultado_expediente,
                        COUNT(*) as cantidad,
                        SUM(COALESCE(precio_total_expediente, 0)) as total_importe
                    FROM rma_maestro 
                    WHERE cliente = ? {where_fecha}
                    GROUP BY resultado_expediente
                    ORDER BY cantidad DESC
                """
                
                cursor.execute(query_resumen, params)
                datos_resumen = cursor.fetchall()
                
                df_resumen = pd.DataFrame(datos_resumen, columns=['Resultado', 'Cantidad', 'Importe Total'])
                df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
                
                # Ajustar el ancho de las columnas automáticamente
                workbook = writer.book
                
                # Ajustar columnas de la hoja Expedientes
                worksheet_expedientes = writer.sheets['Expedientes']
                for column in worksheet_expedientes.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    # Establecer un ancho mínimo y máximo para las columnas
                    adjusted_width = min(max(max_length + 2, 10), 50)
                    worksheet_expedientes.column_dimensions[column_letter].width = adjusted_width
                
                # Ajustar columnas de la hoja Resumen
                worksheet_resumen = writer.sheets['Resumen']
                for column in worksheet_resumen.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    # Establecer un ancho mínimo y máximo para las columnas
                    adjusted_width = min(max(max_length + 2, 12), 40)
                    worksheet_resumen.column_dimensions[column_letter].width = adjusted_width
                
                # Aplicar formato a los headers (primera fila)
                from openpyxl.styles import Font, PatternFill, Alignment
                
                # Formato para headers
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                center_alignment = Alignment(horizontal="center", vertical="center")
                
                # Aplicar formato a headers de Expedientes
                for cell in worksheet_expedientes[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center_alignment
                
                # Aplicar formato a headers de Resumen
                for cell in worksheet_resumen[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center_alignment
                
                # Aplicar formato de moneda a las columnas de importes
                from openpyxl.styles import NamedStyle
                
                # Estilo para moneda
                currency_style = NamedStyle(name="currency_style")
                currency_style.number_format = '€#,##0.00'
                
                # Aplicar formato de moneda a la columna "Precio Total" en Expedientes (columna E)
                for row in range(2, len(df) + 2):
                    cell = worksheet_expedientes[f'E{row}']
                    cell.style = currency_style
                
                # Aplicar formato de moneda a la columna "Importe Total" en Resumen (columna C)
                for row in range(2, len(df_resumen) + 2):
                    cell = worksheet_resumen[f'C{row}']
                    cell.style = currency_style
            
            conn.close()
            
            messagebox.showinfo("Exportación Exitosa", f"Estadísticas exportadas a:\n{nombre_archivo}")
            
        except Exception as e:
            print(f"Error exportando a Excel: {e}")
            messagebox.showerror("Error", f"Error al exportar a Excel:\n{str(e)}")
    
    def crear_item_contacto(self, parent_frame, contacto, cliente_id):
        """Crea un elemento visual para un contacto."""
        contacto_id, nombre, cargo, email, telefono, es_principal, activo = contacto
        
        contacto_frame = ctk.CTkFrame(parent_frame)
        contacto_frame.pack(fill="x", padx=5, pady=5)
        
        info_frame = ctk.CTkFrame(contacto_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=10, pady=10)
        
        # Nombre y cargo
        header_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        header_frame.pack(fill="x")
        
        nombre_text = f"👤 {nombre}"
        if es_principal:
            nombre_text += " ⭐ (Principal)"
        
        ctk.CTkLabel(header_frame, text=nombre_text, 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(side="left")
        
        if cargo:
            ctk.CTkLabel(header_frame, text=f"💼 {cargo}", 
                        font=ctk.CTkFont(size=11), text_color="blue").pack(side="left", padx=(10,0))
        
        # Información de contacto
        if email or telefono:
            contact_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
            contact_frame.pack(fill="x", pady=(5,0))
            
            if email:
                ctk.CTkLabel(contact_frame, text=f"📧 {email}", 
                            font=ctk.CTkFont(size=11)).pack(side="left")
            
            if telefono:
                ctk.CTkLabel(contact_frame, text=f"📞 {telefono}", 
                            font=ctk.CTkFont(size=11)).pack(side="left", padx=(20,0))
        
        # Botones de acción
        btn_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        btn_frame.pack(fill="x", pady=(5,0))
        
        btn_editar = ctk.CTkButton(btn_frame, text="✏️ Editar", 
                                 command=lambda: self.editar_contacto_cliente(contacto_id, cliente_id),
                                 width=80, height=25)
        btn_editar.pack(side="left", padx=(0,5))
        
        if not es_principal:
            btn_principal = ctk.CTkButton(btn_frame, text="⭐ Hacer Principal", 
                                        command=lambda: self.hacer_contacto_principal(contacto_id, cliente_id),
                                        width=120, height=25)
            btn_principal.pack(side="left", padx=(0,5))
    
    def crear_item_rma_historial(self, parent_frame, numero_rma, datos, cliente_id=None):
        """Crea un elemento visual para un RMA en el historial."""
        from lib.cliente_utils import abrir_rma_por_codigo
        
        info = datos['info']  # número, fecha, estado, motivo
        productos = datos['productos']
        
        rma_frame = ctk.CTkFrame(parent_frame)
        rma_frame.pack(fill="x", padx=5, pady=5)
        
        header_frame = ctk.CTkFrame(rma_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=10, pady=10)
        
        # Información principal del RMA
        info_principal = ctk.CTkFrame(header_frame, fg_color="transparent")
        info_principal.pack(fill="x")
        
        label_rma = ctk.CTkLabel(info_principal, text=f"📦 RMA #{numero_rma}", 
                    font=ctk.CTkFont(size=14, weight="bold"), cursor="hand2")
        label_rma.pack(side="left")
        
        # Doble clic para abrir expediente
        def abrir_expediente(e):
            abrir_rma_por_codigo(numero_rma, self.master.conectar_db, self)
        
        label_rma.bind("<Double-Button-1>", abrir_expediente)
        Tooltip(label_rma, "Doble clic para abrir el expediente")
        
        # Estado del RMA
        estado_color = {"Abierto": "orange", "En Proceso": "blue", "Cerrado": "green", "Cancelado": "red"}.get(info[2], "gray")
        ctk.CTkLabel(info_principal, text=f"🏷️ {info[2]}", 
                    font=ctk.CTkFont(size=11), text_color=estado_color).pack(side="right")
        
        # Fecha - manejo robusto
        try:
            fecha_str = str(info[1])
            if ' ' in fecha_str:
                fecha_mostrar = fecha_str.split()[0]  # Solo la parte de la fecha
            else:
                fecha_mostrar = fecha_str  # Si no tiene espacios, usar toda la cadena
        except (IndexError, AttributeError):
            fecha_mostrar = "Sin fecha"
        
        ctk.CTkLabel(info_principal, text=f"📅 {fecha_mostrar}", 
                    font=ctk.CTkFont(size=11), text_color="gray").pack(side="right", padx=(0,10))
        
        # Motivo
        if info[3]:
            motivo_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            motivo_frame.pack(fill="x", pady=(5,0))
            ctk.CTkLabel(motivo_frame, text=f"📝 Motivo: {info[3]}", 
                        font=ctk.CTkFont(size=11)).pack(side="left")
        
        # Productos (si los hay)
        if productos:
            productos_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            productos_frame.pack(fill="x", pady=(5,0))
            
            ctk.CTkLabel(productos_frame, text=f"📦 Productos ({len(productos)}):", 
                        font=ctk.CTkFont(size=11, weight="bold")).pack(anchor="w")
            
            for producto, cantidad, precio, estado in productos[:3]:  # Mostrar máximo 3
                prod_text = f"  • {producto} (Cant: {cantidad}"
                if precio:
                    prod_text += f", €{precio}"
                if estado:
                    prod_text += f", Estado: {estado}"
                prod_text += ")"
                
                ctk.CTkLabel(productos_frame, text=prod_text, 
                            font=ctk.CTkFont(size=10)).pack(anchor="w")
            
            if len(productos) > 3:
                ctk.CTkLabel(productos_frame, text=f"  ... y {len(productos)-3} más", 
                            font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w")
    
    def crear_item_nota(self, parent_frame, nota, parent_window=None):
        """Crea un elemento visual para una nota."""
        # Estructura correcta: nota_id, titulo, contenido, tipo, fecha, usuario, privada
        nota_id, titulo, contenido, tipo, fecha, usuario, privada = nota
        
        nota_frame = ctk.CTkFrame(parent_frame)
        nota_frame.pack(fill="x", padx=5, pady=5)
        
        header_frame = ctk.CTkFrame(nota_frame, fg_color="transparent")
        header_frame.pack(fill="x", padx=10, pady=10)
        
        # Título y tipo
        titulo_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        titulo_frame.pack(fill="x")
        
        titulo_icon = "🔒" if privada else "📝"
        ctk.CTkLabel(titulo_frame, text=f"{titulo_icon} {titulo}", 
                    font=ctk.CTkFont(size=14, weight="bold")).pack(side="left")
        
        tipo_color = {
            "General": "blue", 
            "Incidencia": "red", 
            "Comercial": "green", 
            "Técnica": "orange"
        }.get(tipo, "gray")
        ctk.CTkLabel(titulo_frame, text=f"🏷️ {tipo}", 
                    font=ctk.CTkFont(size=11), text_color=tipo_color).pack(side="right")
        
        # Fecha y usuario
        info_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        info_frame.pack(fill="x", pady=(5,0))
        
        fecha_formateada = fecha.split(' ')[0] if fecha else "Sin fecha"
        ctk.CTkLabel(info_frame, text=f"📅 {fecha_formateada} | 👤 {usuario}", 
                    font=ctk.CTkFont(size=10), text_color="gray").pack(side="left")
        
        # Botones de acción
        botones_frame = ctk.CTkFrame(info_frame, fg_color="transparent")
        botones_frame.pack(side="right")
        
        btn_ver = ctk.CTkButton(botones_frame, text="👁️", width=30, height=25,
                               command=lambda: self.ver_nota_completa(nota))
        btn_ver.pack(side="right", padx=2)
        
        btn_editar = ctk.CTkButton(botones_frame, text="✏️", width=30, height=25,
                                  command=lambda: self.editar_nota(nota_id, parent_window))
        btn_editar.pack(side="right", padx=2)
        
        # Contenido (limitado)
        if contenido:
            contenido_preview = contenido[:200] + "..." if len(contenido) > 200 else contenido
            contenido_frame = ctk.CTkFrame(nota_frame)
            contenido_frame.pack(fill="x", padx=10, pady=(0,10))
            
            ctk.CTkLabel(contenido_frame, text=contenido_preview, 
                        font=ctk.CTkFont(size=11), 
                        wraplength=700, justify="left").pack(anchor="w", padx=10, pady=5)
    
    def mostrar_estadisticas_principales(self, parent_frame, stats):
        """Muestra las estadísticas principales del cliente."""
        # Nueva estructura de stats desde estadisticas_cliente:
        # [0]: cliente_id, [1]: nombre, [2]: tipo_cliente, [3]: activo,
        # [4]: total_rmas, [5]: rmas_completados, [6]: rmas_abiertos,
        # [7]: rmas_pendientes, [8]: rmas_pendientes_autorizar, [9]: rmas_recibidos,
        # [10]: rmas_en_tramite, [11]: tasa_exito, [12]: primer_rma, [13]: ultimo_rma,
        # [14]: productos_diferentes, [15]: rmas_ultimo_mes
        
        ctk.CTkLabel(parent_frame, text="📊 Estadísticas Generales", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(0,10))
        
        stats_frame = ctk.CTkFrame(parent_frame)
        stats_frame.pack(fill="x", pady=(0,20))
        stats_frame.grid_columnconfigure((0,1,2), weight=1)
        
        # Fila 1: RMAs totales, abiertas (NO completados), cerradas (completados)
        ctk.CTkLabel(stats_frame, text=f"📦 Total RMAs\n{stats[4] or 0}", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color="blue").grid(row=0, column=0, padx=10, pady=10)
        
        ctk.CTkLabel(stats_frame, text=f"🟡 RMAs Abiertos\n{stats[6] or 0}", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color="orange").grid(row=0, column=1, padx=10, pady=10)
        
        ctk.CTkLabel(stats_frame, text=f"🟢 RMAs Cerrados\n{stats[5] or 0}", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color="green").grid(row=0, column=2, padx=10, pady=10)
        
        # Fila 2: Tasa de éxito y otros datos
        tasa_exito = stats[11] or 0  # Actualizado al índice correcto
        # Convertir a número si viene como string
        try:
            tasa_exito = float(tasa_exito)
        except (ValueError, TypeError):
            tasa_exito = 0.0
        
        color_tasa = "green" if tasa_exito >= 80 else "orange" if tasa_exito >= 60 else "red"
        
        ctk.CTkLabel(stats_frame, text=f"✅ Tasa de Éxito\n{tasa_exito:.1f}%", 
                    font=ctk.CTkFont(size=14, weight="bold"), 
                    text_color=color_tasa).grid(row=1, column=0, padx=10, pady=10)
        
        if len(stats) > 14 and stats[14]:  # productos_diferentes (índice corregido)
            ctk.CTkLabel(stats_frame, text=f"📦 Total Productos\n{stats[14]}", 
                        font=ctk.CTkFont(size=14, weight="bold"), 
                        text_color="blue").grid(row=1, column=1, padx=10, pady=10)
        
        if len(stats) > 13 and stats[13]:  # ultimo_rma (índice corregido)
            ctk.CTkLabel(stats_frame, text=f"📅 Último RMA\n{stats[13]}", 
                        font=ctk.CTkFont(size=11, weight="bold"), 
                        text_color="gray").grid(row=1, column=2, padx=10, pady=10)
    
    def mostrar_productos_problematicos(self, parent_frame, productos):
        """Muestra los productos más problemáticos del cliente."""
        ctk.CTkLabel(parent_frame, text="⚠️ Productos Más Problemáticos", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", pady=(20,10))
        
        productos_frame = ctk.CTkFrame(parent_frame)
        productos_frame.pack(fill="x", pady=(0,10))
        
        for i, producto in enumerate(productos[:5], 1):  # Top 5
            # producto: cliente_id, producto, total_devoluciones, valor_total
            prod_frame = ctk.CTkFrame(productos_frame, fg_color="transparent")
            prod_frame.pack(fill="x", padx=10, pady=5)
            
            ctk.CTkLabel(prod_frame, text=f"{i}. {producto[1]}", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(side="left")
            
            ctk.CTkLabel(prod_frame, text=f"🔄 {producto[2]} devoluciones", 
                        font=ctk.CTkFont(size=11), text_color="red").pack(side="right")
            
            if producto[3]:  # valor_total
                ctk.CTkLabel(prod_frame, text=f"💰 €{producto[3]:.2f}", 
                            font=ctk.CTkFont(size=11), text_color="gray").pack(side="right", padx=(0,10))
    
    # Funciones auxiliares para gestión de contactos y notas
    def nuevo_contacto_cliente(self, cliente_id):
        """Muestra el formulario para agregar un nuevo contacto."""
        ventana = ctk.CTkToplevel(self)
        ventana.title("Nuevo Contacto")
        ventana.geometry("400x500")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="👤 Nuevo Contacto", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=300)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Nombre (obligatorio)
        ctk.CTkLabel(form_frame, text="Nombre del Contacto *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_nombre = ctk.CTkEntry(form_frame, placeholder_text="Nombre completo del contacto")
        entry_nombre.pack(fill="x", pady=(0,10))
        
        # Cargo
        ctk.CTkLabel(form_frame, text="Cargo", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_cargo = ctk.CTkEntry(form_frame, placeholder_text="Cargo o posición en la empresa")
        entry_cargo.pack(fill="x", pady=(0,10))
        
        # Email
        ctk.CTkLabel(form_frame, text="Email", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_email = ctk.CTkEntry(form_frame, placeholder_text="email@ejemplo.com")
        entry_email.pack(fill="x", pady=(0,10))
        
        # Teléfono
        ctk.CTkLabel(form_frame, text="Teléfono", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_telefono = ctk.CTkEntry(form_frame, placeholder_text="Número de teléfono")
        entry_telefono.pack(fill="x", pady=(0,10))
        
        # Contacto principal
        check_principal = ctk.CTkCheckBox(form_frame, text="🌟 Marcar como contacto principal")
        check_principal.pack(anchor="w", pady=10)
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        def guardar_contacto():
            nombre = entry_nombre.get().strip()
            if not nombre:
                messagebox.showerror("Error", "El nombre del contacto es obligatorio")
                return
            
            datos = {
                'nombre': nombre,
                'cargo': entry_cargo.get().strip(),
                'email': entry_email.get().strip(),
                'telefono': entry_telefono.get().strip(),
                'es_principal': check_principal.get()
            }
            
            if self.crear_contacto_cliente(cliente_id, datos):
                messagebox.showinfo("Éxito", f"Contacto '{nombre}' creado correctamente")
                ventana.destroy()
                # Recargar la pestaña de contactos si está abierta
            else:
                messagebox.showerror("Error", "Error al crear el contacto")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Guardar Contacto", 
                                  command=guardar_contacto,
                                  width=120)
        btn_guardar.pack(side="right")
    
    def editar_contacto_cliente(self, contacto_id, cliente_id):
        """Permite editar un contacto existente."""
        # Obtener datos del contacto
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return
            
            cursor.execute("""
                SELECT nombre, cargo, email, telefono, es_principal
                FROM contactos_cliente 
                WHERE contacto_id = ?
            """, (contacto_id,))
            
            contacto = cursor.fetchone()
            conn.close()
            
            if not contacto:
                messagebox.showerror("Error", "No se pudo cargar el contacto")
                return
                
        except Exception as e:
            messagebox.showerror("Error", f"Error cargando contacto: {e}")
            return
        
        # Crear ventana de edición
        ventana = ctk.CTkToplevel(self)
        ventana.title("Editar Contacto")
        ventana.geometry("400x500")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="✏️ Editar Contacto", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=300)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Nombre (obligatorio)
        ctk.CTkLabel(form_frame, text="Nombre del Contacto *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_nombre = ctk.CTkEntry(form_frame, placeholder_text="Nombre completo del contacto")
        entry_nombre.pack(fill="x", pady=(0,10))
        entry_nombre.insert(0, contacto[0])
        
        # Cargo
        ctk.CTkLabel(form_frame, text="Cargo", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_cargo = ctk.CTkEntry(form_frame, placeholder_text="Cargo o posición en la empresa")
        entry_cargo.pack(fill="x", pady=(0,10))
        entry_cargo.insert(0, contacto[1] or "")
        
        # Email
        ctk.CTkLabel(form_frame, text="Email", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_email = ctk.CTkEntry(form_frame, placeholder_text="email@ejemplo.com")
        entry_email.pack(fill="x", pady=(0,10))
        entry_email.insert(0, contacto[2] or "")
        
        # Teléfono
        ctk.CTkLabel(form_frame, text="Teléfono", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_telefono = ctk.CTkEntry(form_frame, placeholder_text="Número de teléfono")
        entry_telefono.pack(fill="x", pady=(0,10))
        entry_telefono.insert(0, contacto[3] or "")
        
        # Contacto principal
        check_principal = ctk.CTkCheckBox(form_frame, text="🌟 Marcar como contacto principal")
        check_principal.pack(anchor="w", pady=10)
        if contacto[4]:
            check_principal.select()
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        btn_eliminar = ctk.CTkButton(botones_frame, text="🗑️ Eliminar", 
                                   command=lambda: self.eliminar_contacto_cliente(contacto_id, ventana),
                                   width=100, fg_color="red")
        btn_eliminar.pack(side="left")
        
        def actualizar_contacto():
            nombre = entry_nombre.get().strip()
            if not nombre:
                messagebox.showerror("Error", "El nombre del contacto es obligatorio")
                return
            
            try:
                conn, cursor = self.master.conectar_db()
                if not conn: 
                    return
                
                # Si es contacto principal, desmarcar otros
                if check_principal.get():
                    cursor.execute("""
                        UPDATE contactos_cliente 
                        SET es_principal = 0 
                        WHERE cliente_id = ? AND contacto_id != ?
                    """, (cliente_id, contacto_id))
                
                cursor.execute("""
                    UPDATE contactos_cliente 
                    SET nombre = ?, cargo = ?, email = ?, telefono = ?, es_principal = ?,
                        fecha_actualizacion = CURRENT_TIMESTAMP
                    WHERE contacto_id = ?
                """, (nombre, entry_cargo.get().strip(), entry_email.get().strip(),
                      entry_telefono.get().strip(), check_principal.get(), contacto_id))
                
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Éxito", f"Contacto '{nombre}' actualizado correctamente")
                ventana.destroy()
                
            except Exception as e:
                messagebox.showerror("Error", f"Error al actualizar contacto: {e}")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Actualizar", 
                                  command=actualizar_contacto,
                                  width=120)
        btn_guardar.pack(side="right", padx=(10,0))
    
    def hacer_contacto_principal(self, contacto_id, cliente_id):
        """Marca un contacto como principal."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return
            
            # Desmarcar todos los contactos del cliente
            cursor.execute("""
                UPDATE contactos_cliente 
                SET es_principal = 0 
                WHERE cliente_id = ?
            """, (cliente_id,))
            
            # Marcar el contacto seleccionado como principal
            cursor.execute("""
                UPDATE contactos_cliente 
                SET es_principal = 1, fecha_actualizacion = CURRENT_TIMESTAMP
                WHERE contacto_id = ?
            """, (contacto_id,))
            
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Éxito", "Contacto marcado como principal correctamente")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al marcar contacto como principal: {e}")
    
    def eliminar_contacto_cliente(self, contacto_id, ventana_padre):
        """Elimina un contacto del cliente (eliminación suave)."""
        if messagebox.askyesno("Confirmar Eliminación", 
                              "¿Estás seguro de que deseas eliminar este contacto?\n\nEsta acción no se puede deshacer."):
            try:
                conn, cursor = self.master.conectar_db()
                if not conn: 
                    return
                
                cursor.execute("""
                    UPDATE contactos_cliente 
                    SET activo = 0, fecha_actualizacion = CURRENT_TIMESTAMP
                    WHERE contacto_id = ?
                """, (contacto_id,))
                
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Éxito", "Contacto eliminado correctamente")
                ventana_padre.destroy()
                
            except Exception as e:
                messagebox.showerror("Error", f"Error al eliminar contacto: {e}")
    
    def nueva_nota_cliente(self, cliente_id, parent_window=None):
        """Muestra el formulario para agregar una nueva nota."""
        ventana = ctk.CTkToplevel(parent_window or self)
        ventana.title("Nueva Nota")
        ventana.geometry("500x600")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(parent_window or self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="📝 Nueva Nota", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=400)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Título de la nota (obligatorio)
        ctk.CTkLabel(form_frame, text="Título de la Nota *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_titulo = ctk.CTkEntry(form_frame, placeholder_text="Título descriptivo de la nota")
        entry_titulo.pack(fill="x", pady=(0,10))
        
        # Tipo (usar nombres de la base de datos)
        ctk.CTkLabel(form_frame, text="Tipo", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        option_tipo = ctk.CTkOptionMenu(form_frame, values=["General", "Incidencia", "Comercial", "Técnica"])
        option_tipo.set("General")
        option_tipo.pack(fill="x", pady=(0,10))
        
        # Contenido
        ctk.CTkLabel(form_frame, text="Contenido de la Nota *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        text_contenido = ctk.CTkTextbox(form_frame, height=200)
        text_contenido.pack(fill="x", pady=(0,10))
        
        # Usuario (opcional, se puede pre-rellenar)
        ctk.CTkLabel(form_frame, text="Usuario", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_usuario = ctk.CTkEntry(form_frame, placeholder_text="Nombre del usuario que crea la nota")
        entry_usuario.pack(fill="x", pady=(0,10))
        entry_usuario.insert(0, "Usuario")  # Valor por defecto
        
        # Checkbox para nota privada
        check_privada = ctk.CTkCheckBox(form_frame, text="Nota privada (solo visible para administradores)")
        check_privada.pack(anchor="w", pady=10)
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        def guardar_nota():
            titulo = entry_titulo.get().strip()
            contenido = text_contenido.get("1.0", "end-1c").strip()
            
            if not titulo:
                messagebox.showerror("Error", "El título de la nota es obligatorio")
                return
            
            if not contenido:
                messagebox.showerror("Error", "El contenido de la nota es obligatorio")
                return
            
            datos = {
                'titulo': titulo,
                'contenido': contenido,
                'tipo': option_tipo.get(),
                'usuario': entry_usuario.get().strip() or "Usuario",
                'privada': check_privada.get()
            }
            
            if self.crear_nota_cliente(cliente_id, datos):
                messagebox.showinfo("Éxito", f"Nota '{titulo}' creada correctamente")
                ventana.destroy()
                # Si hay ventana padre, refrescar la lista
                if hasattr(parent_window, 'title') and 'Gestión de Notas' in parent_window.title():
                    # Buscar y refrescar la lista de notas en la ventana padre
                    pass
            else:
                messagebox.showerror("Error", "Error al crear la nota")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Guardar Nota", 
                                  command=guardar_nota,
                                  width=150)
        btn_guardar.pack(side="right")
    
    def editar_cliente_directo(self, cliente_id, ventana_padre):
        """Permite editar la información del cliente desde la ficha."""
        # Obtener datos del cliente
        cliente = self.obtener_cliente(cliente_id)
        if not cliente:
            messagebox.showerror("Error", "No se pudo cargar la información del cliente")
            return
        
        # Crear ventana de edición
        ventana = ctk.CTkToplevel(self)
        ventana.title(f"Editar Cliente: {cliente[1]}")
        ventana.geometry("500x600")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="✏️ Editar Cliente", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=400)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Llenar formulario con datos actuales
        ctk.CTkLabel(form_frame, text="Nombre del Cliente *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_nombre = ctk.CTkEntry(form_frame, placeholder_text="Nombre completo del cliente")
        entry_nombre.pack(fill="x", pady=(0,10))
        entry_nombre.insert(0, cliente[1])
        
        ctk.CTkLabel(form_frame, text="Tipo de Cliente", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        option_tipo = ctk.CTkOptionMenu(form_frame, values=["Regular", "Premium", "VIP"])
        option_tipo.set(cliente[2])
        option_tipo.pack(fill="x", pady=(0,10))
        
        ctk.CTkLabel(form_frame, text="Dirección", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_direccion = ctk.CTkEntry(form_frame, placeholder_text="Dirección completa")
        entry_direccion.pack(fill="x", pady=(0,10))
        entry_direccion.insert(0, cliente[3] or "")
        
        ctk.CTkLabel(form_frame, text="Teléfono Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_telefono = ctk.CTkEntry(form_frame, placeholder_text="Teléfono de contacto principal")
        entry_telefono.pack(fill="x", pady=(0,10))
        entry_telefono.insert(0, cliente[4] or "")
        
        ctk.CTkLabel(form_frame, text="Email Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_email = ctk.CTkEntry(form_frame, placeholder_text="Email de contacto principal")
        entry_email.pack(fill="x", pady=(0,10))
        entry_email.insert(0, cliente[5] or "")
        
        ctk.CTkLabel(form_frame, text="Notas Generales", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        text_notas = ctk.CTkTextbox(form_frame, height=80)
        text_notas.pack(fill="x", pady=(0,10))
        text_notas.insert("1.0", cliente[6] or "")
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        def actualizar_cliente():
            nombre = entry_nombre.get().strip()
            if not nombre:
                messagebox.showerror("Error", "El nombre del cliente es obligatorio")
                return
            
            datos = {
                'nombre': nombre,
                'tipo_cliente': option_tipo.get(),
                'direccion': entry_direccion.get().strip(),
                'telefono_principal': entry_telefono.get().strip(),
                'email_principal': entry_email.get().strip(),
                'notas_generales': text_notas.get("1.0", "end-1c").strip()
            }
            
            if self.actualizar_cliente(cliente_id, datos):
                messagebox.showinfo("Éxito", f"Cliente '{nombre}' actualizado correctamente")
                ventana.destroy()
                # Cerrar la ventana padre para que se recargue
                ventana_padre.destroy()
                # Reabrir la ficha con datos actualizados
                self.abrir_ficha_cliente(cliente_id)
            else:
                messagebox.showerror("Error", "Error al actualizar el cliente")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Actualizar Cliente", 
                                  command=actualizar_cliente,
                                  width=140)
        btn_guardar.pack(side="right")
    
    def editar_cliente(self, cliente_id):
        """Permite editar la información básica del cliente."""
        # Obtener datos del cliente
        cliente = self.obtener_cliente(cliente_id)
        if not cliente:
            messagebox.showerror("Error", "No se pudo cargar la información del cliente")
            return
        
        # Crear ventana de edición
        ventana = ctk.CTkToplevel(self)
        ventana.title(f"Editar Cliente: {cliente[1]}")
        ventana.geometry("500x600")
        ventana.resizable(False, False)
        
        # Centrar ventana
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        titulo_frame = ctk.CTkFrame(ventana)
        titulo_frame.pack(fill="x", padx=20, pady=20)
        
        ctk.CTkLabel(titulo_frame, text="✏️ Editar Cliente", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
        
        # Formulario
        form_frame = ctk.CTkScrollableFrame(ventana, height=400)
        form_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Llenar formulario con datos actuales
        ctk.CTkLabel(form_frame, text="Nombre del Cliente *", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_nombre = ctk.CTkEntry(form_frame, placeholder_text="Nombre completo del cliente")
        entry_nombre.pack(fill="x", pady=(0,10))
        entry_nombre.insert(0, cliente[1])
        
        ctk.CTkLabel(form_frame, text="Tipo de Cliente", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        option_tipo = ctk.CTkOptionMenu(form_frame, values=["Regular", "Premium", "VIP"])
        option_tipo.set(cliente[2])
        option_tipo.pack(fill="x", pady=(0,10))
        
        ctk.CTkLabel(form_frame, text="Dirección", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_direccion = ctk.CTkEntry(form_frame, placeholder_text="Dirección completa")
        entry_direccion.pack(fill="x", pady=(0,10))
        entry_direccion.insert(0, cliente[3] or "")
        
        ctk.CTkLabel(form_frame, text="Teléfono Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_telefono = ctk.CTkEntry(form_frame, placeholder_text="Teléfono de contacto principal")
        entry_telefono.pack(fill="x", pady=(0,10))
        entry_telefono.insert(0, cliente[4] or "")
        
        ctk.CTkLabel(form_frame, text="Email Principal", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        entry_email = ctk.CTkEntry(form_frame, placeholder_text="Email de contacto principal")
        entry_email.pack(fill="x", pady=(0,10))
        entry_email.insert(0, cliente[5] or "")
        
        ctk.CTkLabel(form_frame, text="Notas Generales", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
        text_notas = ctk.CTkTextbox(form_frame, height=80)
        text_notas.pack(fill="x", pady=(0,10))
        text_notas.insert("1.0", cliente[6] or "")
        
        # Estado del cliente
        check_activo = ctk.CTkCheckBox(form_frame, text="Cliente activo")
        check_activo.pack(anchor="w", pady=10)
        if cliente[7]:  # cliente[7] es activo
            check_activo.select()
        
        # Botones
        botones_frame = ctk.CTkFrame(ventana)
        botones_frame.pack(fill="x", padx=20, pady=20)
        
        btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                   command=ventana.destroy,
                                   width=100)
        btn_cancelar.pack(side="right", padx=(10,0))
        
        if cliente[7]:  # Solo mostrar eliminar si está activo
            btn_eliminar = ctk.CTkButton(botones_frame, text="🗑️ Desactivar", 
                                       command=lambda: self.confirmar_eliminar_cliente(cliente_id, ventana),
                                       width=100, fg_color="red")
            btn_eliminar.pack(side="left")
        
        def actualizar_cliente():
            nombre = entry_nombre.get().strip()
            if not nombre:
                messagebox.showerror("Error", "El nombre del cliente es obligatorio")
                return
            
            datos = {
                'nombre': nombre,
                'tipo_cliente': option_tipo.get(),
                'direccion': entry_direccion.get().strip(),
                'telefono_principal': entry_telefono.get().strip(),
                'email_principal': entry_email.get().strip(),
                'notas_generales': text_notas.get("1.0", "end-1c").strip()
            }
            
            # Actualizar estado si cambió
            if check_activo.get() != cliente[7]:
                try:
                    conn, cursor = self.master.conectar_db()
                    if conn:
                        cursor.execute("""
                            UPDATE clientes 
                            SET activo = ?, fecha_actualizacion = CURRENT_TIMESTAMP
                            WHERE cliente_id = ?
                        """, (check_activo.get(), cliente_id))
                        conn.commit()
                        conn.close()
                except Exception as e:
                    print(f"Error actualizando estado: {e}")
            
            if self.actualizar_cliente(cliente_id, datos):
                messagebox.showinfo("Éxito", f"Cliente '{nombre}' actualizado correctamente")
                ventana.destroy()
                self.cargar_lista_clientes()  # Recargar lista
            else:
                messagebox.showerror("Error", "Error al actualizar el cliente")
        
        btn_guardar = ctk.CTkButton(botones_frame, text="💾 Actualizar Cliente", 
                                  command=actualizar_cliente,
                                  width=140)
        btn_guardar.pack(side="right")
    
    def confirmar_eliminar_cliente(self, cliente_id, ventana_padre):
        """Confirma y realiza la eliminación suave del cliente."""
        if messagebox.askyesno("Confirmar Desactivación", 
                              "¿Estás seguro de que deseas desactivar este cliente?\n\nEl cliente no se eliminará, solo se marcará como inactivo."):
            if self.eliminar_cliente(cliente_id):
                messagebox.showinfo("Éxito", "Cliente desactivado correctamente")
                ventana_padre.destroy()
                self.cargar_lista_clientes()  # Recargar lista
            else:
                messagebox.showerror("Error", "Error al desactivar el cliente")
    
    def gestionar_notas_cliente(self, cliente_id):
        """Gestiona las notas del cliente - Abre ventana de gestión completa."""
        try:
            # Crear ventana de gestión de notas
            ventana_notas = ctk.CTkToplevel(self)
            ventana_notas.title("Gestión de Notas")
            ventana_notas.geometry("800x600")
            
            # Configurar para permitir minimización
            ventana_notas.resizable(True, True)
            ventana_notas.attributes('-topmost', False)
            ventana_notas.minsize(600, 400)
            # No usar transient para permitir minimización completa
            ventana_notas.focus_set()  # Dar foco sin bloquear
            
            # Forzar aparición al frente (incluso si la principal está maximizada)
            ventana_notas.attributes('-topmost', True)   # Temporalmente al frente
            ventana_notas.lift()
            try:
                ventana_notas.focus_force()
            except:
                pass
            
            def quitar_topmost_ventana_notas():
                try:
                    if ventana_notas.winfo_exists():
                        ventana_notas.attributes('-topmost', False)
                except:
                    pass
            
            ventana_notas.after(500, quitar_topmost_ventana_notas)
            
            # Agregar icono personalizado
            try:
                ventana_notas.iconbitmap("Icono_Ilutrek.ico")
            except Exception:
                pass
            
            # Obtener información del cliente
            cliente = self.obtener_cliente(cliente_id)
            if not cliente:
                messagebox.showerror("Error", "No se pudo obtener información del cliente")
                ventana_notas.destroy()
                return
            
            # Header
            header_frame = ctk.CTkFrame(ventana_notas)
            header_frame.pack(fill="x", padx=10, pady=10)
            
            ctk.CTkLabel(header_frame, text=f"📝 Notas de {cliente[1]}", 
                        font=ctk.CTkFont(size=18, weight="bold")).pack(side="left", padx=10, pady=10)
            
            btn_nueva_nota = ctk.CTkButton(header_frame, text="➕ Nueva Nota", 
                                         command=lambda: self.nueva_nota_cliente(cliente_id, ventana_notas),
                                         width=120)
            btn_nueva_nota.pack(side="right", padx=10, pady=10)
            
            # Frame para filtros
            filtros_frame = ctk.CTkFrame(ventana_notas)
            filtros_frame.pack(fill="x", padx=10, pady=(0,10))
            
            ctk.CTkLabel(filtros_frame, text="Filtrar por tipo:").pack(side="left", padx=10, pady=10)
            
            tipo_filter = ctk.CTkComboBox(filtros_frame, values=["Todos", "General", "Incidencia", "Comercial", "Técnica"])
            tipo_filter.set("Todos")
            tipo_filter.pack(side="left", padx=10, pady=10)
            
            # Lista de notas
            notas_frame = ctk.CTkScrollableFrame(ventana_notas, height=400)
            notas_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            def cargar_notas():
                # Limpiar frame
                for widget in notas_frame.winfo_children():
                    widget.destroy()
                
                # Obtener notas
                filtro_tipo = tipo_filter.get()
                notas = self.obtener_notas_cliente(cliente_id, filtro_tipo if filtro_tipo != "Todos" else None)
                
                if not notas:
                    ctk.CTkLabel(notas_frame, 
                               text="📭 No hay notas que coincidan con el filtro.\nHaz clic en 'Nueva Nota' para agregar una.",
                               font=ctk.CTkFont(size=13)).pack(pady=50)
                else:
                    for nota in notas:
                        self.crear_item_nota(notas_frame, nota, ventana_notas)
            
            # Conectar filtro
            tipo_filter.configure(command=lambda x: cargar_notas())
            
            # Cargar notas inicial
            cargar_notas()
            
            # Botón cerrar
            btn_cerrar = ctk.CTkButton(ventana_notas, text="Cerrar", 
                                     command=ventana_notas.destroy,
                                     width=100)
            btn_cerrar.pack(pady=10)
            
        except Exception as e:
            print(f"❌ Error gestionando notas: {e}")
            messagebox.showerror("Error", f"Error al abrir gestión de notas: {e}")

    # ===============================================
    # FUNCIONES DE BASE DE DATOS PARA CLIENTES
    # ===============================================
    
    def obtener_cliente(self, cliente_id):
        """Obtiene la información completa de un cliente."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return None
            
            cursor.execute("""
                SELECT cliente_id, nombre, tipo_cliente, direccion, telefono_principal,
                       email_principal, notas_generales, activo, fecha_registro, fecha_actualizacion
                FROM clientes 
                WHERE cliente_id = ?
            """, (cliente_id,))
            
            cliente = cursor.fetchone()
            conn.close()
            return cliente
            
        except Exception as e:
            print(f"Error obteniendo cliente: {e}")
            return None
    
    def actualizar_cliente(self, cliente_id, datos):
        """Actualiza la información de un cliente."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return False
            
            cursor.execute("""
                UPDATE clientes 
                SET nombre = ?, tipo_cliente = ?, direccion = ?, telefono_principal = ?,
                    email_principal = ?, notas_generales = ?, fecha_actualizacion = CURRENT_TIMESTAMP
                WHERE cliente_id = ?
            """, (datos['nombre'], datos['tipo_cliente'], datos['direccion'], 
                  datos['telefono_principal'], datos['email_principal'], 
                  datos['notas_generales'], cliente_id))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            print(f"Error actualizando cliente: {e}")
            return False
    
    def eliminar_cliente(self, cliente_id):
        """Marca un cliente como inactivo (eliminación suave)."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return False
            
            cursor.execute("""
                UPDATE clientes 
                SET activo = 0, fecha_actualizacion = CURRENT_TIMESTAMP
                WHERE cliente_id = ?
            """, (cliente_id,))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            print(f"Error eliminando cliente: {e}")
            return False
    
    def obtener_contactos_cliente(self, cliente_id):
        """Obtiene todos los contactos de un cliente."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return []
            
            cursor.execute("""
                SELECT contacto_id, nombre, cargo, email, telefono, es_principal, activo
                FROM contactos_cliente 
                WHERE cliente_id = ? AND activo = 1
                ORDER BY es_principal DESC, nombre
            """, (cliente_id,))
            
            contactos = cursor.fetchall()
            conn.close()
            return contactos
            
        except Exception as e:
            print(f"Error obteniendo contactos: {e}")
            return []
    
    def crear_contacto_cliente(self, cliente_id, datos):
        """Crea un nuevo contacto para un cliente."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return False
            
            # Si es contacto principal, desmarcar otros
            if datos.get('es_principal', False):
                cursor.execute("""
                    UPDATE contactos_cliente 
                    SET es_principal = 0 
                    WHERE cliente_id = ?
                """, (cliente_id,))
            
            cursor.execute("""
                INSERT INTO contactos_cliente (cliente_id, nombre, cargo, email, telefono, es_principal)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (cliente_id, datos['nombre'], datos.get('cargo', ''), 
                  datos.get('email', ''), datos.get('telefono', ''), 
                  datos.get('es_principal', False)))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            print(f"Error creando contacto: {e}")
            return False
    
    def obtener_notas_cliente(self, cliente_id, tipo_filtro=None):
        """Obtiene todas las notas de un cliente."""
        try:
            conn = connect_db()
            cursor = conn.cursor()
            
            # Query base
            query = """
                SELECT nota_id, titulo, contenido, tipo, fecha, usuario, privada
                FROM notas_cliente 
                WHERE cliente_id = ?
            """
            params = [cliente_id]
            
            # Agregar filtro por tipo si se especifica
            if tipo_filtro:
                query += " AND tipo = ?"
                params.append(tipo_filtro)
            
            query += " ORDER BY fecha DESC"
            
            cursor.execute(query, params)
            notas = cursor.fetchall()
            conn.close()
            return notas
            
        except Exception as e:
            print(f"❌ Error obteniendo notas: {e}")
            return []
    
    def crear_nota_cliente(self, cliente_id, datos):
        """Crea una nueva nota para un cliente."""
        try:
            conn = connect_db()
            cursor = conn.cursor()
            
            cursor.execute("""
                INSERT INTO notas_cliente (cliente_id, usuario, tipo, titulo, contenido, privada)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                cliente_id, 
                datos.get('usuario', 'Usuario'),
                datos.get('tipo', 'General'), 
                datos['titulo'], 
                datos['contenido'], 
                datos.get('privada', False)
            ))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            print(f"❌ Error creando nota: {e}")
            return False

    def ver_nota_completa(self, nota):
        """Muestra la nota completa en una ventana emergente."""
        nota_id, titulo, contenido, tipo, fecha, usuario, privada = nota
        
        ventana = ctk.CTkToplevel(self)
        ventana.title(f"Nota: {titulo}")
        ventana.geometry("600x500")
        ventana.transient(self)
        
        # Header
        header_frame = ctk.CTkFrame(ventana)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        titulo_icon = "🔒" if privada else "📝"
        ctk.CTkLabel(header_frame, text=f"{titulo_icon} {titulo}", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(anchor="w", padx=10, pady=10)
        
        # Info
        info_frame = ctk.CTkFrame(ventana)
        info_frame.pack(fill="x", padx=10, pady=(0,10))
        
        ctk.CTkLabel(info_frame, text=f"🏷️ Tipo: {tipo} | 📅 Fecha: {fecha} | 👤 Usuario: {usuario}", 
                    font=ctk.CTkFont(size=12)).pack(anchor="w", padx=10, pady=10)
        
        # Contenido
        contenido_frame = ctk.CTkScrollableFrame(ventana, height=300)
        contenido_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        ctk.CTkLabel(contenido_frame, text=contenido, 
                    font=ctk.CTkFont(size=12), 
                    wraplength=550, justify="left").pack(anchor="w", padx=10, pady=10)
        
        # Botón cerrar
        ctk.CTkButton(ventana, text="Cerrar", command=ventana.destroy).pack(pady=10)

    def editar_nota(self, nota_id, parent_window=None):
        """Permite editar una nota existente."""
        try:
            # Obtener datos actuales de la nota
            conn = connect_db()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM notas_cliente WHERE nota_id = ?", (nota_id,))
            nota_actual = cursor.fetchone()
            conn.close()
            
            if not nota_actual:
                messagebox.showerror("Error", "No se pudo cargar la nota")
                return
            
            # nota_actual: nota_id, cliente_id, usuario, fecha, tipo, titulo, contenido, privada
            _, cliente_id, usuario_actual, _, tipo_actual, titulo_actual, contenido_actual, privada_actual = nota_actual
            
            ventana = ctk.CTkToplevel(parent_window or self)
            ventana.title("Editar Nota")
            ventana.geometry("500x600")
            ventana.transient(parent_window or self)
            ventana.grab_set()
            
            # Formulario similar al de nueva nota
            titulo_frame = ctk.CTkFrame(ventana)
            titulo_frame.pack(fill="x", padx=20, pady=20)
            
            ctk.CTkLabel(titulo_frame, text="📝 Editar Nota", 
                        font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
            
            form_frame = ctk.CTkScrollableFrame(ventana, height=400)
            form_frame.pack(fill="both", expand=True, padx=20, pady=10)
            
            # Título
            ctk.CTkLabel(form_frame, text="Título de la Nota *", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
            entry_titulo = ctk.CTkEntry(form_frame)
            entry_titulo.pack(fill="x", pady=(0,10))
            entry_titulo.insert(0, titulo_actual)
            
            # Tipo
            ctk.CTkLabel(form_frame, text="Tipo", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
            option_tipo = ctk.CTkOptionMenu(form_frame, values=["General", "Incidencia", "Comercial", "Técnica"])
            option_tipo.set(tipo_actual)
            option_tipo.pack(fill="x", pady=(0,10))
            
            # Contenido
            ctk.CTkLabel(form_frame, text="Contenido de la Nota *", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
            text_contenido = ctk.CTkTextbox(form_frame, height=200)
            text_contenido.pack(fill="x", pady=(0,10))
            text_contenido.insert("1.0", contenido_actual)
            
            # Usuario
            ctk.CTkLabel(form_frame, text="Usuario", 
                        font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,2))
            entry_usuario = ctk.CTkEntry(form_frame)
            entry_usuario.pack(fill="x", pady=(0,10))
            entry_usuario.insert(0, usuario_actual)
            
            # Privada
            check_privada = ctk.CTkCheckBox(form_frame, text="Nota privada")
            check_privada.pack(anchor="w", pady=10)
            if privada_actual:
                check_privada.select()
            
            # Botones
            botones_frame = ctk.CTkFrame(ventana)
            botones_frame.pack(fill="x", padx=20, pady=20)
            
            btn_cancelar = ctk.CTkButton(botones_frame, text="❌ Cancelar", 
                                       command=ventana.destroy, width=100)
            btn_cancelar.pack(side="right", padx=(10,0))
            
            def guardar_cambios():
                titulo = entry_titulo.get().strip()
                contenido = text_contenido.get("1.0", "end-1c").strip()
                
                if not titulo or not contenido:
                    messagebox.showerror("Error", "Título y contenido son obligatorios")
                    return
                
                try:
                    conn = connect_db()
                    cursor = conn.cursor()
                    cursor.execute("""
                        UPDATE notas_cliente 
                        SET titulo = ?, contenido = ?, tipo = ?, usuario = ?, privada = ?
                        WHERE nota_id = ?
                    """, (titulo, contenido, option_tipo.get(), 
                          entry_usuario.get().strip(), check_privada.get(), nota_id))
                    
                    conn.commit()
                    conn.close()
                    
                    messagebox.showinfo("Éxito", "Nota actualizada correctamente")
                    ventana.destroy()
                    
                except Exception as e:
                    messagebox.showerror("Error", f"Error al actualizar nota: {e}")
            
            btn_guardar = ctk.CTkButton(botones_frame, text="💾 Guardar Cambios", 
                                      command=guardar_cambios, width=150)
            btn_guardar.pack(side="right")
            
        except Exception as e:
            print(f"❌ Error editando nota: {e}")
            messagebox.showerror("Error", f"Error al editar nota: {e}")
    
    def obtener_estadisticas_cliente(self, cliente_id):
        """Obtiene las estadísticas completas de un cliente."""
        try:
            conn, cursor = self.master.conectar_db()
            if not conn: 
                return None
            
            cursor.execute("""
                SELECT * FROM estadisticas_cliente 
                WHERE cliente_id = ?
            """, (cliente_id,))
            
            stats = cursor.fetchone()
            
            # También obtener productos problemáticos
            cursor.execute("""
                SELECT * FROM productos_problematicos 
                WHERE cliente_id = ?
                ORDER BY total_devoluciones DESC
                LIMIT 10
            """, (cliente_id,))
            
            productos_problematicos = cursor.fetchall()
            
            conn.close()
            
            return {
                'estadisticas': stats,
                'productos_problematicos': productos_problematicos
            }
            
        except Exception as e:
            print(f"Error obteniendo estadísticas: {e}")
            return None

    def mostrar_eliminar_rma(self):
        """Ventana simple para eliminar RMA."""
        ventana = ctk.CTkToplevel(self)
        ventana.title("ELIMINAR RMA")
        ventana.geometry("400x300")
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        ctk.CTkLabel(ventana, text="ELIMINAR RMA", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=20)
        
        # Campo entrada
        ctk.CTkLabel(ventana, text="Código RMA:", font=ctk.CTkFont(size=14)).pack()
        entry = ctk.CTkEntry(ventana, width=200, height=30, font=ctk.CTkFont(size=14))
        entry.pack(pady=10)
        
        # Info
        info = ctk.CTkLabel(ventana, text="", font=ctk.CTkFont(size=12))
        info.pack(pady=10)
        
        # Buscar
        def buscar():
            codigo = entry.get().strip()
            if not codigo:
                info.configure(text="Escribe un código RMA", text_color="orange")
                btn_eliminar.configure(state="disabled")
                return
            
            conn, cursor = self.conectar_db()
            if conn:
                cursor.execute("SELECT cliente, fecha_emision FROM rma_maestro WHERE codigo_rma = ?", (codigo,))
                data = cursor.fetchone()
                conn.close()
                if data:
                    info.configure(text=f"Encontrado: {data[0]} - {data[1]}", text_color="green")
                    btn_eliminar.configure(state="normal")
                else:
                    info.configure(text="RMA no encontrado", text_color="red")
                    btn_eliminar.configure(state="disabled")
        
        ctk.CTkButton(ventana, text="BUSCAR", command=buscar, width=100, height=30).pack(pady=5)
        
        # Eliminar
        def eliminar():
            codigo = entry.get().strip()
            if messagebox.askyesno("Confirmar", f"¿Eliminar RMA {codigo}?"):
                try:
                    conn, cursor = self.conectar_db()
                    cursor.execute("SELECT id FROM rma_maestro WHERE codigo_rma = ?", (codigo,))
                    rma_id = cursor.fetchone()[0]
                    cursor.execute("DELETE FROM rma_historial WHERE rma_id = ?", (rma_id,))
                    cursor.execute("DELETE FROM rma_detalles WHERE rma_id = ?", (rma_id,))
                    cursor.execute("DELETE FROM rma_maestro WHERE id = ?", (rma_id,))
                    conn.commit()
                    conn.close()
                    messagebox.showinfo("OK", f"RMA {codigo} eliminado")
                    ventana.destroy()
                except Exception as e:
                    messagebox.showerror("Error", str(e))
        
        btn_eliminar = ctk.CTkButton(ventana, text="ELIMINAR", command=eliminar, 
                                    fg_color="red", width=150, height=40, state="disabled")
        btn_eliminar.pack(pady=20)
        
        ctk.CTkButton(ventana, text="CANCELAR", command=ventana.destroy, width=100).pack()
        
    def mostrar_generar_numero_manual(self):
        """Ventana simple para crear RMA manual."""
        ventana = ctk.CTkToplevel(self)
        ventana.title("CREAR RMA MANUAL")
        ventana.geometry("400x250")
        ventana.transient(self)
        ventana.grab_set()
        
        # Título
        ctk.CTkLabel(ventana, text="CREAR RMA MANUAL", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=20)
        
        # Campo entrada
        ctk.CTkLabel(ventana, text="Código RMA:", font=ctk.CTkFont(size=14)).pack()
        entry = ctk.CTkEntry(ventana, width=200, height=30, font=ctk.CTkFont(size=14))
        entry.pack(pady=10)
        
        # Info
        info = ctk.CTkLabel(ventana, text="", font=ctk.CTkFont(size=12))
        info.pack(pady=10)
        
        # Validar
        def validar():
            codigo = entry.get().strip().upper()
            if not codigo:
                info.configure(text="Escribe un código RMA", text_color="orange")
                btn_crear.configure(state="disabled")
                return
            
            if not codigo.startswith("RMA25"):
                info.configure(text="Debe empezar con RMA25", text_color="red")
                btn_crear.configure(state="disabled")
                return
            
            conn, cursor = self.conectar_db()
            if conn:
                cursor.execute("SELECT COUNT(*) FROM rma_maestro WHERE codigo_rma = ?", (codigo,))
                existe = cursor.fetchone()[0] > 0
                conn.close()
                if existe:
                    info.configure(text="Ya existe ese código", text_color="red")
                    btn_crear.configure(state="disabled")
                else:
                    info.configure(text="Código válido", text_color="green")
                    btn_crear.configure(state="normal")
        
        ctk.CTkButton(ventana, text="VALIDAR", command=validar, width=100, height=30).pack(pady=5)
        
        # Crear
        def crear():
            codigo = entry.get().strip().upper()
            if messagebox.askyesno("Confirmar", f"¿Crear RMA {codigo}?"):
                try:
                    conn, cursor = self.conectar_db()
                    fecha = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    cursor.execute("""
                        INSERT INTO rma_maestro (codigo_rma, cliente, fecha_emision, creado_por, estado, precio_total_expediente)
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (codigo, "[Manual]", fecha, self.username, "Creado", 0.0))
                    conn.commit()
                    conn.close()
                    messagebox.showinfo("OK", f"RMA {codigo} creado")
                    ventana.destroy()
                except Exception as e:
                    messagebox.showerror("Error", str(e))
        
        btn_crear = ctk.CTkButton(ventana, text="CREAR", command=crear, 
                                 fg_color="green", width=150, height=40, state="disabled")
        btn_crear.pack(pady=20)
        
        ctk.CTkButton(ventana, text="CANCELAR", command=ventana.destroy, width=100).pack()

# ----------------------------------------------------------------------
# 7. EJECUCIÓN DEL PROGRAMA
# ----------------------------------------------------------------------

if __name__ == "__main__":
    # Solo verificar base de datos local si no estamos usando Turso
    turso_url = os.getenv("TURSO_DATABASE_URL")
    turso_token = os.getenv("TURSO_AUTH_TOKEN")
    
    if not (turso_url and turso_token):
        # Solo validar archivo local si no hay configuración de Turso
        if not os.path.exists(DB_NAME):
            print("🚨 Error Crítico: No se encuentra el archivo de base de datos 'rma_app.db'.")
            print("Asegúrate de ejecutar primero 'python db_setup.py'.")
            sys.exit(1)
    else:
        print("🌩️ Usando base de datos Turso cloud")
    
    # Mostrar un splash/spinner de arranque y ejecutar optimize_database en background
    try:
        print("🔧 Iniciando: splash de arranque y optimización en segundo plano...")

        # Nota: no aplicamos aquí el tema del usuario para evitar recargas visuales.
        # Los ajustes de tema se mantienen desactivados por ahora porque los colores
        # están fijados en el theme y la aplicación provocaba una doble renderización.

        # Crear ventana splash simple (tkinter, sin bordes)
        splash = tk.Tk()
        splash.overrideredirect(True)
        splash.configure(bg="white")
        w, h = 420, 140
        sw = splash.winfo_screenwidth()
        sh = splash.winfo_screenheight()
        x = (sw - w) // 2
        y = (sh - h) // 2
        splash.geometry(f"{w}x{h}+{x}+{y}")

        frame = tk.Frame(splash, bg="white")
        frame.pack(fill="both", expand=True)

        label = tk.Label(frame, text="Iniciando Gestor RMA - Expedientes", font=("Segoe UI", 12, "bold"), bg="white")
        label.pack(pady=(20, 6))

        sub = tk.Label(frame, text="Preparando la aplicación...", font=("Segoe UI", 10), bg="white", wraplength=380, justify="left")
        sub.pack(pady=(0, 10))

        try:
            pb = ttk.Progressbar(frame, mode="indeterminate", length=340)
            pb.pack(pady=(0, 12))
            pb.start(12)
        except Exception:
            pb = None

        # Verificar tipo de BD
        import time
        if turso_url and turso_token:
            try:
                sub.config(text="🌩️ Usando base de datos Turso cloud")
                splash.update()
            except Exception:
                pass
            time.sleep(0.3)
        else:
            try:
                sub.config(text="💾 Usando base de datos local SQLite")
                splash.update()
            except Exception:
                pass
            time.sleep(0.3)
        
        # Verificar conexión a Dropbox
        try:
            sub.config(text="Verificando conexión a Dropbox...")
            splash.update()
        except Exception:
            pass
        
        dropbox_connected = False
        try:
            if DROPBOX_APP_KEY and DROPBOX_APP_SECRET:
                if DROPBOX_REFRESH_TOKEN:
                    dbx_test = dropbox.Dropbox(
                        oauth2_refresh_token=DROPBOX_REFRESH_TOKEN,
                        app_key=DROPBOX_APP_KEY,
                        app_secret=DROPBOX_APP_SECRET
                    )
                    dbx_test.users_get_current_account()
                    dropbox_connected = True
                elif DROPBOX_ACCESS_TOKEN and DROPBOX_ACCESS_TOKEN != "tu_access_token_aqui":
                    dbx_test = dropbox.Dropbox(DROPBOX_ACCESS_TOKEN)
                    dbx_test.users_get_current_account()
                    dropbox_connected = True
        except Exception:
            pass
        
        if dropbox_connected:
            try:
                sub.config(text="☁️ Dropbox conectado correctamente")
                splash.update()
            except Exception:
                pass
            time.sleep(0.3)
        else:
            try:
                sub.config(text="📁 Dropbox no disponible - Usando almacenamiento local")
                splash.update()
            except Exception:
                pass
            time.sleep(0.3)
        
        # Iniciar optimización en un hilo daemon
        t = threading.Thread(target=optimize_database, daemon=True)
        t.start()

        # Mostrar mensajes por etapas mientras la optimización corre
        try:
            while t.is_alive():
                # Mostrar etapa de optimización
                try:
                    sub.config(text="🔧 Optimizando la base de datos...")
                except Exception:
                    pass
                try:
                    splash.update()
                except Exception:
                    pass
                time.sleep(0.05)

            # Una vez terminado, indicar carga de la interfaz
            try:
                sub.config(text="✨ Cargando interfaz...")
                splash.update()
            except Exception:
                pass

            # Pequeña pausa para que el usuario vea el estado final
            time.sleep(0.25)

        except KeyboardInterrupt:
            # Permitir salir con Ctrl-C si se interrumpe
            pass

        # Cerramos el splash ANTES de crear la ventana principal para evitar conflictos
        try:
            if pb:
                pb.stop()
        except Exception:
            pass
        try:
            splash.destroy()
        except Exception:
            pass

        # Crear la aplicación principal (ahora que la optimización ha finalizado)
        app = LoginApp()
        app.mainloop()

    except Exception:
        # Si algo falla en el splash, fallback al comportamiento previo
        print("🔧 Optimizando base de datos en segundo plano...")
        t = threading.Thread(target=optimize_database, daemon=True)
        t.start()
        app = LoginApp()
        app.mainloop()